#
# from ffmpeg import Progress
# from ffmpeg.asyncio import FFmpeg
import threading
from email.mime.text import MIMEText
from email.mime.multipart import MIMEMultipart
from email.mime.application import MIMEApplication
from email.header import Header
import smtplib

import random
from MondeVert.SHAINE_MonderVert import Instagram_Posts as IP
import shutil
from docx import Document
import pandas as pd
import pyttsx3
import datetime
import re
from boto3 import Session
from botocore.exceptions import BotoCoreError, ClientError
from contextlib import closing
import os
import sys
from tempfile import gettempdir
from MondeVert.SHAINE_MonderVert.SHAINE_WIZARD_PROMPTS import Stories_For_Audio_Files as SAF
import MondeVert.SHAINE_MonderVert.SHAINE_WIZARD_PROMPTS.User_Prefs as up
from MondeVert.SHAINE_MonderVert.Utilities import DoNotCommit as DNC

from MondeVert.SHAINE_MonderVert import SHAINE as GPT

current_time1 = datetime.datetime.now()
current_time = current_time1.strftime('%m-%d-%Y_%H.%M')



def Check_Folder_Exists(SavePath):
    isExist = os.path.exists(SavePath)
    if not isExist:
        # Create a new directory because it does not exist
        os.makedirs(SavePath)
        print("Created new filePath: " + SavePath)






def speak(self, audio, Add2T=True, voice=''):
    if voice == '':
        voice = self.voice

    if Add2T == True:
        Add2Transcript(self, self.AssistantName + ': ' + audio)
    engine = pyttsx3.init()
    # getter method(gets the current value
    # of engine property)
    voices = engine.getProperty('voices')
    engine.setProperty('voice', voices[voice].id)
    # Method for the speaking of the assistant
    if self.SilentMode == False:
        engine.say(audio + "")
        # Blocks while processing all the currently
        # queued commands
        engine.runAndWait()
    else:
        print('Silent Mode Activated: ' + audio)




def Add2Transcript(self, text2Add):
    transcript_Final = ''
    transcript_Final += text2Add + ' \n'
    current_time1 = datetime.datetime.now()
    current_time2 = current_time1.strftime('%m-%d-%Y_%H.%M.%S')

    data = [(current_time2, text2Add)]
    print(data)
    try:
        df1 = pd.DataFrame(data, columns=["TimeStamp", "Transcript"])
        add2Master2(df1)
    except:
        print('Error could not add to transcript')

def add2Master2(df1):
     f2 = up.getMF2Path()
     df2 = pd.read_excel(f2)
     df3 = pd.concat([df1, df2])
     df3.iloc[:, 1:]
     df3.drop_duplicates()

     # creating a new excel file and save the data
     df3.to_excel(f2, index=False)







def SaveCSV(Text, Title, FolderPath_original,FolderPath_Final):
    current_time1 = datetime.datetime.now()
    current_time2 = current_time1.strftime('%m-%d-%Y_%H.%M.%S')
    f2 = FolderPath_original
    SavePath1 = f2
    invalidCharRemoved = re.sub(r"[^a-zA-Z0-9 ]", "", Title)
    folder = str(invalidCharRemoved)
    if len(folder) > 44:
        folder = folder[0:44]
    Title1 = '\\' + str(invalidCharRemoved) + "_"
    if len(Title1) > 44:
        Title1 = Title1[0:44]
    SavePath2 = SavePath1 + '\\' + folder
    Title2 = SavePath2 + Title1 + current_time2 + '.txt'


    Check_Folder_Exists(FolderPath_Final)



    try:
        df1 = pd.DataFrame(Text, columns=["TimeStamp", "Transcript"])
        print(df1)

        df1.to_csv(Title2 )
        # MondeVert.SaveText(self,df1,'MondeVert Assistant', 'Full Transcript')
        add2Master2(df1)
    except:
        print('Review Error File did not save ')

    return Title2


# ********************************
"""Getting Started Example for Python 2.7+/3.3+"""

# Create a client using the credentials and region defined in the [adminuser]
# section of the AWS credentials file (~/.aws/credentials).

from pydub import AudioSegment
def Combine_Splitsof_audio(AudioFiles_ordered,FilePath, SavePath = up.SavePath,Opening_Sound = up.MondeVertIntro, Voice = random.choices(SAF.Original_List_of_Voices_English)[0]):
    Check_Folder_Exists(SavePath)
    Merged_Audio = Opening_Sound

    counter = 1
    len_Audio = len(AudioFiles_ordered)
    for x in range (0,len_Audio):
        if x ==0:
            print(Opening_Sound)
            #Opening = AudioSegment.from_mp3(Opening_Sound)
            #Merged_Audio = Opening

        else:
            sound2add = AudioSegment.from_mp3(AudioFiles_ordered[x])
            Merged_Audio += sound2add

    try:
    # simple export
         FilePath = Merged_Audio.export(FilePath, format="mp3")
    except:
         print('could not save file audio file')
    #     FilePath = 'No File Saved'

    return FilePath

#on character list Have Narrator Scene (stern and manly) and then you have preferred Narrators to choose from (to do all other narrations)
#Split on Space and :
#also look for Narrator, for now we should mostly ignore whats in the parenthesis so I will work to remove these
#Setting and other parts are Narrator1, and then narrator2
#def SplitScriptbyPart(Script,Character_Voice_Relationship,FileName,SavePath )
#remove (voiceover), otherwise have the system remove anywhere that has the parenthesis as its not important for audio
#(in person) (on the phone)
#Grab the Left most item on the line
#Def Clean_Text_For_Script

#Remove - FADE OUT.
#(Impressed) dip in the voice idk
#(defeated)/(Nervous) make quieter -2/(relieved) quiet but happy

#(Excited).(Urgently)/(upset)  - speed up maybe --> Make it louder

#other parenthesis should be spoken by one of the narrators or just the one

# not sure what to do maybe add a pause or something like that(smiles)
#Replace with onomonapea
#(sighs)
#(Laughs)


#Remove the reduandant parts like out loud and voiceover etc. eventually I can key off this but dont need it
#if line above is NULL or " " then you look for : if no : then ignore the line completely
#if the line says Music: send to GPT to come up with a quick prompt to relate it to the story/scene
#Make sure to add results to a script so we can use it again (Call it on the fly Revisions)
#if line says Narrator, remove parenthesis and just talk for the most part
#If its not a scenario above it is a character so we identify the name prior to the : and after the line before that is NULL or " "
#We eventually try to map the name to a voice and then we use the same voice for the character (for now we can live with random and maybe assign it and remove it from the list)
#We will eventually have the narrator say the lines that are in parenthesis, we can have the words split up or something


def Split_Audio(Text, SavePath, FileName,FilePath = '', OpeningSound = up.MondeVertIntro, Chunk_Limit = 1500, Voice = random.choices(SAF.Original_List_of_Voices_English)[0] ):
    Length_Text = len(Text)
    Chunk_Limit = Chunk_Limit -20
    Text_Chunks = []
    Text_New = Text
    FilePaths = []
    length = len(Text)
    Audio_File_Count= 1
    SavePath_Chunks = SavePath + r'\Audio Chunks'
    Check_Folder_Exists(SavePath)
    Check_Folder_Exists(SavePath_Chunks)
    RunningCheck = 0
    LastPunc = 100
    while RunningCheck < (Length_Text-10) and LastPunc != -1:
        Text_Chunk = Text_New[:Chunk_Limit]
        Text_Chunk_Info_only = Text_Chunk.replace('.','!')
        Text_Chunk_Info_only = Text_Chunk_Info_only.replace('?', '!')

        #print(Text_Chunk_Info_only)
        # print('Last Punc Position')
        # print(LastPunc)
        LastPunc =Text_Chunk_Info_only.rfind("!")

        print('Last Punc Position')
        print(LastPunc)
        Text_Chunk_Final = Text_New[:LastPunc]
        #print(Text_Chunk_Final)
        FilePath = SaveText2Audio(Text = Text_Chunk_Final, SavePath = SavePath_Chunks, FileName=FileName + '_Chunk_' +str(Audio_File_Count), Voice = Voice )
        FilePaths.append(FilePath)
        Text_Chunks.append(Text_Chunk_Final)
        RunningCheck +=len(Text_Chunk_Final)
        Audio_File_Count +=1
        try:
            #length = len(Text_New)
            New_pos = (LastPunc + 1)
            Text_New = Text_New[New_pos:]
            #length = len(Text_New)
        except:
            print('Review Error Triggered, could be just how this has to work')
            #print(Text_New)
            #print(len(Text_New))
    try:
        FilePath2 = Combine_Splitsof_audio(AudioFiles_ordered = FilePaths, FilePath=FilePath,SavePath= SavePath)
    except:
        FilePath2 = 'Could not combine audio'

        print('Error COuld not parse files together, see chunks for now')
    return FilePath2

def CleanFileName(Text):
    Text = re.sub(r"[^a-zA-Z0-9 ]", "", Text)
    return Text
def SaveText2Audio(Text = SAF.Text, SavePath =up.SavePath, FileName="Text2Audio", FilePath = '', Neural='Neural', Mode='Text2Audio', Chunk_Limit = 1500,Voice = random.choices(SAF.Original_List_of_Voices_English)[0]):
    Voice = CleanFileName(Voice)

    Check_Folder_Exists(SavePath)

    FilePaths = ''
    FileName = FileName + '_' + Voice
    FilePath1 = 'No File Saved'
    FilePath = SavePath + '\\' + FileName  + '.mp3'
    print(len(str(Text)))

    #polly = Session(region_name='eu-west-2').client('polly')



    if len(str(Text)) >= Chunk_Limit:
        #try:
        FilePath = Split_Audio(Text = Text,SavePath = SavePath, FileName=FileName, FilePath=FilePath, Chunk_Limit= Chunk_Limit)
        # except:
        #     print('Error trying to create audio file try again later')

    else:
        try:
            session = Session(aws_access_key_id=DNC.aws_access_key_id, aws_secret_access_key=DNC.aws_secret_access_key,
                              region_name='us-west-2')
            polly = session.client("polly")
            response = polly.synthesize_speech(Text=Text, OutputFormat="mp3",
                                               VoiceId=Voice, Engine='neural', )

            # # Request speech synthesis #StartSpeechSynthesisTask #synthesize_speech
            # response = polly.start_speech_synthesis_task(Text=Text, OutputFormat="mp3",
            #                                    VoiceId=Voice, Engine='neural',
            #                                 OutputS3KeyPrefix= "InputAudio",
            #                                 OutputS3BucketName='s3bucketname',)
        except (BotoCoreError, ClientError) as error:
            # The service returned an error, exit gracefully
            print(error)
            sys.exit(-1)

        # Access the audio stream from the response
        if "AudioStream" in response:
            # Note: Closing the stream is important because the service throttles on the
            # number of parallel connections. Here we are using contextlib.closing to
            # ensure the close method of the stream object will be called automatically
            # at the end of the with statement's scope.
            with closing(response["AudioStream"]) as stream:
                output = os.path.join(gettempdir(), FilePath)


                try:
                    # Open a file for writing the output as a binary stream
                    with open(output, "wb") as file:
                        file.write(stream.read())
                        FilePath1 = FilePath
                except IOError as error:
                    # Could not write to file, exit gracefully
                    print(error)
                    sys.exit(-1)

        else:
            # The response didn't contain audio data, exit gracefully
            print("Could not stream audio as file is large, trying alternate ways")
           # Monitor_Audio_File(FileName, FilePath)
            sys.exit(-1)
        #
        # # Play the audio using the platform's default player
        # if sys.platform == "win32":
        #     os.startfile(output)
        # else:
        #     # The following works on macOS and Linux. (Darwin = mac, xdg-open = linux).
        #     opener = "open" if sys.platform == "darwin" else "xdg-open"
        #     subprocess.call([opener, output])
        #


    return FilePath


#
#
# import boto3
# def Monitor_Audio_File(FileName,FilePath):
#
#     session = Session(aws_access_key_id=DNC.aws_access_key_id, aws_secret_access_key=DNC.aws_secret_access_key,
#                       region_name='us-west-2')
#     polly = session.client("polly")
#     s3 = boto3.client('s3', aws_access_key_id=DNC.aws_access_key_id, aws_secret_access_key=DNC.aws_secret_access_key)
#     #somhow download whatever this is lol
#     while FileDone ==False:
#         polly.get_speech_synthesis_task()
#         try:
#             s3.get_object(
#                 #Bucket=self._bucket,
#                 Key=FileName,
#             )
#             FileDone =  True
#         except s3.exceptions.NoSuchKey:
#             FileDone =  False
#             print('File not there, waiting 15 seconds')
#             time.sleep(15)
#     polly.download_file('your_bucket',FileName,FilePath)



def add2Master(df1):
    f2 = up.getMFPath()
    df2 = pd.read_excel(f2)
    df3 = pd.concat([df1, df2])
    df3.drop_duplicates()
    df3.iloc[:, 1:]
    # creating a new excel file and save the data
    df3.to_excel(f2, index=False)

def add2Master2(df1):
    f2 = up.getMF2Path()
    df2 = pd.read_excel(f2)
    df3 = pd.concat([df1, df2])
    df3.iloc[:, 1:]
    df3.drop_duplicates()

    # creating a new excel file and save the data
    df3.to_excel(f2, index=False)

def add2Master3(df1):
    f2 = up.MasterFile3
    df2 = pd.read_excel(f2)
    df3 = pd.concat([df1, df2])
    df3.iloc[:, 1:]
    df3.drop_duplicates()

    # creating a new excel file and save the data
    df3.to_excel(f2, index=False)



def SaveText( df, FileName, tabname):

    text1 = df
    f2 = up.getPath()
    SavePath1 = f2
    invalidCharRemoved = re.sub(r"[^a-zA-Z0-9 ]", "", FileName)
    Filename = '\\' + str(invalidCharRemoved) + "_"
    SavePath2 = SavePath1 + Filename + ".xlsx"
    # SavePath2 = r'D:\ShakeBot Testing\ShaKeBotTest for DA - 12-20-2022.xlsx'
    try:
        with pd.ExcelWriter(SavePath2) as writer:
            text1.to_excel(writer, sheet_name=str(tabname))
            #MondeVert.speak(self, 'File Saved')
            print('File Saved')
    except:
        Exception




def saveTranscript(Transcript,current_time):


    Filename = 'SHAINE Transcript v1'
    Filename = '\\' + Filename + "_"
    f2 = up.getPath()
    SavePath1 = f2
    data = [(current_time, Transcript)]
    print(data)
    try:
        df1 = pd.DataFrame(data, columns=["TimeStamp", "Transcript"])
        # df1 = pd.DataFrame(data)
        df1.to_csv(SavePath1 + Filename + current_time + '.csv')

        # MondeVert.SaveText(self,df1,'MondeVert Assistant', 'Full Transcript')
        add2Master2(df1)
    except:
        print('Review Error')


def Add2MasterLyrics(self, current_time2, Mode, title, ArtistPoetInfo, poem, Tag, SavePath2, Prompts):
    data = [(current_time2, Mode, title, ArtistPoetInfo, poem, Tag, SavePath2, Prompts)]
    df = pd.DataFrame(data, columns=['Date_Time_Added', 'Art_Type', 'Title', 'Poet_Artist_Info', 'Poem_Song_Lyrics',
                                     'Quality', 'Folder_Path', 'Prompts_Used'])
    add2Master3(df)





def SaveCSV(Text, Title, SavePath, current_time = current_time):
    Title2 = ''
    try:
        Title1 = CleanFileName(Title)

        Title1 = '\\' + Title1 + "_"
        if len(Title1) > 44:
            Title1 = Title1[0:44]

        Title2 = SavePath + Title1  + '.txt'

        Check_Folder_Exists(SavePath)
        data1 = [(Text)]
        data = [(current_time, Text)]
        #print(data)
        df = pd.DataFrame(data1)

        df1 = pd.DataFrame(data, columns=["TimeStamp", "Transcript"])
        #print(df1)

        df.to_csv(Title2 )
        # MondeVert.SaveText(self,df1,'MondeVert Assistant', 'Full Transcript')
        add2Master2(df1)
        print('File Saved in the following location: ' + Title2)
    except:
        print('Review Error File did not save ')

    return Title2



def NamePoemSavePoem(self, poem, ArtPaths, Prompts_Used, ArtistPoetInfo, title='', SavePath=up.AI_Poetry_Path,
                     Mode='SHAINE', current_time = current_time ):


    try:
        dfPrompts = ''
        Tag = ''
        if Mode == 'Poem':
            title_p = up.Poem_Title_Prompt
        else:
            title_p = up.Song_Title_Prompt
        try:
            if title == '':
                x = GPT.MondeVert()
                Title = x.Get_Title_GPT(  Text= poem[:1500] )
            else:
                Title = title

        except:
            print('error could not get title')

        if Title == '':
            Title = 'AI - No Name Project'

        Title1 = CleanFileName(Title)

        Title1 = '\\' + str(Title1) + "_"
        if len(Title1) > 60:
            Title1 = Title1[0:60]


        Title2 = SavePath + Title1

        Check_Folder_Exists(SavePath)





        data = Title +  ArtistPoetInfo + poem +  SavePath+  Mode +  Tag
        # print(dat
        data = [(current_time, data)]
        print(data)
        try:
            df1 = pd.DataFrame(data, columns=["TimeStamp", "Transcript"])
            print(df1)
            print(Title2)
            df1.to_csv(Title2 + '.txt')
            Add2MasterLyrics(self, current_time, Mode, title, ArtistPoetInfo, poem, Tag, SavePath,dfPrompts)
            # MondeVert.SaveText(self,df1,'MondeVert Assistant', 'Full Transcript')
            try:
                add2Master2(df1)

            except:
                print('Did not add to master, error')
        except:
            print('Review Error')

            try:
    #Word Doc Creation below, not quite what I need so I do text file too above
                document = Document()
                document.add_heading(Title, 0)
                document.add_heading(ArtistPoetInfo, 1)
                p = document.add_paragraph()
                r = p.add_run()
                # r.add_text(ArtistPoetInfo)
                p = document.add_paragraph()
                r = p.add_run()
                r.add_text(poem)
                p = document.add_paragraph()
                r = p.add_run()
                image_counter = 0
                for i in ArtPaths:
                    image_counter += 1
                    r.add_picture(i)
                    original = i
                    target = SavePath + Title1 + str(image_counter)  + '.png'
                    shutil.copyfile(original, target)

                    try:
                        IP.png2JPG(original, up.PNGPath, Title1, up.PNGPath_Archive, Del=False)

                    except:
                        print('instagram failed, as always!')
                document.add_heading('AI Prompts used: ', 7)

                for i in Prompts_Used:
                    p = document.add_paragraph()
                    r = p.add_run()
                    r.add_text(i)
                    dfPrompts += i + '|'
                document.save(Title2 + '_Details.docx')

                document2 = Document()
                document2.add_heading(Title, 0)
                document2.add_heading(ArtistPoetInfo, 1)

                p = document2.add_paragraph()
                r = p.add_run()
                r.add_text(poem)
                document2.save(Title2 + '.docx')
            except:
                print('Error saving word doc')
    except:
        print('Error did not save file')

    def MultiThread(self, Functions, Args):
        thread_list = []
        print("Start")
        print(Args)

        t = threading.Thread(target=Functions, args=(Args))
        thread_list.append(t)

        # Starts threads
        for thread in thread_list:
            print("New Thread Started")
            thread.start()


        # This blocks the calling thread until the thread whose join() method is called is terminated.
        # From http://docs.python.org/2/library/threading.html#thread-objects
        for thread in thread_list:
            thread.join()

    # Demonstrates that the main process waited for threads to complete
    print("Done")


def send_email_w_attachment_outlook(to, subject, body, filename):
    sender = DNC.OutlookEmail
    receivers = to
    #smtp
    smtpHost = 'smtp.office365.com'
    smtpPort = 587
    password = DNC.Outlookpassword
    subject = "outlook email test"
    # Add the From: and To: headers at the start!
    message = ("From: %s\r\nTo: %s\r\nSubject: %s\r\n\r\n"
           % (sender, ", ".join(receivers), subject))
    message += """This is a test e-mail message."""
    print (message)
    try:
        smtpObj = smtplib.SMTP(smtpHost, smtpPort)
        #smtpObj.set_debuglevel(1)
        smtpObj.ehlo()
        smtpObj.starttls()
        smtpObj.ehlo()
        smtpObj.login(sender,password)
        smtpObj.sendmail(sender, receivers, message)
        smtpObj.quit()
        print ("Successfully sent email")
    except:
        print ("Error: unable to send email")


def send_email_w_attachment_gmail(to, subject, body, filename):
    # create message object
    to = 'sdonovan@mondevert.co; RichardDX44@gmail.com'
    subject = 'New Song for Rich'
    message = MIMEMultipart()
    # add in header
    message['From'] = Header(up.user)
    message['To'] = Header(to)
    message['Subject'] = Header(subject)
    # attach message body as MIMEText
    message.attach(MIMEText(body, 'plain', 'utf-8'))

    # locate and attach desired attachments

    for n in filename:
        att_name = os.path.basename(filename[n])
        _f = open(filename[n], 'rb')
        att = MIMEApplication(_f.read(), _subtype="txt")
        _f.close()
        att.add_header('Content-Disposition', 'attachment', filename=att_name)
        message.attach(att)

    # setup email server
    server = smtplib.SMTP_SSL(DNC.host, DNC.port)
    server.login(DNC.gmail_user, DNC.gmail_pass)

    # send email and quit server
    server.sendmail(DNC.user, to, message.as_string())
    server.quit()


def send_email_no_attachment_gmail(to, subject, body):
    # create message object
    to = 'sdonovan@mondevert.co'
    subject = 'New Song for Rich'

    message = MIMEMultipart()

    # add in header
    message['From'] = Header(up.user)
    message['To'] = Header(to)
    message['Subject'] = Header(subject)

    # attach message body as MIMEText
    message.attach(MIMEText(body, 'plain', 'utf-8'))

    # setup email server
    server = smtplib.SMTP_SSL(DNC.host, DNC.port)
    server.login(DNC.user, DNC.gmail_pass)

    # send email and quit server
    server.sendmail(DNC.user, to, message.as_string())
    server.quit()
