import glob
import Stories_For_Audio_Files as SAF
import datetime
import pandas as pd
import wikipedia
global Record
import sys
import random
from secrets import randbelow
import Instagram_Posts as IP
import Long_User_Prompts as lup
import ShakesBot as s
import Common_Sayings as cs
from threading import Event
import User_Prefs as up
from gingerit.gingerit import GingerIt
import threading
import numpy
from docx import Document
import re
# from exceptions import PendingDeprecationWarning
import requests
import cloudscraper
import tensorflow
import selenium
Record = ''
#import urllib2
import webbrowser
from fpdf import FPDF
import User_Prefs as up
import Common_Utilities as cu
import os
import platform
import subprocess
import requests
import random
import string
import openai
import pyttsx3
# Import the speech recognition library
import speech_recognition as sr
from dotenv import load_dotenv
from DoNotCommit import API_Key
import DoNotCommit as DNC
import atexit
from docx import Document
from docx.shared import Inches

import threading
import shutil

import smtplib
import os
from email.mime.text import MIMEText
from email.mime.multipart import MIMEMultipart
from email.mime.application import MIMEApplication
from email.header import Header

# from email_config import gmail_pass, user, host, port


URL = "..."  # noqa
API_KEY = "..."


# make a transcript class that takes down all of my requests and saves to a text file
# have specific command for the assistant to save in a certain folder (always add date & time to the file
# create things like to do lists
# ideas
# stories and dialogue
# the key is being able to store it properly and sytemically access as well
# every program I create should be executeable through the assistant


# Urgent to do
# Make it a class
# have transcribe and also proofread (later) add method where in the middle of dictating you can add titles so it gets properly makred eventually
# bullets etc. parenthesis
# Have it so every pause puts a new line in between so its easy to pull out pauses etc.
# Most important is to have a mode where it callse poembot and reads the lines to me.
# .... also I would like for it to be a mini drama where I get multiple voices and characters that read
# the output etc eventually learn how to animate/make a nice script.
# create a series so eventually we might have a way to feed in characters to poem bot so its not just random like it is now lol

# 12/18/2022 get the assistant able to open the respective web browsers
# set up mouse so it has shortcuts
# be able to run on both mac and also the pc
# import the site view bot
# update the youtube bot so it hides ips properly and also so it loops through videos to not draw attention
#    (make it so it starts off slow then has pockets of 100 here and there big days occasionally and find out how to make it work properly
# have assistant able to run those on command (buy wifi adapter so I can have it running whenever & get new monitor for that too


# Update transcribe commands (Review file so far,#Have a separate function to make the title (different than subject)
##Also I should be able to mark the file as important based on results (add tag)
# Make a program that goes through and corrects the script into a new tab (puts subject as start of the word blocks (with other data in parenthesis for tagging)
# Also make it so that the poems do not overwrite but get concatenated together
# Make a theatre method that uses multithreading to do stuff
# overall threads should allow this to run faster
# Make it so I can update all meta data and also add more metadata
# Goal, mark as header (prior or next maybe)




# def MultiThread2(Args):
#     command_array = Args
#     number_of_commands = len(command_array)
#     def run_the_command(index):
#         exec (command_array[index])
#
#     threads = []
#     for i in range(number_of_commands):
#         try:
#             t = threading.Thread(target=run_the_command, args=(i,))
#             t.start()
#             threads.append(t)
#
#         except:
#             print('Error unable to pull the value from the ARGs parameter')
#

def MultiThread(self, Functions, Args):
    thread_list = []
    print("Start")
    print(Args)

    t = threading.Thread(target=Functions, args=(Args))
    thread_list.append(t)

    # Starts threads
    for thread in thread_list:
        print("New Thread Started")
        thread.start()


    # This blocks the calling thread until the thread whose join() method is called is terminated.
    # From http://docs.python.org/2/library/threading.html#thread-objects
    for thread in thread_list:
        thread.join()

    # Demonstrates that the main process waited for threads to complete
    print("Done")


def send_email_w_attachment_outlook(to, subject, body, filename):
    sender = DNC.OutlookEmail
    receivers = to
    #smtp
    smtpHost = 'smtp.office365.com'
    smtpPort = 587
    password = "youremailpassword"
    subject = "outlook email test"
    # Add the From: and To: headers at the start!
    message = ("From: %s\r\nTo: %s\r\nSubject: %s\r\n\r\n"
           % (sender, ", ".join(receivers), subject))
    message += """This is a test e-mail message."""
    print (message)
    try:
        smtpObj = smtplib.SMTP(smtpHost, smtpPort)
        #smtpObj.set_debuglevel(1)
        smtpObj.ehlo()
        smtpObj.starttls()
        smtpObj.ehlo()
        smtpObj.login(sender,password)
        smtpObj.sendmail(sender, receivers, message)
        smtpObj.quit()
        print ("Successfully sent email")
    except:
        print ("Error: unable to send email")


def send_email_w_attachment_gmail(to, subject, body, filename):
    # create message object
    to = 'sdonovan@mondevert.co; RichardDX44@gmail.com'
    subject = 'New Song for Rich'
    message = MIMEMultipart()
    # add in header
    message['From'] = Header(up.user)
    message['To'] = Header(to)
    message['Subject'] = Header(subject)
    # attach message body as MIMEText
    message.attach(MIMEText(body, 'plain', 'utf-8'))

    # locate and attach desired attachments

    for n in filename:
        att_name = os.path.basename(filename[n])
        _f = open(filename[n], 'rb')
        att = MIMEApplication(_f.read(), _subtype="txt")
        _f.close()
        att.add_header('Content-Disposition', 'attachment', filename=att_name)
        message.attach(att)

    # setup email server
    server = smtplib.SMTP_SSL(DNC.host, DNC.port)
    server.login(DNC.gmail_user, DNC.gmail_pass)

    # send email and quit server
    server.sendmail(DNC.user, to, message.as_string())
    server.quit()


def send_email_no_attachment_gmail(to, subject, body):
    # create message object
    to = 'sdonovan@mondevert.com'
    subject = 'New Song for Rich'

    message = MIMEMultipart()

    # add in header
    message['From'] = Header(up.user)
    message['To'] = Header(to)
    message['Subject'] = Header(subject)

    # attach message body as MIMEText
    message.attach(MIMEText(body, 'plain', 'utf-8'))

    # setup email server
    server = smtplib.SMTP_SSL(DNC.host, DNC.port)
    server.login(DNC.user, DNC.gmail_pass)

    # send email and quit server
    server.sendmail(DNC.user, to, message.as_string())
    server.quit()


class MondeVert():
    def __init__(self, voice=4, language_settings=1):
        self.voice = voice
        self.language_settings = language_settings
        self.AssistantName = up.getAssistantName()
        self.UserName = up.getUserName()
        self.transcript_Final = ''
        self.API_Key = API_Key
        self.Correction_Comment = ''
        self.AI_Corrected_Content = ''
        self.SilentMode = False
        self.Mode = 'SHAINE booted up'
        global xVoice
        global voice_set
        voice_set = self.voice
        xVoice = 1

    def Open_Web(url):
        url = url
        webbrowser.open_new_tab(url)

    def speak(self, audio, Add2T=True, voice=''):
        if voice == '':
            voice = self.voice

        if Add2T == True:
            MondeVert.Add2Transcript(self, self.AssistantName + ': ' + audio)
        engine = pyttsx3.init()
        # getter method(gets the current value
        # of engine property)
        voices = engine.getProperty('voices')
        engine.setProperty('voice', voices[voice].id)
        # Method for the speaking of the assistant
        if self.SilentMode == False:
            engine.say(audio + "")
            # Blocks while processing all the currently
            # queued commands
            engine.runAndWait()
        else:
            print('Silent Mode Activated: ' + audio)

    def speakSweet(self, audio, Add2T=True):
        if Add2T == True:
            MondeVert.Add2Transcript(self, self.AssistantName + ': ' + audio)
        engine = pyttsx3.init()
        # getter method(gets the current value
        # of engine property)
        voices = engine.getProperty('voices')
        engine.setProperty('voice', voices[4].id)
        # Method for the speaking of the assistant
        if self.SilentMode == False:
            engine.say(audio + "")
            # Blocks while processing all the currently
            # queued commands
            engine.runAndWait()
        else:
            print('Silent Mode Activated: ' + audio)

        # def RandomCharacters

    def testvoices(self):
        # import pyttsx3
        engine = pyttsx3.init()
        voices = engine.getProperty('voices')
        index = 0
        for voice in voices:
            print(f'index-> {index} -- {voice.name}')
            engine.setProperty('voice', voices[index].id)
            engine.say("Shane do you like this voice more than the last one you busy little bee turkish" + "")

            index += 1
        engine.runAndWait()

    def Setvoices(self, Quick=False):

        if Quick == False:
            engine = pyttsx3.init()
            voices = engine.getProperty('voices')
            index = 0
            indexCount = 0
            # load_dotenv()
            # Set up the OpenAI API client :
            openai.api_key = API_Key

            prompt = 'Write a short paragraph (less than 44 words) that uses different pronunciations and complex words while being warm and welcoming to a person named Shane'
            completions = openai.Completion.create(
                engine="text-davinci-003",
                prompt=prompt,
                max_tokens=2500,
                n=1,
                stop=None,
                temperature=0.5,
            )
            speak1 = 'Do you want to set this voice as active?'

            # Print the generated text
            message = completions.choices[0].text
            print(message)
            MondeVert.Add2Transcript(self, text2Add=message)
            set1 = False

            for voice in voices:
                indexCount += 1

            while set1 == False:
                if index <= indexCount:
                    print(f'index-> {index} -- {voices[index].name}')
                    engine.setProperty('voice', voices[index].id)
                    engine.say(message)
                    #MondeVert.speak(self, speak1)
                    Query = MondeVert.getUserResponse(self)

                    if "yes" in Query or "set" in Query or ("make" in Query and "active" in Query) or (
                            "set" in Query and "active" in Query):
                        self.voice = index
                        voice_set = self.voice
                        xVoice = index
                        #MondeVert.speak(self, 'Voice Set')
                        set1 = True

                    index += 1
                    engine.runAndWait()
                else:
                    MondeVert.Setvoices(self, Quick=True)

        else:
            self.voice = int(input('What Voice do you want to make active?'))
            #MondeVert.speak(self, 'Voice Set')

        ####################################################################################################################################################################################
        ## Below are more utlities that help me to do repetitive tasks quicker
        ####################################################################################################################################################################################

    def createDF(self):
        test = [('', '', '', '', '', '', '', '', '', '', '', '', '', '', '')]
        self.transcript = pd.DataFrame(test, columns=['Entry #', 'Date', 'Subject', 'Type/Marker', 'Priority',
                                                      'Original_Text', 'Edited_Text', 'AI_Corrected_Text',
                                                      'AI_Correction_Comment', 'Note', 'AI_Content',
                                                      'AI_Correction_Content', 'File_Name', 'Added to text File',
                                                      'Completed/Added to blog'])

    def adddata2DF(self):
        current_time1 = datetime.datetime.now()
        xx = current_time1.strftime('%m/%d/%Y %H:%M:%S')
        row = (self.entry - 1)
        self.transcript.loc[row] = [str(self.entry), xx, self.subject, self.Type, self.Significance, self.query1,
                                    self.new_Query, self.AI_Corrected_Text, self.Correction_Comment, self.Notes,
                                    self.Words, self.AI_Corrected_Content, self.FileName, self.Added2TextFile,
                                    self.Completed]

    def ActivateSilentMode(self):
        self.SilentMode = True

    def ActivateLoudMode(self):
        self.SilentMode = False

    def saveTranscript(self):

        current_time1 = datetime.datetime.now()
        current_time2 = current_time1.strftime('%m-%d-%Y_%H.%M.%S')
        Filename = 'SHAINE Transcript v1'
        Filename = '\\' + Filename + "_"
        f2 = up.getPath()
        SavePath1 = f2
        data = [(current_time2, self.transcript_Final)]
        print(data)
        try:
            df1 = pd.DataFrame(data, columns=["TimeStamp", "Transcript"])
            # df1 = pd.DataFrame(data)
            df1.to_csv(SavePath1 + Filename + current_time2 + '.csv')

            # MondeVert.SaveText(self,df1,'MondeVert Assistant', 'Full Transcript')
            MondeVert.add2Master2(df1)
        except:
            print('Review Error')

    def Add2Transcript(self, text2Add):
        self.transcript_Final += text2Add + ' \n'

    def SaveText(self, df, FileName, tabname):
        current_time1 = datetime.datetime.now()
        current_time2 = current_time1.strftime('%m-%d-%Y_%H.%M.%S')
        text1 = df
        f2 = up.getPath()
        SavePath1 = f2
        invalidCharRemoved = re.sub(r"[^a-zA-Z0-9 ]", "", FileName)
        Filename = '\\' + str(invalidCharRemoved) + "_"
        SavePath2 = SavePath1 + Filename + current_time2 + ".xlsx"
        # SavePath2 = r'D:\ShakeBot Testing\ShaKeBotTest for DA - 12-20-2022.xlsx'
        try:
            with pd.ExcelWriter(SavePath2) as writer:
                text1.to_excel(writer, sheet_name=str(tabname))
                #MondeVert.speak(self, 'File Saved')
                print('File Saved')
        except:
            Exception
        # MondeVert.speak(self,'Error Saving')
        # print ('Error Saving')


    #2 options for better transcribe tool 1). Make my own where it transcribes the whole thing (this is probably best), I can have a multithread to print or output either way the live text so I can see what I am saying and it will make a rough correction on the fly.
    #Have the ability to kick off certain processes like sending emails, scheduling meetings. Post instagram (auto or manual (basic or advance), Have the ability to run the entire tool by voice activation, further you can decide if its fully on the fly by talking to AI or if you want to direct it to certain inputs and pick a file for another. Outline mode, explain mode etc.)
    #if we can run everything, keep recording and outputting the Text to speech I think we are in a sweet spot for the actual AI. There should be a noice supression piece eventually to remove ambient noise.
    # you should be able to train model so it recognize your voice over other sounds (this tech exists 100% see siri alexa teams somehow etc.)
    #Ask walt about how to get better sounding voice models for text to speech?? Look up model how to train your voice to be used

    def cleanText(self, text, SHAINE_Clean = False):
        # result = ''
        parser = GingerIt()
        try:
            if int(len(text)) <= 4999:
                # print(len(text))
                result = pd.DataFrame(parser.parse(text))

                # result.drop_duplicates()
            else:
                result = pd.DataFrame('', '', '')
                self.Notes = 'Text is Large so may need to review the parts that may have been cut by mistake'
                for i in range(0, (len(text) // 5000) + 1):
                    texttemp = text[0 + (5000 * i):4999 + (5000 * i)]
                    if len(texttemp) > 0:
                        result11 = pd.DataFrame(parser.parse(text))
                        result1 = [result1 + texttemp]
                result = pd.DataFrame(result1, columns='result')

            if len(result.loc[0, 'result']) > 0:
                # print(result.loc[:,'result'].values)
                rr = str(result.loc[0, 'result'])
                self.Correction_Comment += 'Corrections Made.'
            else:
                rr = text
                self.Correction_Comment += 'No Corrections Needed.'



        except:
            print('Text was not properly cleaned by GingerIT: ' + text)
            # print(text)
            rr = text
            self.Correction_Comment += 'No Corrections Needed (Caused an error and did not return anything).'

        # print('Sending Text to SHAINE to clean it up ' + text)
        # if(SHAINE_Clean ==True):
        #     try:
        #         if int(len(text)) <= 4999:
        #             # print(len(text))
        #             result = pd.DataFrame(parser.parse(text))
        #             # result.drop_duplicates()
        #         else:
        #             result = pd.DataFrame('', '', '')
        #             self.Notes = 'Text is Large so may need to review the parts that may have been cut by mistake'
        #             for i in range(0, (len(text) // 5000) + 1):
        #                 texttemp = text[0 + (5000 * i):4999 + (5000 * i)]
        #                 if len(texttemp) > 0:
        #                     print("Current SubSet of Text being cleaned by SHAINE: " + texttemp)
        #                     #This is where I use
        #                     #result11 = pd.DataFrame(parser.parse(text))
        #                     result1 = [result1 + texttemp]
        #             result = pd.DataFrame(result1, columns='result')
        #     except:
        #
        #         return rr

        return rr

    def makeQuickPoem(self):
        x = s.PoemBot(1, 1, 1, 1, 40, 3)
        x.ReloadModel("model.h5")
        x.setupdata()
        self.Words = x.shakesbot_DA_Make_Script(size=300)
        # self.Words[0:100]
        print(self.Words)
        self.Words = MondeVert.cleanText(self, self.Words)
        MondeVert.speakSweet(self, self.Words)

    def StartThread(self):
        trd1 = threading.Thread(target=MondeVert.Make_Script2(self))
        self.w[i] = trd1.start()
        x = MondeVert.Make_Script2(self)
        print(x)

        self.threads.append(trd1)

    def Make_Script(self):
        size = 1
        self.threads = []
        self.w = numpy.empty(size, dtype=str)
        for self.i in range(0, size):
            MondeVert.StartThread(self)
            print(self.w[i])
            # globals()[f"trd{i}"] = threading.Thread(target=MondeVert.Make_Script2(self))
            # w[i] = globals()[f"trd{i}"].start()

        for x in self.threads:
            x.join()

        for ii in range(0, size):
            self.Words += str(self.w[ii])

        self.AI_Corrected_Content = MondeVert.cleanText(self, self.Words)

    def Make_Script2(self):
        pb = s.PoemBot(1, 1, 1, 1, 40, 3)
        pb.ReloadModel('model.h5')
        pb.setupdata()
        w1 = pb.shakesbot_DA_Make_Script()
        MondeVert.speakSweet(self, w1)
        # print (w1)
        return w1

    def generateGreek(self):
        w2 = self.pb.generateGreek() + ' \n' + ' \n'
        return w2

    def add2Master(df1):
        f2 = up.getMFPath()
        df2 = pd.read_excel(f2)
        df3 = pd.concat([df1, df2])
        df3.drop_duplicates()
        df3.iloc[:, 1:]
        # creating a new excel file and save the data
        df3.to_excel(f2, index=False)

    def add2Master2(df1):
        f2 = up.getMF2Path()
        df2 = pd.read_excel(f2)
        df3 = pd.concat([df1, df2])
        df3.iloc[:, 1:]
        df3.drop_duplicates()

        # creating a new excel file and save the data
        df3.to_excel(f2, index=False)

    def add2Master3(df1):
        f2 = up.MasterFile3
        df2 = pd.read_excel(f2)
        df3 = pd.concat([df1, df2])
        df3.iloc[:, 1:]
        df3.drop_duplicates()

        # creating a new excel file and save the data
        df3.to_excel(f2, index=False)

    def ChatGPTDA(self, temp=0.5, MakeArt=False, Prompt='', UseArtPrompt=False, ConfirmBot=True):
        MondeVert.Add2Transcript(self, ' \n')
        Prior_Mode = self.Mode
        self.Mode = 'Chat GPT Mode'
        self.message = ''
        MondeVert.Add2Transcript(self, text2Add=(Prior_Mode + ' - ' + self.Mode + ':'))
        MondeVert.Add2Transcript(self, ' \n')
        prompt = ''
        self.promptB = False
        # Import the OpenAI library
        # Set up the OpenAI API client :
        openai.api_key = API_Key
        # Record the audio
        #MondeVert.speak(self, " Chat GPT is running!", voice=4)

        if UseArtPrompt == True:
            MakeArtPrompt = "Provide me with a prompt to share with artificial inteligence that will create a unique and visually pleasing work of art. "
        else:
            MakeArtPrompt = ''

        if Prompt == '':
            sStop1 = False
        else:
            sStop1 = True

        sStop = False
        # check if the reply contains 'yes'
        while sStop == False:
            # Record the audio

            MondeVert.Add2Transcript(self, MakeArtPrompt)

            # engine.runAndWait()
            if sStop1 == True:
                prompt = MakeArtPrompt + Prompt
                print(prompt)
            else:
                prompt = MakeArtPrompt + MondeVert.getUserResponse(self, pause=1.44)
                if ('STOP' in prompt.upper() and (
                        'LISTEN' in prompt.upper())) or 'QUIT' == prompt.upper() or ' QUIT ' == prompt.upper() or (
                        'END' in prompt.upper() and ('CONVERSA' in prompt.upper())):
                    sStop = True
                    # print ('Its supposed to quit here')
                    continue

            # Generate text
            completions = openai.Completion.create(
                engine="text-davinci-003",
                prompt=prompt,
                max_tokens=2444,
                n=1,
                stop=None,
                temperature=temp,
            )
            self.message = completions.choices[0].text
            print(self.message)

            if MakeArt == False and (
                    'MAKE' in prompt.upper() or 'ART' in prompt.upper()) and 'PROMPT' in prompt.upper() and sStop1 == False:
                MondeVert.ConfirmBOT(self, self.message)
                if self.promptB == True:
                    MondeVert.makeArt(self, self.message)
                    continue
                else:
                    continue

            if MakeArt == True or sStop1 == True:
                if ConfirmBot == True:
                    MondeVert.ConfirmBOT(self, self.message)
                else:
                    MondeVert.speakSweet(self, self.message)
                    self.promptB = True

                if self.promptB == True:
                    sStop = True
                    continue
                else:

                    continue
            else:
                MondeVert.speakSweet(self, self.message)

        return self.message

    ############################################################################################################################################################################################################################
    ############################################################################################################################################################################################################################
    ################Above is code I did not write, below is code that allows user to update values of respective values
    ############################################################################################################################################################################################################################
    ############################################################################################################################################################################################################################
    def getfilename(self):
        self.FileName = MondeVert.getdata(self, 'FileName')

    def getsubject(self):
        self.subject = MondeVert.getdata(self, 'subject')

    def getDictation_Type(self):
        self.Dictation_Type = MondeVert.getdata(self, 'Type')

    def getSignificance(self):
        self.Significance = MondeVert.getdata(self, 'Significance')

    def getNote(self):
        self.Notes += '|' + MondeVert.getdata(self, 'Notes')

    def getdata(self, Field):

        #MondeVert.speak(self, 'what is the ' + Field)
        print('Please confirm the ' + Field)
        s = True
        Query = ''
        while (s == True):
            Query = MondeVert.getUserResponse(self)
            Query2 = ''
            WaitforResponse = False
            while (WaitforResponse == False):
                #MondeVert.speak(self, Query + ' Is that correct?')
                Query2 = MondeVert.getUserResponse(self)
                if "yes" in Query2 or "correct" in Query2 or "ya" in Query2 or "yeah" in Query2:
                    #MondeVert.speak(self, 'Thanks for confirming the ' + Field)
                    s = False
                    WaitforResponse = True
                    continue
                if "no" in Query2 or "wrong" in Query2 or "not it" in Query2 or "nope" in Query2:
                    #MondeVert.speak(self, 'ok  please say the ' + Field + ' you want to set')
                    print('ok  please say the ' + Field + ' you want to set')
                    WaitforResponse = True
                    continue
                else:
                    #MondeVert.speak(self, 'Lets try that again what is the ' + Field)
                    Event().wait(1)
                    continue
        return Query

    ############################################################################################################################################################################################################################
    ############################################################################################################################################################################################################################
    ################ below is code that allow me to confirm with the user the inputs are correct
    ############################################################################################################################################################################################################################
    ############################################################################################################################################################################################################################
    def takeCommand(self):
        Query = MondeVert.getUserResponse(self, Response="Action Requested")
        return Query



    def ConfirmBOT(self, message):
        #MondeVert.speak(self, 'Chat GPT responded with: ' + message + ' do you want to use this prompt?', voice=4)
        query = MondeVert.getUserResponse(self)

        if 'yes' in query or 'correct' in query or 'yeah' in query or 'submit' in query:
            self.promptB = True

        if 'edit' in query or 'change' in query and ('prompt' in query or 'words' in query):
            MondeVert.editBotPrompt(self, message)

    def transcribe_Build_Query(self, pause=3.13):
        Query = MondeVert.getUserResponse(self)
        return Query

    def transcribe_Build_Query_Pause(self):
        Speak1 = 'What Tool do you want to use?'
        Print1 = 'Options: ChatGPT (ChatBot), Transcribe Mode , Make To Do List, Change Subject, Edit, Poem, Stop Recording --> Build others here'
        #MondeVert.speak(self, Speak1)
        print(Print1)
        Query = MondeVert.getUserResponse(self, Response="Tool Selected")
        return Query

    ############################################################################################################################################################################################################################
    ############################################################################################################################################################################################################################
    ############################################################################################################################################################################################################################
    ##########################                below is ChatGPT                  ##############################################################################################################
    ############################################################################################################################################################################################################################
    ############################################################################################################################################################################################################################




    def makeArt(self, Prompt=''):

        MondeVert.Add2Transcript(self, ' \n')
        Prior_Mode = self.Mode
        self.Mode = 'AI Art Mode'
        MondeVert.Add2Transcript(self, text2Add=(Prior_Mode + ' - ' + self.Mode + ':'))
        MondeVert.Add2Transcript(self, ' \n')
        self.promptB = False

        Promptc = 'True'

        if Prompt == '':
            Promptc = 'True'
        else:
            Promptc = 'False'

        if Prompt == '':
            print('Prompt is NULL: ' + Promptc)

            while self.promptB == False:
                AIspeak = 'What do you want me to draw?'

                #MondeVert.speak(self, AIspeak)

                print(AIspeak)
                prompt = MondeVert.getUserResponse(self, Response="Art Prompted", pause=1)

                if 'cancel' in prompt or 'stop' in prompt or 'quit' in prompt or 'done' in prompt or 'exit' in prompt:
                    MondeVert.saveTranscript(self)
                    self.promptB == True
                    continue

                MondeVert.ConfirmUR(self, prompt)
                print()
        else:
            prompt = Prompt
            print(prompt)
            MondeVert.Add2Transcript(self, '(Prompt Fed Directly into Function)')

        # Set up the OpenAI API client :
        openai.api_key = API_Key
        print("Sending to OpenAI...")
        print()

        response = openai.Image.create(
            prompt=prompt,
            n=1,
            size="1024x1024"
        )

        image_url = response['data'][0]['url']

        print(f"Image URL: {image_url}")
        print()

        # URL of the image to be downloaded is defined as image_url
        r = requests.get(image_url)  # create HTTP response object

        # send a HTTP request to the server and save
        # the HTTP response in a response object called r
        FileName = prompt
        print('Length of File Name: ' + str(len(FileName)))
        if len(FileName) >= 150:
            FileName = FileName[0:150]
            print('Length of File Name: ' + str(len(FileName)))

        fname = f"{''.join([c for c in FileName.strip().replace(' ', '_') if c.isalnum() or c == '_'])}.png"

        if not os.path.exists("../Utilities/images"):
            os.mkdir("../Utilities/images")

        if os.path.isfile(os.path.join(up.AI_Art_Path, '\'', fname)):
            fname = fname.split(".")[0] + f".{''.join(random.choice(string.ascii_letters) for x in range(5))}.png"


        fname_only = fname
        fname = os.path.join(up.AI_Art_Path, fname)
        print(f"Filename: {fname}")
        with open(fname, 'wb') as f:

            # Saving received content as a png file in
            # binary format

            # write the contents of the response (r.content)
            # to a new file in binary mode.
            f.write(r.content)

        if platform.system() == 'Darwin':  # macOS
            subprocess.call(('open', fname))
        elif platform.system() == 'Windows':  # Windows
            os.startfile(fname)
        else:  # linux variants
            subprocess.call(('xdg-open', fname))

        #this line saves the
        #IP.png2JPG(fname, up.PNGPath,fname_only, up.PNGPath_Archive, Del = False )



        return fname

    # Define a function that sends a message to ChatGPT
    def chat_query(self, prompt):
        model_engine = "text-davinci-003"
        completions = openai.Completion.create(
            engine=model_engine,
            prompt=prompt,
            max_tokens=2048,
            n=1,
            temperature=0.5,
        )

        message = completions.choices[0].text
        return message

    # Define a function that handles the conversation
    def conversation_handler(self, prompt):
        # Send the prompt to ChatGPT
        load_dotenv()
        # Set up the OpenAI API client :
        openai.api_key = API_Key
        response = MondeVert.chat_query(prompt)
        print(f"MondeVert: {response}")

        # End the conversation if ChatGPT says goodbye
        if "goodbye" in response.lower():
            return

        # Otherwise, get user input and continue the conversation
        prompt = input("User: ")
        MondeVert.conversation_handler(prompt)




    ############################################################################################################################################################################################################################
    ############################################################################################################################################################################################################################
    ############################################################################################################################################################################################################################
    ################Long Code where I give Digital Assistant commands##    (user menus)
    ############################################################################################################################################################################################################################
    ############################################################################################################################################################################################################################

    def RunChatGPT(self):
       #this will look very diff in the future The digital assistant will pass in certain key words to run different functions, but really this should exist without digital assist this is auxilary to anything I run it will be way easier to maintain
        dn = 1

    # Provide me with a copy of the short bio you come up with and then using that I want you to role play you are the respective artist and write a short 1 page story about something that would happen in the life of the respective artist you described. It can be about any topic but it should feel authentic to the life of the person you described."}
    def Make_a_Song(self, Mode='Random Song'):
        ArtPaths = []
        openai.api_key = API_Key

        crazy = round((randbelow(520000) + 170000) / 100000, 0)
        crazy = crazy / 10
        print(crazy)

        response = openai.ChatCompletion.create(
            model="gpt-3.5-turbo",
            messages=[
                {"role": "system", "content": up.system_Text},
                {"role": "user", "content": up.ArtistBio_SongArtist}
            ], temperature=crazy
        )

        Artist_Bio = response.choices[0].message.content

        Song1 = openai.ChatCompletion.create(
            model="gpt-3.5-turbo",
            messages=[
                {"role": "system", "content": up.system_Text},
                {"role": "user", "content": up.RolePlay_SongArtist + Artist_Bio},
                {"role": "user", "content": up.Song_Prompt_SongArtists},
                {"role": "user", "content": up.Title_SongArtists},
                {"role": "user", "content": up.Samples_SongArtists},
                {"role": "user", "content": up.Tune_SongArtists}

            ]
            , temperature=crazy
        )
        Song = Song1.choices[0].message.content

        Art_Prompt1 = openai.ChatCompletion.create(
            model="gpt-3.5-turbo",
            messages=[
                {"role": "system", "content": up.system_Text},
                {"role": "user", "content": up.RolePlay_SongArtist + Artist_Bio},
                {"role": "user", "content": up.ArtPrompt_SongArtist}
            ], temperature=crazy
        )

        Art_Prompt = Art_Prompt1.choices[0].message.content

        print(Artist_Bio)
        #MondeVert.speak(self, Artist_Bio)
        print(Song)
        #MondeVert.speak(self, Song)
        print(Art_Prompt)
        #MondeVert.speak(self, Art_Prompt)

        # Art2 = MondeVert.ChatGPTDA(self, temp=crazy, MakeArt=True, Prompt=(Art_Prompt),ConfirmBot=False)
        try:
            ArtPath2 = MondeVert.makeArt(self, Art_Prompt)
            ArtPaths.append(ArtPath2)

        except:
            ArtPath2 = ''

        Title = MondeVert.ChatGPTDA(self, temp=crazy, MakeArt=True, Prompt=(up.Song_Title_Prompt + Song),
                                        ConfirmBot=False)
        prompt = "Artist info:" + Artist_Bio + "Work of Art Inspiration:" + Art_Prompt + "Song:" + Song
        Prompts_Used = [
            up.system_Text + up.RolePlay_SongArtist + "AI Created a Persona shown as the Artist Bio Above" + up.Song_prompt + up.ArtPrompt_SongArtist]
        ArtistPoetInfo = 'Lyrics Written By: ' + up.Song_Writer + '      (' + 'Artwork by: ' + up.AI_ArtistName + ')'
        MondeVert.NamePoemSavePoem(self, prompt, ArtPaths, Prompts_Used, ArtistPoetInfo, title=Title,
                                       FolderPath=up.AI_Music_Path + '\\' + Mode, Mode='Song Lyrics' + Mode)

    def Make_a_SongRR(self, System = up.system_Text, Role= up.RolePlay_SongArtist, Background=up.ArtistBio_SongArtist, Task = up.Song_Prompt_SongArtists,Special=up.Samples_SongArtists, Format='', Mode='Random Song'):
        ArtPaths = []
        openai.api_key = API_Key

        crazy = round((randbelow(520000) + 170000) / 100000, 0)
        crazy = crazy / 10
        print(crazy)

        crazy += .2
        if crazy < .4:
            crazy = .6
        if crazy > .9:
            crazy = .7

        # Title = "Richies Music"
        current_time1 = datetime.datetime.now()
        current_time2 = current_time1.strftime('%m-%d-%Y_%H.%M.%S')
        # f2 = FolderPath
        f2 = up.SavePath
        SavePath1 = f2

        response = openai.ChatCompletion.create(
            model="gpt-3.5-turbo",
            messages=[
                {"role": "system", "content": up.system_TextRR},
                {"role": "user", "content": up.ArtistBio_SongArtistRR + up.Artist_Bio_DetailsRR}
            ], temperature=crazy
        )

        Artist_Bio = response.choices[0].message.content

        # print('Artist_Bio: ' + Artist_Bio)
        print(up.breakupOutput2)

        Song1 = openai.ChatCompletion.create(
            model="gpt-3.5-turbo",
            messages=[
                {"role": "system", "content": up.system_TextRR},
                {"role": "user", "content": up.RolePlay_SongArtistRR + Artist_Bio},
                {"role": "user", "content": up.Song_Prompt_SongArtistsRR + up.Song_Subject},
                {"role": "user", "content": up.Song_Format_Prompt + up.SongFormat},
                # {"role": "user", "content": up.Title_SongArtistsRR},
                # {"role": "user", "content": up.Samples_SongArtistsRR},
                # {"role": "user", "content": up.Tune_SongArtistsRR}

            ]
            , temperature=crazy
        )
        Song = Song1.choices[0].message.content

        # Create original details
        response = openai.ChatCompletion.create(
            model="gpt-3.5-turbo",
            messages=[
                {"role": "system", "content": up.system_TextRR},
                {"role": "user", "content": up.RolePlay_SongArtistRR + Artist_Bio},
                {"role": "user", "content": up.Title_SongArtistsRR + Song}

            ], temperature=crazy
        )

        Title = response.choices[0].message.content

        invalidCharRemoved = re.sub(r"[^a-zA-Z0-9 ]", "", Title)
        folder = str(invalidCharRemoved)
        if len(folder) > 44:
            folder = folder[0:44]
        Title1 = '\\' + str(invalidCharRemoved) + "_"
        if len(Title1) > 44:
            Title1 = Title1[0:44]
        SavePath2 = SavePath1 + '\\' + folder
        Title2 = SavePath2 + Title1 + current_time2

        # Create original details
        response = openai.ChatCompletion.create(
            model="gpt-3.5-turbo",
            messages=[
                {"role": "system", "content": up.system_TextRR},
                {"role": "user", "content": up.RolePlay_SongArtistRR + Artist_Bio},
                {"role": "user", "content": up.Samples_SongArtistsRR2 + Song}

            ], temperature=crazy
        )

        Samples2 = response.choices[0].message.content

        # Create original details
        response = openai.ChatCompletion.create(
            model="gpt-3.5-turbo",
            messages=[
                {"role": "system", "content": up.system_TextRR},
                {"role": "user", "content": up.RolePlay_SongArtistRR + Artist_Bio},
                {"role": "user", "content": up.ReWrite_Song + Song}

            ], temperature=crazy
        )

        ReWrite = response.choices[0].message.content

        # Create original details
        response = openai.ChatCompletion.create(
            model="gpt-3.5-turbo",
            messages=[
                {"role": "system", "content": up.system_TextDJ},
                {"role": "user", "content": up.RolePlay_SongArtistRR + Artist_Bio},
                {"role": "user", "content": Samples2},
                {"role": "user", "content": up.ExplainTheBeat + ReWrite}

            ], temperature=crazy
        )

        DJMondeVert = response.choices[0].message.content

        print(up.breakupOutput2)

        #MondeVert.speak(self, "DJ MondeVert Work complete yo")
        data = """Artist_Bio: """ + Artist_Bio + ": " + print(up.breakupOutput2) + """ Potential Samples: """ + Samples2  + print(up.breakupOutput2) +  """Title: """ + Title + + print(up.breakupOutput2) + """Song: """ + Song + + print(up.breakupOutput2) +"""Song 2.0: """ + ReWrite + print(up.breakupOutput2) + """DJMondeVert: """ + DJMondeVert

        print(data)

        #MondeVert.speak(self, "Saving the files round 1")
        # print(data)

        data = [(current_time2, data)]
        print(data)
        try:
            df1 = pd.DataFrame(data, columns=["TimeStamp", "Transcript"])
            # print(df1)
            # print(Title2)
            filename = Title2
            MondeVert.SaveCSV(data, filename, SavePath2)

            # MondeVert.SaveText(self,df1,'MondeVert Assistant', 'Full Transcript')
            MondeVert.add2Master2(df1)
        except:
            print('Review Error File did not save ')

        try:
            send_email_w_attachment_gmail(up.to, up.subject, Song, filename)
        except:
            print('email not send, its possible file was not created')

        #MondeVert.speak(self, "Making the art, painting yo picture")
        Art_Prompt1 = openai.ChatCompletion.create(
            model="gpt-3.5-turbo",
            messages=[
                {"role": "system", "content": up.system_TextRR},
                {"role": "user", "content": up.RolePlay_SongArtistRR + Artist_Bio},
                {"role": "user", "content": up.ArtPrompt_SongArtistRR}
            ], temperature=crazy
        )

        Art_Prompt = Art_Prompt1.choices[0].message.content
        print("Work of Art Inspiration:" + Art_Prompt)

        # print(Artist_Bio)
        # MondeVert.speak(self,Artist_Bio)
        # print(Song)
        # MondeVert.speak(self, Song)
        # print(Art_Prompt)
        # MondeVert.speak(self, Art_Prompt)

        # Art2 = MondeVert.ChatGPTDA(self, temp=crazy, MakeArt=True, Prompt=(Art_Prompt),ConfirmBot=False)
        try:
            ArtPath2 = MondeVert.makeArt(self, Art_Prompt)
            ArtPaths.append(ArtPath2)

        except:
            ArtPath2 = ''

        #MondeVert.speak(self, "Saving the files round 2")
        prompt = data + "Work of Art Inspiration:" + Art_Prompt
        Prompts_Used = [
            up.system_Text + up.RolePlay_SongArtist + "AI Created a Persona shown as the Artist Bio Above" + up.Song_prompt + up.ArtPrompt_SongArtist]
        ArtistPoetInfo = 'Lyrics Written By: ' + up.Song_Writer + '      (' + 'Artwork by: ' + up.AI_ArtistName + ')'
        MondeVert.NamePoemSavePoem(self, prompt, ArtPaths, Prompts_Used, ArtistPoetInfo, title=Title,
                                       FolderPath=up.AI_Music_Path + '\\' + Mode, Mode='Song Lyrics' + Mode)



    def Manual_Audio_File(self, SavePath = up.SavePath,FilePath = r"A:\MondeVert Productions\SHAINE - MondeVert AI Assistant\AI Tasks\MondeVert_Audio_Video_Story\Echoes of the Heart_Miniseries.txt", Text_override = '',Voice = random.choices(SAF.Original_List_of_Voices_English)[0], FileName = 'SHAINE Testing'):
    #I can likely use this for all of the ChatGPT conversations
        #Text  = pd.read_csv(FilePath)

        current_time1 = datetime.datetime.now()
        current_time2 = current_time1.strftime('%m-%d-%Y_%H.%M')
        Voice_fix = cu.CleanFileName(Voice)
        if Text_override =='':
            with open(FilePath) as f:
                Text = f.read()

        else:
            Text = Text_override

        FilePathnew= FilePath[:-4] + '_'+Voice_fix + '_'+ current_time2 +  '.mp3'
        NewTempPath = cu.SaveText2Audio(Text, SavePath=SavePath, FileName=FileName, FilePath=FilePathnew)
        print(NewTempPath)
        try:
            shutil.move(NewTempPath,FilePathnew)
            print('Audio File Saved in the following location: ' + FilePathnew)
            print('Audio File created with Voice: ' + str(Voice))
        except:
            print(NewTempPath)
    def Basic_GPT_Query(self,   Line2_Role  , Line3_Format,Line4_Task,Special = '',Line1_System_Rule = up.system_TextJoaT, crazy = .5, Subject= '', Outline = ''):#use this to create art style for the work

        if Subject != '':
            Line2_Role = Line2_Role + """Your role and subject matter expertise should fit the following Subject and or style and mood in the {Text} provided by the user Text:###""" + Subject + """###"""




        KeepGoing = False
        KillSwitch = 0
        while KeepGoing == False and KillSwitch < 8:
            try:

                Line1_System_Rule=Line1_System_Rule+ Line2_Role
                Line2_Role =Line2_Role + Special

                # This is for the result if you let the AI describe project and details and then make the response
                response = openai.ChatCompletion.create(
                    model="gpt-3.5-turbo",
                    messages=[
                        {"role": "system", "content": Line1_System_Rule},
                        {"role": "user", "content": Line2_Role },
                        {"role": "user", "content": Line3_Format},
                        {"role": "user", "content": Line4_Task },
                    ]
                    , temperature=crazy
                )
                GPT_Response = str(response.choices[0].message.content)

                KeepGoing = True
            except:
                print(' Error ChatGPT failed, trying to rerun prompt now.... if this happens too many times we will kill the script')
                KillSwitch += 1
                if KillSwitch == 9:
                    print('could not create a writer persona, redoing it now')
                    GPT_Response = MondeVert.Basic_GPT_Query(self)
                continue

            Full_User_Prompt = """User Inputs to Chat GPT: 
            1). """ + Line1_System_Rule + """
            2).""" + Line2_Role+ """
            3).""" + Line3_Format+ """
            4).""" + Line4_Task
            #print(up.breakupOutput)
            return GPT_Response



    def summarize_art_style_for_short_story(self,Role, Task, Artist_Persona, outline,writer_persona,Format = lup.Art_Summary_short_story_format, crazy = .5 ):
        Line2_Role = Role
        Line3_Format = Format
        Line4_Task = """Using the following {Artist_Persona},{outline} and {writer_persona} complete the table previously provided with information that aligns with the intersection of all 3 data points""" + """
        Artist Persona:###""" +Artist_Persona +"""###
        outline:  ###""" +outline +"""###
        Writer Persona:  ###""" +writer_persona +"""###"""


        Art_Style_details =  MondeVert.Basic_GPT_Query(self, Line2_Role = Line2_Role, Line4_Task = Line4_Task, Line3_Format= Line3_Format)
        #print(up.breakupOutput2)
        return Art_Style_details
    # use this to create characters art
#ned to run the prompt one first then this one (same goes for by scene too, this will just create the art ultimately
    def Create_Art_from_String_to_list(self, Art_Prompt, SavePath, FileName, Art_Style_details,Key_Word, Key_Char = ':',Delimiter = '|', crazy = .5):
        dummy =1
        ArtPaths = []
        if dummy ==1:
            try:
                Art_PromptCharo = Art_Prompt
            except:
                print('Error Trying to make original Painting')
                Art_PromptCharo =  Art_Prompt

            Art_PromptCharoo = Art_PromptCharo
            # ArtPrompt_Mondevert = + Result
            # ArtPrompt_Mondevert = ArtFormatAdvance + Result

            try:
                if Delimiter in Art_PromptCharo:
                    Art_PromptCharo = Art_PromptCharo.split(Delimiter)
                else:
                    Art_PromptCharo = Art_PromptCharo.replace(Key_Word + Key_Char, Delimiter)
                    Art_PromptCharo = Art_PromptCharo.replace(Key_Word, Delimiter)

                if Delimiter in Art_PromptCharo:
                    Art_PromptCharo = Art_PromptCharo.split(Delimiter)
                else:
                    Art_PromptCharo = Art_PromptCharo.replace(Key_Char, Delimiter)
                    Art_PromptCharo = Art_PromptCharo.split(Delimiter)

                # Test_Skel1 = Result.replace("Page:", "#")
                # Test_Skel1 = Test_Skel1.replace("Page", "#")

               # print('Art_PromptCharo')
                #print(Art_PromptCharo)



            except:
                print('Could not properly split out the string for multiple illustrations')



#before sending as art I should do one more clean up of the language/prompt to make it best

            KillSwitch = 0
            KeepGoing = False
            while KeepGoing == False and KillSwitch < 6:
                try:
                    Art_PromptCharoo + 'Revised Prompts: '
                    Characters = []
                    x = len(Art_PromptCharo)
                    for c in range(0, x):
                        # print(c)
                        #Do the change here before setting the value or keep both
                        Characters.append(Art_PromptCharo[c])
                        Clean_Prompt_User_GPT_Input = MondeVert.Basic_GPT_Query(self,   Line2_Role = lup.Clean_Role_after_Delimit + c , Line3_Format = lup.Clean_Format_after_Delimit,Line4_Task= lup.Clean_Task_after_Delimit ,Special = 'Use the following Text for the styles to embody/artist persona to embody###  ' + Art_Style_details+ c + '###',Line1_System_Rule = up.system_Text_Art, crazy = .5, Subject= '')
                        Characters.append(Art_PromptCharo[Clean_Prompt_User_GPT_Input])

                        Art_PromptCharoo +=Clean_Prompt_User_GPT_Input


                    #print('Characters')
                    #print(Characters)

                    for c in Characters:
                        print(' ')
                        print('-------------------------------------------')
                        print('-------------------------------------------')
                        print("Character: " +  c)
                        print('-------------------------------------------')
                        # print("Character Art1: " + up.Character_Art1 + c)
                        ArtPath2 = MondeVert.makeArt(self, Art_PromptCharo + c)
                        ArtPaths.append(ArtPath2)
                        #print(ArtPath2)


                except:
                    print('Error Trying to create multiple paintings')
                    KillSwitch += 1
                KeepGoing = True

        return ArtPaths #after its returned I append this to the folder with all relevant data
    # Make sure to start a master tracker with this information, use the writer persona as part of the data stored (TimeStamp, WriterPersona, Outline, Format, and/or put the prompt together for reference


    #use this to create characters
    def Create_Characters_Short_Story(self,  Task, Special, Format, Outline, Persona, Role= lup.Short_Story_Role, crazy = .5):

        if Persona =='':
            MondeVert.Writer_Persona_Short_Story(self, Role = Role )
#Make sure to start a master tracker with this information, use the writer persona as part of the data stored (TimeStamp, WriterPersona, Outline, Format, and/or put the prompt together for reference
        Character_Personas = MondeVert.Basic_GPT_Query(self, Line2_Role=Role, Line4_Task=Task, Line3_Format=Format, Special=Special)

        #make Character Art prompts for later
        Character_Art_Prompts_Main = MondeVert.Make_Art_prompts(self, Line2_Role=lup.Persona_artist_Role, Line3_Format=lup.Character_Art_Format_Main,
                         Line4_Task=lup.Character_Art_Task_Main, Special='', Line1_System_Rule=up.system_Text_Art,
                         Art_Style_details='Imitate a random artist and/or art style pick a random subject', crazy=.5,
                         Subject=Character_Personas)




        Character_Art_Prompts_Minor = MondeVert.Make_Art_prompts(self, Line2_Role=lup.Persona_artist_Role, Line3_Format=lup.Character_Art_Format_Main,
                         Line4_Task=lup.Character_Art_Task_Main, Special='', Line1_System_Rule=up.system_Text_Art,
                         Art_Style_details='Imitate a random artist and/or art style pick a random subject', crazy=.5,
                         Subject=Character_Personas)

        self.Character_Art_Prompts_Main = Character_Art_Prompts_Main
        self.Character_Art_Prompts_Minor = Character_Art_Prompts_Minor

        return Character_Personas


    def Writer_Persona_Short_Story(self, Role=lup.Persona_Role, Task=lup.Persona_Background,
                                   Format=lup.Persona_Format, Special=lup.Persona_Special,
                                   Subject='', crazy=.5):
        if Subject != '':
            Role = Role + """Your role and subject matter expertise should fit the following Subject and or style and mood in the {Text} provided by the user Text:###""" + Subject + """###"""

        Writer_Persona = MondeVert.Basic_GPT_Query(self,Line2_Role = Role,Line4_Task= Task, Line3_Format = Format, Special = Special, crazy = crazy)
        #Save a csv of this info in ssavepath
        #Add to Master tracker of artists and writer personas
        return Writer_Persona

    def Artist_Persona_Short_Story(self,Role = lup.Persona_artist_Role, Task = lup.Persona_artist_Background,Format = lup.Persona_artist_Format, Special = lup.Persona_artist_Special, Writer_persona = '', crazy = .5):
        if Writer_persona !='':
            Special = Special + """You should fit the style and mood provided by the following Outline:###""" + Writer_persona + """###"""

            Artist_Persona = MondeVert.Basic_GPT_Query(self,Line2_Role = Role,Line4_Task= Task, Line3_Format = Format, Special = Special, crazy = crazy)
            if Artist_Persona =='':
                Artist_Persona = 'Same artist that works for the simpsons, make it in simpson style art'
            #print(up.breakupOutput)
            return Artist_Persona


    def Make_Art_prompts(self,   Line2_Role = lup.Persona_artist_Role , Line3_Format = lup.Persona_artist_Format ,Line4_Task = lup.Persona_artist_Task,Special = '',Line1_System_Rule = up.system_Text_Art, Art_Style_details = 'Imitate a random artist and/or art style pick a random subject', crazy = .5, Subject = 'Choose a random artist and style and make an interesting prompt up'):
        Art_Prompt_With_Delimiter = MondeVert.Basic_GPT_Query(self, Line2_Role = Line2_Role+ Art_Style_details, Line3_Format = Line3_Format, Line4_Task = Line4_Task, Special='', Line1_System_Rule=Line1_System_Rule,crazy=crazy, Subject=   '')
        return Art_Prompt_With_Delimiter

    #run this once for characters then do it for the scenes --> for characters need the respective info (see other functions)
    #--> for scenes try to set it up too
    #--> also set up random info
    #lastly dont forget to make audio file


    def Quick_Title(self, Text, Line1_System_Rule = up.system_TextJoaT, crazy = .5):
        KeepGoing = False
        KillSwitch = 0
        while KeepGoing == False and KillSwitch < 8:
            try:

                Line1_System_Rule = Line1_System_Rule
                Line2_Task_Text = """Extract the Title from the following data, do not include 'Title:' or 'Title' in the title. Text: ###""" + Text + "###"

                # This is for the result if you let the AI describe project and details and then make the response
                response = openai.ChatCompletion.create(
                    model="gpt-3.5-turbo",
                    messages=[
                        {"role": "system", "content": Line1_System_Rule},
                        {"role": "user", "content": Line2_Task_Text}
                    ]
                    , temperature=crazy
                )
                Title = response.choices[0].message.content

                KeepGoing = True
            except:
                print(' Error ChatGPT failed, trying to rerun prompt now.... if this happens too many times we will kill the script')
                KillSwitch += 1
                if KillSwitch == 9:
                    print('could not create a writer persona, redoing it now')
                    Title = 'Unknown title - Too many errors'
                continue
        return Title

    def MondeVert_Audio_Video_Story(self, System='', Role='', Background='', Task='', Special='', Title='SHAINE Project Fail Safe', Mode='Audio_Video_Story',                 Logic_AI=0, Format='', SavePath=up.AI_AudioBook_Path):

        Episode_Count = 6
        ArtPaths = []
        openai.api_key = API_Key
        Art_PromptCharoo = ''
        Art_PromptCharo = ''
        Art_PromptForDALLE = ''
        Episode = []
        # ******************************************************************
        Episode_Outlines = ''





        Outline_Main = 'N/A not ready to test this yet'
        Outline_All_Episodes = 'N/A not ready to test this yet'
        ScreenPlay = 'N/A not ready to test this yet'
        Outline_Episode_Specific = 'N/A not ready to test this yet'
        Title = 'N/A not ready to test this yet'
        Project_Description = 'N/A not ready to test this yet'
        Result = 'N/A not ready to test this yet'
        Result_AI = 'N/A not ready to test this yet'
        AIPrompt = 'N/A not ready to test this yet'
        Writer_Persona = 'N/A not ready to test this yet'
        Artist_Persona = 'N/A not ready to test this yet'
        Art_Prompts = 'N/A not ready to test this yet'
        Episode_Outline = []


        Titleo = Title

        if Logic_AI == 0:
            crazy = round((randbelow(520000) + 170000) / 100000, 0)
            crazy = crazy / 10
            print('Crazy Rating: ' + str(crazy))
            crazy += .2
            if crazy < .4:
                crazy = .6
            if crazy > .9:
                crazy = .7
        else:
            crazy = Logic_AI


        current_time1 = datetime.datetime.now()
        current_time2 = current_time1.strftime('%m-%d-%Y_%H.%M.%S')
        # f2 = FolderPath
        f2 = up.AI_AudioBook_Path
        f2 = SavePath
        SavePath1 = f2





        Format0 = up.Test_Format0
        KillSwitch = 0

        Format_f = Format
        if Mode == 'skit':
            Format_f = Format0

        if Title != '':
            Background = Title + Background

        tempFormat = Format
        if Mode == 'PictureBook':
            Format = up.Test_Format_PictureBook_outline



#Call Writer persona function
        Writer_Persona  = MondeVert.Writer_Persona_Short_Story(self, Role=lup.Persona_Role, Task=lup.Persona_Background,
                                   Format=lup.Persona_Format2, Special=lup.Persona_Special,
                                   Subject='', crazy=crazy)

        print('Created Writer Persona')
# Call Artist persona function
        Artist_Persona = MondeVert.Artist_Persona_Short_Story(self,Role = lup.Persona_artist_Role, Task = lup.Persona_artist_Background,Format = lup.Persona_artist_Format, Special = lup.Persona_artist_Special, Writer_persona = Writer_Persona, crazy = crazy)

        print('Created Artist Persona')
        # Call outline - Main function
        Outline_Main = MondeVert.Basic_GPT_Query(self,   Line2_Role = 'Take on the following writer persona:' + Writer_Persona , Line3_Format= lup.Short_Story_Outline_Format,Line4_Task = lup.Short_Story_Outline_Task , Special = '',Line1_System_Rule = up.system_TextJoaT, crazy = crazy, Subject= '')
        print('Created Main Outline')

        #print(up.breakupOutput)
        #print('User_Input_Outline_Main: ' + User_Input_Outline_Main)
        #print(up.breakupOutput)
        #Call Art Theme Function
        Art_Style_For_Story = MondeVert.summarize_art_style_for_short_story(self,Role, Task, Artist_Persona, Outline_Main,Writer_Persona,Format = lup.Art_Summary_short_story_format, crazy = crazy )
        print('Created Artist style summary')

#Call Character function
        Characters = MondeVert.Create_Characters_Short_Story(self, Role = lup.Characters_Role, Task = lup.Characters_Task, Special = lup.Characters_Special, Format =  lup.Characters_Format, Outline = Outline_Main, Persona = Writer_Persona, crazy = crazy)
        # print(up.breakupOutput)
        # print('User_Input_Character: ' +  User_Input_Character)
        # print(up.breakupOutput)

        print('Created Character Personas')
 #Call Outline - All Episodes
        Outline_ALL_Episodes = MondeVert.Basic_GPT_Query(self,   Line2_Role = 'Take on the following writer persona:' + Writer_Persona , Line3_Format= lup.Short_Story_Season_Outline_Format,Line4_Task = lup.Short_Story_Season_Outline_Task + """Use the following High level outline as a source for the new task you are working on Outline:###""" + Outline_Main + '###',Special = '',Line1_System_Rule = up.system_TextJoaT, crazy = crazy, Subject= '', Outline = Outline_Main)

        print('Created All Espisode Outline')
#Call Outline - Episode by episode
        Pilot_Short_story_fix = ''
        for_num = Episode_Count + 1
        for episode_num in range(1,for_num):




            if episode_num ==1:
                Pilot_Short_story_fix = lup.Pilot_Short_story_fix
            else:
                Pilot_Short_story_fix = ''
            print('Created Outline for Episode ' + str(episode_num))
            Outline_Episodes_by_episode = MondeVert.Basic_GPT_Query(self,Line2_Role='Take on the following writer persona:' + Writer_Persona,Line3_Format=lup.Short_Story_Episode_Outline_Format,Line4_Task=lup.Short_Story_Episode_Outline_Task + """Use the following High level outline as a source for the new task you are working on Outline:###""" + Outline_ALL_Episodes + '###',Special=lup.Short_Story_Episode_Outline_Special + str(episode_num) + "### " + Pilot_Short_story_fix,Line1_System_Rule=up.system_TextJoaT,crazy=crazy,Subject='',Outline=Outline_ALL_Episodes)
            Episode_Outline.append(Outline_Episodes_by_episode)
            Episode_Story = MondeVert.Basic_GPT_Query(self,Line2_Role='Take on the following writer persona:' + Writer_Persona,Line3_Format=lup.Short_Story_Format,Line4_Task=lup.Short_Story_Task + """Use the following Episode-specific-Detailed  outline as a source for the new task you are working on Outline:###""" + Outline_Episodes_by_episode + '###',Special=lup.Short_Story_Special + str(episode_num) + "### " + Pilot_Short_story_fix,Line1_System_Rule=up.system_TextJoaT,crazy=crazy,Subject='',Outline=Outline_Episodes_by_episode)
            Episode.append(Episode_Story)

            # print(up.breakupOutput)
            # print('User_Input_AllEpisodes_Episodes_by_episode: ' + User_Input_AllEpisodes_Episodes_by_episode)
            # print(up.breakupOutput)

#******************************************************************

        #print(up.breakupOutput)
# ************************************************************************************************************************************
# ************************************************************************************************************************************

        try:
            if Title == '':
                # Title = MondeVert.ChatGPTDA(self, temp=0.5, MakeArt=True, Prompt=( Titleo+ '"' + Background + '"'),ConfirmBot=False)
                Title = Titleo
        except:
            print('could not come up with title review')
            Title = Mode
#********************************************************************************************************************************
#********************************************************************************************************************************

     # Below is where I make character art can be done after all the other work so I set it up that way in case it fails I still save other work.

        Title = MondeVert.Quick_Title(self, Outline_Main)
        Title = re.sub(r"[^a-zA-Z0-9 ]", "", Title)
        FileName = Title + current_time2
        SavePath = SavePath1 + '\\' + Title
        FileName_Char_Main = Title + '_Main Characters_' + current_time2
        FileName_Char_Minor = Title + '_Minor Characters_' + current_time2
        # Call Character art prompt

        ScreenPlay = ''
        for x in range(0, Episode_Count):
            Episode_Outlines += Episode_Outline[x]
            ScreenPlay += Episode[x]
        # MondeVert.speak(self, "MondeVert Project Work complete yo")

        print(up.breakupOutput2)

        #Save all of the userinputs for reference later UserInputs =
        #Save Details and Have voice speak it out

        # print(Title)
        # print(Writer_Persona)
        # print(Artist_Persona)
        # print(Art_Style_For_Story)
        # print(Outline_Main)
        # print(Characters)
        # print(self.Character_Art_Prompts_Minor)
        # print(self.Character_Art_Prompts_Major)
        # print(Outline_All_Episodes)
        # print(Outline_Episode_Specific)
        # print(ScreenPlay)

        data = """Title: """ + Title + up.breakupOutput2 + """Writer Bio: """ + Writer_Persona + up.breakupOutput2 + """Artist Bio: """ + Artist_Persona + Art_Style_For_Story + up.breakupOutput2 + """Outline_Main: """ + Outline_Main + up.breakupOutput2 +"""Characters: """+ Characters + up.breakupOutput2 + """Character Descriptions: """+ self.Character_Art_Prompts_Main + up.breakupOutput2 +  + self.Character_Art_Prompts_Minor +  """     Outline_All_Episodes: """ + Outline_All_Episodes + up.breakupOutput2 + """Outline_Episode_Specific: """ + Outline_Episode_Specific + up.breakupOutput2 + """ScreenPlay: """ + ScreenPlay
        AudioBook_Text =  """Title: """ + Title + up.breakupOutput2 + ScreenPlay
        Outline_and_Details = """Title: """ + Title + up.breakupOutput2 + """Writer Bio: """ + Writer_Persona + up.breakupOutput2 + """Artist Bio: """ + Artist_Persona + Art_Style_For_Story + up.breakupOutput2 + """Outline_Main: """ + Outline_Main + up.breakupOutput2 +"""Characters: """+ Characters + up.breakupOutput2 + """Character Descriptions:    Main Characters - """+ self.Character_Art_Prompts_Main + up.breakupOutput2 +  """ | Minor Characters - """ + self.Character_Art_Prompts_Minor +  """     Outline_All_Episodes: """ + Outline_All_Episodes + up.breakupOutput2 + """Outline_Episode_Specific: """ + Outline_Episode_Specific + up.breakupOutput2
        print('Output for : ' + Mode)
        print(data)

        print(up.breakupOutput2)
        # print(up.breakupOutput2)
        # print(up.breakupOutput2)
        print('End of creation, saving files')


        try:
            for i in range (1,Episode_Count+1):
                FileName_Audio = Title + '_Episode_' + i + current_time2
                Voice = random.choices(SAF.Original_List_of_Voices_English)
                cu.SaveText2Audio(SavePath=SavePath, FileName=FileName_Audio, Voice=Voice, Neural='Neural', Mode='Text2Audio')
        except:
            print('Audio Files not recorded error occurred')



#This is just the characters, need to do this for scenes etc too
        ArtPaths1 = MondeVert.Create_Art_from_String_to_list(self, self.Character_Art_Prompts_Main, SavePath,
                                                             FileName_Char_Main, Art_Style_For_Story,
                                                             Key_Word='Character', Key_Char=':', Delimiter='|',
                                                             crazy=crazy)
        ArtPaths.append(ArtPaths1)
        ArtPaths1 = MondeVert.Create_Art_from_String_to_list(self, self.Character_Art_Prompts_Minor, SavePath,
                                                             FileName_Char_Minor,
                                                             Art_Style_For_Story, Key_Word='Character', Key_Char=':',
                                                             Delimiter='|',
                                                             crazy=crazy)
        ArtPaths.append(ArtPaths1)
        #print(up.breakupOutput)


        # MondeVert.speak(self, "Saving the files round 1")
        # print(data)

        data1 = [(data)]

        try:
            # df1 = pd.DataFrame(data1, columns=['Text'])

            filename = Title
            MondeVert.SaveCSV(data, filename, SavePath)

        except:
            print('Review Error File did not save ')

        try:
            send_email_w_attachment_gmail(up.to, up.subject, Project_Description, SavePath + filename)
            print("Gmail email sent")
            send_email_w_attachment_outlook(up.to, up.subject, Project_Description, SavePath + filename)
            print("Outlook email sent")
        except:
            print('email not send, its possible file was not created')

        # Art changes based on mode for now I will switch only when the mode calls for it otherwise default


        prompt = data + "Work of Art Inspiration(s):" + Art_PromptForDALLE + "     Work of Art Inspiration(scene by scene):" + Art_PromptCharoo
        Prompts_Used = [
            up.system_Text + up.RolePlay_SongArtist + "AI Created a Persona shown as the Artist Bio Above" + up.Song_prompt + up.ArtPrompt_SongArtist]
        ArtistPoetInfo = 'Shane Donovan - SHAINE - MondeVert'
        MondeVert.NamePoemSavePoem(self, prompt, ArtPaths, Prompts_Used, ArtistPoetInfo, title=Title,
                                   FolderPath=SavePath1 + '\\' + Mode, Mode=Mode)
        KillSwitch = 7

    # MondeVert.speak(self, ReWrite )
    def MondeVertAuto(self, Mode='AutoSocial', Role = '', System = ''):
        if Mode == 'AutoSocial':
            Caption = MondeVert.MondeVertTask(self, System=up.system_TextJoaT, Role=up.Test_Role_SMA, Background=up.Test_Background_SMA,Task=up.Test_Task_SMA, Special=up.Test_Special_SMA, Format=up.Test_Format_SMA, Title = up.Test_Title_SMA,Mode='AutoSocial')
        return Caption + up.Instagram_Adds
    def MondeVertMenu_up(self, Mode='Basic', Role = '', System = ''):
        self.Mode = Mode

        if Mode == 'ScreenPlay':
            MondeVert.Make_a_ScreenPlay(self,System = up.system_Text_ScreenPlay0, Mode='ScreenPlay')
        elif Mode == 'Music':
            MondeVert.Make_a_SongRR(self,System = up.system_TextRR, Mode='Music')
        elif Mode == 'Skit':
            MondeVert.MondeVertTask(self,System = up.system_TextJoaT,Role=up.Test_Role_Skit, Background=up.Test_Background_Skit, Task = up.Test_Task_Skit,Special=up.Test_Special_Skit, Format=up.Test_Format_Skit, Mode='Skit')
        elif Mode == 'Basic':
            MondeVert.MondeVertTask(self,System = up.system_TextJoaT,Role=up.Test_Role_Explain, Background=up.Test_Background_Explain, Task = up.Test_Task_Explain,Special=up.Test_Special_Explain, Format=up.Test_Format_Explain,Title = up.Test_Title_Explain, Mode='Basic', Logic_AI=.5)
        elif Mode == 'Advance':
            MondeVert.MondeVertTask(self,System = up.system_TextJoaT,Role=up.Test_Role_Explain, Background=up.Test_Background_Explain, Task = up.Test_Task_Explain, Special=up.Test_Special_Explain, Format=up.Test_Format_Explain, Mode='Advance')
        elif Mode == 'Interview':
            MondeVert.MondeVertTask(self,System = up.system_TextJoaT,Role=up.Test_Role_Explain, Background=up.Test_Background_Explain, Task = up.Test_Task_Explain,Special=up.Test_Special_Explain, Format=up.Test_Format_Explain, Mode='Interview')
        elif Mode == 'Social':
            MondeVert.MondeVertTask(self, System=up.system_TextJoaT, Role=up.Test_Role_SM, Background=up.Test_Background_SM,Task=up.Test_Task_SM, Special=up.Test_Special_SM, Format=up.Test_Format_SM, Mode='Social')
        elif Mode == 'RealEstate':
            MondeVert.MondeVertTask(self, System=up.system_TextJoaT, Role=up.Test_Role_RE,Background=up.Test_Background_RE, Task=up.Test_Task_RE, Special=up.Test_Special_RE,Format=up.Test_Format_RE, Mode='RealEstate')
        elif Mode == 'PictureBook':
            MondeVert.MondeVertTask(self, System=up.system_TextJoaT, Role=up.Test_Role_PictureBook,Background=up.Test_Background_PictureBook, Task=up.Test_Task_PictureBook, Special=up.Test_Special_PictureBook,Format=up.Test_Format_PictureBook, Mode='PictureBook', ArtPrompt= up.MondeVert_ArtPrompt,ArtFormat = up.MondeVert_ArtFormat, AdvanceArtPrompt = up.MondeVert_ArtPrompt_PictureBook , AdvanceArtFormat= up.MondeVert_ArtFormat_PictureBook)
        elif Mode == 'JobDescription':
            MondeVert.MondeVertTask(self, System=up.system_TextJoaT, Role=lup.Test_Role_Summarize_JD,
                                    Background=lup.Test_Background_Summarize_JD, Task=lup.Test_Task_Summarize_JD,
                                    Special=lup.Test_Special_Summarize_JD, Format=lup.Test_Format_Summarize_JD, Mode='JobDescription')

        elif Mode == 'Resume':
            MondeVert.MondeVertTask(self, System=up.system_TextJoaT, Role=lup.Test_Role_Resume,
                                    Background=lup.Test_Background_Resume, Task=lup.Test_Task_Resume,
                                    Special=lup.Test_Special_Resume, Format=lup.Test_Format_Resume, Mode='Resume')

        elif Mode == 'Resume_Review':
            MondeVert.MondeVertTask(self, System=up.system_TextJoaT, Role=lup.Test_Role_Resume_Review,
                                    Background=lup.Test_Background_Resume_Review, Task=lup.Test_Task_Resume_Review,
                                    Special=lup.Test_Special_Resume_Review, Format=lup.Test_Format_Resume_Review, Mode='Resume_Review')




        elif Mode == 'Resume_Consolidate_Old':
            MondeVert.MondeVertTask(self, System=up.system_TextJoaT, Role=lup.Test_Role_Resume_old,
                                    Background=lup.Test_Background_Resume_old, Task=lup.Test_Task_Resume_old,
                                    Special=lup.Test_Format_Resume_old, Format=lup.Test_Format_Resume_old, Mode='Resume_Consolidate_Old')



        elif Mode == 'Resume_Combine_old_new':
            MondeVert.MondeVertTask(self, System=up.system_TextJoaT, Role=lup.Test_Role_Resume_old_New_Combo,
                                    Background=lup.Test_Background_Resume_old_New_Combo, Task=lup.Test_Task_Resume_old_New_Combo,
                                    Special=lup.Test_Format_Resume_old_New_Combo, Format=lup.Test_Format_Resume_old_New_Combo, Mode='Resume_Combine_old_new')

        elif Mode == 'LinkedIn':
            MondeVert.MondeVertTask(self, System=up.system_TextJoaT, Role=lup.Test_Role_LinkedIn,
                                    Background=lup.Test_Background_LinkedIn, Task=lup.Test_Task_LinkedIn,
                                    Special=lup.Test_Format_LinkedIn, Format=lup.Test_Format_LinkedIn, Mode='LinkedIn')



        elif Mode == 'MondeVert_Audio_Video_Story':
           MondeVert.MondeVert_Audio_Video_Story(self)



        elif Mode == 'Freelance_Services':
            MondeVert.MondeVertTask(self, System=up.system_TextJoaT, Role=lup.Test_Role_Resume,
                                    Background=lup.Test_Background_Summarize_FL, Task=lup.Test_Task_Summarize_FL,
                                    Special=lup.Test_Special_Summarize_FL, Format=lup.Test_Format_Summarize_FL, Mode='Freelance_Services')


        elif Mode == 'Create_Persona_Writer':
            MondeVert.MondeVertTask(self, System=up.system_TextJoaT, Role=lup.Persona_Role,
                                    Background=lup.Persona_Background, Task=lup.Persona_Task,
                                    Special=lup.Persona_Special, Format=lup.Persona_Format, Mode='Create_Persona_Writer')




        elif Mode == 'ReSearch':
            MondeVert.MondeVertTask(self, System=up.system_TextJoaT, Role=up.Test_Role_ReSearch,Background=up.Test_Background_ReSearch, Task=up.Test_Task_ReSearch, Special=up.Test_Special_ReSearch,Format=up.Test_Format_ReSearch, Mode='Research')
        elif Mode == 'ReSearch2':
            MondeVert.MondeVertTask(self, System=up.system_TextJoaT, Role=up.Test_Role_ReSearch,
                                    Background=up.Test_Background_ReSearch, Task=up.Test_Task_ReSearch2,
                                    Special=up.Test_Special_ReSearch, Format=up.Test_Format_ReSearch2, Mode='Research2')
        elif Mode == 'ReSearch3':
            MondeVert.MondeVertTask(self, System=up.system_TextJoaT, Role=up.Test_Role_ReSearch,
                                    Background=up.Test_Background_ReSearch, Task=up.Test_Task_ReSearch3,
                                    Special=up.Test_Special_ReSearch, Format=up.Test_Format_ReSearch3, Mode='Research3')

        elif Mode == 'Book':
            MondeVert.Make_a_ScreenPlay(self, Mode='ScreenPlay')
        # elif Mode == 'Series':
        #     MondeVert.Series(self, Mode='Series', Episodes= 10, Seasons = 1, Audience = 'Kids')

        #atexit.register(MondeVert.saveTranscript(self))




    def MondeVertTask(self, System = '', Role='', Background='', Task='',Special='', Title='', Mode='Basic', Logic_AI=0,Format = '',  SavePath = up.AI_Task_Path, ArtPrompt= up.MondeVert_ArtPrompt,ArtFormat = up.MondeVert_ArtFormat, AdvanceArtPrompt = up.MondeVert_ArtPrompt_PictureBook , AdvanceArtFormat= up.MondeVert_ArtFormat_PictureBook, ART_System = up.system_Text_Art):
        ArtPaths = []
        openai.api_key = API_Key
        Art_PromptCharoo = ''
        Art_PromptCharo = ''
        Art_PromptForDALLE = ''

        Titleo = Title

        if Logic_AI == 0:
            crazy = round((randbelow(520000) + 170000) / 100000, 0)
            crazy = crazy / 10
            print(crazy)
            crazy += .2
            if crazy < .4:
                crazy = .6
            if crazy > .9:
                crazy = .7
        else:
            crazy = Logic_AI

        # Title = "Richies Music"
        current_time1 = datetime.datetime.now()
        current_time2 = current_time1.strftime('%m-%d-%Y_%H.%M.%S')
        # f2 = FolderPath
        f2 = up.SavePath
        f2 = SavePath
        SavePath1 = f2



        Project_Description = 'N/A not ready to test this yet'
        Result = 'N/A not ready to test this yet'
        Result_AI = 'N/A not ready to test this yet'
        AIPrompt = 'N/A not ready to test this yet'
        Result_Combine = 'N/A not ready to test this yet'
        Bitter_Critic = 'N/A not ready to test this yet'


        Task2 = up.Test_Task_Critic
        Format2 = up.Test_Format_Critic
        Special2 = up.Test_Special_Critic
        Role2 = up.Test_Role_Critic

        Format0 = up.Test_Format0
        KillSwitch = 0

        Format_f = Format
        if Mode =='skit':
            Format_f = Format0

        if Title != '':
            Background = Title + Background

        tempFormat = Format
        if Mode == 'PictureBook':
            Format = up.Test_Format_PictureBook_outline

        while KillSwitch < 7:
            KeepGoing = False
            while KeepGoing ==False and KillSwitch < 6:
                # dummy = 1
                # if dummy ==1:
                try:# This is for the result when you ask AI to summarize the project
                    response = openai.ChatCompletion.create(
                        model="gpt-3.5-turbo",
                        messages=[
                            {"role": "system", "content": up.system_TextJoaT + Role},
                            {"role": "user", "content": Role},
                            {"role": "user", "content": Format_f},
                            {"role": "user", "content": Special},
                            {"role": "user", "content": Background}
                        ], temperature=crazy
                    )

                    Project_Description = response.choices[0].message.content

                    if Title != '':
                        Project_Description = Title + Project_Description
                    else:
                        Project_Description = Project_Description

                    KeepGoing = True

                except:

                    print('1). Error ChatGPT failed, trying to rerun prompt now.... if this happens too many times we will kill the script')
                    KillSwitch +=1
                    if KillSwitch == 6:
                        Project_Description = Background
                    continue

                Format = tempFormat

                print('************************************************************************************************')

#This was misguided lol but I think getting rid of prompt for now is ok eventually maybe this will cause an issue, going to be fully reworked soon
            if 'RESUME' in Mode.upper():
                Resume_Fix =  Project_Description
            else:
                Resume_Fix = up.AI_Background_Prompt + Project_Description

            KeepGoing = False
            KillSwitch = 0
            while KeepGoing == False and KillSwitch < 6:
                try:
                    # This is for the result if you let the AI describe project and details and then make the response
                    response = openai.ChatCompletion.create(
                        model="gpt-3.5-turbo",
                        messages=[
                            {"role": "system", "content": up.system_TextJoaT + Role},
                            {"role": "user", "content": up.Role_Play_Prompt + Role},
                            {"role": "user", "content":  Format},
                            {"role": "user", "content": up.USER_SPECIAL_Prompt + Special},
                            {"role": "user", "content": Resume_Fix},
                            {"role": "user", "content": up.USER_Task_Prompt + Task},
                        ]
                        , temperature=crazy
                    )
                    Result = response.choices[0].message.content

                    KeepGoing = True
                except:
                    print('2). Error ChatGPT failed, trying to rerun prompt now.... if this happens too many times we will kill the script')
                    KillSwitch += 1
                    continue



                print(up.breakupOutput)

                KillSwitch = 0
                while KeepGoing == False and KillSwitch < 6:
                    try:
                        response = openai.ChatCompletion.create(
                            model="gpt-3.5-turbo",
                            messages=[
                                {"role": "system", "content": up.system_TextJoaT + Role},
                                {"role": "user", "content": up.Role_Play_Prompt + Role},
                                {"role": "user", "content": up.UserRequest + Format + Special + Task},
                                # {"role": "user", "content": up.USER_SPECIAL_Prompt + Special},
                                #{"role": "user", "content": up.AI_Background_Prompt + Project_Description},
                                {"role": "user",
                                 "content": up.CombineBothResults + "Result 1: " + Result + "Result 2: " + Result_AI},
                                #{"role": "user", "content": up.USER_Task_Prompt + Task},
                            ]
                            , temperature=crazy
                        )
                        Result_Combine = response.choices[0].message.content
                        KeepGoing = True
                    except:
                        print('5). Error ChatGPT failed, trying to rerun prompt now.... if this happens too many times we will kill the script')
                        #KillSwitch += 1

                        try:
                            Result2 = openai.ChatCompletion.create(
                                model="gpt-3.5-turbo",
                                messages=[
                                    {"role": "system", "content": up.system_TextJoaT + Role},
                                    {"role": "user", "content": up.Role_Play_Prompt + Role},
                                    {"role": "user", "content": up.UserRequest + Format + Special},
                                    # {"role": "user", "content": up.USER_SPECIAL_Prompt + Special},
                                    #{"role": "user", "content": Project_Description},
                                    {"role": "user",
                                     "content": up.CombineBothResults + "Result 1: " + Result},
                                    #{"role": "user", "content": up.USER_Task_Prompt + Task},
                                ]
                                , temperature=crazy
                            )
                            Result_Combine = Result2.choices[0].message.content
                            print("Only revised v1")
                            KeepGoing = True
                        except:
                            print("could not combine 2 versions")
                            print( '6). Error ChatGPT failed, trying to rerun prompt now.... if this happens too many times we will kill the script')
                            KillSwitch += 1
                            continue




                    # This is funny content where the AI reviews the youtube script
                KillSwitch =0
                KeepGoing = False
                while KeepGoing == False and KillSwitch < 6:

                    try:
                        if   Mode =='skit' or Mode =='skit'or Mode =='skit'or Mode =='PictureBook':
                            # Create original details
                            response = openai.ChatCompletion.create(
                                model="gpt-3.5-turbo",
                                messages=[
                                    {"role": "system", "content": up.system_TextJoaT + Role2},
                                    {"role": "user", "content": up.Role_Play_Prompt + Role2},
                                    {"role": "user", "content": up.UserRequest + Format2 + Special2 + Project_Description + Task2 + Title},
                                    {"role": "user", "content": up.USER_SPECIAL_Prompt + Special2},
                                    {"role": "user", "content": up.Test_Background_Critic + Result_Combine},
                                    #{"role": "user", "content": up.USER_Task_Prompt + Task2},

                                ], temperature=crazy
                            )

                            Bitter_Critic = response.choices[0].message.content

                        KeepGoing = True
                        continue
                    except:
                        print('8). Error ChatGPT failed, trying to rerun prompt now.... if this happens too many times we will kill the script')
                        KillSwitch += 1
                        continue

                KillSwitch = 0
                KeepGoing = False
                while KeepGoing == False and KillSwitch < 6:
                    try:
                        if Title == '':
                            Title = MondeVert.Get_Title_GPT(self, Project_Description = Project_Description)
                            KeepGoing = True
                        else:
                            KeepGoing =True
                    except:
                        try:
                            Title = MondeVert.Get_Title_GPT(self, Project_Description=Background)
                            print('First try did not work for title, but went through second time')
                        except:
                            print(Mode)
                            print (Title)
                            print('7). Error ChatGPT failed, trying to rerun prompt now.... if this happens too many times we will kill the script')
                            KillSwitch += 1
                            continue


            try:
                if Title == '':
                    #Title = MondeVert.ChatGPTDA(self, temp=0.5, MakeArt=True, Prompt=( Titleo+ '"' + Background + '"'),ConfirmBot=False)
                    Title = Titleo
            except:
                print('could not come up with title review')
                Title = Mode


            #this should be a utility
            invalidCharRemoved = re.sub(r"[^a-zA-Z0-9 ]", "", Title)
            folder = Mode
            if len(folder) > 44:
                folder = folder[0:44]
            Title1 = '\\' + str(invalidCharRemoved) + "_"
            if len(Title1) > 44:
                Title1 = Title1[0:44]
            SavePath2 = SavePath1 + '\\' + folder



            print('************************************************************************************************')
            print('************************************************************************************************')
            print('************************************************************************************************')

            #MondeVert.speak(self, "MondeVert Project Work complete yo")

            data = """Project_Description: """ + Project_Description + up.breakupOutput2 + """Title: """ + Title  +up.breakupOutput2 + """     Result: """ + Result +                up.breakupOutput2 + """Results 2.0 (reworked a second time): """ + Result_Combine +                up.breakupOutput2 + """Bitter_Critic: """ + Bitter_Critic
            print('Output for : '+ Mode)
            print(data)

            #MondeVert.speak(self, "Saving the files round 1")
            # print(data)

            data1 = [(data)]

            try:
                # df1 = pd.DataFrame(data1, columns=['Text'])

                filename = Title1
                MondeVert.SaveCSV(data, filename, SavePath2)
                FilePath_email = cu.SaveText2Audio(SavePath=SavePath2, FileName=filename, Voice = random.choices(SAF.Original_List_of_Voices_English)[0], Neural='Neural',Mode=Mode)

            except:
                print('Review Error File did not save ')

            try:
                send_email_w_attachment_gmail(up.to, up.subject, Project_Description, FilePath_email)
                print("Gmail email sent")
                send_email_w_attachment_outlook(up.to, up.subject, Project_Description,  FilePath_email)
                print("Outlook email sent")
            except:
                print('email not send, its possible file was not created')



            #Art changes based on mode for now I will switch only when the mode calls for it otherwise default
            if Mode =='PictureBook':
                ArtPrompt_Mondevert = ArtPrompt + Result
                # ArtFormat_Mondevert = ArtPrompt + Result
            else:
                ArtPrompt_Mondevert = ArtPrompt + Project_Description
                # ArtFormat_Mondevert = ArtFormat

            #ArtPrompt_Mondevert0 = up.MondeVert_ArtFormat

            #MondeVert.speak(self, "Making the art, painting yo picture")

            crazyArt = crazy
            crazyArt = .4
            try:
                Art_PromptForDALLE = MondeVert.GPTArt(self, crazy, prompt=AIPrompt, Plot=ArtPrompt_Mondevert,ArtFormat=ArtFormat,sys_prompt=ART_System)
            except:
                Art_PromptForDALLE = 'Make a work of art based on the following text:' + Result

            # Art_Prompt1 = openai.ChatCompletion.create(
            #     model="gpt-3.5-turbo",
            #     messages=[
            #         {"role": "system", "content": up.system_Text_Art},
            #         {"role": "user", "content": ArtPrompt_Mondevert},
            #         {"role": "user", "content": ArtFormat},
            #     ], temperature=crazyArt
            # )
            #
            # Art_PromptForDALLE = Art_Prompt1.choices[0].message.content
            print("Work of Art Inspiration:" + Art_PromptForDALLE)

            try:
                ArtPath2 = MondeVert.makeArt(self, Art_PromptForDALLE)
                ArtPaths.append(ArtPath2)

            except:
                ArtPath2 = ''

            #MondeVert.speak(self, "Saving the files round 2")

            try:
                if '|' in Art_PromptForDALLE:
                    Art_PromptForDALLE_Simple = Art_PromptForDALLE.split('|')
                else:
                    Art_PromptForDALLE_Simple = Art_PromptForDALLE.replace("Subject:", "|")
                    Art_PromptForDALLE_Simple = Art_PromptForDALLE.replace("Subject", "|")

                Art_PromptForDALLE_Simple = Art_PromptForDALLE_Simple[:-1]

            except:
                print('Error did not Properly split the list')



            if Mode == 'PictureBook' or Mode == 'script':
                try:
                    Art_PromptCharo = MondeVert.GPTArt(self, crazy, prompt = AdvanceArtPrompt , Plot =   c, ArtFormat = AdvanceArtFormat,
                                                  sys_prompt=up.system_Text_Art_PictureBook)


                except:
                    print('Error Trying to make original Painting')
                    Art_PromptCharo = 'Make a work of art based on the following text:' + Art_PromptForDALLE

                Art_PromptCharoo = Art_PromptCharo
                # ArtPrompt_Mondevert = + Result
                # ArtPrompt_Mondevert = ArtFormatAdvance + Result

                try:
                    if '|' in Art_PromptCharo:
                        Art_PromptCharo = Art_PromptCharo.split('|')
                    else:
                        Art_PromptCharo = Art_PromptCharo.replace("Page:", "|")
                        Art_PromptCharo = Art_PromptCharo.replace("Page", "|")


                    if '|' in Art_PromptCharo:
                        Art_PromptCharo = Art_PromptCharo.split('|')
                    else:
                        Art_PromptCharo = Art_PromptCharo.replace(":", "|")
                        Art_PromptCharo = Art_PromptCharo.split('|')

                    # Test_Skel1 = Result.replace("Page:", "#")
                    # Test_Skel1 = Test_Skel1.replace("Page", "#")

                    print('Art_PromptCharo')
                    print(Art_PromptCharo)



                except:
                    print('Could not properly split out the string for multiple illistrations')


                KillSwitch =0
                KeepGoing = False
                while KeepGoing == False and KillSwitch < 6:
                    try:
                        Characters = []
                        x = len(Art_PromptCharo)
                        for c in range(0, x):
                            # print(c)
                            Characters.append(Art_PromptCharo[c])

                        print('Characters')
                        print(Characters)

                        for c in Characters:
                            print(' ')
                            print('-------------------------------------------')
                            print('-------------------------------------------')
                            print("Character: " + Art_PromptForDALLE + c)
                            print('-------------------------------------------')
                            #print("Character Art1: " + up.Character_Art1 + c)
                            ArtPath2 = MondeVert.makeArt(self, Art_PromptForDALLE + c)
                            ArtPaths.append(ArtPath2)
                            print(ArtPath2)


                    except:
                        print('Error Trying to create multiple paintings')
                        KillSwitch += 1
                    KeepGoing = True



            Art_PromptCharo

            prompt = data + "Work of Art Inspiration(s):" + Art_PromptForDALLE + "     Work of Art Inspiration(scene by scene):"+ Art_PromptCharoo
            Prompts_Used = [
                up.system_Text + up.RolePlay_SongArtist + "AI Created a Persona shown as the Artist Bio Above" + up.Song_prompt + up.ArtPrompt_SongArtist]
            ArtistPoetInfo = 'Lyrics Written By: ' + up.Song_Writer + '      (' + 'Artwork by: ' + up.AI_ArtistName + ')'
            MondeVert.NamePoemSavePoem(self, prompt, ArtPaths, Prompts_Used, ArtistPoetInfo, title=Title,
                                           FolderPath=SavePath1+'\\' + Mode, Mode= Mode)
            KillSwitch = 7




    def GPTSummary(self, crazy='', prompt=up.GPT_Summary, Plot='',
                   sys_prompt=up.system_Text_ScreenPlay0 + up.system_Text_ScreenPlay3):

        openai.api_key = API_Key

        if crazy != '':
            crazy = round((randbelow(520000) + 170000) / 100000, 0)
            crazy = crazy / 10

        print(crazy)

        response = openai.ChatCompletion.create(
            model="gpt-3.5-turbo",
            messages=[
                {"role": "system", "content": sys_prompt},
                {"role": "user", "content": prompt + Plot}

            ], temperature=crazy
        )

        Skeleton_Story = response.choices[0].message.content
        return Skeleton_Story

    # Every Scene I should come up with pictures to show (find things to trigger a picture getting created and multi-thread it so it does not stop anything, I should have it nice and output within the App if possible
    # Also I should have the ability to assign a voice to a given character and keep that for the audio transcript/recording. Also provide the Raw Words.
    # Snatch prequel
    # Big Lewbowski Prequel and Sequel
    # Something about Mary Sequel (add my own ideas)
    # Ancient egyption culture and other cultures
    # #(have the different technologies showing how they use sound and frequencies to move giant stones underground. Large caves and artifical light makes the world seem like another planet,
    # #have it seem like a distant planet etc, also have some of the extinct animals and animals/plants we never heard of. Describe landscape in vivid detail and make the reader fall in love with the lore.
    # Twist is they end up being underground and when humans are dealing with global warming we break in and accidentally become slaves.
    #   eventually we try to become a unified people as they realize we made the same mistakes as them
    #    #Write the MetaVerse Elon Musk Story -- This is going to be a video game too(like pokemon)
    def GPTArt(self, crazy='', prompt=up.ArtPrompt_ScreenPlay , Plot='Pick a random subject and medium go wild and make it exciting, beautiful and shocking' ,ArtFormat = up.MondeVert_ArtFormat,
               sys_prompt=up.system_Text_Artist_ScreenPlay):
        openai.api_key = API_Key

        Art_Prompt1 = openai.ChatCompletion.create(
            model="gpt-3.5-turbo",
            messages=[
                {"role": "system", "content": sys_prompt},
                {"role": "user", "content":  prompt + Plot},
                {"role": "user", "content": ArtFormat}
            ], temperature=crazy
        )

        GPTARTPROMPT = Art_Prompt1.choices[0].message.content
        return GPTARTPROMPT

    #
    # $$$$$Savetime

    # Give artist a profile and save it for use - make them have a set style that is repeated for the entire book (will be great for gothic and kid stories)

    # 1). Adapt for Movie/Book/Short Story/Indie Film
    #     2). Enhance the tool so you have option to pass in subject (by Default choose a random genre
    # 3). Need to build a google search version where I ask open ended questions and it summarizes and provides useful links for me to explore
    # 4). Coding version of this and start to use it
    # 5). Update Rap one for me and make it indie etc.
    # 6). create a version that does advertisements
    # 7). Animation
    # 8). Take all of the notes I add and make a story
    # 9). Make a tool that starts with basic idea, makes song.... book, short story, Movie
    # 10. Make tool post to instagram and twitter (do this asap and start doing it)
    # 11). Have a folder of pictures that it goes through and 3 times a day or something posts to my sites. Then it moves to another folder. Take poetry and maybe post it with all nuts and bolts ready
    # 12). For sampler, maybe build a tool to go out and download songs from itunes when it suggests it so I can mess around
    # 13). Tabs in and out (edit sounds to be other instruments)
    # 14). Edit the voice output so I can text to speech sound better
    # 15). make a play out of the voices we already have??
    # 16 get picture working with comma separated list.

    def Make_a_ScreenPlay(self, System = '',Mode='ScreenPlay'):
        Summary_Episodes = []
        ScreenPlay = []
        i = 0
        data = []
        Outline_Next2_Ref2 = ''
        Outline_Details = ''
        Plot = ''
        Add_IN = ''
        Summary_Summer_90s = ''
        filename = []
        filename1 = []
        Outline_Next2 = ''
        Outline_Next = ''
        self.Mode = 'ScreenPlay'
        ArtPaths = []
        openai.api_key = API_Key
        Plot = ""
        crazy = round((randbelow(520000) + 170000) / 100000, 0)
        crazy = crazy / 10
        if crazy < .3:
            crazy += .1

        print('Crazy Factor: ' + str(crazy))

        UserInput = False

        if UserInput == False:
            Subject = up.RandomGenre()
        else:
            Subject = UserInput
        print('Subject: ' + Subject)

        # Create original details
        response = openai.ChatCompletion.create(
            model="gpt-3.5-turbo",
            messages=[
                {"role": "system", "content": up.system_Text_ScreenPlay0 + up.system_Text_ScreenPlay2},
                {"role": "user", "content": up.Direction_ScreenPlay_Detail},
                {"role": "user", "content": up.Direction_ScreenPlay_Detail2},
                {"role": "user", "content": up.Direction_ScreenPlay_Detail3 + Subject},
                {"role": "user", "content": up.Setting_ScreenPlay},
                {"role": "user", "content": up.Characters_ScreenPlay},
                # {"role": "user", "content": up.Title_ScreenPlay}

            ], temperature=crazy
        )

        Skeleton_Story1 = response.choices[0].message.content

        # Create original details
        response = openai.ChatCompletion.create(
            model="gpt-3.5-turbo",
            messages=[
                {"role": "system", "content": up.system_Text_ScreenPlay0 + up.system_Text_ScreenPlay2},
                {"role": "user", "content": up.Title_ScreenPlay + Skeleton_Story1}

            ], temperature=crazy
        )

        Skeleton_Story2 = response.choices[0].message.content

        Skeleton_Story = Skeleton_Story2 + Skeleton_Story1

        print('***********************************************************************************')
        print(Skeleton_Story)
        print('***********************************************************************************')

        # MondeVert.speak(self, Skeleton_Story)

        # Create full outline
        response = openai.ChatCompletion.create(
            model="gpt-3.5-turbo",
            messages=[
                {"role": "system", "content": up.system_Text_ScreenPlay0 + up.system_Text_ScreenPlay4},
                {"role": "user", "content": up.Create_Outline1 + Skeleton_Story},

                # {"role": "user", "content": up.Title_ScreenPlay}

            ], temperature=crazy
        )

        Outline_All = response.choices[0].message.content

        print(Outline_All)

        print('***********************************************************************************')
        print('***********************************************************************************')

        # This is how I set up the prompts for the show, I could randomize when to activate red herrings etc based on Add-In Feature
        Direction_ScreenPlay = []
        Direction_ScreenPlay.append(up.Direction_ScreenPlay_Pilot)
        Direction_ScreenPlay.append(up.Direction_ScreenPlayMiddle)
        Direction_ScreenPlay.append(up.Direction_ScreenPlayMiddleLate)
        Direction_ScreenPlay.append(up.Direction_ScreenPlayMiddleLate)
        Direction_ScreenPlay.append(up.Direction_ScreenPlayFinal)

        ScreenPlay_Final = ''
        Episode_Count = 5
        for i in range(0, Episode_Count):
            Episode_Direction = Direction_ScreenPlay[i]

            Add_IN = ''
            # if i == 1:
            #     Add_IN = up.RedHerring_ScreenPlay
            # elif i == 2:
            #     Add_IN = up.UnsungHero + up.Twist_ScreenPlay
            # elif i == 3:
            #     Add_IN = up.Mystery_Character_Past + up.VillanOrigin

            # Prior to writing an episode we should create the outline for the episode based on

            if i != 0:

                # Create outline for next 2 episodes
                response = openai.ChatCompletion.create(
                    model="gpt-3.5-turbo",
                    messages=[
                        {"role": "system", "content": up.system_Text_ScreenPlay0 + up.system_Text_ScreenPlay4},
                        {"role": "user", "content": up.Compare_Detail_v_Summary1 + Plot},
                        {"role": "user", "content": up.Compare_Detail_v_Summary2 + Outline_All},
                        {"role": "user", "content": up.Compare_Detail_v_Summary3},

                        # {"role": "user", "content": up.Title_ScreenPlay}

                    ], temperature=crazy
                )

                Outline_Next2 = response.choices[0].message.content

                # Store this to maybe use later as a way to keep story moving together well

                # Revise
                response = openai.ChatCompletion.create(
                    model="gpt-3.5-turbo",
                    messages=[
                        {"role": "system", "content": up.system_Text_ScreenPlay0 + up.system_Text_ScreenPlay4},
                        {"role": "user", "content": up.Compare_Detail_v_Summary_Outline_New + Outline_Next2},
                        {"role": "user", "content": up.Compare_Detail_v_Summary_Revise},
                        {"role": "user", "content": up.Compare_Detail_v_Summary_Summary_New + PriorScreenPlay},
                        # {"role": "user", "content": up.Title_ScreenPlay}

                    ], temperature=crazy
                )

                Outline_Next = response.choices[0].message.content

            else:
                # if i = 0 want to create an outline for next episode

                # Create outline for next 2 episodes
                response = openai.ChatCompletion.create(
                    model="gpt-3.5-turbo",
                    messages=[
                        {"role": "system", "content": up.system_Text_ScreenPlay0 + up.system_Text_ScreenPlay4},
                        {"role": "user", "content": up.Compare_Detail_v_Summary2 + Outline_All},
                        {"role": "user", "content": up.Outline_Pilot}

                        # {"role": "user", "content": up.Title_ScreenPlay}

                    ], temperature=crazy
                )

                Outline_Next = response.choices[0].message.content

            print('  ')
            print('  ')
            print('  ')
            # print('***********************************************************')
            # print('***********************************************************')
            print('***********************************************************')
            print('***********************************************************')
            print('***********************************************************')
            print('  ')
            print('  ')
            print('  ')
            print('Episode ' + str(i + 1))
            Add_IN = ''

            Song1 = openai.ChatCompletion.create(
                model="gpt-3.5-turbo",
                messages=[
                    {"role": "user", "content": up.system_Text_ScreenPlay + up.system_Text_ScreenPlay1},
                    {"role": "system", "content": up.Tie_In_ScreenPlay + Skeleton_Story},
                    {"role": "system", "content": up.Tie_In_ScreenPlay2 + Outline_Next},
                    {"role": "system", "content": up.Direction_ScreenPlayALL},
                    {"role": "system", "content": up.Direction_ScreenPlayALL2},
                    {"role": "system", "content": up.Direction_ScreenPlayALL3},
                    ##--------------------Common Above/Unique Below----------
                    # {"role": "system", "content":  Episode_Direction+ Add_IN},
                    {"role": "system", "content": Episode_Direction},
                    ##--------------------Dialog----------
                    {"role": "system", "content": up.Dialog_ScreenPlay}
                ], temperature=crazy)
            ScreenPlay0 = Song1.choices[0].message.content
            ScreenPlay.append(ScreenPlay0)

            print(ScreenPlay[i])

            Summary_Episode1 = MondeVert.GPTSummary(self, crazy, up.Summarize_ScreenPlay, ScreenPlay[i])
            Summary_Episodes.append(Summary_Episode1)

            PriorSummary = Summary_Episodes[i]
            PriorScreenPlay = ScreenPlay[i]
            PriorNext2 = Outline_Next2
            PriorOutline = Outline_Next

            # print('***********************************************************')
            # print("Summary Episode " + str(i + 1))
            # print(Summary_Episodes[i])
            Plot += Summary_Episodes[i]
            ScreenPlay += "EPISODE " + str(i + 1) + ScreenPlay[i]
            Outline_Details = Skeleton_Story + Outline_All
            Outline_Next2_Ref2 += Outline_Next2

        Title = MondeVert.Get_Title_GPT(self,Plot)

        current_time1 = datetime.datetime.now()
        current_time2 = current_time1.strftime('%m-%d-%Y_%H.%M.%S')
        # f2 = FolderPath
        f2 = up.SavePath
        SavePath1 = f2
        invalidCharRemoved = re.sub(r"[^a-zA-Z0-9 ]", "", Title)
        folder = Mode
        if len(folder) > 44:
            folder = folder[0:44]
        Title1 = r"\\" +  str(invalidCharRemoved) + "_"
        Title11 =  str(invalidCharRemoved) + "_"
        if len(Title1) > 44:
            Title1 = Title1[0:44]
        SavePath2 = SavePath1 + r"\\"+ folder

        filename1.append('Screen Play - ' + Title11)
        filename1.append('Details_Outline - ' + Title11)
        filename1.append('OutlineNext2_Ref - ' + Title11)
        filename1.append('Summaries - ' + Title11)

        print('Review the filnames:')
        print(filename)

        data.append([ScreenPlay])
        data.append([Outline_Details])
        data.append([Outline_Next2_Ref2])
        data.append([Plot])

        for x in range(0, 4):
            # print(data)
            try:

                filename = MondeVert.SaveCSV(data[x], filename1[x], SavePath2)
                # MondeVert.SaveText(self,df1,'MondeVert Assistant', 'Full Transcript')

                # MondeVert.add2Master2(df1)
            except:
                print('Review Error File did not save ')

        try:
            send_email_w_attachment_gmail(up.to, up.subject, Skeleton_Story, filename)
            print('Gmail Email Sent ')
            send_email_w_attachment_outlook(up.to, up.subject, Skeleton_Story, filename)
            print('Outlook Email Sent ')

        except:
            print('email not send, its possible file was not created')

        try:
            Remove_l = (-1) * len(Summary_Episodes[4] + Summary_Episodes[5])
            Plot_NoSpoiler = Plot[Remove_l:]
            print(Plot_NoSpoiler)
        except:
            print('error on no spoiler')
        #MondeVert.speak(self, ScreenPlay_Final)

        try:
            Summary_Summer_90s = MondeVert.GPTSummary(self, crazy, up.GPT_90s_Summary, Plot_NoSpoiler)
        except:
            print('error on 90s trailer')

            print('***********************************************************')
            print('***********************************************************')
            print('***********************************************************')
            print('***********************************************************')
            print("90s Movie Preview: ")
            print(Summary_Summer_90s)
            #MondeVert.speak(self, Summary_Summer_90s)

            print('***********************************************************')
            print('***********************************************************')
            print('***********************************************************')
            print('***********************************************************')
            print("Character Art prompts ")

            # ''
            # if this does not work I can make replace that text with the char so it works
            # Test_Skel1 = Skeleton_Story.replace(":", "")

            print('Skeleton_Story')
            print(Skeleton_Story)


            if '|' in Skeleton_Story:
                Test_Skel = Skeleton_Story.split('|')
            else:
                Test_Skel = Skeleton_Story.replace("Characters:", "|")
                Test_Skel = Test_Skel.replace("Characters", "|")
                Test_Skel = Test_Skel.split('|')




            print('Test_Skel')
            print(Test_Skel)

            # Test_Skel = Test_Skel.split('')
            # print('Test_Skel')
            # print(Test_Skel)
            Details = Test_Skel[0]




            Characters = []
            x = len(Test_Skel)
            for c in range(0, x):
                # print(c)
                if c != 0:
                    Characters.append(Test_Skel[c])

            print('Characters')
            print(Characters)

            for c in Characters:
                print(' ')
                print('-------------------------------------------')
                print('-------------------------------------------')
                # print("Character: " + c)
                print('-------------------------------------------')
                print("Character Art1: " + up.Character_Art1 + c)
                Art_PromptChar = MondeVert.GPTArt(self, crazy, up.Character_Art2 + c, Plot = Skeleton_Story)

                print("Character Art2: " + up.Character_Art2 + c)
                Art_PromptChar2 = MondeVert.GPTArt(self, crazy, up.Character_Art2 + c, Plot = Skeleton_Story)

            try:
                #MondeVert.speak(self, Art_PromptChar)
                print(Art_PromptChar)
                ArtPath2 = MondeVert.makeArt(self, Art_PromptChar)
                ArtPaths.append(ArtPath2)
            except:
                dn = ''

            try:
                #Art_PromptChar2)
                print(Art_PromptChar2)
                ArtPath2 = MondeVert.makeArt(self, Art_PromptChar2)
                ArtPaths.append(ArtPath2)
            except:
                dn = ''

            print('***********************************************************')
            print('***********************************************************')
            print('***********************************************************')
            print('***********************************************************')
            print("Art prompts ")

            try:
                Art_Prompt = MondeVert.GPTArt(self, crazy,up.direction_Text_Artist_ScreenPlay + up.ArtPrompt_ScreenPlay,
                                                  Skeleton_Story,
                                                  sys_prompt=up.system_Text_Artist_ScreenPlay + up.system_Text_ScreenPlay_Art1)
                #Art_Prompt)
                print('***********************************************************')
                print("Art prompt Movie Poster ")
                print(Art_Prompt)
                try:
                    ArtPath2 = MondeVert.makeArt(self, Art_Prompt)
                    ArtPaths.append(ArtPath2)
                except:
                    d = ''
            except:
                ArtPath2 = ''
            # Art_Prompt = Art_Prompt11.choices[0].message.content

            for x in range(0, 5):

                Summary_Episode = Summary_Episodes[x]

                try:
                    Art_Prompt1 = MondeVert.GPTArt(self, crazy,
                                                       up.direction_Text_Artist_ScreenPlay + up.ArtPrompt_ScreenPlay_Scene,
                                                       Summary_Episode,
                                                       sys_prompt=up.system_Text_Artist_ScreenPlay + up.system_Text_ScreenPlay_Art1 + up.system_Text_ScreenPlay_Art)
                except:
                    print('Error did not work but kept going')

                try:

                    # Test_Skel1 = Skeleton_Story.replace("Art Prompts:", "")
                    # Test_Skel1 = Test_Skel1.replace("Art Prompts", "#")
                    # ''

                    if '|' in Skeleton_Story:
                        Test_Skel = Skeleton_Story.split('|')
                    else:
                        Test_Skel = Skeleton_Story.replace("Art Prompt:", "|")
                        Test_Skel = Test_Skel.replace("Art Prompt", "|")
                        Test_Skel = Test_Skel.split('|')


                    #Test_Skel = Art_Prompt1.split('#')

                    Art_Prompts_pre = []
                    Art_Prompts = []
                    x = len(Test_Skel)
                    for c in range(0, x):
                        # print(c)
                        if c != 0:
                            Art_Prompts_pre.append(Test_Skel[c])

                    print(Art_Prompt1)
                    Test_Skel = Art_Prompt1.split(';')
                    print("Test_Skel: " + Test_Skel)

                    x = len(Test_Skel)
                    for c in range(0, x):
                        Art_Prompts.append(Test_Skel[c])

                    for ap in Art_Prompts:
                        #MondeVert.speak(self, ap)
                        print('***********************************************************')
                        print("Art prompt Movie Poster ")
                        print(ap)
                        ArtPath2 = MondeVert.makeArt(self, ap)
                        ArtPaths.append(ArtPath2)

                except:
                    print("New stuff failed yo")

            # Art_Prompt2 = MondeVert.GPTArt(self, crazy, up.direction_Text_Artist_ScreenPlay + up.ArtPrompt_ScreenPlay_Scene, Summary_Episode2, sys_prompt= up.system_Text_Artist_ScreenPlay + up.system_Text_ScreenPlay_Art1 +up.system_Text_ScreenPlay_Art)
            # Art_Prompt3 = MondeVert.GPTArt(self, crazy, up.direction_Text_Artist_ScreenPlay + up.ArtPrompt_ScreenPlay_Scene, Summary_Episode3, sys_prompt= up.system_Text_Artist_ScreenPlay + up.system_Text_ScreenPlay_Art1 +up.system_Text_ScreenPlay_Art)
            # Art_Prompt4 = MondeVert.GPTArt(self, crazy, up.direction_Text_Artist_ScreenPlay + up.ArtPrompt_ScreenPlay_Scene, Summary_Episode4, sys_prompt= up.system_Text_Artist_ScreenPlay + up.system_Text_ScreenPlay_Art1 +up.system_Text_ScreenPlay_Art)
            # Art_Prompt5 = MondeVert.GPTArt(self, crazy, up.direction_Text_Artist_ScreenPlay + up.ArtPrompt_ScreenPlay_Scene, Summary_EpisodeFinal,sys_prompt= up.system_Text_Artist_ScreenPlay + up.system_Text_ScreenPlay_Art1 +up.system_Text_ScreenPlay_Art)
            #

        # ScreenPlay = "EPISODE I      :" +  ScreenPlay[0] +"               EPISODE II: " +  ScreenPlay2 + "       EPISODE III: " + ScreenPlay3 +"         EPISODE IV:" +  ScreenPlay4+"         EPISODE V:" +  ScreenPlayFinal

        # Title = 'SD - MondeVert Productions - Classified'

        if len(Title) > 45:
            Title = Title[0:45]

        ArtPrompt_Combined = ''

        for f in Art_Prompt:
                ArtPrompt_Combined +=f

        prompt = " 90s Summer Preview: " + Summary_Summer_90s + " Details:" + Skeleton_Story + 'Summary: ' + Plot + "Work of Art Inspiration:" + ArtPrompt_Combined + "      Song:" + ScreenPlay
        Prompts_Used = [
            up.system_Text + up.RolePlay_SongArtist + "AI Created a Persona shown as the Artist Bio Above" + up.Song_prompt + up.ArtPrompt_SongArtist]
        ArtistPoetInfo = 'Lyrics Written By: ' + up.Song_Writer + '      (' + 'Artwork by: ' + up.AI_ArtistName + ')'
        MondeVert.NamePoemSavePoem(self, prompt, ArtPaths, Prompts_Used, ArtistPoetInfo, title=Title,
                                       FolderPath=up.AI_Screen_Plays + '\\' + Mode, Mode='Screen Play' + Mode)

    def Add2MasterLyrics(self, current_time2, Mode, title, ArtistPoetInfo, poem, Tag, SavePath2, Prompts):
        data = [(current_time2, Mode, title, ArtistPoetInfo, poem, Tag, SavePath2, Prompts)]
        df = pd.DataFrame(data, columns=['Date_Time_Added', 'Art_Type', 'Title', 'Poet_Artist_Info', 'Poem_Song_Lyrics',
                                         'Quality', 'Folder_Path', 'Prompts_Used'])
        MondeVert.add2Master3(df)

    def quickArt(self):
        prompt = MondeVert.ChatGPTDA(self, MakeArt=True, Prompt=(up.QuickArt))

        ArtPaths = []
        try:
            ArtPath = MondeVert.makeArt(self, prompt)
            ArtPaths.append(ArtPath)
        except:
            #MondeVert.speak(self, 'Could not make the Art due to an error')
            print( 'Could not make the Art due to an error')

        Prompts_Used = [str('Quick_Art_Prompt: ' + up.QuickArt)]
        ArtistPoetInfo = 'Written By: ' + up.Bot_Name + '      (' + 'Artwork by: ' + up.AI_ArtistName + ')'
        MondeVert.NamePoemSavePoem(self, prompt, ArtPaths, Prompts_Used, ArtistPoetInfo, title='Quick_Art')





    def SaveCSV(Text, Title, FolderPath):
        current_time1 = datetime.datetime.now()
        current_time2 = current_time1.strftime('%m-%d-%Y_%H.%M.%S')
        f2 = FolderPath
        SavePath1 = f2
        invalidCharRemoved = re.sub(r"[^a-zA-Z0-9 ]", "", Title)
        folder = str(invalidCharRemoved)
        if len(folder) > 44:
            folder = folder[0:44]
        Title1 = '\\' + str(invalidCharRemoved) + "_"
        if len(Title1) > 44:
            Title1 = Title1[0:44]
        SavePath2 = SavePath1 + '\\' + folder
        Title2 = SavePath2 + Title1 + current_time2 + '.txt'

        isExist = os.path.exists(SavePath1)
        if not isExist:
            # Create a new directory because it does not exist
            os.makedirs(SavePath1)
            print("Created new filePath: " + SavePath1)

        isExist = os.path.exists(SavePath2)
        if not isExist:
            # Create a new directory because it does not exist
            os.makedirs(SavePath2)
            print("Created new filePath: " + SavePath2)
        data = [(current_time2, Text)]
        print(data)
        try:
            df1 = pd.DataFrame(data, columns=["TimeStamp", "Transcript"])
            print(df1)

            df1.to_csv(Title2 )
            # MondeVert.SaveText(self,df1,'MondeVert Assistant', 'Full Transcript')
            MondeVert.add2Master2(df1)
        except:
            print('Review Error File did not save ')

        return Title2



    def Get_Title_GPT(self,  Text, Role = '',mood = '', bio = '', crazy = .5):

        # Create original details
        response = openai.ChatCompletion.create(
            model="gpt-3.5-turbo",
            messages=[
                {"role": "system", "content": up.system_TextJoaT},
                #{"role": "user", "content": up.Role_Play_Prompt + Role},
                {"role": "user", "content": up.MondeVert_Title + Text}

            ], temperature=crazy
        )

        Title = response.choices[0].message.content


        Title = Title.replace("Title:", "")
        Title = Title.replace("Title", "")
        return Title

    def NamePoemSavePoem(self, poem, ArtPaths, Prompts_Used, ArtistPoetInfo, title='', FolderPath=up.AI_Poetry_Path,
                         Mode='Poem'):
        dfPrompts = ''
        Tag = ''
        if Mode == 'Poem':
            title_p = up.Poem_Title_Prompt
        else:
            title_p = up.Song_Title_Prompt
        try:
            if title == '':
                Title = MondeVert.Get_Title_GPT(self,  Text= poem )
            else:
                Title = title

        except:
            print('error could not get title')

        if Title == '':
            Title = 'AI - No Name Project'

        current_time1 = datetime.datetime.now()
        current_time2 = current_time1.strftime('%m-%d-%Y_%H.%M.%S')
        f2 = FolderPath
        SavePath1 = f2
        invalidCharRemoved = re.sub(r"[^a-zA-Z0-9 ]", "", Title)
        folder = str(re.sub(r"[^a-zA-Z0-9 ]", "", Mode))
        if len(folder) > 44:
            folder = folder[0:44]
        Title1 = '\\' + str(invalidCharRemoved) + "_"
        if len(Title1) > 60:
            Title1 = Title1[0:60]
        SavePath2 = SavePath1 + '\\' + folder
        Title2 = SavePath2 + Title1 + current_time2

#not using for now
        # Tag = ""
        # Tag += MondeVert.YayorNay(self)
        # Title2 = Title2 + Tag

        isExist = os.path.exists(SavePath1)
        if not isExist:
            # Create a new directory because it does not exist
            os.makedirs(SavePath1)

        isExist = os.path.exists(SavePath2)
        if not isExist:
            # Create a new directory because it does not exist
            os.makedirs(SavePath2)


#Adds to master files and also creates its own CSV file


        data = Title +  ArtistPoetInfo + poem +  FolderPath+  Mode +  Tag
        # print(dat
        data = [(current_time2, data)]
        print(data)
        try:
            df1 = pd.DataFrame(data, columns=["TimeStamp", "Transcript"])
            print(df1)
            print(Title2)
            df1.to_csv(Title2 + '.txt')
            MondeVert.Add2MasterLyrics(self, current_time2, Mode, title, ArtistPoetInfo, poem, Tag, SavePath2,dfPrompts)
            # MondeVert.SaveText(self,df1,'MondeVert Assistant', 'Full Transcript')
            try:
                MondeVert.add2Master2(df1)

            except:
                print('Did not add to master, error')
        except:
            print('Review Error')



            try:
    #Word Doc Creation below, not quite what I need so I do text file too above
                document = Document()
                document.add_heading(Title, 0)
                document.add_heading(ArtistPoetInfo, 1)
                p = document.add_paragraph()
                r = p.add_run()
                # r.add_text(ArtistPoetInfo)
                p = document.add_paragraph()
                r = p.add_run()
                r.add_text(poem)
                p = document.add_paragraph()
                r = p.add_run()
                image_counter = 0
                for i in ArtPaths:
                    image_counter += 1
                    r.add_picture(i)
                    original = i
                    target = SavePath2 + Title1 + str(image_counter) + current_time2 + '.png'
                    shutil.copyfile(original, target)


                    IP.png2JPG(original, up.PNGPath, Title1, up.PNGPath_Archive, Del=False)

                document.add_heading('AI Prompts used: ', 7)

                for i in Prompts_Used:
                    p = document.add_paragraph()
                    r = p.add_run()
                    r.add_text(i)
                    dfPrompts += i + '|'
                document.save(Title2 + '_Details.docx')

                document2 = Document()
                document2.add_heading(Title, 0)
                document2.add_heading(ArtistPoetInfo, 1)

                p = document2.add_paragraph()
                r = p.add_run()
                r.add_text(poem)
                document2.save(Title2 + '.docx')
            except:
                print('Error saving word doc')
    # Blog Make Art Live mode



#test this to see how to generate images from another one
    def animate(self, image, Prompt):
        response = openai.Image.create_edit(
            image=open("A:\AI Art\Approved Quality Art\Favorites\Art_Blunts_Using_the_classic_style_of_Pablo_Picasso_create_a_unique_work_of_art_that_celebrates_the_creative_combination_of_Snoop_Dogg_and_Backwood.png", "rb"),
            mask=open("mask.png", "rb"),
            prompt="A sunlit indoor lounge area with a pool containing a flamingo",
            n=1,
            size="1024x1024"
        )
        image_url = response['data'][0]['url']



#
# from multiprocessing.dummy import Pool as ThreadPool
#
# # urls = [
# #   'http://www.python.org',
# #   'http://www.python.org/about/',
# #   'http://www.onlamp.com/pub/a/python/2003/04/17/metaclasses.html',
# #   'http://www.python.org/doc/',
# #   'http://www.python.org/download/',
# #   'http://www.python.org/getit/',
# #   'http://www.python.org/community/',
# #   'https://wiki.python.org/moin/',
# # ]
#
# # Make the Pool of workers
# pool = ThreadPool(3)
#
# # Open the URLs in their own threads
# # and return the results
# results = pool.map(urllib2.urlopen, urls)
#
# # Close the pool and wait for the work to finish
# pool.close()
# pool.join()
#









# def run_the_command(index, command_array):
#     try:
#         command_array = ...
#         number_of_commands = len(command_array)
#
#         exec
#         command_array[index]
#     except:
#         print('Error but powering through it')

#
# def run_the_command(index):
#     exec
#     command_array[index]


#def SHAINEBootUP(Functions, Args = ['Basic'], MultiT = False, PostInsta = False):
def SHAINEBootUP(Args = 'Basic'):
        dummy = 1
        if dummy == 1:
        #try:
            x = MondeVert()
            x.MondeVertMenu_up(Mode = Args)
        # except:
        #     print('Error unable to pull the value from the ARGs parameter')

if __name__ == '__main__':
    # main method for executing
    # the functions
    Record = ''




    #args = ['JobDescription']
    #args = ['Resume_Consolidate_Old']
    #args = ['Resume_Consolidate_Old','Resume', 'Resume_Combine_old_new']
    #args = ['Resume_Combine_old_new']

    #args = ['ReSearch3', 'ReSearch',  'ReSearch2']
    #args = ['ReSearch2']
    #args = ['PictureBook']
    #args = ['Freelance_Services', 'JobDescription','Resume']
    #args = ['Freelance_Services']
    #args = ['Resume']

    #args = ['LinkedIn', 'Create_Persona_Writer', 'Resume_Review', 'Basic']
    args = ['MondeVert_Audio_Video_Story']
    args.append(args)
    args.append(args)



    number_of_commands = len(args)

    threads = []
    for i in range(number_of_commands):
        ArgX = args[i]
        #Dummy line
        i2 = 1

        if i2 == 1:
        #try:

            print("Start Thread " + str(i))
            t = threading.Thread(target=SHAINEBootUP, args = (ArgX,)).start()
            threads.append(t)

        # except:
        #     print('Error - Could not start new thread')

 # Wait for all of them to finish
 #    for t in threads:
 #        t.join()
    print("Job Complete!")




# IP.upload_pictures()
    #Image1 = x.animate(image= r"A:\AI Art\DALLE_please_create_an_art_piece_that_showcases_a_luxurious_and_modern_real_estate_property_The_property_should_have_a_stunning_view_and_be_surroun.png", Prompt = "Add a Stick Figure with long blonde hair to the following image")
    #Image1 = x.Take_query(image = Image1, Prompt= 'Make a variation where everything is the same but you put the stick figure in a slightly different position')
    # Testing
    # x.Make_a_ScreenPlay_Testing()
    # x.Make_a_ScreenPlay()
    # x.saveTranscript()




