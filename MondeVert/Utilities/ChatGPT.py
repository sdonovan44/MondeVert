import glob
import datetime
import pandas as pd
import wikipedia
global Record
import sys
import random
from secrets import randbelow
import Instagram_Posts as IP
import ShakesBot as s
import Common_Sayings as cs
from threading import Event
import User_Prefs as up
from gingerit.gingerit import GingerIt
import threading
import numpy
from docx import Document
import re
# from exceptions import PendingDeprecationWarning
import requests
import cloudscraper
import tensorflow
import selenium
Record = ''
#import urllib2
import webbrowser
from fpdf import FPDF
import User_Prefs as up
import os
import platform
import subprocess
import requests
import random
import string
import openai
import pyttsx3
# Import the speech recognition library
import speech_recognition as sr
from dotenv import load_dotenv
from DoNotCommit import API_Key
import DoNotCommit as DNC
import atexit
from docx import Document
from docx.shared import Inches

import shutil

import smtplib
import os
from email.mime.text import MIMEText
from email.mime.multipart import MIMEMultipart
from email.mime.application import MIMEApplication
from email.header import Header

# from email_config import gmail_pass, user, host, port


URL = "..."  # noqa
API_KEY = "..."


# make a transcript class that takes down all of my requests and saves to a text file
# have specific command for the assistant to save in a certain folder (always add date & time to the file
# create things like to do lists
# ideas
# stories and dialogue
# the key is being able to store it properly and sytemically access as well
# every program I create should be executeable through the assistant


# Urgent to do
# Make it a class
# have transcribe and also proofread (later) add method where in the middle of dictating you can add titles so it gets properly makred eventually
# bullets etc. parenthesis
# Have it so every pause puts a new line in between so its easy to pull out pauses etc.
# Most important is to have a mode where it callse poembot and reads the lines to me.
# .... also I would like for it to be a mini drama where I get multiple voices and characters that read
# the output etc eventually learn how to animate/make a nice script.
# create a series so eventually we might have a way to feed in characters to poem bot so its not just random like it is now lol

# 12/18/2022 get the assistant able to open the respective web browsers
# set up mouse so it has shortcuts
# be able to run on both mac and also the pc
# import the site view bot
# update the youtube bot so it hides ips properly and also so it loops through videos to not draw attention
#    (make it so it starts off slow then has pockets of 100 here and there big days occasionally and find out how to make it work properly
# have assistant able to run those on command (buy wifi adapter so I can have it running whenever & get new monitor for that too


# Update transcribe commands (Review file so far,#Have a separate function to make the title (different than subject)
##Also I should be able to mark the file as important based on results (add tag)
# Make a program that goes through and corrects the script into a new tab (puts subject as start of the word blocks (with other data in parenthesis for tagging)
# Also make it so that the poems do not overwrite but get concatenated together
# Make a theatre method that uses multithreading to do stuff
# overall threads should allow this to run faster
# Make it so I can update all meta data and also add more metadata
# Goal, mark as header (prior or next maybe)






def send_email_w_attachment_outlook(to, subject, body, filename):
    sender = DNC.OutlookEmail
    receivers = to
    #smtp
    smtpHost = 'smtp.office365.com'
    smtpPort = 587
    password = "youremailpassword"
    subject = "outlook email test"
    # Add the From: and To: headers at the start!
    message = ("From: %s\r\nTo: %s\r\nSubject: %s\r\n\r\n"
           % (sender, ", ".join(receivers), subject))
    message += """This is a test e-mail message."""
    print (message)
    try:
        smtpObj = smtplib.SMTP(smtpHost, smtpPort)
        #smtpObj.set_debuglevel(1)
        smtpObj.ehlo()
        smtpObj.starttls()
        smtpObj.ehlo()
        smtpObj.login(sender,password)
        smtpObj.sendmail(sender, receivers, message)
        smtpObj.quit()
        print ("Successfully sent email")
    except:
        print ("Error: unable to send email")


def send_email_w_attachment_gmail(to, subject, body, filename):
    # create message object
    to = 'sdonovan@mondevert.co; RichardDX44@gmail.com'
    subject = 'New Song for Rich'
    message = MIMEMultipart()
    # add in header
    message['From'] = Header(up.user)
    message['To'] = Header(to)
    message['Subject'] = Header(subject)
    # attach message body as MIMEText
    message.attach(MIMEText(body, 'plain', 'utf-8'))

    # locate and attach desired attachments

    for n in filename:
        att_name = os.path.basename(filename[n])
        _f = open(filename[n], 'rb')
        att = MIMEApplication(_f.read(), _subtype="txt")
        _f.close()
        att.add_header('Content-Disposition', 'attachment', filename=att_name)
        message.attach(att)

    # setup email server
    server = smtplib.SMTP_SSL(DNC.host, DNC.port)
    server.login(DNC.gmail_user, DNC.gmail_pass)

    # send email and quit server
    server.sendmail(DNC.user, to, message.as_string())
    server.quit()


def send_email_no_attachment_gmail(to, subject, body):
    # create message object
    to = 'sdonovan@mondevert.com'
    subject = 'New Song for Rich'

    message = MIMEMultipart()

    # add in header
    message['From'] = Header(up.user)
    message['To'] = Header(to)
    message['Subject'] = Header(subject)

    # attach message body as MIMEText
    message.attach(MIMEText(body, 'plain', 'utf-8'))

    # setup email server
    server = smtplib.SMTP_SSL(DNC.host, DNC.port)
    server.login(DNC.user, DNC.gmail_pass)

    # send email and quit server
    server.sendmail(DNC.user, to, message.as_string())
    server.quit()


class MondeVert():
    def __init__(self, voice=4, language_settings=1):
        self.voice = voice
        self.language_settings = language_settings
        self.AssistantName = up.getAssistantName()
        self.UserName = up.getUserName()
        self.transcript_Final = ''
        self.API_Key = API_Key
        self.Correction_Comment = ''
        self.AI_Corrected_Content = ''
        self.SilentMode = False
        global xVoice
        global voice_set
        voice_set = self.voice
        xVoice = 1

    def Open_Web(url):
        url = url
        webbrowser.open_new_tab(url)

    def speak(self, audio, Add2T=True, voice=''):
        if voice == '':
            voice = self.voice

        if Add2T == True:
            MondeVert.Add2Transcript(self, self.AssistantName + ': ' + audio)
        engine = pyttsx3.init()
        # getter method(gets the current value
        # of engine property)
        voices = engine.getProperty('voices')
        engine.setProperty('voice', voices[voice].id)
        # Method for the speaking of the assistant
        if self.SilentMode == False:
            engine.say(audio + "")
            # Blocks while processing all the currently
            # queued commands
            engine.runAndWait()
        else:
            print('Silent Mode Activated: ' + audio)

    def speakSweet(self, audio, Add2T=True):
        if Add2T == True:
            MondeVert.Add2Transcript(self, self.AssistantName + ': ' + audio)
        engine = pyttsx3.init()
        # getter method(gets the current value
        # of engine property)
        voices = engine.getProperty('voices')
        engine.setProperty('voice', voices[4].id)
        # Method for the speaking of the assistant
        if self.SilentMode == False:
            engine.say(audio + "")
            # Blocks while processing all the currently
            # queued commands
            engine.runAndWait()
        else:
            print('Silent Mode Activated: ' + audio)

        # def RandomCharacters

    def testvoices(self):
        # import pyttsx3
        engine = pyttsx3.init()
        voices = engine.getProperty('voices')
        index = 0
        for voice in voices:
            print(f'index-> {index} -- {voice.name}')
            engine.setProperty('voice', voices[index].id)
            engine.say("Shane do you like this voice more than the last one you busy little bee turkish" + "")

            index += 1
        engine.runAndWait()

    def Setvoices(self, Quick=False):

        if Quick == False:
            engine = pyttsx3.init()
            voices = engine.getProperty('voices')
            index = 0
            indexCount = 0
            # load_dotenv()
            # Set up the OpenAI API client :
            openai.api_key = API_Key

            prompt = 'Write a short paragraph (less than 44 words) that uses different pronunciations and complex words while being warm and welcoming to a person named Shane'
            completions = openai.Completion.create(
                engine="text-davinci-003",
                prompt=prompt,
                max_tokens=2500,
                n=1,
                stop=None,
                temperature=0.5,
            )
            speak1 = 'Do you want to set this voice as active?'

            # Print the generated text
            message = completions.choices[0].text
            print(message)
            MondeVert.Add2Transcript(self, text2Add=message)
            set1 = False

            for voice in voices:
                indexCount += 1

            while set1 == False:
                if index <= indexCount:
                    print(f'index-> {index} -- {voices[index].name}')
                    engine.setProperty('voice', voices[index].id)
                    engine.say(message)
                    MondeVert.speak(self, speak1)
                    Query = MondeVert.getUserResponse(self)

                    if "yes" in Query or "set" in Query or ("make" in Query and "active" in Query) or (
                            "set" in Query and "active" in Query):
                        self.voice = index
                        voice_set = self.voice
                        xVoice = index
                        MondeVert.speak(self, 'Voice Set')
                        set1 = True

                    index += 1
                    engine.runAndWait()
                else:
                    MondeVert.Setvoices(self, Quick=True)

        else:
            self.voice = int(input('What Voice do you want to make active?'))
            MondeVert.speak(self, 'Voice Set')

        ####################################################################################################################################################################################
        ## Below are more utlities that help me to do repetitive tasks quicker
        ####################################################################################################################################################################################

    def createDF(self):
        test = [('', '', '', '', '', '', '', '', '', '', '', '', '', '', '')]
        self.transcript = pd.DataFrame(test, columns=['Entry #', 'Date', 'Subject', 'Type/Marker', 'Priority',
                                                      'Original_Text', 'Edited_Text', 'AI_Corrected_Text',
                                                      'AI_Correction_Comment', 'Note', 'AI_Content',
                                                      'AI_Correction_Content', 'File_Name', 'Added to text File',
                                                      'Completed/Added to blog'])

    def adddata2DF(self):
        current_time1 = datetime.datetime.now()
        xx = current_time1.strftime('%m/%d/%Y %H:%M:%S')
        row = (self.entry - 1)
        self.transcript.loc[row] = [str(self.entry), xx, self.subject, self.Type, self.Significance, self.query1,
                                    self.new_Query, self.AI_Corrected_Text, self.Correction_Comment, self.Notes,
                                    self.Words, self.AI_Corrected_Content, self.FileName, self.Added2TextFile,
                                    self.Completed]

    def ActivateSilentMode(self):
        self.SilentMode = True

    def ActivateLoudMode(self):
        self.SilentMode = False

    def saveTranscript(self):

        current_time1 = datetime.datetime.now()
        current_time2 = current_time1.strftime('%m-%d-%Y_%H.%M.%S')
        Filename = 'MondeVert Assistant'
        Filename = '\\' + Filename + "_"
        f2 = up.getPath()
        SavePath1 = f2
        data = [(current_time2, self.transcript_Final)]
        print(data)
        try:
            df1 = pd.DataFrame(data, columns=["TimeStamp", "Transcript"])
            # df1 = pd.DataFrame(data)
            df1.to_csv(SavePath1 + Filename + current_time2 + '.csv')

            # MondeVert.SaveText(self,df1,'MondeVert Assistant', 'Full Transcript')
            MondeVert.add2Master2(df1)
        except:
            print('Review Error')

    def Add2Transcript(self, text2Add):
        self.transcript_Final += text2Add + ' \n'

    def SaveText(self, df, FileName, tabname):
        current_time1 = datetime.datetime.now()
        current_time2 = current_time1.strftime('%m-%d-%Y_%H.%M.%S')
        text1 = df
        f2 = up.getPath()
        SavePath1 = f2
        invalidCharRemoved = re.sub(r"[^a-zA-Z0-9 ]", "", FileName)
        Filename = '\\' + str(invalidCharRemoved) + "_"
        SavePath2 = SavePath1 + Filename + current_time2 + ".xlsx"
        # SavePath2 = r'D:\ShakeBot Testing\ShaKeBotTest for DA - 12-20-2022.xlsx'
        try:
            with pd.ExcelWriter(SavePath2) as writer:
                text1.to_excel(writer, sheet_name=str(tabname))
                MondeVert.speak(self, 'File Saved')
                print('File Saved')
        except:
            Exception
        # MondeVert.speak(self,'Error Saving')
        # print ('Error Saving')

    def cleanText(self, text):
        # result = ''
        parser = GingerIt()
        try:
            if int(len(text)) <= 4999:
                # print(len(text))
                result = pd.DataFrame(parser.parse(text))

                # result.drop_duplicates()
            else:
                result = pd.DataFrame('', '', '')
                self.Notes = 'Text is Large so may need to review the parts that may have been cut by mistake'
                for i in range(0, (len(text) // 5000) + 1):
                    texttemp = text[0 + (5000 * i):4999 + (5000 * i)]
                    if len(texttemp) > 0:
                        result11 = pd.DataFrame(parser.parse(text))
                        result1 = [result1 + texttemp]
                result = pd.DataFrame(result1, columns='result')

            if len(result.loc[0, 'result']) > 0:
                # print(result.loc[:,'result'].values)
                rr = str(result.loc[0, 'result'])
                self.Correction_Comment += 'Corrections Made.'
            else:
                rr = text
                self.Correction_Comment += 'No Corrections Needed.'


        except:
            print('Text was not properly cleaned: ' + text)
            # print(text)
            rr = text
            self.Correction_Comment += 'No Corrections Needed (Caused an error and did not return anything).'
            return rr

        return rr

    def makeQuickPoem(self):
        x = s.PoemBot(1, 1, 1, 1, 40, 3)
        x.ReloadModel("model.h5")
        x.setupdata()
        self.Words = x.shakesbot_DA_Make_Script(size=300)
        # self.Words[0:100]
        print(self.Words)
        self.Words = MondeVert.cleanText(self, self.Words)
        MondeVert.speakSweet(self, self.Words)

    def StartThread(self):
        trd1 = threading.Thread(target=MondeVert.Make_Script2(self))
        self.w[i] = trd1.start()
        x = MondeVert.Make_Script2(self)
        print(x)

        self.threads.append(trd1)

    def Make_Script(self):
        size = 1
        self.threads = []
        self.w = numpy.empty(size, dtype=str)
        for self.i in range(0, size):
            MondeVert.StartThread(self)
            print(self.w[i])
            # globals()[f"trd{i}"] = threading.Thread(target=MondeVert.Make_Script2(self))
            # w[i] = globals()[f"trd{i}"].start()

        for x in self.threads:
            x.join()

        for ii in range(0, size):
            self.Words += str(self.w[ii])

        self.AI_Corrected_Content = MondeVert.cleanText(self, self.Words)

    def Make_Script2(self):
        pb = s.PoemBot(1, 1, 1, 1, 40, 3)
        pb.ReloadModel('model.h5')
        pb.setupdata()
        w1 = pb.shakesbot_DA_Make_Script()
        MondeVert.speakSweet(self, w1)
        # print (w1)
        return w1

    def generateGreek(self):
        w2 = self.pb.generateGreek() + ' \n' + ' \n'
        return w2

    def add2Master(df1):
        f2 = up.getMFPath()
        df2 = pd.read_excel(f2)
        df3 = pd.concat([df1, df2])
        df3.drop_duplicates()
        df3.iloc[:, 1:]
        # creating a new excel file and save the data
        df3.to_excel(f2, index=False)

    def add2Master2(df1):
        f2 = up.getMF2Path()
        df2 = pd.read_excel(f2)
        df3 = pd.concat([df1, df2])
        df3.iloc[:, 1:]
        df3.drop_duplicates()

        # creating a new excel file and save the data
        df3.to_excel(f2, index=False)

    def add2Master3(df1):
        f2 = up.MasterFile3
        df2 = pd.read_excel(f2)
        df3 = pd.concat([df1, df2])
        df3.iloc[:, 1:]
        df3.drop_duplicates()

        # creating a new excel file and save the data
        df3.to_excel(f2, index=False)

    def CompileBlog(self):
        f2 = up.getMFPath()
        df2 = pd.read_excel(f2)
        df2 = df2.loc[df2['Type/Marker'] != 'To Do list']
        df2.sort_values(by=['Date', 'Subject'], ascending=False)
        # need to add more stuff here for if it has notes or other tags etc
        for i in range(0, len(df2)):
            if df2['Priority'].loc[i] == '**AI content attached':
                df2['Combined'] = df2['AI_Corrected_Text'] + '\\n' + '\\n' + 'Note: AI Generated content: ' + df2[
                    'AI_Content']
            else:
                df2['Combined'] = df2['AI_Corrected_Text'] + '\\n' + '\\n'

        df2['Combined_Edit'] = MondeVert.cleanText(self, str(df2['Combined']))

        text1 = df2['Combined'].str.cat(sep='\\n \n')
        text2 = df2['Combined_Edit'].str.cat(sep='\n \n')
        text3 = MondeVert.cleanText(self, str(df2['text1']))
        text4 = MondeVert.cleanText(self, str(df2['text2']))

        topics = df2['Subject'].str.cat(sep='\n \n')
        notes = df2['Note'].str.cat(sep='\n \n')
        StartDate = df2['Date'].aggregate(['min']).strftime('%m-%d-%Y')
        End1Date = df2['Date'].aggregate(['min']).strftime('%m-%d-%Y')

        current_time2 = datetime.datetime.now().strftime('%m-%d-%Y')
        document = Document()
        document.add_heading(
            'MondeVert Blog --> ' + current_time2 + '\n' + 'Topics: ' + topics + 'Notes: ' + notes + ' ' + StartDate + 'to' + End1Date,
            level=1)
        document.add_paragraph(text1)
        document.save('MondeVert Blog (Raw) --> ' + current_time2 + '.docx')

        document = Document(text2)
        document.add_heading(
            'MondeVert Blog --> ' + current_time2 + '\n' + 'Topics: ' + topics + 'Notes: ' + notes + ' ' + StartDate + 'to' + End1Date,
            level=1)
        document.add_paragraph()
        document.save('MondeVert Blog (Polished) --> ' + current_time2 + '.docx')

        document = Document(text3)
        document.add_heading(
            'MondeVert Blog --> ' + current_time2 + '\n' + 'Topics: ' + topics + 'Notes: ' + notes + ' ' + StartDate + 'to' + End1Date,
            level=1)
        document.add_paragraph()
        document.save('MondeVert Blog (Polished2) --> ' + current_time2 + '.docx')

        document = Document(text4)
        document.add_heading(
            'MondeVert Blog --> ' + current_time2 + '\n' + 'Topics: ' + topics + 'Notes: ' + notes + ' ' + StartDate + 'to' + End1Date,
            level=1)
        document.add_paragraph()
        document.save('MondeVert Blog (Polished3) --> ' + current_time2 + '.docx')

    def compileFiles(self):
        f1 = up.getPath()
        files = glob.glob(f1 + '*.xlsx')
        tempDF = ''
        tempDF2 = ''
        df1 = pd.DataFrame()
        for fp in files:
            df1 = pd.DataFrame(pd.read_excel(fp))
            df1.drop_duplicates()
            r = len(df1.index)
            # print(r)

            for x in df1.columns:
                if x == 'Original_Text':
                    xs = 'Original_Text'
                elif x == 'Original_Wording':
                    xs = 'Original_Wording'
                else:
                    xx2 = ''

            for i in range(r):
                x = str(df1.loc[i, xs])

                try:
                    tempDF = MondeVert.cleanText(self, x)

                except:
                    xx2 = ''
                    # print('Error - Review somewhere')
                    # print(len(x))
                    # print(x)
                    # print('End of Error Details')
                else:
                    df1.loc[i, 'AI_Corrected_Text'] = tempDF
                    df1.loc[i, 'AI_Correction_Comment'] = self.Correction_Comment
            MondeVert.add2Master(df1)
            del df1



    ############################################################################################################################################################################################################################
    ############################################################################################################################################################################################################################
    ################Above is code I did not write, below is code that allows user to update values of respective values
    ############################################################################################################################################################################################################################
    ############################################################################################################################################################################################################################
    def getfilename(self):
        self.FileName = MondeVert.getdata(self, 'FileName')

    def getsubject(self):
        self.subject = MondeVert.getdata(self, 'subject')

    def getDictation_Type(self):
        self.Dictation_Type = MondeVert.getdata(self, 'Type')

    def getSignificance(self):
        self.Significance = MondeVert.getdata(self, 'Significance')

    def getNote(self):
        self.Notes += '|' + MondeVert.getdata(self, 'Notes')

    def getdata(self, Field):

        MondeVert.speak(self, 'what is the ' + Field)
        print('Please confirm the ' + Field)
        s = True
        Query = ''
        while (s == True):
            Query = MondeVert.getUserResponse(self)
            Query2 = ''
            WaitforResponse = False
            while (WaitforResponse == False):
                MondeVert.speak(self, Query + ' Is that correct?')
                Query2 = MondeVert.getUserResponse(self)
                if "yes" in Query2 or "correct" in Query2 or "ya" in Query2 or "yeah" in Query2:
                    MondeVert.speak(self, 'Thanks for confirming the ' + Field)
                    s = False
                    WaitforResponse = True
                    continue
                if "no" in Query2 or "wrong" in Query2 or "not it" in Query2 or "nope" in Query2:
                    MondeVert.speak(self, 'ok  please say the ' + Field + ' you want to set')
                    print('ok  please say the ' + Field + ' you want to set')
                    WaitforResponse = True
                    continue
                else:
                    MondeVert.speak(self, 'Lets try that again what is the ' + Field)
                    Event().wait(1)
                    continue
        return Query

    ############################################################################################################################################################################################################################
    ############################################################################################################################################################################################################################
    ################ below is code that allow me to confirm with the user the inputs are correct
    ############################################################################################################################################################################################################################
    ############################################################################################################################################################################################################################
    def takeCommand(self):
        Query = MondeVert.getUserResponse(self, Response="Action Requested")
        return Query



    def ConfirmBOT(self, message):
        MondeVert.speak(self, 'Chat GPT responded with: ' + message + ' do you want to use this prompt?', voice=4)
        query = MondeVert.getUserResponse(self)

        if 'yes' in query or 'correct' in query or 'yeah' in query or 'submit' in query:
            self.promptB = True

        if 'edit' in query or 'change' in query and ('prompt' in query or 'words' in query):
            MondeVert.editBotPrompt(self, message)

    def transcribe_Build_Query(self, pause=3.13):
        Query = MondeVert.getUserResponse(self)
        return Query

    def transcribe_Build_Query_Pause(self):
        Speak1 = 'What Tool do you want to use?'
        Print1 = 'Options: ChatGPT (ChatBot), Transcribe Mode , Make To Do List, Change Subject, Edit, Poem, Stop Recording --> Build others here'
        MondeVert.speak(self, Speak1)
        print(Print1)
        Query = MondeVert.getUserResponse(self, Response="Tool Selected")
        return Query

    ############################################################################################################################################################################################################################
    ############################################################################################################################################################################################################################
    ############################################################################################################################################################################################################################
    ##########################                below is ChatGPT                  ##############################################################################################################
    ############################################################################################################################################################################################################################
    ############################################################################################################################################################################################################################




    def makeArt(self, Prompt=''):

        MondeVert.Add2Transcript(self, ' \n')
        Prior_Mode = self.Mode
        self.Mode = 'AI Art Mode'
        MondeVert.Add2Transcript(self, text2Add=(Prior_Mode + ' - ' + self.Mode + ':'))
        MondeVert.Add2Transcript(self, ' \n')
        self.promptB = False

        Promptc = 'True'

        if Prompt == '':
            Promptc = 'True'
        else:
            Promptc = 'False'

        if Prompt == '':
            print('Prompt is NULL: ' + Promptc)

            while self.promptB == False:
                AIspeak = 'What do you want me to draw?'

                MondeVert.speak(self, AIspeak)

                print(AIspeak)
                prompt = MondeVert.getUserResponse(self, Response="Art Prompted", pause=1)

                if 'cancel' in prompt or 'stop' in prompt or 'quit' in prompt or 'done' in prompt or 'exit' in prompt:
                    MondeVert.saveTranscript(self)
                    self.promptB == True
                    continue

                MondeVert.ConfirmUR(self, prompt)
                print()
        else:
            prompt = Prompt
            print(prompt)
            MondeVert.Add2Transcript(self, '(Prompt Fed Directly into Function)')

        # Set up the OpenAI API client :
        openai.api_key = API_Key
        print("Sending to OpenAI...")
        print()

        response = openai.Image.create(
            prompt=prompt,
            n=1,
            size="1024x1024"
        )

        image_url = response['data'][0]['url']

        print(f"Image URL: {image_url}")
        print()

        # URL of the image to be downloaded is defined as image_url
        r = requests.get(image_url)  # create HTTP response object

        # send a HTTP request to the server and save
        # the HTTP response in a response object called r
        FileName = prompt
        print('Length of File Name: ' + str(len(FileName)))
        if len(FileName) >= 150:
            FileName = FileName[0:150]
            print('Length of File Name: ' + str(len(FileName)))

        fname = f"{''.join([c for c in FileName.strip().replace(' ', '_') if c.isalnum() or c == '_'])}.png"

        if not os.path.exists("images"):
            os.mkdir("images")

        if os.path.isfile(os.path.join(up.AI_Art_Path, '\'', fname)):
            fname = fname.split(".")[0] + f".{''.join(random.choice(string.ascii_letters) for x in range(5))}.png"

        fname = os.path.join(up.AI_Art_Path, fname)
        print(f"Filename: {fname}")
        with open(fname, 'wb') as f:

            # Saving received content as a png file in
            # binary format

            # write the contents of the response (r.content)
            # to a new file in binary mode.
            f.write(r.content)

        if platform.system() == 'Darwin':  # macOS
            subprocess.call(('open', fname))
        elif platform.system() == 'Windows':  # Windows
            os.startfile(fname)
        else:  # linux variants
            subprocess.call(('xdg-open', fname))

        return fname

    # Define a function that sends a message to ChatGPT
    def chat_query(self, prompt):
        model_engine = "text-davinci-003"
        completions = openai.Completion.create(
            engine=model_engine,
            prompt=prompt,
            max_tokens=2048,
            n=1,
            temperature=0.5,
        )

        message = completions.choices[0].text
        return message

    # Define a function that handles the conversation
    def conversation_handler(self, prompt):
        # Send the prompt to ChatGPT
        load_dotenv()
        # Set up the OpenAI API client :
        openai.api_key = API_Key
        response = MondeVert.chat_query(prompt)
        print(f"MondeVert: {response}")

        # End the conversation if ChatGPT says goodbye
        if "goodbye" in response.lower():
            return

        # Otherwise, get user input and continue the conversation
        prompt = input("User: ")
        MondeVert.conversation_handler(prompt)




    ############################################################################################################################################################################################################################
    ############################################################################################################################################################################################################################
    ############################################################################################################################################################################################################################
    ################Long Code where I give Digital Assistant commands##    (user menus)
    ############################################################################################################################################################################################################################
    ############################################################################################################################################################################################################################

    def RunChatGPT(self):
       #this will look very diff in the future The digital assistant will pass in certain key words to run different functions, but really this should exist without digital assist this is auxilary to anything I run it will be way easier to maintain
        dn = 1

    # Provide me with a copy of the short bio you come up with and then using that I want you to role play you are the respective artist and write a short 1 page story about something that would happen in the life of the respective artist you described. It can be about any topic but it should feel authentic to the life of the person you described."}
    def Make_a_Song(self, Mode='Random Song'):
        ArtPaths = []
        openai.api_key = API_Key

        crazy = round((randbelow(520000) + 170000) / 100000, 0)
        crazy = crazy / 10
        print(crazy)

        response = openai.ChatCompletion.create(
            model="gpt-3.5-turbo",
            messages=[
                {"role": "system", "content": up.system_Text},
                {"role": "user", "content": up.ArtistBio_SongArtist}
            ], temperature=crazy
        )

        Artist_Bio = response.choices[0].message.content

        Song1 = openai.ChatCompletion.create(
            model="gpt-3.5-turbo",
            messages=[
                {"role": "system", "content": up.system_Text},
                {"role": "user", "content": up.RolePlay_SongArtist + Artist_Bio},
                {"role": "user", "content": up.Song_Prompt_SongArtists},
                {"role": "user", "content": up.Title_SongArtists},
                {"role": "user", "content": up.Samples_SongArtists},
                {"role": "user", "content": up.Tune_SongArtists}

            ]
            , temperature=crazy
        )
        Song = Song1.choices[0].message.content

        Art_Prompt1 = openai.ChatCompletion.create(
            model="gpt-3.5-turbo",
            messages=[
                {"role": "system", "content": up.system_Text},
                {"role": "user", "content": up.RolePlay_SongArtist + Artist_Bio},
                {"role": "user", "content": up.ArtPrompt_SongArtist}
            ], temperature=crazy
        )

        Art_Prompt = Art_Prompt1.choices[0].message.content

        print(Artist_Bio)
        MondeVert.speak(self, Artist_Bio)
        print(Song)
        MondeVert.speak(self, Song)
        print(Art_Prompt)
        MondeVert.speak(self, Art_Prompt)

        # Art2 = MondeVert.ChatGPTDA(self, temp=crazy, MakeArt=True, Prompt=(Art_Prompt),ConfirmBot=False)
        try:
            ArtPath2 = MondeVert.makeArt(self, Art_Prompt)
            ArtPaths.append(ArtPath2)

        except:
            ArtPath2 = ''

        Title = MondeVert.ChatGPTDA(self, temp=crazy, MakeArt=True, Prompt=(up.Song_Title_Prompt + Song),
                                        ConfirmBot=False)
        prompt = "Artist info:" + Artist_Bio + "Work of Art Inspiration:" + Art_Prompt + "Song:" + Song
        Prompts_Used = [
            up.system_Text + up.RolePlay_SongArtist + "AI Created a Persona shown as the Artist Bio Above" + up.Song_prompt + up.ArtPrompt_SongArtist]
        ArtistPoetInfo = 'Lyrics Written By: ' + up.Song_Writer + '      (' + 'Artwork by: ' + up.AI_ArtistName + ')'
        MondeVert.NamePoemSavePoem(self, prompt, ArtPaths, Prompts_Used, ArtistPoetInfo, title=Title,
                                       FolderPath=up.AI_Music_Path + '\\' + Mode, ArtType='Song Lyrics' + Mode)

    def Make_a_SongRR(self, System = up.system_Text, Role= up.RolePlay_SongArtist, Background=up.ArtistBio_SongArtist, Task = up.Song_Prompt_SongArtists,Special=up.Samples_SongArtists, Format='', Mode='Random Song'):
        ArtPaths = []
        openai.api_key = API_Key

        crazy = round((randbelow(520000) + 170000) / 100000, 0)
        crazy = crazy / 10
        print(crazy)

        crazy += .2
        if crazy < .4:
            crazy = .6
        if crazy > .9:
            crazy = .7

        # Title = "Richies Music"
        current_time1 = datetime.datetime.now()
        current_time2 = current_time1.strftime('%m-%d-%Y_%H.%M.%S')
        # f2 = FolderPath
        f2 = up.SavePath
        SavePath1 = f2

        response = openai.ChatCompletion.create(
            model="gpt-3.5-turbo",
            messages=[
                {"role": "system", "content": up.system_TextRR},
                {"role": "user", "content": up.ArtistBio_SongArtistRR + up.Artist_Bio_DetailsRR}
            ], temperature=crazy
        )

        Artist_Bio = response.choices[0].message.content

        # print('Artist_Bio: ' + Artist_Bio)
        print('************************************************************************************************')

        Song1 = openai.ChatCompletion.create(
            model="gpt-3.5-turbo",
            messages=[
                {"role": "system", "content": up.system_TextRR},
                {"role": "user", "content": up.RolePlay_SongArtistRR + Artist_Bio},
                {"role": "user", "content": up.Song_Prompt_SongArtistsRR + up.Song_Subject},
                {"role": "user", "content": up.Song_Format_Prompt + up.SongFormat},
                # {"role": "user", "content": up.Title_SongArtistsRR},
                # {"role": "user", "content": up.Samples_SongArtistsRR},
                # {"role": "user", "content": up.Tune_SongArtistsRR}

            ]
            , temperature=crazy
        )
        Song = Song1.choices[0].message.content

        # Create original details
        response = openai.ChatCompletion.create(
            model="gpt-3.5-turbo",
            messages=[
                {"role": "system", "content": up.system_TextRR},
                {"role": "user", "content": up.RolePlay_SongArtistRR + Artist_Bio},
                {"role": "user", "content": up.Title_SongArtistsRR + Song}

            ], temperature=crazy
        )

        Title = response.choices[0].message.content

        invalidCharRemoved = re.sub(r"[^a-zA-Z0-9 ]", "", Title)
        folder = str(invalidCharRemoved)
        if len(folder) > 44:
            folder = folder[0:44]
        Title1 = '\\' + str(invalidCharRemoved) + "_"
        if len(Title1) > 44:
            Title1 = Title1[0:44]
        SavePath2 = SavePath1 + '\\' + folder
        Title2 = SavePath2 + Title1 + current_time2

        # Create original details
        response = openai.ChatCompletion.create(
            model="gpt-3.5-turbo",
            messages=[
                {"role": "system", "content": up.system_TextRR},
                {"role": "user", "content": up.RolePlay_SongArtistRR + Artist_Bio},
                {"role": "user", "content": up.Samples_SongArtistsRR2 + Song}

            ], temperature=crazy
        )

        Samples2 = response.choices[0].message.content

        # Create original details
        response = openai.ChatCompletion.create(
            model="gpt-3.5-turbo",
            messages=[
                {"role": "system", "content": up.system_TextRR},
                {"role": "user", "content": up.RolePlay_SongArtistRR + Artist_Bio},
                {"role": "user", "content": up.ReWrite_Song + Song}

            ], temperature=crazy
        )

        ReWrite = response.choices[0].message.content

        # Create original details
        response = openai.ChatCompletion.create(
            model="gpt-3.5-turbo",
            messages=[
                {"role": "system", "content": up.system_TextDJ},
                {"role": "user", "content": up.RolePlay_SongArtistRR + Artist_Bio},
                {"role": "user", "content": Samples2},
                {"role": "user", "content": up.ExplainTheBeat + ReWrite}

            ], temperature=crazy
        )

        DJMondeVert = response.choices[0].message.content

        print('************************************************************************************************')
        print('************************************************************************************************')
        print('************************************************************************************************')

        MondeVert.speak(self, "DJ MondeVert Work complete yo")
        data = """Artist_Bio: """ + Artist_Bio + """:                                                                                      

        *************************************************
        *************************************************


        Potential Samples: """ + Samples2 + """


        *************************************************
        *************************************************


        Title: """ + Title + """



        *************************************************
        *************************************************



        Song: """ + Song + """ 


        *************************************************
        *************************************************




        Song 2.0: """ + ReWrite + """


        *************************************************
        *************************************************




        DJMondeVert: """ + DJMondeVert + """
        *************************************************
        *************************************************

        """

        print(data)

        MondeVert.speak(self, "Saving the files round 1")
        # print(data)

        data1 = [(data)]

        try:
            df1 = pd.DataFrame(data1, columns='Text')
            # print(df1)
            # print(Title2)
            filename = Title2
            MondeVert.SaveCSV(data, filename, SavePath2)

            # MondeVert.SaveText(self,df1,'MondeVert Assistant', 'Full Transcript')
            MondeVert.add2Master2(df1)
        except:
            print('Review Error File did not save ')

        try:
            send_email_w_attachment_gmail(up.to, up.subject, Song, filename)
        except:
            print('email not send, its possible file was not created')

        MondeVert.speak(self, "Making the art, painting yo picture")
        Art_Prompt1 = openai.ChatCompletion.create(
            model="gpt-3.5-turbo",
            messages=[
                {"role": "system", "content": up.system_TextRR},
                {"role": "user", "content": up.RolePlay_SongArtistRR + Artist_Bio},
                {"role": "user", "content": up.ArtPrompt_SongArtistRR}
            ], temperature=crazy
        )

        Art_Prompt = Art_Prompt1.choices[0].message.content
        print("Work of Art Inspiration:" + Art_Prompt)

        # print(Artist_Bio)
        # MondeVert.speak(self,Artist_Bio)
        # print(Song)
        # MondeVert.speak(self, Song)
        # print(Art_Prompt)
        # MondeVert.speak(self, Art_Prompt)

        # Art2 = MondeVert.ChatGPTDA(self, temp=crazy, MakeArt=True, Prompt=(Art_Prompt),ConfirmBot=False)
        try:
            ArtPath2 = MondeVert.makeArt(self, Art_Prompt)
            ArtPaths.append(ArtPath2)

        except:
            ArtPath2 = ''

        MondeVert.speak(self, "Saving the files round 2")
        prompt = data + "Work of Art Inspiration:" + Art_Prompt
        Prompts_Used = [
            up.system_Text + up.RolePlay_SongArtist + "AI Created a Persona shown as the Artist Bio Above" + up.Song_prompt + up.ArtPrompt_SongArtist]
        ArtistPoetInfo = 'Lyrics Written By: ' + up.Song_Writer + '      (' + 'Artwork by: ' + up.AI_ArtistName + ')'
        MondeVert.NamePoemSavePoem(self, prompt, ArtPaths, Prompts_Used, ArtistPoetInfo, title=Title,
                                       FolderPath=up.AI_Music_Path + '\\' + Mode, ArtType='Song Lyrics' + Mode)

        # MondeVert.speak(self, ReWrite )
    def MondeVertAuto(self, Mode='AutoSocial', Role = '', System = ''):
        if Mode == 'AutoSocial':
            Caption = MondeVert.MondeVertTask(self, System=up.system_TextJoaT, Role=up.Test_Role_SMA, Background=up.Test_Background_SMA,Task=up.Test_Task_SMA, Special=up.Test_Special_SMA, Format=up.Test_Format_SMA, Title = up.Test_Title_SMA,Mode='AutoSocial')
        return Caption + up.Instagram_Adds
    def MondeVertMenu(self, Mode='Basic', Role = '', System = ''):

        if Mode == 'ScreenPlay':
                    MondeVert.Make_a_ScreenPlay(self,System = up.system_Text_ScreenPlay0, Mode='ScreenPlay')
        elif Mode == 'Music':
            MondeVert.Make_a_ScreenPlay(self,System = up.system_TextRR, Mode='Make_a_SongRR')
        elif Mode == 'Skit':
            MondeVert.MondeVertTask(self,System = up.system_TextJoaT,Role=up.Test_Role, Background=up.Test_Background, Task = up.Test_Task,Special=up.Test_Special, Format=up.Test_Format, Mode='Skit')
        elif Mode == 'Basic':
            MondeVert.MondeVertTask(self,System = up.system_TextJoaT,Role=up.Test_Role, Background=up.Test_Background, Task = up.Test_Task,Special=up.Test_Special, Format=up.Test_Format, Mode='Basic', Logic_AI=.5)
        elif Mode == 'Advance':
            MondeVert.MondeVertTask(self,System = up.system_TextJoaT,Role=up.Test_Role, Background=up.Test_Background, Task = up.Test_Task, Special=up.Test_Special, Format=up.Test_Format, Mode='Advance')
        elif Mode == 'Interview':
            MondeVert.MondeVertTask(self,System = up.system_TextJoaT,Role=up.Test_Role, Background=up.Test_Background, Task = up.Test_Task,Special=up.Test_Special, Format=up.Test_Format, Mode='Interview')
        elif Mode == 'Social':
            MondeVert.MondeVertTask(self, System=up.system_TextJoaT, Role=up.Test_Role, Background=up.Test_Background_SM,Task=up.Test_Task_SM, Special=up.Test_Special_SM, Format=up.Test_Format_SM, Mode='Social')
        elif Mode == 'Book':
            MondeVert.Make_a_ScreenPlay(self, Mode='ScreenPlay')
        elif Mode == 'Series':
            MondeVert.Series(self, Mode='Series', Episodes= 10, Seasons = 1, Audience = 'Kids')








    def MondeVertTask(self, System = '', Role='', Background='', Task='',FinalOutput='Text with proper MLA formatting, do not provide made up information only facts (unless user asks for fiction)',
                      Special='', Title='', Mode='Basic', Logic_AI=0, skit = False,Format = '', Review_Critic = False):
        ArtPaths = []
        openai.api_key = API_Key
        if Logic_AI == 0:
            crazy = round((randbelow(520000) + 170000) / 100000, 0)
            crazy = crazy / 10
            print(crazy)
            crazy += .2
            if crazy < .4:
                crazy = .6
            if crazy > .9:
                crazy = .7
        else:
            crazy = Logic_AI

        # Title = "Richies Music"
        current_time1 = datetime.datetime.now()
        current_time2 = current_time1.strftime('%m-%d-%Y_%H.%M.%S')
        # f2 = FolderPath
        f2 = up.SavePath
        SavePath1 = f2



        Project_Description = 'N/A not ready to test this yet'
        Result = 'N/A not ready to test this yet'
        Result_AI = 'N/A not ready to test this yet'
        AIPrompt = 'N/A not ready to test this yet'
        Result_Combine = 'N/A not ready to test this yet'
        Bitter_Critic = 'N/A not ready to test this yet'
        Title = up.Test_Title




        Task2 = up.Test_Task_Critic
        Format2 = up.Test_Format_Critic
        Special2 = up.Test_Special_Critic
        Role2 = up.Test_Role_Critic

        Format0 = up.Test_Format0


        Format_f = Format
        if Mode =='skit':
            Format_f = Format0

        if Title != '':
            Background = Title + up.Test_Background
        else:
            Background = up.Test_Background

        # This is for the result when you ask AI to summarize the project
        response = openai.ChatCompletion.create(
            model="gpt-3.5-turbo",
            messages=[
                {"role": "system", "content": up.system_TextJoaT + Role},
                {"role": "user", "content": up.Role_Play_Prompt + Role},
                {"role": "user", "content": up.USER_Format_Prompt + Format_f},
                {"role": "user", "content": up.USER_SPECIAL_Prompt + Special},
                {"role": "user", "content": up.USER_Background_Prompt + Background}
            ], temperature=crazy
        )

        Project_Description = response.choices[0].message.content

        if Title != '':
            Project_Description = Title + Project_Description
        else:
            Project_Description = Project_Description

        print('************************************************************************************************')

        # This is for the result if you let the AI describe project and details and then make the response
        Result1 = openai.ChatCompletion.create(
            model="gpt-3.5-turbo",
            messages=[
                {"role": "system", "content": up.system_TextJoaT + Role},
                {"role": "user", "content": up.Role_Play_Prompt + Role},
                {"role": "user", "content": up.USER_Format_Prompt + Format},
                {"role": "user", "content": up.USER_SPECIAL_Prompt + Special},
                {"role": "user", "content": up.AI_Background_Prompt + Project_Description},
                {"role": "user", "content": up.USER_Task_Prompt + Task},
            ]
            , temperature=crazy
        )
        Result = Result1.choices[0].message.content

        response = openai.ChatCompletion.create(
            model="gpt-3.5-turbo",
            messages=[
                {"role": "system", "content": up.system_TextJoaT + Role},
                {"role": "user", "content": up.Role_Play_Prompt + Role},
                # {"role": "user", "content": up.USER_Format_Prompt + Format},
                {"role": "user", "content": up.USER_SPECIAL_Prompt + Special},
                {"role": "user", "content": up.AI_Background_Prompt + Project_Description},
                {"role": "user", "content": up.USER_Task_Prompt + Task},
            ], temperature=crazy
        )

        AIPrompt = response.choices[0].message.content

        # print('Artist_Bio: ' + Artist_Bio)
        print('************************************************************************************************')

        Result1 = openai.ChatCompletion.create(
            model="gpt-3.5-turbo",
            messages=[
                {"role": "system", "content": up.system_TextJoaT + Role},
                {"role": "user", "content": up.Role_Play_Prompt + Role},
                {"role": "user", "content": up.USER_Format_Prompt + Format},
                {"role": "user", "content": up.USER_SPECIAL_Prompt + Special},
                {"role": "user", "content": up.AI_Background_Prompt + Project_Description},
                {"role": "user", "content": up.USER_Task_Prompt + Task},
            ]
            , temperature=crazy
        )
        Result_AI = Result1.choices[0].message.content

        print('************************************************************************************************')

        try:
            Result2 = openai.ChatCompletion.create(
                model="gpt-3.5-turbo",
                messages=[
                    {"role": "system", "content": up.system_TextJoaT + Role},
                    {"role": "user", "content": up.Role_Play_Prompt + Role},
                    {"role": "user", "content": up.UserRequest + Format + Special + Project_Description + Task},
                    # {"role": "user", "content": up.USER_SPECIAL_Prompt + Special},
                    {"role": "user", "content": up.AI_Background_Prompt + Project_Description},
                    {"role": "user",
                     "content": up.CombineBothResults + "Result 1: " + Result + "Result 2: " + Result_AI},
                    {"role": "user", "content": up.USER_Task_Prompt + Task},
                ]
                , temperature=crazy
            )
            Result_Combine = Result1.choices[0].message.content

        except:
            try:
                Result2 = openai.ChatCompletion.create(
                    model="gpt-3.5-turbo",
                    messages=[
                        {"role": "system", "content": up.system_TextJoaT + Role},
                        {"role": "user", "content": up.Role_Play_Prompt + Role},
                        {"role": "user", "content": up.UserRequest + Format + Special + Project_Description + Task},
                        # {"role": "user", "content": up.USER_SPECIAL_Prompt + Special},
                        {"role": "user", "content": up.AI_Background_Prompt + Project_Description},
                        {"role": "user",
                         "content": up.CombineBothResults + "Result 1: " + Result},
                        {"role": "user", "content": up.USER_Task_Prompt + Task},
                    ]
                    , temperature=crazy
                )
                Result_Combine = Result1.choices[0].message.content
                print("Only revised v1")
            except:
                print("could not combine 2 versions")
                dn = 55

        if Title == '':
            Title = MondeVert.Get_Title_GPT(self)



        # This is funny content where the AI reviews the youtube script
        if   Mode =='skit':
            # Create original details
            response = openai.ChatCompletion.create(
                model="gpt-3.5-turbo",
                messages=[
                    {"role": "system", "content": up.system_TextJoaT + Role2},
                    {"role": "user", "content": up.Role_Play_Prompt + Role2},
                    {"role": "user", "content": up.UserRequest + Format2 + Special2 + Project_Description + Task2},
                    {"role": "user", "content": up.USER_SPECIAL_Prompt + Special2},
                    {"role": "user", "content": up.AI_Background_Prompt + Result_Combine},
                    {"role": "user", "content": up.USER_Task_Prompt + Task2},

                ], temperature=crazy
            )

            Bitter_Critic = response.choices[0].message.content

        invalidCharRemoved = re.sub(r"[^a-zA-Z0-9 ]", "", Title)
        folder = str(invalidCharRemoved)
        if len(folder) > 44:
            folder = folder[0:44]
        Title1 = '\\' + str(invalidCharRemoved) + "_"
        if len(Title1) > 44:
            Title1 = Title1[0:44]
        SavePath2 = SavePath1 + '\\' + folder
        Title2 = SavePath2 + Title1 + current_time2


        print('************************************************************************************************')
        print('************************************************************************************************')
        print('************************************************************************************************')

        MondeVert.speak(self, "MondeVert Project Work complete yo")
        data = """Project_Description: """ + Project_Description + """:                                                                                      

        *************************************************
        *************************************************


        Title: """ + Title + """



        *************************************************
        *************************************************



        Result: """ + Result + """



        *************************************************
        *************************************************



        AIPrompt: """ + AIPrompt + """



        *************************************************
        *************************************************



        Result_AI: """ + Result_AI + """



        *************************************************
        *************************************************



        Result_Combine: """ + Result_Combine + """



        *************************************************
        *************************************************



        Bitter_Critic: """ + Bitter_Critic

        print(data)

        MondeVert.speak(self, "Saving the files round 1")
        # print(data)

        data1 = [(data)]

        try:
            df1 = pd.DataFrame(data1, columns=['Text'])

            filename = Title2
            MondeVert.SaveCSV(data, filename, SavePath2)

        except:
            print('Review Error File did not save ')

        try:
            send_email_w_attachment_gmail(up.to, up.subject, Project_Description, filename)
            print("Gmail email sent")
            send_email_w_attachment_outlook(up.to, up.subject, Project_Description, filename)
            print("Outlook email sent")
        except:
            print('email not send, its possible file was not created')

        MondeVert.speak(self, "Making the art, painting yo picture")
        Art_Prompt1 = openai.ChatCompletion.create(
            model="gpt-3.5-turbo",
            messages=[
                {"role": "system", "content": up.system_Text_Art},
                {"role": "user", "content": up.MondeVert_ArtPrompt + Project_Description},
            ], temperature=crazy
        )

        Art_Prompt = Art_Prompt1.choices[0].message.content
        print("Work of Art Inspiration:" + Art_Prompt)

        try:
            ArtPath2 = MondeVert.makeArt(self, Art_Prompt)
            ArtPaths.append(ArtPath2)

        except:
            ArtPath2 = ''

        MondeVert.speak(self, "Saving the files round 2")



        prompt = data + "Work of Art Inspiration:" + Art_Prompt
        Prompts_Used = [
            up.system_Text + up.RolePlay_SongArtist + "AI Created a Persona shown as the Artist Bio Above" + up.Song_prompt + up.ArtPrompt_SongArtist]
        ArtistPoetInfo = 'Lyrics Written By: ' + up.Song_Writer + '      (' + 'Artwork by: ' + up.AI_ArtistName + ')'
        MondeVert.NamePoemSavePoem(self, prompt, ArtPaths, Prompts_Used, ArtistPoetInfo, title=Title,
                                       FolderPath=up.AI_Music_Path + '\\' + Mode, ArtType='Song Lyrics' + Mode)

    def GPTSummary(self, crazy='', prompt=up.GPT_Summary, Plot='',
                   sys_prompt=up.system_Text_ScreenPlay0 + up.system_Text_ScreenPlay3):

        openai.api_key = API_Key

        if crazy != '':
            crazy = round((randbelow(520000) + 170000) / 100000, 0)
            crazy = crazy / 10

        print(crazy)

        response = openai.ChatCompletion.create(
            model="gpt-3.5-turbo",
            messages=[
                {"role": "system", "content": sys_prompt},
                {"role": "user", "content": prompt + Plot}

            ], temperature=crazy
        )

        Skeleton_Story = response.choices[0].message.content
        return Skeleton_Story

    # Every Scene I should come up with pictures to show (find things to trigger a picture getting created and multi-thread it so it does not stop anything, I should have it nice and output within the App if possible
    # Also I should have the ability to assign a voice to a given character and keep that for the audio transcript/recording. Also provide the Raw Words.
    # Snatch prequel
    # Big Lewbowski Prequel and Sequel
    # Something about Mary Sequel (add my own ideas)
    # Ancient egyption culture and other cultures
    # #(have the different technologies showing how they use sound and frequencies to move giant stones underground. Large caves and artifical light makes the world seem like another planet,
    # #have it seem like a distant planet etc, also have some of the extinct animals and animals/plants we never heard of. Describe landscape in vivid detail and make the reader fall in love with the lore.
    # Twist is they end up being underground and when humans are dealing with global warming we break in and accidentally become slaves.
    #   eventually we try to become a unified people as they realize we made the same mistakes as them
    #    #Write the MetaVerse Elon Musk Story -- This is going to be a video game too(like pokemon)
    def GPTArt(self, crazy='', prompt=up.ArtPrompt_ScreenPlay + up.direction_Text_Artist_ScreenPlay, Plot='',
               sys_prompt=up.system_Text_Artist_ScreenPlay):
        openai.api_key = API_Key

        Art_Prompt1 = openai.ChatCompletion.create(
            model="gpt-3.5-turbo",
            messages=[
                {"role": "system", "content": sys_prompt},
                {"role": "user", "content": prompt + Plot}
            ], temperature=crazy
        )

        GPTARTPROMPT = Art_Prompt1.choices[0].message.content
        return GPTARTPROMPT

    #
    # $$$$$Savetime

    # Give artist a profile and save it for use - make them have a set style that is repeated for the entire book (will be great for gothic and kid stories)

    # 1). Adapt for Movie/Book/Short Story/Indie Film
    #     2). Enhance the tool so you have option to pass in subject (by Default choose a random genre
    # 3). Need to build a google search version where I ask open ended questions and it summarizes and provides useful links for me to explore
    # 4). Coding version of this and start to use it
    # 5). Update Rap one for me and make it indie etc.
    # 6). create a version that does advertisements
    # 7). Animation
    # 8). Take all of the notes I add and make a story
    # 9). Make a tool that starts with basic idea, makes song.... book, short story, Movie
    # 10. Make tool post to instagram and twitter (do this asap and start doing it)
    # 11). Have a folder of pictures that it goes through and 3 times a day or something posts to my sites. Then it moves to another folder. Take poetry and maybe post it with all nuts and bolts ready
    # 12). For sampler, maybe build a tool to go out and download songs from itunes when it suggests it so I can mess around
    # 13). Tabs in and out (edit sounds to be other instruments)
    # 14). Edit the voice output so I can text to speech sound better
    # 15). make a play out of the voices we already have??
    # 16 get picture working with comma separated list.

    def Make_a_ScreenPlay(self, System = '',Mode='ScreenPlay'):
        Summary_Episodes = []
        ScreenPlay = []
        i = 0
        data = []
        Outline_Next2_Ref2 = ''
        Outline_Details = ''
        Plot = ''
        Add_IN = ''
        Summary_Summer_90s = ''
        filename = []
        filename1 = []
        Outline_Next2 = ''
        Outline_Next = ''
        self.Mode = 'ScreenPlay'
        ArtPaths = []
        openai.api_key = API_Key
        Plot = ""
        crazy = round((randbelow(520000) + 170000) / 100000, 0)
        crazy = crazy / 10
        if crazy < .3:
            crazy += .1

        print('Crazy Factor: ' + str(crazy))

        UserInput = False

        if UserInput == False:
            Subject = up.RandomGenre()
        else:
            Subject = UserInput
        print('Subject: ' + Subject)

        # Create original details
        response = openai.ChatCompletion.create(
            model="gpt-3.5-turbo",
            messages=[
                {"role": "system", "content": up.system_Text_ScreenPlay0 + up.system_Text_ScreenPlay2},
                {"role": "user", "content": up.Direction_ScreenPlay_Detail},
                {"role": "user", "content": up.Direction_ScreenPlay_Detail2},
                {"role": "user", "content": up.Direction_ScreenPlay_Detail3 + Subject},
                {"role": "user", "content": up.Setting_ScreenPlay},
                {"role": "user", "content": up.Characters_ScreenPlay},
                # {"role": "user", "content": up.Title_ScreenPlay}

            ], temperature=crazy
        )

        Skeleton_Story1 = response.choices[0].message.content

        # Create original details
        response = openai.ChatCompletion.create(
            model="gpt-3.5-turbo",
            messages=[
                {"role": "system", "content": up.system_Text_ScreenPlay0 + up.system_Text_ScreenPlay2},
                {"role": "user", "content": up.Title_ScreenPlay + Skeleton_Story1}

            ], temperature=crazy
        )

        Skeleton_Story2 = response.choices[0].message.content

        Skeleton_Story = Skeleton_Story2 + Skeleton_Story1

        print('***********************************************************************************')
        print(Skeleton_Story)
        print('***********************************************************************************')

        # MondeVert.speak(self, Skeleton_Story)

        # Create full outline
        response = openai.ChatCompletion.create(
            model="gpt-3.5-turbo",
            messages=[
                {"role": "system", "content": up.system_Text_ScreenPlay0 + up.system_Text_ScreenPlay4},
                {"role": "user", "content": up.Create_Outline1 + Skeleton_Story},

                # {"role": "user", "content": up.Title_ScreenPlay}

            ], temperature=crazy
        )

        Outline_All = response.choices[0].message.content

        print(Outline_All)

        print('***********************************************************************************')
        print('***********************************************************************************')

        # This is how I set up the prompts for the show, I could randomize when to activate red herrings etc based on Add-In Feature
        Direction_ScreenPlay = []
        Direction_ScreenPlay.append(up.Direction_ScreenPlay_Pilot)
        Direction_ScreenPlay.append(up.Direction_ScreenPlayMiddle)
        Direction_ScreenPlay.append(up.Direction_ScreenPlayMiddleLate)
        Direction_ScreenPlay.append(up.Direction_ScreenPlayMiddleLate)
        Direction_ScreenPlay.append(up.Direction_ScreenPlayFinal)

        ScreenPlay_Final = ''
        Episode_Count = 5
        for i in range(0, Episode_Count):
            Episode_Direction = Direction_ScreenPlay[i]

            if i == 1:
                Add_IN = up.RedHerring_ScreenPlay
            elif i == 2:
                Add_IN = up.UnsungHero + up.Twist_ScreenPlay
            elif i == 3:
                Add_IN = up.Mystery_Character_Past + up.VillanOrigin

            # Prior to writing an episode we should create the outline for the episode based on

            if i != 0:

                # Create outline for next 2 episodes
                response = openai.ChatCompletion.create(
                    model="gpt-3.5-turbo",
                    messages=[
                        {"role": "system", "content": up.system_Text_ScreenPlay0 + up.system_Text_ScreenPlay4},
                        {"role": "user", "content": up.Compare_Detail_v_Summary1 + Plot},
                        {"role": "user", "content": up.Compare_Detail_v_Summary2 + Outline_All},
                        {"role": "user", "content": up.Compare_Detail_v_Summary3},

                        # {"role": "user", "content": up.Title_ScreenPlay}

                    ], temperature=crazy
                )

                Outline_Next2 = response.choices[0].message.content

                # Store this to maybe use later as a way to keep story moving together well

                # Revise
                response = openai.ChatCompletion.create(
                    model="gpt-3.5-turbo",
                    messages=[
                        {"role": "system", "content": up.system_Text_ScreenPlay0 + up.system_Text_ScreenPlay4},
                        {"role": "user", "content": up.Compare_Detail_v_Summary_Outline_New + Outline_Next2},
                        {"role": "user", "content": up.Compare_Detail_v_Summary_Revise},
                        {"role": "user", "content": up.Compare_Detail_v_Summary_Summary_New + PriorScreenPlay},
                        # {"role": "user", "content": up.Title_ScreenPlay}

                    ], temperature=crazy
                )

                Outline_Next = response.choices[0].message.content

            else:
                # if i = 0 want to create an outline for next episode

                # Create outline for next 2 episodes
                response = openai.ChatCompletion.create(
                    model="gpt-3.5-turbo",
                    messages=[
                        {"role": "system", "content": up.system_Text_ScreenPlay0 + up.system_Text_ScreenPlay4},
                        {"role": "user", "content": up.Compare_Detail_v_Summary2 + Outline_All},
                        {"role": "user", "content": up.Outline_Pilot}

                        # {"role": "user", "content": up.Title_ScreenPlay}

                    ], temperature=crazy
                )

                Outline_Next = response.choices[0].message.content

            print('  ')
            print('  ')
            print('  ')
            # print('***********************************************************')
            # print('***********************************************************')
            print('***********************************************************')
            print('***********************************************************')
            print('***********************************************************')
            print('  ')
            print('  ')
            print('  ')
            print('Episode ' + str(i + 1))
            Add_IN = ''

            Song1 = openai.ChatCompletion.create(
                model="gpt-3.5-turbo",
                messages=[
                    {"role": "user", "content": up.system_Text_ScreenPlay + up.system_Text_ScreenPlay1},
                    {"role": "system", "content": up.Tie_In_ScreenPlay + Skeleton_Story},
                    {"role": "system", "content": up.Tie_In_ScreenPlay2 + Outline_Next},
                    {"role": "system", "content": up.Direction_ScreenPlayALL},
                    {"role": "system", "content": up.Direction_ScreenPlayALL2},
                    {"role": "system", "content": up.Direction_ScreenPlayALL3},
                    ##--------------------Common Above/Unique Below----------
                    # {"role": "system", "content":  Episode_Direction+ Add_IN},
                    {"role": "system", "content": Episode_Direction},
                    ##--------------------Dialog----------
                    {"role": "system", "content": up.Dialog_ScreenPlay}
                ], temperature=crazy)

            ScreenPlay0 = Song1.choices[0].message.content
            ScreenPlay.append(ScreenPlay0)

            print(ScreenPlay[i])

            Summary_Episode1 = MondeVert.GPTSummary(self, crazy, up.Summarize_ScreenPlay, ScreenPlay[i])
            Summary_Episodes.append(Summary_Episode1)

            PriorSummary = Summary_Episodes[i]
            PriorScreenPlay = ScreenPlay[i]
            PriorNext2 = Outline_Next2
            PriorOutline = Outline_Next

            # print('***********************************************************')
            # print("Summary Episode " + str(i + 1))
            # print(Summary_Episodes[i])
            Plot += Summary_Episodes[i]
            ScreenPlay += "EPISODE " + str(i + 1) + ScreenPlay[i]
            Outline_Details = Skeleton_Story + Outline_All
            Outline_Next2_Ref2 += Outline_Next2

        Title = MondeVert.Get_Title_GPT(self,Plot)

        current_time1 = datetime.datetime.now()
        current_time2 = current_time1.strftime('%m-%d-%Y_%H.%M.%S')
        # f2 = FolderPath
        f2 = up.SavePath
        SavePath1 = f2
        invalidCharRemoved = re.sub(r"[^a-zA-Z0-9 ]", "", Title)
        folder = str(invalidCharRemoved)
        if len(folder) > 44:
            folder = folder[0:44]
        Title1 = '\\' + str(invalidCharRemoved) + "_"
        if len(Title1) > 44:
            Title1 = Title1[0:44]
        SavePath2 = SavePath1 + '\\' + folder

        filename1.append('Screen Play - ' + Title1)
        filename1.append('Details_Outline - ' + Title1)
        filename1.append('OutlineNext2_Ref - ' + Title1)
        filename1.append('Summaries - ' + Title1)

        print('Review the filnames:')
        print(filename)

        data.append([ScreenPlay])
        data.append([Outline_Details])
        data.append([Outline_Next2_Ref2])
        data.append([Plot])

        for x in range(0, 4):
            # print(data)
            try:

                MondeVert.SaveCSV(data[x], filename1[x], SavePath2)
                # MondeVert.SaveText(self,df1,'MondeVert Assistant', 'Full Transcript')

                # MondeVert.add2Master2(df1)
            except:
                print('Review Error File did not save ')

        try:
            send_email_w_attachment_gmail(up.to, up.subject, Skeleton_Story, filename)
            print('Gmail Email Sent ')
            send_email_w_attachment_outlook(up.to, up.subject, Skeleton_Story, filename)
            print('Outlook Email Sent ')

        except:
            print('email not send, its possible file was not created')

        try:
            Remove_l = (-1) * len(Summary_Episodes[4] + Summary_Episodes[5])
            Plot_NoSpoiler = Plot[Remove_l:]
            print(Plot_NoSpoiler)
        except:
            print('error on no spoiler')
        MondeVert.speak(self, ScreenPlay_Final)

        try:
            Summary_Summer_90s = MondeVert.GPTSummary(self, crazy, up.GPT_90s_Summary, Plot_NoSpoiler)
        except:
            print('error on 90s trailer')

            print('***********************************************************')
            print('***********************************************************')
            print('***********************************************************')
            print('***********************************************************')
            print("90s Movie Preview: ")
            print(Summary_Summer_90s)
            MondeVert.speak(self, Summary_Summer_90s)

            print('***********************************************************')
            print('***********************************************************')
            print('***********************************************************')
            print('***********************************************************')
            print("Character Art prompts ")

            # ''
            # if this does not work I can make replace that text with the char so it works
            # Test_Skel1 = Skeleton_Story.replace(":", "")

            print('Skeleton_Story')
            print(Skeleton_Story)

            Test_Skel1 = Skeleton_Story.replace("Characters:", "#")
            Test_Skel1 = Test_Skel1.replace("Characters", "#")

            print('Test_Skel')
            print(Test_Skel1)

            Test_Skel = Test_Skel1.split('#')
            print('Test_Skel')
            print(Test_Skel)
            Details = Test_Skel[0]

            Characters = []
            x = len(Test_Skel)
            for c in range(0, x):
                # print(c)
                if c != 0:
                    Characters.append(Test_Skel[c])

            print('Characters')
            print(Characters)

            for c in Characters:
                print(' ')
                print('-------------------------------------------')
                print('-------------------------------------------')
                # print("Character: " + c)
                print('-------------------------------------------')
                print("Character Art1: " + up.Character_Art1 + c)
                Art_PromptChar = MondeVert.GPTArt(self, crazy, up.Character_Art2 + c, Skeleton_Story)

                print("Character Art2: " + up.Character_Art2 + c)
                Art_PromptChar2 = MondeVert.GPTArt(self, crazy, up.Character_Art2 + c, Skeleton_Story)

            try:
                MondeVert.speak(self, Art_PromptChar)
                print(Art_PromptChar)
                ArtPath2 = MondeVert.makeArt(self, Art_PromptChar)
                ArtPaths.append(ArtPath2)
            except:
                dn = ''

            try:
                MondeVert.speak(self, Art_PromptChar2)
                print(Art_PromptChar2)
                ArtPath2 = MondeVert.makeArt(self, Art_PromptChar2)
                ArtPaths.append(ArtPath2)
            except:
                dn = ''

            print('***********************************************************')
            print('***********************************************************')
            print('***********************************************************')
            print('***********************************************************')
            print("Art prompts ")

            try:
                Art_Prompt = MondeVert.GPTArt(self, crazy,
                                                  up.direction_Text_Artist_ScreenPlay + up.ArtPrompt_ScreenPlay,
                                                  Skeleton_Story,
                                                  sys_prompt=up.system_Text_Artist_ScreenPlay + up.system_Text_ScreenPlay_Art1)
                MondeVert.speak(self, Art_Prompt)
                print('***********************************************************')
                print("Art prompt Movie Poster ")
                print(Art_Prompt)
                try:
                    ArtPath2 = MondeVert.makeArt(self, Art_Prompt)
                    ArtPaths.append(ArtPath2)
                except:
                    d = ''
            except:
                ArtPath2 = ''
            # Art_Prompt = Art_Prompt11.choices[0].message.content

            for x in range(0, 5):

                Summary_Episode = Summary_Episodes[x]

                try:
                    Art_Prompt1 = MondeVert.GPTArt(self, crazy,
                                                       up.direction_Text_Artist_ScreenPlay + up.ArtPrompt_ScreenPlay_Scene,
                                                       Summary_Episode,
                                                       sys_prompt=up.system_Text_Artist_ScreenPlay + up.system_Text_ScreenPlay_Art1 + up.system_Text_ScreenPlay_Art)
                except:
                    print('Error did not work but kept going')

                try:

                    Test_Skel1 = Skeleton_Story.replace("Art Prompts:", "")
                    Test_Skel1 = Test_Skel1.replace("Art Prompts", "#")
                    # ''
                    Test_Skel = Art_Prompt1.split('#')

                    Art_Prompts_pre = []
                    Art_Prompts = []
                    x = len(Test_Skel)
                    for c in range(0, x):
                        # print(c)
                        if c != 0:
                            Art_Prompts_pre.append(Test_Skel[c])

                    print(Art_Prompt1)
                    Test_Skel = Art_Prompt1.split(';')
                    print("Test_Skel: " + Test_Skel)

                    x = len(Test_Skel)
                    for c in range(0, x):
                        Art_Prompts.append(Test_Skel[c])

                    for ap in Art_Prompts:
                        MondeVert.speak(self, ap)
                        print('***********************************************************')
                        print("Art prompt Movie Poster ")
                        print(ap)
                        ArtPath2 = MondeVert.makeArt(self, ap)
                        ArtPaths.append(ArtPath2)

                except:
                    print("New stuff failed yo")

            # Art_Prompt2 = MondeVert.GPTArt(self, crazy, up.direction_Text_Artist_ScreenPlay + up.ArtPrompt_ScreenPlay_Scene, Summary_Episode2, sys_prompt= up.system_Text_Artist_ScreenPlay + up.system_Text_ScreenPlay_Art1 +up.system_Text_ScreenPlay_Art)
            # Art_Prompt3 = MondeVert.GPTArt(self, crazy, up.direction_Text_Artist_ScreenPlay + up.ArtPrompt_ScreenPlay_Scene, Summary_Episode3, sys_prompt= up.system_Text_Artist_ScreenPlay + up.system_Text_ScreenPlay_Art1 +up.system_Text_ScreenPlay_Art)
            # Art_Prompt4 = MondeVert.GPTArt(self, crazy, up.direction_Text_Artist_ScreenPlay + up.ArtPrompt_ScreenPlay_Scene, Summary_Episode4, sys_prompt= up.system_Text_Artist_ScreenPlay + up.system_Text_ScreenPlay_Art1 +up.system_Text_ScreenPlay_Art)
            # Art_Prompt5 = MondeVert.GPTArt(self, crazy, up.direction_Text_Artist_ScreenPlay + up.ArtPrompt_ScreenPlay_Scene, Summary_EpisodeFinal,sys_prompt= up.system_Text_Artist_ScreenPlay + up.system_Text_ScreenPlay_Art1 +up.system_Text_ScreenPlay_Art)
            #

        # ScreenPlay = "EPISODE I      :" +  ScreenPlay[0] +"               EPISODE II: " +  ScreenPlay2 + "       EPISODE III: " + ScreenPlay3 +"         EPISODE IV:" +  ScreenPlay4+"         EPISODE V:" +  ScreenPlayFinal

        # Title = 'SD - MondeVert Productions - Classified'

        if len(Title) > 45:
            Title = Title[0:45]

        prompt = " 90s Summer Preview: " + Summary_Summer_90s + " Details:" + Skeleton_Story + 'Summary: ' + Plot + "Work of Art Inspiration:" + Art_Prompt + "      Song:" + ScreenPlay
        Prompts_Used = [
            up.system_Text + up.RolePlay_SongArtist + "AI Created a Persona shown as the Artist Bio Above" + up.Song_prompt + up.ArtPrompt_SongArtist]
        ArtistPoetInfo = 'Lyrics Written By: ' + up.Song_Writer + '      (' + 'Artwork by: ' + up.AI_ArtistName + ')'
        MondeVert.NamePoemSavePoem(self, prompt, ArtPaths, Prompts_Used, ArtistPoetInfo, title=Title,
                                       FolderPath=up.AI_Screen_Plays + '\\' + Mode, ArtType='Screen Play' + Mode)

    def Add2MasterLyrics(self, current_time2, ArtType, title, ArtistPoetInfo, poem, Tag, SavePath2, Prompts):
        data = [(current_time2, ArtType, title, ArtistPoetInfo, poem, Tag, SavePath2, Prompts)]
        df = pd.DataFrame(data, columns=['Date_Time_Added', 'Art_Type', 'Title', 'Poet_Artist_Info', 'Poem_Song_Lyrics',
                                         'Quality', 'Folder_Path', 'Prompts_Used'])
        MondeVert.add2Master3(df)

    def quickArt(self):
        prompt = MondeVert.ChatGPTDA(self, MakeArt=True, Prompt=(up.QuickArt))

        ArtPaths = []
        try:
            ArtPath = MondeVert.makeArt(self, prompt)
            ArtPaths.append(ArtPath)
        except:
            MondeVert.speak(self, 'Could not make the Art due to an error')

        Prompts_Used = [str('Quick_Art_Prompt: ' + up.QuickArt)]
        ArtistPoetInfo = 'Written By: ' + up.Bot_Name + '      (' + 'Artwork by: ' + up.AI_ArtistName + ')'
        MondeVert.NamePoemSavePoem(self, prompt, ArtPaths, Prompts_Used, ArtistPoetInfo, title='Quick_Art')





    def SaveCSV(Text, Title, FolderPath):
        current_time1 = datetime.datetime.now()
        current_time2 = current_time1.strftime('%m-%d-%Y_%H.%M.%S')
        f2 = FolderPath
        SavePath1 = f2
        invalidCharRemoved = re.sub(r"[^a-zA-Z0-9 ]", "", Title)
        folder = str(invalidCharRemoved)
        if len(folder) > 44:
            folder = folder[0:44]
        Title1 = '\\' + str(invalidCharRemoved) + "_"
        if len(Title1) > 44:
            Title1 = Title1[0:44]
        SavePath2 = SavePath1 + '\\' + folder
        Title2 = SavePath2 + Title1 + current_time2 + '.txt'

        isExist = os.path.exists(SavePath1)
        if not isExist:
            # Create a new directory because it does not exist
            os.makedirs(SavePath1)
            print("Created new filePath: " + SavePath1)

        isExist = os.path.exists(SavePath2)
        if not isExist:
            # Create a new directory because it does not exist
            os.makedirs(SavePath2)
            print("Created new filePath: " + SavePath2)
        data = [(Text)]
        # print(data)
        try:

            print(Title2)
            df1 = pd.DataFrame(data, columns=['Text'])
            print(df1)

            df1.to_csv(Title2 )
            # MondeVert.SaveText(self,df1,'MondeVert Assistant', 'Full Transcript')
            MondeVert.add2Master2(df1)
        except:
            print('Review Error File did not save ')





    def Get_Title_GPT(self,  Text, Role,mood = '', bio = '', crazy = .5):

        # Create original details
        response = openai.ChatCompletion.create(
            model="gpt-3.5-turbo",
            messages=[
                {"role": "system", "content": up.system_TextJoaT},
                {"role": "user", "content": up.Role_Play_Prompt + Role},
                {"role": "user", "content": up.MondeVert_Title + Text}

            ], temperature=crazy
        )

        Title = response.choices[0].message.content


        Title = Title.replace("Title:", "")
        Title = Title.replace("Title", "")
        return Title

    def NamePoemSavePoem(self, poem, ArtPaths, Prompts_Used, ArtistPoetInfo, title='', FolderPath=up.AI_Poetry_Path,
                         ArtType='Poem'):
        dfPrompts = ''
        Tag = ''
        if ArtType == 'Poem':
            title_p = up.Poem_Title_Prompt
        else:
            title_p = up.Song_Title_Prompt

        if title == '':
            Title = MondeVert.Get_Title_GPT(self,  text= poem )
        else:
            Title = title

        current_time1 = datetime.datetime.now()
        current_time2 = current_time1.strftime('%m-%d-%Y_%H.%M.%S')
        f2 = FolderPath
        SavePath1 = f2
        invalidCharRemoved = re.sub(r"[^a-zA-Z0-9 ]", "", Title)
        folder = str(invalidCharRemoved)
        if len(folder) > 44:
            folder = folder[0:44]
        Title1 = '\\' + str(invalidCharRemoved) + "_"
        if len(Title1) > 44:
            Title1 = Title1[0:44]
        SavePath2 = SavePath1 + '\\' + folder
        Title2 = SavePath2 + Title1 + current_time2

#not using for now
        # Tag = ""
        # Tag += MondeVert.YayorNay(self)
        # Title2 = Title2 + Tag

        isExist = os.path.exists(SavePath1)
        if not isExist:
            # Create a new directory because it does not exist
            os.makedirs(SavePath1)

        isExist = os.path.exists(SavePath2)
        if not isExist:
            # Create a new directory because it does not exist
            os.makedirs(SavePath2)


#Adds to master files and also creates its own CSV file


        data = [(Title, ArtistPoetInfo, poem, FolderPath, ArtType, Tag)]
        # print(data)
        try:
            df1 = pd.DataFrame(data, columns=['Title', 'Artist Info', 'ScreenPlay', 'FolderPath', 'Category',
                                              'Self-Critique'])
            print(df1)
            print(Title2)
            df1.to_csv(Title2 + '.txt')
            MondeVert.Add2MasterLyrics(self, current_time2, ArtType, title, ArtistPoetInfo, poem, Tag, SavePath2,                                           dfPrompts)
            # MondeVert.SaveText(self,df1,'MondeVert Assistant', 'Full Transcript')
            try:
                MondeVert.add2Master2(df1)

            except:
                print('Did not add to master, error')
        except:
            print('Review Error')


#Word Doc Creation below, not quite what I need so I do text file too above
        document = Document()
        document.add_heading(Title, 0)
        document.add_heading(ArtistPoetInfo, 1)
        p = document.add_paragraph()
        r = p.add_run()
        # r.add_text(ArtistPoetInfo)
        p = document.add_paragraph()
        r = p.add_run()
        r.add_text(poem)
        p = document.add_paragraph()
        r = p.add_run()
        image_counter = 0
        for i in ArtPaths:
            image_counter += 1
            r.add_picture(i)
            original = i
            target = SavePath2 + Title1 + str(image_counter) + current_time2 + '.png'
            shutil.copyfile(original, target)

        document.add_heading('AI Prompts used: ', 7)

        for i in Prompts_Used:
            p = document.add_paragraph()
            r = p.add_run()
            r.add_text(i)
            dfPrompts += i + '|'
        document.save(Title2 + '_Details.docx')

        document2 = Document()
        document2.add_heading(Title, 0)
        document2.add_heading(ArtistPoetInfo, 1)

        p = document2.add_paragraph()
        r = p.add_run()
        r.add_text(poem)
        document2.save(Title2 + '.docx')

    # Blog Make Art Live mode



#test this to see how to generate images from another one
    def animate(self, image, Prompt):
        response = openai.Image.create_edit(
            image=open("A:\AI Art\Approved Quality Art\Favorites\Art_Blunts_Using_the_classic_style_of_Pablo_Picasso_create_a_unique_work_of_art_that_celebrates_the_creative_combination_of_Snoop_Dogg_and_Backwood.png", "rb"),
            mask=open("mask.png", "rb"),
            prompt="A sunlit indoor lounge area with a pool containing a flamingo",
            n=1,
            size="1024x1024"
        )
        image_url = response['data'][0]['url']



#
# from multiprocessing.dummy import Pool as ThreadPool
#
# # urls = [
# #   'http://www.python.org',
# #   'http://www.python.org/about/',
# #   'http://www.onlamp.com/pub/a/python/2003/04/17/metaclasses.html',
# #   'http://www.python.org/doc/',
# #   'http://www.python.org/download/',
# #   'http://www.python.org/getit/',
# #   'http://www.python.org/community/',
# #   'https://wiki.python.org/moin/',
# # ]
#
# # Make the Pool of workers
# pool = ThreadPool(3)
#
# # Open the URLs in their own threads
# # and return the results
# results = pool.map(urllib2.urlopen, urls)
#
# # Close the pool and wait for the work to finish
# pool.close()
# pool.join()
#





    def MultiThread(self,Functions,Args):

        thread_list = []
        print("Start")

        t = threading.Thread(target=MondeVert.MondeVertMenu, args=(Args))
        thread_list.append(t)




        # Starts threads
        for thread in thread_list:
            thread.start()
            print("New Thread Started")

        # This blocks the calling thread until the thread whose join() method is called is terminated.
        # From http://docs.python.org/2/library/threading.html#thread-objects
        for thread in thread_list:
            thread.join()

    # Demonstrates that the main process waited for threads to complete
        print ( "Done")



if __name__ == '__main__':
    # main method for executing
    # the functions
    Record = ''

    x = MondeVert()

    Args = ['Basic','Social','']

    Functions = [x.MondeVertMenu(),x.MondeVertMenu(),IP.upload_pictures()]
    x.MultiThread(Functions,Args)


    # x.MondeVertMenu(Mode='Basic')
    #
    # x.MondeVertMenu(Mode='Social')
    #
    # IP.upload_pictures()


    #x.MondeVertMenu(Role=up.Test_Role, Background=up.Test_Background, Special=up.Test_Special, Format=up.Test_Format, Mode='Music')

    # x.MondeVertMenu(Role=up.Test_Role, Background=up.Test_Background, Special=up.Test_Special, Format=up.Test_Format, Mode='Skit')
    # x.MondeVertMenu(Role=up.Test_Role, Background=up.Test_Background, Special=up.Test_Special, Format=up.Test_Format, Mode='Book')
    # x.MondeVertMenu(Role=up.Test_Role, Background=up.Test_Background, Special=up.Test_Special, Format=up.Test_Format, Mode='Series')
    #
    # x.MondeVertMenu(Role=up.Test_Role, Background=up.Test_Background, Special=up.Test_Special, Format=up.Test_Format, Mode='Art')
    # x.MondeVertMenu(Role=up.Test_Role, Background=up.Test_Background, Special=up.Test_Special, Format=up.Test_Format, Mode='Advance')


    #Image1 = x.animate(image= r"A:\AI Art\DALLE_please_create_an_art_piece_that_showcases_a_luxurious_and_modern_real_estate_property_The_property_should_have_a_stunning_view_and_be_surroun.png", Prompt = "Add a Stick Figure with long blonde hair to the following image")
    #Image1 = x.Take_query(image = Image1, Prompt= 'Make a variation where everything is the same but you put the stick figure in a slightly different position')

    # Testing
    # x.Make_a_ScreenPlay_Testing()
    # x.Make_a_ScreenPlay()

    # x.saveTranscript()
    atexit.register(x.saveTranscript)


