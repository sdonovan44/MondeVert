import glob
import datetime
import pandas as pd
import wikipedia
global Record
import sys
import ShakesBot as s
import Common_Sayings as cs
from threading import Event
import User_Prefs as up
from gingerit.gingerit import GingerIt
import threading
import numpy
from docx import Document
#from exceptions import PendingDeprecationWarning
import requests
import cloudscraper
import tensorflow
import selenium

Record = ''
import webbrowser


import User_Prefs as up
import os
import platform
import subprocess

import requests
import random
import string
import openai
import pyttsx3
# Import the speech recognition library
import speech_recognition as sr
from dotenv import load_dotenv



URL = "..."  # noqa
API_KEY = "..."






        #make a transcript class that takes down all of my requests and saves to a text file
        #have specific command for the assistant to save in a certain folder (always add date & time to the file
        #create things like to do lists
        #ideas
        #stories and dialogue
        #the key is being able to store it properly and sytemically access as well
        #every program I create should be executeable through the assistant



        #Urgent to do
        #Make it a class
        #have transcribe and also proofread (later) add method where in the middle of dictating you can add titles so it gets properly makred eventually
        # bullets etc. parenthesis
        #Have it so every pause puts a new line in between so its easy to pull out pauses etc.
        #Most important is to have a mode where it callse poembot and reads the lines to me.
        # .... also I would like for it to be a mini drama where I get multiple voices and characters that read
        # the output etc eventually learn how to animate/make a nice script.
        # create a series so eventually we might have a way to feed in characters to poem bot so its not just random like it is now lol

        #12/18/2022 get the assistant able to open the respective web browsers
        #set up mouse so it has shortcuts
        #be able to run on both mac and also the pc
        #import the site view bot
        #update the youtube bot so it hides ips properly and also so it loops through videos to not draw attention
        #    (make it so it starts off slow then has pockets of 100 here and there big days occasionally and find out how to make it work properly
        #have assistant able to run those on command (buy wifi adapter so I can have it running whenever & get new monitor for that too




        #Update transcribe commands (Review file so far,#Have a separate function to make the title (different than subject)
        ##Also I should be able to mark the file as important based on results (add tag)
        #Make a program that goes through and corrects the script into a new tab (puts subject as start of the word blocks (with other data in parenthesis for tagging)
        #Also make it so that the poems do not overwrite but get concatenated together
        #Make a theatre method that uses multithreading to do stuff
        #overall threads should allow this to run faster
        #Make it so I can update all meta data and also add more metadata
        #Goal, mark as header (prior or next maybe)


#
# class GingerIt(object):
#     def __init__(self):
#         self.url = URL
#         self.api_key = API_KEY
#         self.api_version = "2.0"
#         self.lang = "US"
#
#     def parse(self, text, verify=True):
#         session = cloudscraper.create_scraper()
#         request = session.get(
#             self.url,
#             params={
#                 "lang": self.lang,
#                 "apiKey": self.api_key,
#                 "clientVersion": self.api_version,
#                 "text": text,
#             },
#             verify=verify,
#         )
#         data = request.json()
#         return self._process_data(text, data)


class DigitalAssist():
    def __init__(self, voice = 1, language_settings=1):
        self.voice = voice
        self.language_settings = language_settings
        self.AssistantName = up.getAssistantName()
        self.UserName = up.UserName()

    def Open_Web(url):
        url = url
        webbrowser.open_new_tab(url)


    def speak(self,audio):
        engine = pyttsx3.init()
        # getter method(gets the current value
        # of engine property)
        voices = engine.getProperty('voices')
        engine.setProperty('voice', voices[self.voice].id)
        # Method for the speaking of the assistant
        engine.say(audio + "")
        # Blocks while processing all the currently
        # queued commands
        engine.runAndWait()

    def speakSweet(audio):
        engine = pyttsx3.init()
        # getter method(gets the current value
        # of engine property)
        voices = engine.getProperty('voices')
        engine.setProperty('voice', voices[15].id)
        # Method for the speaking of the assistant
        engine.say(audio + "")
        # Blocks while processing all the currently
        # queued commands
        engine.runAndWait()


        #def RandomCharacters

    def testvoices(self):
        #import pyttsx3
        engine = pyttsx3.init()
        voices = engine.getProperty('voices')
        index = 0
        for voice in voices:

            print(f'index-> {index} -- {voice.name}')
            engine.setProperty('voice', voices[index].id)
            engine.say("Shane do you like this voice more than the last one you busy little bee turkish" + "")

            index += 1
        engine.runAndWait()



    def Setvoices(self):
        #import pyttsx3
        engine = pyttsx3.init()
        voices = engine.getProperty('voices')
        index = 0

        prompt = 'Write a short paragraph that uses different prononciations and complex words while being warm and welcoming to a person named Shane'
        completions = openai.Completion.create(
            engine="text-davinci-002",
            prompt=prompt,
            max_tokens=2048,
            n=1,
            stop=None,
            temperature=0.5,
        )
        speak1 = 'Do you want to set this voice as active?'

        # Print the generated text
        message = completions.choices[0].text
        DigitalAssist.Add2Transcript(message)

        for voice in voices:



            print(f'index-> {index} -- {voice.name}')
            engine.setProperty('voice', voices[index].id)
            engine.say(message)
            engine.say(speak1)
            Query = DigitalAssist.getUserResponse()
            DigitalAssist.Add2Transcript(speak1 + ' Voice: ' + index)
            DigitalAssist.Add2Transcript(Query)

            if "yes" in Query or "set" in Query or ("make" in Query and "active" in Query) or  ("set" in Query and "active" in Query):
                self.voice = index
                continue


            index += 1
        engine.runAndWait()



    def tellDay(self):
        # This function is for telling the
        # day of the week
        day = datetime.datetime.today().weekday() + 1

        # this line tells us about the number
        # that will help us in telling the day
        Day_dict = {1: 'Monday', 2: 'Tuesday',
                    3: 'Wednesday', 4: 'Thursday',
                    5: 'Friday', 6: 'Saturday',
                    7: 'Sunday'}

        if day in Day_dict.keys():
            day_of_the_week = Day_dict[day]
            print(day_of_the_week)
            DigitalAssist.speak("The day is " + day_of_the_week)


    def tellTime():
        # This method will give the time
        time = str(datetime.datetime.now())
        print(time)
        hour = time[11:13]
        min = time[14:16]
        t1 = cs.DA_Time1()
        t2 = cs.DA_Time2()
        t3 = cs.DA_Time3()
        DigitalAssist.speak( t1 + hour + t2+ min + t3)


    def Hello():
        # This function is for when the assistant
        # is called it will say hello and then
        # take query
        r = cs.DAgreet()
        print (r)
        DigitalAssist.speak(r)

    def takeCommand(self):
        Query = DigitalAssist.getUserResponse(Response = "Action Requested")

        return Query

############################################################################################################################################################################################################################
################Above is code I did not write, below is ChatGPT
############################################################################################################################################################################################################################

    def makeArt(self):
        print()
        prompt = input("Text prompt for image generation (e.g. 'A blue eagle flying over a desert at sunset'): ")
        print()

        # Set up the OpenAI API client :
        openai.api_key = up.getAPIKey()
        print("Sending to OpenAI...")
        print()

        response = openai.Image.create(
            prompt=prompt,
            n=1,
            size="1024x1024"
        )

        image_url = response['data'][0]['url']

        print(f"Image URL: {image_url}")
        print()

        # URL of the image to be downloaded is defined as image_url
        r = requests.get(image_url)  # create HTTP response object

        # send a HTTP request to the server and save
        # the HTTP response in a response object called r
        fname = f"{''.join([c for c in prompt.strip().replace(' ', '_') if c.isalnum() or c == '_'])}.png"

        if not os.path.exists("images"):
            os.mkdir("images")

        if os.path.isfile(os.path.join("images", fname)):
            fname = fname.split(".")[0] + f".{''.join(random.choice(string.ascii_letters) for x in range(5))}.png"

        fname = os.path.join("images", fname)
        print(f"Filename: {fname}")
        with open(fname, 'wb') as f:

            # Saving received content as a png file in
            # binary format

            # write the contents of the response (r.content)
            # to a new file in binary mode.
            f.write(r.content)

        if platform.system() == 'Darwin':  # macOS
            subprocess.call(('open', fname))
        elif platform.system() == 'Windows':  # Windows
            os.startfile(fname)
        else:  # linux variants
            subprocess.call(('xdg-open', fname))

    # Define a function that sends a message to ChatGPT
    def chat_query(self,prompt):
        model_engine = "text-davinci-003"
        completions = openai.Completion.create(
            engine=model_engine,
            prompt=prompt,
            max_tokens=2048,
            n=1,
            temperature=0.5,
        )

        message = completions.choices[0].text
        return message

    # Define a function that handles the conversation
    def conversation_handler(self,prompt):
        # Send the prompt to ChatGPT

        load_dotenv()
        # Set up the OpenAI API client :
        openai.api_key = up.getAPIKey()
        response = DigitalAssist.chat_query(prompt)
        print(f"ChatGPT: {response}")

        # End the conversation if ChatGPT says goodbye
        if "goodbye" in response.lower():
            return

        # Otherwise, get user input and continue the conversation
        prompt = input("You: ")
        DigitalAssist.conversation_handler(prompt)

    def ChatGPTDA(self):
        # Import the OpenAI library

        # Set up the OpenAI API client :
        openai.api_key = up.getAPIKey()

        # Set up the recognizer
        r = sr.Recognizer()

        # Import the text to speech synthesis library

        engine = pyttsx3.init()
        engine.setProperty('rate', 180)  # setting up new voice rate
        voices = engine.getProperty('voices')
        engine.setProperty('voice', voices[2].id)  # changing index, changes voices. 1 for female

        # Record the audio
        engine.say(" Do you have any questions to ask Chat GPT?")
        engine.runAndWait()
        print('---------------------Please speak now, recording start-----------------------')
        with sr.Microphone() as source:
            audio = r.listen(source)
        print('---------------------Recording finished---------------------------------------')

        # Convert the audio to text
        prompt = r.recognize_google(audio)
        print(prompt)

        # check if the reply contains 'yes'
        while 'YES' in prompt.upper() or 'QUESTION' in prompt.upper():
            # Record the audio
            engine.say(" Please ask your question and get help from Chat GPT?")
            engine.runAndWait()
            print('---------------------Please speak now, recording start-----------------------')
            with sr.Microphone() as source:
                audio = r.listen(source)
            print('---------------------Recording finished---------------------------------------')

            # Convert the audio to text
            prompt = r.recognize_google(audio)

            # Generate text
            completions = openai.Completion.create(
                engine="text-davinci-002",
                prompt=prompt,
                max_tokens=2048,
                n=1,
                stop=None,
                temperature=0.5,
            )

            # Print the generated text
            message = completions.choices[0].text
            print(message)

            # generate the feedback
            engine.say(message)
            engine.runAndWait()
            # engine.stop()

            # Record the audio
            engine.say(" Any other questions to ask Chat GPT?")
            engine.runAndWait()
            print('---------------------Please speak now, recording start-----------------------')
            with sr.Microphone() as source:
                audio = r.listen(source)
            print('---------------------Recording finished---------------------------------------')

            # Convert the audio to text
            prompt = r.recognize_google(audio)
            print(prompt)

    ############################################################################################################################################################################################################################
    ################Above  is ChatGPT (I did not write original code but adapted for my use case)
    ############################################################################################################################################################################################################################

    def Add2Transcript(self, text2Add):
        self.transcript_Final += text2Add + ' \n'

    def getUserResponse(self, pause =.5, Response = 'You Said' ):

        r = sr.Recognizer()
        with sr.Microphone() as source:
            print('Listening...')
            r.pause_threshold = pause
            audio = r.listen(source)
            try:
                Query = r.recognize_google(audio, language='en-in')
                print(Response+ ": ", Query)

            except Exception as e:
                print(e)
                #uk2 = cs.DA_Unknown_Command_Serious()
                #DigitalAssist.speak(uk2)
                print("Please repeat Words not understood...")
                Query = DigitalAssist.getUserResponse()
                #Event().wait(2)

            return Query



    def transcribe_Build_Query(self, pause = 3.13):

        Query = DigitalAssist.getUserResponse()

        return Query


    def transcribe_Build_Query_Pause(self):
        s = True
        while (s == True):
            Speak1 = 'What Tool do you want to use?'
            Print1 = 'Options: Change Subject, Edit, Poem, Stop Recording --> Build others here'
            DigitalAssist.speak(Speak1)
            print(Print1)
            Query = DigitalAssist.getUserResponse(Response="Tool Selected")
            DigitalAssist.Add2Transcript(Speak1)
        return Query




    def getdata(self, Field):

        DigitalAssist.speak('what is the ' + Field)
        print('Please confirm the ' + Field)
        s = True
        Query = ''
        while (s == True):

            Query = DigitalAssist.getUserResponse()

            Query2 = ''
            WaitforResponse = False
            while (WaitforResponse == False):
                DigitalAssist.speak(Query + ' Is that correct?')

                Query2 = DigitalAssist.getUserResponse()


                    if  "yes" in Query2 or "correct" in Query2 or "ya" in Query2 or "yeah" in Query2:
                        DigitalAssist.speak('Thanks for confirming the '+ Field )
                        s = False
                        WaitforResponse = True
                        continue
                    if "no" in Query2 or "wrong" in Query2 or "not it" in Query2 or "nope" in Query2:
                        DigitalAssist.speak('ok  please say the ' + Field+ ' you want to set' )
                        print('ok  please say the ' + Field + ' you want to set' )
                        WaitforResponse = True
                        continue
                    else:
                        DigitalAssist.speak('Lets try that again what is the ' +Field)
                        Event().wait(1)
                        continue
            return Query


    def editentry(self, Query1):
        s = True
        while (s == True):
            DigitalAssist.speak('I am ready to make the edit please proceed')
            print('ready to make the edit please proceed')
            print('Original Entry: ' + Query1)
            Query2 = DigitalAssist.getUserResponse()

            ss = True
            while (ss == True):
                print(Query2)
                DigitalAssist.speak(Query2 + 'Are you satisfied with your edits?')
                print('Options: Yes/Correct, Not yet/Redo, Add as Note, Combine, Cancel')

                Query3 = DigitalAssist.getUserResponse()


                if  "yes" in Query3 or "correct" in Query3 or "ya" in Query3 or "yeah" in Query3 or "looks good" in Query3 or "keep new" in Query3 or "new one" in Query3:
                    DigitalAssist.speak('Thanks for confirming, I will make the updates you provided')
                    print('Thanks for confirming, I will make the updates you provided')
                    Query_Final = Query2
                    s = False
                    ss = False


                elif  "not yet" in Query3 or 'redo'  in Query3 or 'try again'  in Query3 or 'again'  in Query3 or 'one more' in Query3:
                    DigitalAssist.speak('ok, lets go again')
                    print('ok, lets go again')
                    ss = False
                    continue

                elif "add as note" in Query3 or "include note" in Query3 or "as note" in Query3  :
                    DigitalAssist.speak('Thanks for confirming, I will add this as a note')
                    print('Thanks for confirming, I will add this as a note')
                    Query_Final = ''
                    self.Notes = Query2
                    #when this is a class I can update the df for now I will add it to the end same as when I say combine them
                    s = False
                    ss = False

                elif "add to original" in Query3 or "both" in Query3 or 'combine' in Query3 or 'combination' in Query3 or 'merge' in Query3 or 'together' in Query3:
                    DigitalAssist.speak('Thanks for confirming, I will join the two thoughts')
                    print('Thanks for confirming, I will join the two thoughts')
                    Query_Final = Query1 + ' ' + Query2
                    # when this is a clas I can update the df for now I will add it to the end same as when I say combine them
                    s = False
                    ss = False

                elif 'cancel'   in Query3  or 'end' in Query3 or 'stop'  in Query3 or  'stop trying' in Query3:
                    DigitalAssist.speak('Thanks for confirming, I will remove these edits')
                    print('Thanks for confirming, I will remove these edits')
                    Query_Final = ''
                    s = False
                    ss = False

                else:
                    Event().wait(1)


        return Query_Final







    def getfilename(self):
        self.FileName = DigitalAssist.getdata(self, 'FileName')



    def  getsubject(self):
        self.subject = DigitalAssist.getdata(self, 'subject')



    def  getDictation_Type(self):
        self.Dictation_Type  =  DigitalAssist.getdata(self, 'Type')



    def  getSignificance(self):
        self.Significance =  DigitalAssist.getdata(self, 'Significance')


    def  getNote(self):
        self.Notes += '|' + DigitalAssist.getdata(self, 'Notes')

############################################################################################################################################################################################################################
################Long Code where I give Digital Assistant commands##
############################################################################################################################################################################################################################


    def Take_query(self):
        DigitalAssist.Hello()

        self.Mode = 'Main Menu'
        DigitalAssist.Add2Transcript(self.mode + ':')
        s2 = True
        while (s2 == True):
            query = DigitalAssist.takeCommand(self).lower()
            DigitalAssist.Add2Transcript(query)
            if "company dashboard" in query or "open dashboard" in query :
                DigitalAssist.speak("Opening MondDay Vert's Dasboard, good luck with your work!")

                # in the open method we just to give the link
                # of the website and it automatically open
                # it in your default browser
                Mondevert = "http://mondeVert.co"
                godaddy = 'https://dashboard.godaddy.com/'
                Mail = 'https://outlook.office.com/mail/'
                Square_store = 'https://square.online/app/home/'
                DigitalAssist.Open_Web(Mail)
                DigitalAssist.Open_Web(Square_store)
                DigitalAssist.Open_Web(Mondevert)
                DigitalAssist.Open_Web(godaddy)
                continue

            elif "pay bills" in query or "pay credit cards" in query:
                DigitalAssist.speak("Bills Suck man, at least you are getting them done, Hope you have a nice day!")

                Liberty_Bay_CC = "https://www.myaccountaccess.com/onlineCard/login.do"
                Home_Depot_CC = 'https://citiretailservices.citibankonline.com/RSnextgen/svc/launch/index.action?siteId=PLCN_HOMEDEPOT#signon'
                Macys_CC = 'https://www.macys.com/my-credit/gateway/guest'
                Insurance = 'https://customer.concordgroupinsurance.com/ccp/login'
                National_Grid = 'https://login.nationalgridus.com/loginnationalgridus.onmicrosoft.com/oauth2/v2.0/authorize?cancel_uri=https%3A%2F%2Fmyaccount.nationalgrid.com%2Fservices%2Fauth%2Fsso%2FNGP_SignIn_MA_Electric_Home&client_id=88d004b4-3d39-4599-b410-093849907ee5&customer_type=home&p=B2C_1A_UWP_NationalGrid_convert_merge_signin&redirect_uri=https%3A%2F%2Fmyaccount.nationalgrid.com%2Fservices%2Fauthcallback%2FNGP_SignIn_MA_Electric_Home&region=masselec&response_type=code&scope=https%3A%2F%2Flogin.nationalgridus.com%2Fapi%2Fread+openid+profile+email+offline_access&state=CAAAAYUokN2YMDAwMDAwMDAwMDAwMDAwAAAA8Iy1JaCJ8hTkDIfqZ5i6KE2oIrQFv5bPH9zBI_P-5wWki9riiDhgLKrNlnVD4FeR38QBEzYH17vRm_CbsR8msY3J3fEDKFqtBsup_Bp9KwM-Flb38WYF6YqeqejsI-Iuj-yWpXfajw6-QKG_DMoOBtkqcL8zubs4c8KgKiuF-yIk3yoa1_vp_KqVAuBw0DSopTJrDwsYAZN5qW9zaUmwyW-wk6zbTNddqu24EJObkHkfEGKHEKQwrqArl40WukP3dxbKBfj7sSMhdqjTPa0ido8%3D&switch_uri=https%3A%2F%2Fmyaccount.nationalgrid.com%2Fs%2Flogin'
                LLBean_CC = 'https://citiretailservices.citibankonline.com/RSauth/signon?pageName=signon&siteId=PLCN_LLBEAN&langId=en_US'
                Amazon_CC = 'https://secure07ea.chase.com/web/auth/#/logon/logon/chaseOnline'
                DigitalAssist.Open_Web(Liberty_Bay_CC)
                DigitalAssist.Open_Web(Home_Depot_CC)
                DigitalAssist.Open_Web(Macys_CC)
                DigitalAssist.Open_Web(Insurance)
                DigitalAssist.Open_Web(National_Grid)
                DigitalAssist.Open_Web(LLBean_CC)
                DigitalAssist.Open_Web(Amazon_CC)
                continue


            elif ("chat" in query and "bot" in query) or  ("gpt" in query and "chat" in query) or  ("smart" in query and "assist" in query) or  ("extra" in query and "help" in query):
                DigitalAssist.speak("High Tech!")
                DigitalAssist.ChatGPTDA()
                continue

            elif ("make" in query and "art" in query) or ("make" in query and "picture" in query) or ("ai" in query and "art" in query) or ("art" in query and "mode" in query):
                DigitalAssist.speak("High Tech and sheek!")
                DigitalAssist.makeArt()
                continue




            elif "stop" in query or "all set" in query or "quit" in query or "m done" in query or "thank" in query or "peace out" in query:
                DigitalAssist.speak("Happy to help, Peace out brothah man")
                s2 = False

            elif ("record" in query or "listen up" in query or "can you right" in query or "record mode"  in query  or  "transcribe" in query or "could you write" in query or "transcribe" in query or "can you write" in query) and 'stop' not in query :
                DigitalAssist.speakSweet("Dictated but not Red Mode Activated")
                DigitalAssist.transcribe(self)
                continue


            elif ( "blog" in query )and ("create" in query or "together" in query or "put"in query or "make"in query ) or ("add" in query and "blog" in query):
                DigitalAssist.speakSweet("Mon Day Vert is lucky to have a great CEO this is going to be a great post!")
                DigitalAssist.CompileBlog(self)
                DigitalAssist.speakSweet("Blog Consolidation Complete Shane D")
                continue

            elif (("consolidate" in query or "compile" in query  or "collect" in query )and ("file" in query or "folder" in query or "list"in query )) or ("add" in query and "master" in query):
                DigitalAssist.speakSweet("Consolidation in progress")
                DigitalAssist.compileFiles(self)
                DigitalAssist.speakSweet("Consolidation Complete my dude")
                continue

            elif "make list" in query or " list" in query or "to do" in query or "tracker" in query or "checklist" in query  or  "make a plan" in query or "quick note" in query  or "quick plan" in query or "take note" in query:
                #DigitalAssist.speak("Sure thing, One moment while I prepare")
                DigitalAssist.makelist(self)
                continue


            elif "poetic" in query or "swoon me" in query or "me a poem" in query  or  "poem" in query or "while I think you" in query or "poetry will help me think" in query or "write me a poem" in query or "we need some inspiration" in query or "some art" in query or "be creative" in query or "inspire me" in query:
                x22 = s.PoemBot(1, 1, 1, 1, 40, 3)
                x22.ReloadModel("model.h5")
                x22.setupdata()
                self.Words = x22.shakesbot_DA()
                print (self.Words)
                DigitalAssist.speakSweet(self.Words)
                continue




            elif ("play" in query or "script" in query) and ("make" in query or "write" in query or "create" in query or "mode" in query):
                # Call for the save file name after listing all of the subjects to the user
                self.Words = ''
                self.Correction_Comment = ''
                self.Type = 'AI Script Break'
                self.Significance = '**AI content attached'
                self.FileName = 'AI ShakeScript'
                self.Subject = 'AI ShakeScript'
                DigitalAssist.speak('I like your style, my man')

                DigitalAssist.Make_Script(self)

                DigitalAssist.speakSweet(self.Words)
                s3 = pd.DataFrame()
                s3['Script'] = [self.Words]
                s3['Script - Corrected'] = [self.AI_Corrected_Content]
                s3['Correction Comment'] = [self.Correction_Comment]
                DigitalAssist.SaveText(s3, self.FileName, self.Subject)
                continue

            elif ("go" in query or 'try' in query or "mode" in query) and ("greek" in query or "myth" in query ):
                # Call for the save file name after listing all of the subjects to the user
                self.Words = ''
                self.Type = 'Greek Script Break'
                self.Significance = '**AI content attached'
                self.FileName = 'AI Going Greek'
                self.Subject = 'AI Going Greek'
                DigitalAssist.speak('I like your style, my man... Dont eat too much tzatziki!')
                self.pb = s.PoemBot(1, 1, 1, 1, 40, 3)
                self.pb.ReloadModel("model_Greek.h5")
                self.pb.setupdata()
                size = 3
                trd = numpy.empty(size, dtype=str)
                w = numpy.empty(size, dtype=str)
                for i in range(0, size - 1):
                    globals()[f"trd{i}"] = threading.Thread(target=DigitalAssist.generateGreek(self))
                    w[i] = globals()[f"trd{i}"].start()  # starting the thread i

                for i in range(0, size - 1):
                    threading.active_count()
                    threading.enumerate()
                    if i == 0:
                        self.Words += w[i]
                    else:
                        i2 = i - 1
                        globals()[f"trd{i2}"].join()
                        globals()[f"trd{i}"].join()
                        self.Words += w[i]
                        print(self.Words)

                DigitalAssist.speak(self.Words)
                s3 = pd.DataFrame()
                s3['Script'] = [self.Words]
                DigitalAssist.SaveText(s3, self.FileName, self.subject)
                continue



            elif "social media" in query or "marketing" in query:
                DigitalAssist.speak("Will do")

                # in the open method we just to give the link
                # of the website and it automatically open
                # it in your default browser
                Facebook = "https://www.facebook.com/"
                Insta = 'https://www.instagram.com/'
                TikTok = 'https://www.tiktok.com/'
                Twitter = 'https://twitter.com/home'
                DigitalAssist.Open_Web(Facebook)
                DigitalAssist.Open_Web(TikTok)
                DigitalAssist.Open_Web(Twitter)
                DigitalAssist.Open_Web(Insta)
                continue

            elif "open google" in query:
                DigitalAssist.speak("Opening Google ")
                url = "http://google.com"
                DigitalAssist.Open_Web(url)
                continue



            elif "preview" in query or "show case" in query or "showcase" in query or "your" in query or "voice" in query:
                DigitalAssist.s ##Fix here
                continue
            elif "s the date" in query:
                DigitalAssist.tellDay(self)
                continue

            elif "the time" in query or "what time" in query or "clock mode" in query:
                DigitalAssist.tellTime(self)
                continue

            # this will exit and terminate the program
            elif "bye" in query:
                DigitalAssist.speak("Peace Out!")
                s2 = False

            elif "from wikipedia" in query or "wikipedia" in query:
                DigitalAssist.speak("Checking the wikipedia ")
                query = query.replace("wikipedia", "")
                result = wikipedia.summary(query, sentences=4)
                DigitalAssist.speak("According to wikipedia")
                DigitalAssist.speak(result)



            elif "tell me your name" in query or  "your name" in query or "introduce yourself" in query or "introduce your self" in query:
                DigitalAssist.speak("I am Big Master Funk the fourth also known as Brick top, AND i AM  Your personal desktop Assistant")


            elif "shut up" in query:
                DigitalAssist.speak("my bad I am tripping, I will leave you be Shane D")
                s2 = False

            else:
                #uk = cs.DA_Unknown_Command()
                #DigitalAssist.speak(uk)
                #print(uk)
                #Event().wait(1)
                continue




    def transcribe(self):
        self.entry = 1

        DigitalAssist.createDF(self)
        self.subject = 'Default'
        self.FileName = 'Default'
        s2 = True
        DigitalAssist.getsubject(self)
        self.visualList = 'Subject: ' + self.subject + ' \n'
        self.visualList2 = 'Subject*: ' + self.subject + ' \n'
        DigitalAssist.speak("Please prepare for recording")
        print("Please prepare for recording...")
        while (s2 ==True):
            xx = ''
            self.query1 = ''
            self. new_Query = ''
            self.Type = ''
            self.Significance = ''
            self.Words = ''
            self.Notes = ''
            self.Added2TextFile = ''
            self.Completed = ''
            print(self.visualList)
            print(self.visualList2)
            self.query1 = DigitalAssist.transcribe_Build_Query(self,1.3).lower()
            self.query1 = self.query1 + '.'
            self.Type = 'Original'
            self.Significance = 'Default'



            if "stop recording" in self.query1 or "m all set" in self.query1 or "end list" in self.query1 or "stop listening" in self.query1 or "save this file" in self.query1 or "save this recording" in self.query1:
                # Call for the save file name after listing all of the subjects to the user
                DigitalAssist.getfilename(self)
                DigitalAssist.SaveText(self.transcript, self.FileName, self.subject)
                DigitalAssist.add2Master(self.transcript)
                s2 = False
                sys.exit()
                continue


            elif (("brick top" in self.query1 or "bricktop" in self.query1 or "show me" in self.query1 or "extra" in self.query1 or "additional" in self.query1) and ("option" in self.query1 or  "tool" in self.query1)) or ("edit mode" in self.query1) or  ("need to" in self.query1 and "make edit" in self.query1):
                s3 = True
                while (s3 == True):
                #this is where you say next steps (This can be a separate function)
                    query2 = DigitalAssist.transcribe_Build_Query_Pause(self).lower()
                    print(query2)

                    if  "continue" in query2 or "just thinking" in query2 or "yes" in query2 or "chill" in query2 or "just thinking" in query2 or "keep going" in query2:
                        DigitalAssist.speak('no worries, still here')
                        s3 = False
                        #continue



                    elif "change subject" in self.query1 or "update subject" in self.query1 or "switch category" in self.query1 or "change category" in self.query1:
                        # transcript2[self.subject] = [str(self.visualList)]
                        # DigitalAssist.SaveText(self.transcript2, self.FileName + 'v2' + self.subject, self.subject)
                        DigitalAssist.getsubject(self)
                        s3 = False
                        #continue


                    elif  "change" in query2 or "fix" in query2 or "edit" in query2 or "fix that" in query2 or "not quite" in query2 or "review last" in query2 or "back to me" in query2 or "read that back" in query2:
                        # Call for the save file name after listing all of the subjects to the user
                        self.Type = 'Edited'
                        self.Significance = 'Review edits'
                        self.new_Query = DigitalAssist.editentry(self,self.query1)
                        s3 = False
                        #continue


                    elif  "poem" in query2 or "while I think you" in query2 or "poetry will help me think" in query2 or "write me a poem" in query2 or "we need some inspiration" in query2 or "some art" in query2 or "be creative" in query2 or "inspire me" in query2:
                        # Call for the save file name after listing all of the subjects to the user
                        self.Type = 'Poetry Break'
                        self.Significance = '**AI content attached'
                        DigitalAssist.speakSweet('I like your style, Shane you are brilliant and beautiful')
                        x = s.PoemBot(1, 1, 1, 1, 40, 3)
                        x.ReloadModel("model.h5")
                        x.setupdata()
                        self.Words += x.shakesbot_DA()
                        print (self.Words)
                        DigitalAssist.speakSweet(self.Words)
                        s3 = False
                        #continue



                    elif ( "play" in query2 or "script" in query2 )and ("make" in query2 or "write" in query2 or "create" in query2 or "mode" in query2 ):
                        # Call for the save file name after listing all of the subjects to the user
                        self.Type = 'AI Script Break'
                        self.Significance = '**AI content attached'

                        DigitalAssist.speak('I like your style, my man')


                        DigitalAssist.Make_Script(self)
                        DigitalAssist.speakSweet(self.Words)

                        s3 = False
                        #continue
                    elif ("go" in query or 'try' in query2 or "mode" in query2) and ("greek" in query2 or "myth" in query2):
                        # Call for the save file name after listing all of the subjects to the user
                        self.Words = ''
                        self.Type = 'Greek Script Break'
                        self.Significance = '**AI content attached'
                        self.FileName = 'AI Going Greek'
                        self.Subject = 'AI Going Greek'
                        DigitalAssist.speak('I like your style, my man... Dont eat too much tzatziki!')
                        self.pb = s.PoemBot(1, 1, 1, 1, 40, 3)
                        self.pb.ReloadModel("model_Greek.h5")
                        self.pb.setupdata()
                        size = 3
                        trd = numpy.empty(size, dtype=str)
                        w = numpy.empty(size, dtype=str)
                        for i in range(0, size - 1):
                            globals()[f"trd{i}"] = threading.Thread(target=DigitalAssist.generateGreek(self))
                            w[i] = globals()[f"trd{i}"].start()  # starting the thread i

                        for i in range(0, size - 1):
                            threading.active_count()
                            threading.enumerate()
                            if i == 0:
                                self.Words += w[i]
                            else:
                                i2 = i - 1
                                globals()[f"trd{i2}"].join()
                                globals()[f"trd{i}"].join()
                                self.Words += w[i]
                                print(self.Words)

                        DigitalAssist.speak(self.Words)

                        continue
                    elif "stop recording" in query2 or "m all set" in query2 or "no thanks" in query2 or "stop recording" in query2 or "save" in query2 or "no" in query2:
                        #Call for the save file name after listing all of the subjects to the user
                        print(self.subject)
                        DigitalAssist.getfilename(self)
                        DigitalAssist.SaveText(self.transcript, self.FileName, self.subject)
                        DigitalAssist.add2Master(self.transcript)
                        s3 = False
                        s2 = False
                        continue


                    elif query2 == 'none':
                        zz = 1
                        #Event().wait(1)

                    else:
                        zz = 1
                        #Event().wait(1)

            elif self.query1 == 'none':
                #Event().wait(1)
                continue
            else:

                self.AI_Corrected_Text = DigitalAssist.cleanText(self, self.query1)
                DigitalAssist.adddata2DF(self)
                self.visualList += str(self.entry) + '). ' + self.query1 + ' \n'
                self.visualList2 += str(self.entry) + '). ' + self.AI_Corrected_Text + ' \n'
                self.entry += 1

            #Do the dataframe stuff here


    def StartThread(self):
        trd1 = threading.Thread(target=DigitalAssist.Make_Script2(self))
        self.w[i] = trd1.start()
        x = DigitalAssist.Make_Script2(self)
        print(x)

        self.threads.append(trd1)

    def Make_Script(self):
        size = 1
        self.threads = []
        self.w = numpy.empty(size,dtype=str)
        for self.i in range(0, size):
            DigitalAssist.StartThread(self)
            print(self.w[i])
            # globals()[f"trd{i}"] = threading.Thread(target=DigitalAssist.Make_Script2(self))
            # w[i] = globals()[f"trd{i}"].start()

        for x in self.threads:
            x.join()

        for ii in range(0, size):
            self.Words += str(self.w[ii])

        self.AI_Corrected_Content = DigitalAssist.cleanText(self, self.Words)



    def Make_Script2(self):
        pb = s.PoemBot(1, 1, 1, 1, 40, 3)
        pb.ReloadModel('model.h5')
        pb.setupdata()
        w1 = pb.shakesbot_DA_Make_Script()
        DigitalAssist.speakSweet(w1)
        #print (w1)
        return w1

    def generateGreek(self):
        w2 = self.pb.generateGreek() + ' \n' + ' \n'
        return w2



    def makelist(self):
         self.entry = 1
         self.transcript2 = pd.DataFrame()
         #self.transcript = pd.DataFrame()
         DigitalAssist.createDF(self)
         self.subject = 'To Do'
         self.FileName = 'To Do '
         s2 = True
         DigitalAssist.getsubject(self)
         DigitalAssist.speak("Please prepare for recording")
         print("Please prepare for recording...")
         self.visualList = ''
         self.visualList2 = ''
         self.Type = 'To Do list'
         self.Significance = 'High'
         while (s2 == True):
             self.xx = ''
             self.query1 = ''
             self.new_Query = ''
             self.Words = ''
             self.Notes = ''
             self.Added2TextFile = ''
             self.Completed = ''
             print(self.visualList)
             print(self.visualList2)
             self.query1 = DigitalAssist.transcribe_Build_Query(self,1.3).lower()
             self.query1 = self.query1 + '.'




             if "stop recording" in self.query1 or "m all set" in self.query1 or "end list" in self.query1 or "stop listening" in self.query1 or "save this file" in self.query1 or "save this recording" in self.query1:
                 # Call for the save file name after listing all of the subjects to the user
                 DigitalAssist.SaveText(self.transcript,self.FileName,self.subject)
                 DigitalAssist.add2Master(self.transcript)
                 s2 = False
                 sys.exit()
                 continue

             elif "change subject" in self.query1 or "update subject" in self.query1 or "switch category" in self.query1 or"change category" in self.query1:
                 #transcript2[self.subject] = [str(self.visualList)]
                 #DigitalAssist.SaveText(self.transcript2, self.FileName + 'v2' + self.subject, self.subject)
                 DigitalAssist.getsubject(self)

                 continue
             elif  self.query1 == 'none' :
                 continue

             else:
                 self.AI_Corrected_Text = DigitalAssist.cleanText(self, self.query1)
                 DigitalAssist.adddata2DF(self)
                 self.visualList += str(self.entry) + '). ' + self.query1 + ' \n'
                 self.visualList2 += str(self.entry) + '). ' + self.AI_Corrected_Text + ' \n'
                 self.entry += 1






####################################################################################################################################################################################
## Below are more utlities that help me to do repetitive tasks quicker
####################################################################################################################################################################################
    def createDF(self):
        test = [('', '', '', '', '', '', '', '', '', '', '', '','','','')]
        self.transcript = pd.DataFrame(test, columns=['Entry #', 'Date', 'Subject', 'Type/Marker', 'Priority', 'Original_Text','Edited_Text','AI_Corrected_Text','AI_Correction_Comment', 'Note' ,'AI_Content','AI_Correction_Content','File_Name','Added to text File','Completed/Added to blog'])


    def adddata2DF(self):
        current_time1 = datetime.datetime.now()
        xx = current_time1.strftime('%m/%d/%Y %H:%M:%S')
        row = (self.entry - 1)
        self.transcript.loc[row]= [str(self.entry), xx, self.subject, self.Type, self.Significance, self.query1, self.new_Query,self.AI_Corrected_Text, self.Correction_Comment, self.Notes, self.Words,self.AI_Corrected_Content, self.FileName, self.Added2TextFile, self.Completed]

    def SaveText(df,FileName,tabname):
        current_time1 = datetime.datetime.now()
        current_time2 = current_time1.strftime('%m-%d-%Y_%H.%M.%S')
        text1 = df
        f2 = up.getPath()
        SavePath1 = f2
        Filename = '\\' + FileName + "_"
        SavePath2 = SavePath1 + Filename + current_time2 + ".xlsx"
        #SavePath2 = r'D:\ShakeBot Testing\ShaKeBotTest for DA - 12-20-2022.xlsx'
        try:
            with pd.ExcelWriter(SavePath2) as writer:
                text1.to_excel(writer, sheet_name=str(tabname))
                DigitalAssist.speak('File Saved')
                print('File Saved')
        except: Exception
        #DigitalAssist.speak('Error Saving')
        #print ('Error Saving')


    def cleanText(self,text):
        #result = ''
        parser = GingerIt()
        try:
            if int(len(text)) <= 4999:

                    #print(len(text))
                    result = pd.DataFrame(parser.parse(text))

                #result.drop_duplicates()
            else:
                result = pd.DataFrame('','','')
                self.Notes = 'Text is Large so may need to review the parts that may have been cut by mistake'
                for i in range(0, (len(text)//5000) +1):
                    texttemp = text[0+(5000*i):4999+(5000*i)]
                    if len(texttemp) >0:
                        result11 = pd.DataFrame(parser.parse(text))
                        result1 = [result1  + texttemp+ '\\n'+ '\\n']
                result = pd.DataFrame(result1, columns = 'result')

            try:
                if str(result.loc[0,'result']) != '':
                #print(result.loc[:,'result'].values)
                    rr = str(result.loc[0,'result'])
                    self.Correction_Comment += 'Corrections Made.'
            except:
                rr = str(result.loc[0, 'text'])
                #print(result.loc[:, 'text'].values)
                self.Correction_Comment += 'No Corrections Needed.'
            #rr = result.loc[:,:]
                #if rr == '':
                    #print(result.values)
        except:
            print('Text was not properly cleaned: ' + text)
            #print(text)
            rr = text
            self.Correction_Comment += 'No Corrections Needed (Caused an error and did not return anything).'
        finally:
                return rr





    def CompileBlog(self):
        f2 = up.getMFPath()
        df2 = pd.read_excel(f2)
        df2 =  df2.loc[df2['Type/Marker']!='To Do list']
        df2.sort_values(by=['Date','Subject'], ascending=False)
       #need to add more stuff here for if it has notes or other tags etc
        for i  in range(0,len(df2)):
            if df2['Priority'].loc[i] == '**AI content attached':
                df2['Combined'] = df2['AI_Corrected_Text'] + '\\n' +'\\n' + 'Note: AI Generated content: ' +df2['AI_Content']
            else:
                df2['Combined']= df2['AI_Corrected_Text'] + '\\n' +'\\n'

        df2['Combined_Edit']  = DigitalAssist.cleanText(self, str(df2['Combined']))

        text1 = df2['Combined'].str.cat(sep='\\n \n')
        text2 = df2['Combined_Edit'].str.cat(sep='\n \n')
        text3= DigitalAssist.cleanText(self, str(df2['text1']))
        text4 = DigitalAssist.cleanText(self, str(df2['text2']))

        topics = df2['Subject'].str.cat(sep='\n \n')
        notes = df2['Note'].str.cat(sep='\n \n')
        StartDate = df2['Date'].aggregate(['min']).strftime('%m-%d-%Y')
        End1Date = df2['Date'].aggregate(['min']).strftime('%m-%d-%Y')


        current_time2 = datetime.datetime.now().strftime('%m-%d-%Y')
        document = Document()
        document.add_heading('MondeVert Blog --> ' + current_time2 + '\n' + 'Topics: ' + topics + 'Notes: ' + notes + ' ' + StartDate + 'to' + End1Date, level=1)
        document.add_paragraph(text1)
        document.save('MondeVert Blog (Raw) --> ' + current_time2 + '.docx')

        document = Document(text2)
        document.add_heading('MondeVert Blog --> ' + current_time2 + '\n' + 'Topics: ' + topics + 'Notes: ' + notes + ' ' + StartDate + 'to' + End1Date, level=1)
        document.add_paragraph()
        document.save('MondeVert Blog (Polished) --> ' + current_time2 + '.docx')



        document = Document(text3)
        document.add_heading('MondeVert Blog --> ' + current_time2 + '\n' + 'Topics: ' + topics + 'Notes: ' + notes + ' ' + StartDate + 'to' + End1Date, level=1)
        document.add_paragraph()
        document.save('MondeVert Blog (Polished2) --> ' + current_time2 + '.docx')




        document = Document(text4)
        document.add_heading('MondeVert Blog --> ' + current_time2 + '\n' + 'Topics: ' + topics + 'Notes: ' + notes + ' ' + StartDate + 'to' + End1Date, level=1)
        document.add_paragraph()
        document.save('MondeVert Blog (Polished3) --> ' + current_time2 + '.docx')



    def compileFiles(self):
        f1 = up.getPath()
        files = glob.glob(f1 + '*.xlsx')
        tempDF = ''
        tempDF2 = ''
        df1 = pd.DataFrame()
        for fp in files:
            df1 = pd.DataFrame(pd.read_excel(fp))
            df1.drop_duplicates()
            r = len(df1.index)
            #print(r)

            for x in df1.columns:
                if x =='Original_Text':
                    xs = 'Original_Text'
                elif x =='Original_Wording':
                    xs = 'Original_Wording'
                else:
                    xx2 = ''


            for i in range(r):
                x = str(df1.loc[i, xs])

                try:
                    tempDF =  DigitalAssist.cleanText(self,x)

                except:
                    xx2 = ''
                    #print('Error - Review somewhere')
                    #print(len(x))
                    #print(x)
                    #print('End of Error Details')
                else:
                    df1.loc[i, 'AI_Corrected_Text'] = tempDF
                    df1.loc[i, 'AI_Correction_Comment'] = self.Correction_Comment
            DigitalAssist.add2Master(df1)
            del df1





    def add2Master(df1):
        f2 = up.getMFPath()
        df2 = pd.read_excel(f2)
        df3 = pd.concat([df1, df2])
        df3.drop_duplicates()
        df3.iloc[:,1:]
        # creating a new excel file and save the data
        df3.to_excel(f2, index=False)




if __name__ == '__main__':
    # main method for executing
    # the functions
    Record = ''
    x = DigitalAssist(1)
    x.Take_query()
