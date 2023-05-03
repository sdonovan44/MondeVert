
import glob
import datetime
import pandas as pd
import wikipedia
global Record
import sys
import ShakesBot as s
import Common_Sayings as cs
from threading import Event
import User_Prefs as up
from gingerit.gingerit import GingerIt
import threading
import numpy
from docx import Document
import re
#from exceptions import PendingDeprecationWarning
import requests
import cloudscraper
import tensorflow
import selenium

Record = ''
import webbrowser
from fpdf import FPDF

import User_Prefs as up
import os
import platform
import subprocess

import requests
import random
import string
import openai
import pyttsx3
# Import the speech recognition library
import speech_recognition as sr
from dotenv import load_dotenv
from DoNotCommit import API_Key


import atexit

from docx import Document
from docx.shared import Inches

import shutil





URL = "..."  # noqa
API_KEY = "..."






        #make a transcript class that takes down all of my requests and saves to a text file
        #have specific command for the assistant to save in a certain folder (always add date & time to the file
        #create things like to do lists
        #ideas
        #stories and dialogue
        #the key is being able to store it properly and sytemically access as well
        #every program I create should be executeable through the assistant



        #Urgent to do
        #Make it a class
        #have transcribe and also proofread (later) add method where in the middle of dictating you can add titles so it gets properly makred eventually
        # bullets etc. parenthesis
        #Have it so every pause puts a new line in between so its easy to pull out pauses etc.
        #Most important is to have a mode where it callse poembot and reads the lines to me.
        # .... also I would like for it to be a mini drama where I get multiple voices and characters that read
        # the output etc eventually learn how to animate/make a nice script.
        # create a series so eventually we might have a way to feed in characters to poem bot so its not just random like it is now lol

        #12/18/2022 get the assistant able to open the respective web browsers
        #set up mouse so it has shortcuts
        #be able to run on both mac and also the pc
        #import the site view bot
        #update the youtube bot so it hides ips properly and also so it loops through videos to not draw attention
        #    (make it so it starts off slow then has pockets of 100 here and there big days occasionally and find out how to make it work properly
        #have assistant able to run those on command (buy wifi adapter so I can have it running whenever & get new monitor for that too




        #Update transcribe commands (Review file so far,#Have a separate function to make the title (different than subject)
        ##Also I should be able to mark the file as important based on results (add tag)
        #Make a program that goes through and corrects the script into a new tab (puts subject as start of the word blocks (with other data in parenthesis for tagging)
        #Also make it so that the poems do not overwrite but get concatenated together
        #Make a theatre method that uses multithreading to do stuff
        #overall threads should allow this to run faster
        #Make it so I can update all meta data and also add more metadata
        #Goal, mark as header (prior or next maybe)


#
# class GingerIt(object):
#     def __init__(self):
#         self.url = URL
#         self.api_key = API_KEY
#         self.api_version = "2.0"
#         self.lang = "US"
#
#     def parse(self, text, verify=True):
#         session = cloudscraper.create_scraper()
#         request = session.get(
#             self.url,
#             params={
#                 "lang": self.lang,
#                 "apiKey": self.api_key,
#                 "clientVersion": self.api_version,
#                 "text": text,
#             },
#             verify=verify,
#         )
#         data = request.json()
#         return self._process_data(text, data)


class DigitalAssist():
    def __init__(self, voice = 3, language_settings=1):
        self.voice = voice
        self.language_settings = language_settings
        self.AssistantName = up.getAssistantName()
        self.UserName = up.getUserName()
        self.transcript_Final = ''
        self.API_Key = API_Key
        self.Correction_Comment = ''
        self.AI_Corrected_Content = ''
        self.SilentMode = False
        global xVoice
        global voice_set
        voice_set = self.voice
        xVoice = 1

    def Open_Web(url):
        url = url
        webbrowser.open_new_tab(url)
    def speak(self,audio, Add2T = True, voice = ''):
        if voice == '':
            voice = self.voice

        if Add2T == True:
            DigitalAssist.Add2Transcript(self, self.AssistantName + ': ' + audio)
        engine = pyttsx3.init()
        # getter method(gets the current value
        # of engine property)
        voices = engine.getProperty('voices')
        engine.setProperty('voice', voices[voice].id)
        # Method for the speaking of the assistant
        if self.SilentMode == False:
            engine.say(audio + "")
            # Blocks while processing all the currently
            # queued commands
            engine.runAndWait()
        else:
            print('Silent Mode Activated: ' + audio)

    def speakSweet(self,audio, Add2T = True):
        if Add2T == True:
            DigitalAssist.Add2Transcript(self, self.AssistantName + ': ' + audio)
        engine = pyttsx3.init()
        # getter method(gets the current value
        # of engine property)
        voices = engine.getProperty('voices')
        engine.setProperty('voice', voices[4].id)
        # Method for the speaking of the assistant
        if self.SilentMode == False:
            engine.say(audio + "")
            # Blocks while processing all the currently
            # queued commands
            engine.runAndWait()
        else:
            print('Silent Mode Activated: ' + audio)

        #def RandomCharacters

    def testvoices(self):
        #import pyttsx3
        engine = pyttsx3.init()
        voices = engine.getProperty('voices')
        index = 0
        for voice in voices:

            print(f'index-> {index} -- {voice.name}')
            engine.setProperty('voice', voices[index].id)
            engine.say("Shane do you like this voice more than the last one you busy little bee turkish" + "")

            index += 1
        engine.runAndWait()

    def Setvoices(self, Quick = False):



        if Quick ==False:
            engine = pyttsx3.init()
            voices = engine.getProperty('voices')
            index = 0
            indexCount = 0
            #load_dotenv()
            # Set up the OpenAI API client :
            openai.api_key = API_Key

            prompt = 'Write a short paragraph (less than 44 words) that uses different pronunciations and complex words while being warm and welcoming to a person named Shane'
            completions = openai.Completion.create(
                engine="text-davinci-003",
                prompt=prompt,
                max_tokens=2500,
                n=1,
                stop=None,
                temperature=0.5,
            )
            speak1 = 'Do you want to set this voice as active?'

            # Print the generated text
            message = completions.choices[0].text
            print(message)
            DigitalAssist.Add2Transcript(self,text2Add=message)
            set1 = False

            for voice in voices:
                indexCount +=1

            while set1 == False:
                if index <=indexCount:
                    print(f'index-> {index} -- {voices[index].name}')
                    engine.setProperty('voice', voices[index].id)
                    engine.say(message)
                    DigitalAssist.speak(self,speak1)
                    Query = DigitalAssist.getUserResponse(self)


                    if "yes" in Query or "set" in Query or ("make" in Query and "active" in Query) or  ("set" in Query and "active" in Query):
                        self.voice = index
                        voice_set = self.voice
                        xVoice = index
                        DigitalAssist.speak(self, 'Voice Set')
                        set1 = True



                    index += 1
                    engine.runAndWait()
                else:
                    DigitalAssist.Setvoices(self,Quick = True)

        else:
            self.voice = int(input('What Voice do you want to make active?'))
            DigitalAssist.speak(self, 'Voice Set')

        ####################################################################################################################################################################################
        ## Below are more utlities that help me to do repetitive tasks quicker
        ####################################################################################################################################################################################
    def createDF(self):
        test = [('', '', '', '', '', '', '', '', '', '', '', '', '', '', '')]
        self.transcript = pd.DataFrame(test, columns=['Entry #', 'Date', 'Subject', 'Type/Marker', 'Priority',
                                                      'Original_Text', 'Edited_Text', 'AI_Corrected_Text',
                                                      'AI_Correction_Comment', 'Note', 'AI_Content',
                                                      'AI_Correction_Content', 'File_Name', 'Added to text File',
                                                      'Completed/Added to blog'])

    def adddata2DF(self):
        current_time1 = datetime.datetime.now()
        xx = current_time1.strftime('%m/%d/%Y %H:%M:%S')
        row = (self.entry - 1)
        self.transcript.loc[row] = [str(self.entry), xx, self.subject, self.Type, self.Significance, self.query1,
                                    self.new_Query, self.AI_Corrected_Text, self.Correction_Comment, self.Notes,
                                    self.Words, self.AI_Corrected_Content, self.FileName, self.Added2TextFile,
                                    self.Completed]


    def ActivateSilentMode(self):
        self.SilentMode = True
    def ActivateLoudMode(self):
        self.SilentMode = False

    def saveTranscript(self):

        current_time1 = datetime.datetime.now()
        current_time2 = current_time1.strftime('%m-%d-%Y_%H.%M.%S')
        Filename = 'MondeVert Assistant'
        Filename = '\\' + Filename + "_"
        f2 = up.getPath()
        SavePath1 = f2
        data = [current_time2,self.transcript_Final]
        df1 = pd.DataFrame(data,columns = ["TimeStamp","Transcript"])
        df1.to_csv(SavePath1 + Filename + current_time2 + '.csv')
        # DigitalAssist.SaveText(self,df1,'MondeVert Assistant', 'Full Transcript')
        DigitalAssist.add2Master2(df1)


    def Add2Transcript(self, text2Add):
        self.transcript_Final += text2Add + ' \n'

    def SaveText(self, df, FileName, tabname):
        current_time1 = datetime.datetime.now()
        current_time2 = current_time1.strftime('%m-%d-%Y_%H.%M.%S')
        text1 = df
        f2 = up.getPath()
        SavePath1 = f2
        invalidCharRemoved = re.sub(r"[^a-zA-Z0-9 ]", "", FileName)
        Filename = '\\' + str(invalidCharRemoved) + "_"
        SavePath2 = SavePath1 + Filename + current_time2 + ".xlsx"
        # SavePath2 = r'D:\ShakeBot Testing\ShaKeBotTest for DA - 12-20-2022.xlsx'
        try:
            with pd.ExcelWriter(SavePath2) as writer:
                text1.to_excel(writer, sheet_name=str(tabname))
                DigitalAssist.speak(self, 'File Saved')
                print('File Saved')
        except:
            Exception
        # DigitalAssist.speak(self,'Error Saving')
        # print ('Error Saving')
    def cleanText(self, text):
        # result = ''
        parser = GingerIt()
        try:
            if int(len(text)) <= 4999:
                # print(len(text))
                result = pd.DataFrame(parser.parse(text))

                # result.drop_duplicates()
            else:
                result = pd.DataFrame('', '', '')
                self.Notes = 'Text is Large so may need to review the parts that may have been cut by mistake'
                for i in range(0, (len(text) // 5000) + 1):
                    texttemp = text[0 + (5000 * i):4999 + (5000 * i)]
                    if len(texttemp) > 0:
                        result11 = pd.DataFrame(parser.parse(text))
                        result1 = [result1 + texttemp + '\\n' + '\\n']
                result = pd.DataFrame(result1, columns='result')

            if len(result.loc[0, 'result']) > 0:
                # print(result.loc[:,'result'].values)
                rr = str(result.loc[0, 'result'])
                self.Correction_Comment += 'Corrections Made.'
            else:
                rr = text
                self.Correction_Comment += 'No Corrections Needed.'


        except:
            print('Text was not properly cleaned: ' + text)
            # print(text)
            rr = text
            self.Correction_Comment += 'No Corrections Needed (Caused an error and did not return anything).'
            return rr

        return rr
    def makeQuickPoem(self):
        x = s.PoemBot(1, 1, 1, 1, 40, 3)
        x.ReloadModel("model.h5")
        x.setupdata()
        self.Words = x.shakesbot_DA_Make_Script(size=300)
        # self.Words[0:100]
        print(self.Words)
        self.Words = DigitalAssist.cleanText(self, self.Words)
        DigitalAssist.speakSweet(self, self.Words)
    def StartThread(self):
        trd1 = threading.Thread(target=DigitalAssist.Make_Script2(self))
        self.w[i] = trd1.start()
        x = DigitalAssist.Make_Script2(self)
        print(x)

        self.threads.append(trd1)
    def Make_Script(self):
        size = 1
        self.threads = []
        self.w = numpy.empty(size, dtype=str)
        for self.i in range(0, size):
            DigitalAssist.StartThread(self)
            print(self.w[i])
            # globals()[f"trd{i}"] = threading.Thread(target=DigitalAssist.Make_Script2(self))
            # w[i] = globals()[f"trd{i}"].start()

        for x in self.threads:
            x.join()

        for ii in range(0, size):
            self.Words += str(self.w[ii])

        self.AI_Corrected_Content = DigitalAssist.cleanText(self, self.Words)
    def Make_Script2(self):
        pb = s.PoemBot(1, 1, 1, 1, 40, 3)
        pb.ReloadModel('model.h5')
        pb.setupdata()
        w1 = pb.shakesbot_DA_Make_Script()
        DigitalAssist.speakSweet(self, w1)
        # print (w1)
        return w1
    def generateGreek(self):
        w2 = self.pb.generateGreek() + ' \n' + ' \n'
        return w2
    def add2Master(df1):
        f2 = up.getMFPath()
        df2 = pd.read_excel(f2)
        df3 = pd.concat([df1, df2])
        df3.drop_duplicates()
        df3.iloc[:, 1:]
        # creating a new excel file and save the data
        df3.to_excel(f2, index=False)
    def add2Master2(df1):
        f2 = up.getMF2Path()
        df2 = pd.read_excel(f2)
        df3 = pd.concat([df1, df2])
        df3.iloc[:, 1:]
        df3.drop_duplicates()

        # creating a new excel file and save the data
        df3.to_excel(f2, index=False)




    def add2Master3(df1):
        f2 = up.MasterFile3
        df2 = pd.read_excel(f2)
        df3 = pd.concat([df1, df2])
        df3.iloc[:, 1:]
        df3.drop_duplicates()

        # creating a new excel file and save the data
        df3.to_excel(f2, index=False)
    def CompileBlog(self):
        f2 = up.getMFPath()
        df2 = pd.read_excel(f2)
        df2 = df2.loc[df2['Type/Marker'] != 'To Do list']
        df2.sort_values(by=['Date', 'Subject'], ascending=False)
        # need to add more stuff here for if it has notes or other tags etc
        for i in range(0, len(df2)):
            if df2['Priority'].loc[i] == '**AI content attached':
                df2['Combined'] = df2['AI_Corrected_Text'] + '\\n' + '\\n' + 'Note: AI Generated content: ' + df2[
                    'AI_Content']
            else:
                df2['Combined'] = df2['AI_Corrected_Text'] + '\\n' + '\\n'

        df2['Combined_Edit'] = DigitalAssist.cleanText(self, str(df2['Combined']))

        text1 = df2['Combined'].str.cat(sep='\\n \n')
        text2 = df2['Combined_Edit'].str.cat(sep='\n \n')
        text3 = DigitalAssist.cleanText(self, str(df2['text1']))
        text4 = DigitalAssist.cleanText(self, str(df2['text2']))

        topics = df2['Subject'].str.cat(sep='\n \n')
        notes = df2['Note'].str.cat(sep='\n \n')
        StartDate = df2['Date'].aggregate(['min']).strftime('%m-%d-%Y')
        End1Date = df2['Date'].aggregate(['min']).strftime('%m-%d-%Y')

        current_time2 = datetime.datetime.now().strftime('%m-%d-%Y')
        document = Document()
        document.add_heading(
            'MondeVert Blog --> ' + current_time2 + '\n' + 'Topics: ' + topics + 'Notes: ' + notes + ' ' + StartDate + 'to' + End1Date,
            level=1)
        document.add_paragraph(text1)
        document.save('MondeVert Blog (Raw) --> ' + current_time2 + '.docx')

        document = Document(text2)
        document.add_heading(
            'MondeVert Blog --> ' + current_time2 + '\n' + 'Topics: ' + topics + 'Notes: ' + notes + ' ' + StartDate + 'to' + End1Date,
            level=1)
        document.add_paragraph()
        document.save('MondeVert Blog (Polished) --> ' + current_time2 + '.docx')

        document = Document(text3)
        document.add_heading(
            'MondeVert Blog --> ' + current_time2 + '\n' + 'Topics: ' + topics + 'Notes: ' + notes + ' ' + StartDate + 'to' + End1Date,
            level=1)
        document.add_paragraph()
        document.save('MondeVert Blog (Polished2) --> ' + current_time2 + '.docx')

        document = Document(text4)
        document.add_heading(
            'MondeVert Blog --> ' + current_time2 + '\n' + 'Topics: ' + topics + 'Notes: ' + notes + ' ' + StartDate + 'to' + End1Date,
            level=1)
        document.add_paragraph()
        document.save('MondeVert Blog (Polished3) --> ' + current_time2 + '.docx')
    def compileFiles(self):
        f1 = up.getPath()
        files = glob.glob(f1 + '*.xlsx')
        tempDF = ''
        tempDF2 = ''
        df1 = pd.DataFrame()
        for fp in files:
            df1 = pd.DataFrame(pd.read_excel(fp))
            df1.drop_duplicates()
            r = len(df1.index)
            # print(r)

            for x in df1.columns:
                if x == 'Original_Text':
                    xs = 'Original_Text'
                elif x == 'Original_Wording':
                    xs = 'Original_Wording'
                else:
                    xx2 = ''

            for i in range(r):
                x = str(df1.loc[i, xs])

                try:
                    tempDF = DigitalAssist.cleanText(self, x)

                except:
                    xx2 = ''
                    # print('Error - Review somewhere')
                    # print(len(x))
                    # print(x)
                    # print('End of Error Details')
                else:
                    df1.loc[i, 'AI_Corrected_Text'] = tempDF
                    df1.loc[i, 'AI_Correction_Comment'] = self.Correction_Comment
            DigitalAssist.add2Master(df1)
            del df1

    ############################################################################################################################################################################################################################
    ############################################################################################################################################################################################################################
    ################code below is really not used but its how I started
    ############################################################################################################################################################################################################################
    ############################################################################################################################################################################################################################
    def tellDay(self):
        # This function is for telling the
        # day of the week
        day = datetime.datetime.today().weekday() + 1

        # this line tells us about the number
        # that will help us in telling the day
        Day_dict = {1: 'Monday', 2: 'Tuesday',
                    3: 'Wednesday', 4: 'Thursday',
                    5: 'Friday', 6: 'Saturday',
                    7: 'Sunday'}

        if day in Day_dict.keys():
            day_of_the_week = Day_dict[day]
            print(day_of_the_week)
            DigitalAssist.speak(self,"The day is " + day_of_the_week)
    def tellTime(self):
        # This method will give the time
        time = str(datetime.datetime.now())
        print(time)
        hour = time[11:13]
        min = time[14:16]
        t1 = cs.DA_Time1()
        t2 = cs.DA_Time2()
        t3 = cs.DA_Time3()
        DigitalAssist.speak(self, t1 + hour + t2+ min + t3)
    def Hello(self):
        # This function is for when the assistant
        # is called it will say hello and then
        # take query
        r = cs.DAgreet()
        print (r)
        DigitalAssist.speak(self,r)
    ############################################################################################################################################################################################################################
    ############################################################################################################################################################################################################################
    ################Above is code I did not write, below is code that allows user to update values of respective values
    ############################################################################################################################################################################################################################
    ############################################################################################################################################################################################################################
    def getfilename(self):
        self.FileName = DigitalAssist.getdata(self, 'FileName')
    def  getsubject(self):
        self.subject = DigitalAssist.getdata(self, 'subject')
    def  getDictation_Type(self):
        self.Dictation_Type  =  DigitalAssist.getdata(self, 'Type')
    def  getSignificance(self):
        self.Significance =  DigitalAssist.getdata(self, 'Significance')
    def  getNote(self):
        self.Notes += '|' + DigitalAssist.getdata(self, 'Notes')

    def getdata(self, Field):

        DigitalAssist.speak(self, 'what is the ' + Field)
        print('Please confirm the ' + Field)
        s = True
        Query = ''
        while (s == True):
            Query = DigitalAssist.getUserResponse(self)
            Query2 = ''
            WaitforResponse = False
            while (WaitforResponse == False):
                DigitalAssist.speak(self, Query + ' Is that correct?')
                Query2 = DigitalAssist.getUserResponse(self)
                if "yes" in Query2 or "correct" in Query2 or "ya" in Query2 or "yeah" in Query2:
                    DigitalAssist.speak(self, 'Thanks for confirming the ' + Field)
                    s = False
                    WaitforResponse = True
                    continue
                if "no" in Query2 or "wrong" in Query2 or "not it" in Query2 or "nope" in Query2:
                    DigitalAssist.speak(self, 'ok  please say the ' + Field + ' you want to set')
                    print('ok  please say the ' + Field + ' you want to set')
                    WaitforResponse = True
                    continue
                else:
                    DigitalAssist.speak(self, 'Lets try that again what is the ' + Field)
                    Event().wait(1)
                    continue
        return Query

############################################################################################################################################################################################################################
############################################################################################################################################################################################################################
################ below is code that allow me to confirm with the user the inputs are correct
############################################################################################################################################################################################################################
############################################################################################################################################################################################################################
    def takeCommand(self):
        Query = DigitalAssist.getUserResponse(self,Response = "Action Requested")
        return Query


    def getUserResponse(self, pause =.5, Response = 'You Said' ):
        r = sr.Recognizer()
        with sr.Microphone() as source:
            print('Listening...')
            r.pause_threshold = pause
            audio = r.listen(source)
            try:
                Query = r.recognize_google(audio, language='en-in')
                DigitalAssist.Add2Transcript(self,self.UserName + ': '+Query)
                print(Response+ ": ", Query)

            except Exception as e:
                print(e)
                print("Please repeat Words not understood...")
                Query = DigitalAssist.getUserResponse(self)
                #Event().wait(2)
            return Query





    def ConfirmUR(self,prompt):
        DigitalAssist.speak(self,'You Said ' + prompt + ' is this correct?')
        query = DigitalAssist.getUserResponse(self)

        if 'yes' in query or 'correct' in query or 'yeah' in query or 'submit' in query:
            self.promptB =True

    def editBotPrompt(self, message):
        DigitalAssist.speak(self, 'Edit Mode Activated', voice = 4)
        DigitalAssist.speak(self,'What do you want to use for a prompt?',voice = 4)
        query = DigitalAssist.getUserResponse(self)

        if 'cancel'in query:
            r = message
            #Do nothing
        else:
            DigitalAssist.ConfirmBOT(self, query)
            if self.promptB ==True:
                self.message = query


    def ConfirmBOT(self,message):
        DigitalAssist.speak(self,'Chat GPT responded with: ' + message + ' do you want to use this prompt?',voice = 4)
        query = DigitalAssist.getUserResponse(self)

        if 'yes' in query or 'correct' in query or 'yeah' in query or 'submit' in query:
            self.promptB =True

        if 'edit' in query or 'change' in query and ('prompt' in query or 'words' in query):
            DigitalAssist.editBotPrompt(self, message)
    def transcribe_Build_Query(self, pause = 3.13):
        Query = DigitalAssist.getUserResponse(self)
        return Query
    def transcribe_Build_Query_Pause(self):
        Speak1 = 'What Tool do you want to use?'
        Print1 = 'Options: ChatGPT (ChatBot), Transcribe Mode , Make To Do List, Change Subject, Edit, Poem, Stop Recording --> Build others here'
        DigitalAssist.speak(self,Speak1)
        print(Print1)
        Query = DigitalAssist.getUserResponse(self,Response="Tool Selected")
        return Query

    ############################################################################################################################################################################################################################
    ############################################################################################################################################################################################################################
    ############################################################################################################################################################################################################################
    ##########################                below is ChatGPT                  ##############################################################################################################
    ############################################################################################################################################################################################################################
    ############################################################################################################################################################################################################################

    def makeArt(self, Prompt = ''):

        DigitalAssist.Add2Transcript(self, ' \n')
        Prior_Mode = self.Mode
        self.Mode = 'AI Art Mode'
        DigitalAssist.Add2Transcript(self,text2Add= ( Prior_Mode + ' - ' + self.Mode + ':'))
        DigitalAssist.Add2Transcript(self, ' \n')
        self.promptB = False

        Promptc = 'True'

        if Prompt =='':
           Promptc = 'True'
        else:
            Promptc = 'False'



        if Prompt == '':
            print('Prompt is NULL: ' + Promptc)


            while self.promptB ==False:
                AIspeak = 'What do you want me to draw?'

                DigitalAssist.speak(self,AIspeak)

                print(AIspeak)
                prompt  = DigitalAssist.getUserResponse(self,Response = "Art Prompted", pause= 1)

                if 'cancel' in prompt or 'stop' in prompt or 'quit' in prompt or 'done' in prompt or 'exit' in prompt:
                    DigitalAssist.saveTranscript(self)
                    self.promptB == True
                    continue

                DigitalAssist.ConfirmUR(self,prompt)
                print()
        else:
            prompt = Prompt
            print(prompt)
            DigitalAssist.Add2Transcript(self,'(Prompt Fed Directly into Function)')

        # Set up the OpenAI API client :
        openai.api_key = API_Key
        print("Sending to OpenAI...")
        print()

        response = openai.Image.create(
            prompt=prompt,
            n=1,
            size="1024x1024"
        )

        image_url = response['data'][0]['url']

        print(f"Image URL: {image_url}")
        print()

        # URL of the image to be downloaded is defined as image_url
        r = requests.get(image_url)  # create HTTP response object

        # send a HTTP request to the server and save
        # the HTTP response in a response object called r
        FileName = prompt
        print('Length of File Name: ' + str(len(FileName)))
        if len(FileName) >= 150:
            FileName=FileName[0:150]
            print('Length of File Name: ' + str(len(FileName)))

        fname = f"{''.join([c for c in FileName.strip().replace(' ', '_') if c.isalnum() or c == '_'])}.png"

        if not os.path.exists("images"):
            os.mkdir("images")

        if os.path.isfile(os.path.join(up.AI_Art_Path,'\'', fname)):
            fname = fname.split(".")[0] + f".{''.join(random.choice(string.ascii_letters) for x in range(5))}.png"

        fname = os.path.join(up.AI_Art_Path, fname)
        print(f"Filename: {fname}")
        with open(fname, 'wb') as f:

            # Saving received content as a png file in
            # binary format

            # write the contents of the response (r.content)
            # to a new file in binary mode.
            f.write(r.content)

        if platform.system() == 'Darwin':  # macOS
            subprocess.call(('open', fname))
        elif platform.system() == 'Windows':  # Windows
            os.startfile(fname)
        else:  # linux variants
            subprocess.call(('xdg-open', fname))

        return fname



    # Define a function that sends a message to ChatGPT
    def chat_query(self,prompt):
        model_engine = "text-davinci-003"
        completions = openai.Completion.create(
            engine=model_engine,
            prompt=prompt,
            max_tokens=2048,
            n=1,
            temperature=0.5,
        )

        message = completions.choices[0].text
        return message

    # Define a function that handles the conversation
    def conversation_handler(self,prompt):
        # Send the prompt to ChatGPT
        load_dotenv()
        # Set up the OpenAI API client :
        openai.api_key = API_Key
        response = DigitalAssist.chat_query(prompt)
        print(f"ChatGPT: {response}")

        # End the conversation if ChatGPT says goodbye
        if "goodbye" in response.lower():
            return

        # Otherwise, get user input and continue the conversation
        prompt = input("You: ")
        DigitalAssist.conversation_handler(prompt)

    def ChatGPTDA(self, temp = 0.5,MakeArt = False, Prompt = '', UseArtPrompt = False, ConfirmBot = True):
        DigitalAssist.Add2Transcript(self, ' \n')
        Prior_Mode = self.Mode
        self.Mode = 'Chat GPT Mode'
        self.message = ''
        DigitalAssist.Add2Transcript(self,text2Add= ( Prior_Mode + ' - ' + self.Mode + ':'))
        DigitalAssist.Add2Transcript(self, ' \n')
        prompt = ''
        self.promptB = False
        # Import the OpenAI library
        # Set up the OpenAI API client :
        openai.api_key = API_Key
        # Record the audio
        DigitalAssist.speak(self, " Chat GPT is running!",voice = 4)


        if UseArtPrompt ==True:
            MakeArtPrompt = "Provide me with a prompt to share with artificial inteligence that will create a unique and visually pleasing work of art. "
        else:
            MakeArtPrompt = ''


        if Prompt == '':
            sStop1 = False
        else:
            sStop1 = True


        sStop = False
        # check if the reply contains 'yes'
        while sStop == False:
            # Record the audio


            DigitalAssist.Add2Transcript(self, MakeArtPrompt)

            # engine.runAndWait()
            if sStop1 == True:
                prompt = MakeArtPrompt + Prompt
                print(prompt)
            else:
                prompt =  MakeArtPrompt + DigitalAssist.getUserResponse(self, pause=1.44)
                if ('STOP' in prompt.upper() and (
                        'LISTEN' in prompt.upper())) or 'QUIT' == prompt.upper() or ' QUIT ' == prompt.upper() or (
                        'END' in prompt.upper() and ('CONVERSA' in prompt.upper())):
                    sStop = True
                    # print ('Its supposed to quit here')
                    continue


            # Generate text
            completions = openai.Completion.create(
                engine="text-davinci-003",
                prompt=prompt,
                max_tokens=2444,
                n=1,
                stop=None,
                temperature=temp,
            )
            self.message = completions.choices[0].text
            print(self.message)



            if  MakeArt == False and ('MAKE' in prompt.upper() or 'ART' in prompt.upper()) and 'PROMPT' in prompt.upper() and sStop1 == False:
                DigitalAssist.ConfirmBOT(self, self.message)
                if self.promptB == True:
                    DigitalAssist.makeArt(self,self.message)
                    continue
                else:
                    continue


            if MakeArt == True or sStop1 == True:
                if ConfirmBot == True:
                    DigitalAssist.ConfirmBOT(self, self.message)
                else:
                    DigitalAssist.speakSweet(self, self.message)
                    self.promptB = True

                if self.promptB == True:
                    sStop = True
                    continue
                else:

                    continue
            else:
                DigitalAssist.speakSweet(self, self.message)

        return self.message

    ############################################################################################################################################################################################################################
    ############################################################################################################################################################################################################################
    ############################################################################################################################################################################################################################
    ################Above  is ChatGPT (I did not write original code but adapted for my use case)
    ############################################################################################################################################################################################################################
    ############################################################################################################################################################################################################################
    ############################################################################################################################################################################################################################



    def editentry(self, Query1):
        DigitalAssist.Add2Transcript(self, ' \n')
        Prior_Mode = self.Mode
        self.Mode = 'Edit Mode'
        DigitalAssist.Add2Transcript(self,text2Add= ( Prior_Mode + ' - ' + self.Mode + ':'))
        s = True
        while (s == True):
            DigitalAssist.speak(self,'I am ready to make the edit please proceed')
            print('ready to make the edit please proceed')
            print('Original Entry: ' + Query1)
            Query2 = DigitalAssist.getUserResponse()

            ss = True
            while (ss == True):
                print(Query2)
                DigitalAssist.speak(self,Query2 + 'Are you satisfied with your edits?')
                print('Options: Yes/Correct, Not yet/Redo, Add as Note, Combine, Cancel')

                Query3 = DigitalAssist.getUserResponse(self)


                if  "yes" in Query3 or "correct" in Query3 or "ya" in Query3 or "yeah" in Query3 or "looks good" in Query3 or "keep new" in Query3 or "new one" in Query3:
                    DigitalAssist.speak(self,'Thanks for confirming, I will make the updates you provided')
                    print('Thanks for confirming, I will make the updates you provided')
                    Query_Final = Query2
                    s = False
                    ss = False


                elif  "not yet" in Query3 or 'redo'  in Query3 or 'try again'  in Query3 or 'again'  in Query3 or 'one more' in Query3:
                    DigitalAssist.speak(self,'ok, lets go again')
                    print('ok, lets go again')
                    ss = False
                    continue

                elif "add as note" in Query3 or "include note" in Query3 or "as note" in Query3  :
                    DigitalAssist.speak(self,'Thanks for confirming, I will add this as a note')
                    print('Thanks for confirming, I will add this as a note')
                    Query_Final = ''
                    self.Notes = Query2
                    #when this is a class I can update the df for now I will add it to the end same as when I say combine them
                    s = False
                    ss = False

                elif "add to original" in Query3 or "both" in Query3 or 'combine' in Query3 or 'combination' in Query3 or 'merge' in Query3 or 'together' in Query3:
                    DigitalAssist.speak(self,'Thanks for confirming, I will join the two thoughts')
                    print('Thanks for confirming, I will join the two thoughts')
                    Query_Final = Query1 + ' ' + Query2
                    # when this is a clas I can update the df for now I will add it to the end same as when I say combine them
                    s = False
                    ss = False

                elif 'cancel'   in Query3  or 'end' in Query3 or 'stop'  in Query3 or  'stop trying' in Query3:
                    DigitalAssist.speak(self,'Thanks for confirming, I will remove these edits')
                    print('Thanks for confirming, I will remove these edits')
                    Query_Final = ''
                    s = False
                    ss = False

                else:
                    Event().wait(1)


        return Query_Final

############################################################################################################################################################################################################################
############################################################################################################################################################################################################################
############################################################################################################################################################################################################################
################Long Code where I give Digital Assistant commands##    (user menus)
############################################################################################################################################################################################################################
############################################################################################################################################################################################################################





    def RunChatGPT(self):
        self.Mode = 'Chat GPT -  Menu'

        DigitalAssist.speak(self, self.Mode + ' What Chat GPT mode do you want to use?')

        print( self.Mode + ' What Chat GPT mode do you want to use?')

        s2 = True
        while (s2 == True):
            DigitalAssist.Add2Transcript(self,' \n')
            query = DigitalAssist.takeCommand(self).lower()


            #Say Quick Art or Auto Art or Basic Art
            if 'normal'  in query or 'conv'  in query or 'reg' in query or  'talk'  in query :
                prompt = DigitalAssist.ChatGPTDA(self)
                DigitalAssist.RunChatGPT(self)
                s2 = False

            elif "art" in query and ("basic" in query or "auto" in query or "quick" in query):
                DigitalAssist.quickArt(self)
                DigitalAssist.RunChatGPT(self)
                s2 = False

            # Say Prompted Art or super Art or advanced Art
            elif "art" in query and ( "advance" in query or  "super" in query or  "open" in query or ( "with" in query and   "prompt" in query)):
                prompt = DigitalAssist.ChatGPTDA(self, MakeArt=True, UseArtPrompt = True)

                ArtPath =  DigitalAssist.makeArt(self, prompt)
                DigitalAssist.RunChatGPT(self)
                s2 = False

            elif 'shake' not in query and ("inspire" in query or "unique" in query or ("own" in query) or ("version" in query and "3" in query) or ('quick' in query and 'poem' in query ) or  ( "poem" in query or 'poetry' in query)):
                #make program to ask how spicy to make it lol
                Poem_Type = random.choices(up.Random_Poem)
                Poem_Type = Poem_Type[0]
                DigitalAssist.Quick_Poem_v1(self, GPTPrompt=Poem_Type)
                s2 = False
                DigitalAssist.RunChatGPT(self)


            elif "shake" in query and ("poem" in query or "poet" in query  or ("version" in query and "1" in query)):
                DigitalAssist.Shake_Poem_v1(self)
                s2 = False
                DigitalAssist.RunChatGPT(self)

            elif ("poem" in query or 'poetry' in query or 'shake' in query) and ("inspire" in query or "unique" in query or ("make" in query or "own" in query) or ("version" in query and "2" in query)):
                DigitalAssist.Shake_Poem_v2(self)
                s2 = False
                DigitalAssist.RunChatGPT(self)


            elif ("sunday" in query and ('scary' in query or 'scaries' in query)) or ("sunday" in query and ( "poem" in query or "story" in query) ):
                DigitalAssist.SundayScary_Poem(self)
                s2 = False
                DigitalAssist.RunChatGPT(self)


            elif (("live" in query and 'art in query') or ('podcast' in query or 'stream' in query) ):
                DigitalAssist.MakeArtLive(self)
                s2 = False
                DigitalAssist.RunChatGPT(self)





            elif ('chorus' in query or 'song' in query) or ("brick" in query and ( "sing" in query or "dj" in query)) :


                Song_Genre = random.choices(up.Random_Song_Genre_List)
                Song_Genre = Song_Genre[0]
                Song_Genre = 'v3'
                #Song_Genre = 'techno'


                if 'v1' in  Song_Genre:
                    DigitalAssist.Make_a_Song(self)
                    DigitalAssist.speak(self, 'Make a quick song Version 1 Complete')
                elif 'v2' in Song_Genre:
                    DigitalAssist.Make_a_Chorus(self)
                    DigitalAssist.speak(self, 'Make a quick song Version 2 Complete')
                elif 'v3' in Song_Genre:
                    DigitalAssist.Make_a_Song(self)
                    DigitalAssist.speak(self, 'Make a quick song Version 3 Complete')

                elif 'Techno' in Song_Genre:
                    DigitalAssist.Make_a_Chorus(self, Mode='Techno')
                    DigitalAssist.speak(self, 'Techno Lyrics Complete')
                    DigitalAssist.Sampler(self, Mode='Techno')
                    DigitalAssist.speak(self, 'Techno Samples Provided Complete', voice=9)
                elif Song_Genre =='SadRap':
                    DigitalAssist.Make_a_Rap(self, Mode='SadRap')
                    DigitalAssist.speak(self, 'Sad Rap Lyrics Complete')
                elif Song_Genre == 'Reggae':
                    DigitalAssist.Make_a_Rap(self, Mode='Raggae')
                    DigitalAssist.speak(self, 'Reggae Lyrics Complete')
                elif Song_Genre == 'Rap':
                    DigitalAssist.Make_a_Rap(self, Mode='Rap')
                    DigitalAssist.speak(self, 'Rap Lyrics Complete')
                else:
                    DigitalAssist.Make_a_Rap(self, Mode='Raggae')
                    DigitalAssist.speak(self, 'Reggae Lyrics Complete')


                s2 = False
                DigitalAssist.RunChatGPT(self)



            elif 'basic' in query or 'shake' in query:
                DigitalAssist.makeQuickPoem(self)
                DigitalAssist.makeArt(self, self.Words)
                s2 = False
                DigitalAssist.RunChatGPT(self)


            elif 'dad' in query or 'timmy' in query or ' guy' in query   or 'tim ' in query or (('simon' in query or 'old man' in query or 'timmy d' in query ) and 'help' in query) or ('timmy' in query and 'd' in query):
                DigitalAssist.TimmyDMode(self)
                s2 = False
                DigitalAssist.RunChatGPT(self)

            elif 'blog' in query or 'website post' in query  or 'post' in query :
                Blog_Topic = random.choices(up.Random_Blog_Topic)

                Blog_Topic = Blog_Topic[0]


                #DigitalAssist.makeBlogPost(self, GPTprompt=Blog_Topic)
                #DigitalAssist.makeBlogPost(self, GPTprompt=up.RandomTopic)
                DigitalAssist.makeBlogPost(self, GPTprompt=up.CorrectText)




                s2 = False
                DigitalAssist.RunChatGPT(self)


#Make a program that makes a short background for a character, also have a duo/relationship of some sort and make an interesting scene for a play/screenplay

            elif 'wiki' in query or 'wikipedia' in query  or (('pen name' in query or 'page' in query or 'mode' in query ) and 'wiki' in query) :
                DigitalAssist.Wiki4PenNames(self)
                s2 = False
                DigitalAssist.RunChatGPT(self)


            elif 'cancel' in query or 'stop' in query or 'quit'in query or 'done' in query or  'exit' in query:
                DigitalAssist.saveTranscript(self)
                sys.exit()
                s2 = False

            else:
                print('Please try again, ')
                continue

    # def Shake_Art_v1(self, prompt)

    def TimmyDMode(self):
        current_time1 = datetime.datetime.now()
        current_time2 = current_time1.strftime('%m-%d-%Y_%H.%M.%S')
        f2 = up.getPath()
        SavePath1 = f2
        FileName = '5 men on a Bridge - Timmy D Convo & Cover Art '
        Filename = '\\' + FileName + "_"
        Title = SavePath1 + Filename + current_time2

        Script1 = DigitalAssist.ChatGPTDA(self, temp=0.5, MakeArt=False, Prompt=(up.Timmy_D_Dialogue1),ConfirmBot = False)
        Script2 = DigitalAssist.ChatGPTDA(self, temp=0.5, MakeArt=False, Prompt=(up.Timmy_D_Dialogue2), ConfirmBot = False)
        CoverArt1 = DigitalAssist.ChatGPTDA(self, temp=0.5, MakeArt=False, Prompt=(up.Timmy_D_Cover_Art_Prompt1), ConfirmBot = False)
        CoverArtPath1 = DigitalAssist.makeArt(self, CoverArt1)
        CoverArt2 = DigitalAssist.ChatGPTDA(self, temp=0.5, MakeArt=False, Prompt=(up.Timmy_D_Cover_Art_Prompt2), ConfirmBot = False)
        CoverArtPath2 = DigitalAssist.makeArt(self, CoverArt2)
        CoverArtPath3 = DigitalAssist.makeArt(self, up.Timmy_D_Cover_Art_Prompt_Direct_to_ArtBot)
        ArtPaths = [CoverArtPath1, CoverArtPath2,CoverArtPath3]

        document = Document()
        document.add_heading('5 men on a Bridge', 0)
        document.add_heading('Script 1', 4)
        p = document.add_paragraph()
        r = p.add_run()
        r.add_text(Script1)
        document.add_heading('Potential Cover Art', 4)
        for i in ArtPaths:
            r.add_picture(i)
        document.add_heading('Script 2', 4)
        p = document.add_paragraph()
        r = p.add_run()
        r.add_text(Script2)
        document.save(Title+'.docx')


#for reference but not currently used
        # pdf = FPDF()
        # pdf.add_page()
        # pdf.set_xy(0, 0)
        # pdf.set_font('times', size=12.0)
        # pdf.cell( align='L', w=0, txt=Script1, border=0)
        # pdf.cell(align='L', w=0, txt= '\n' + '\n' + '\n' , border=0)
        # pdf.cell(align='L', w=0, txt= '\n' + '\n' + '\n' , border=0)
        # pdf.cell( align='L', w=0, txt=Script2, border=0)
        #
        # for i in ArtPaths:
        #     pdf.image(i)  # , x=10, y=8, w=100)
        #
        #
        # pdf.output(Title + '.pdf', 'F')


    def Add2MasterLyrics(self ,current_time2, ArtType, title, ArtistPoetInfo, poem, Tag, SavePath2, Prompts):
        data = [(current_time2, ArtType, title, ArtistPoetInfo, poem, Tag, SavePath2, Prompts)]
        df = pd.DataFrame(data,columns = ['Date_Time_Added','Art_Type','Title','Poet_Artist_Info','Poem_Song_Lyrics','Quality','Folder_Path', 'Prompts_Used'])
        DigitalAssist.add2Master3(df)



    def quickArt(self):
        prompt = DigitalAssist.ChatGPTDA(self, MakeArt=True, Prompt=(up.QuickArt))

        ArtPaths = []
        try:
            ArtPath = DigitalAssist.makeArt(self, prompt)
            ArtPaths.append(ArtPath)
        except:
            DigitalAssist.speak(self, 'Could not make the Art due to an error')

        Prompts_Used = [str('Quick_Art_Prompt: ' + up.QuickArt)]
        ArtistPoetInfo ='Written By: ' + up.Bot_Name + '      (' + 'Artwork by: ' + up.AI_ArtistName + ')'
        DigitalAssist.NamePoemSavePoem(self, prompt, ArtPaths,Prompts_Used, ArtistPoetInfo, title = 'Quick_Art')


    def Quick_Poem_v1(self, GPTPrompt = up.Random_Poem, ARTPrompt = up.Poem_Art_2_prompt):
        prompt = DigitalAssist.ChatGPTDA(self, temp=0.6, MakeArt=True, Prompt=(GPTPrompt),ConfirmBot = False)
        ArtPaths = []
        try:
            ArtPath = DigitalAssist.makeArt(self, prompt)
            ArtPaths.append(ArtPath)
        except:
            DigitalAssist.speak(self, 'Could not make the Art due to an error')

        Art2 = DigitalAssist.ChatGPTDA(self, temp=0.5, MakeArt=True, Prompt=(ARTPrompt+ '"' + prompt + '"'),ConfirmBot = False)

        try:
            ArtPath2 = DigitalAssist.makeArt(self, Art2)
            ArtPaths.append(ArtPath2)
        except:
            DigitalAssist.speak(self, 'Could not make the Art due to an error')

        Prompts_Used = [str('Quick_Poem_prompt: ' + GPTPrompt),'Poem_Art_1_prompt: Art Generated using poem directly', str('Prompt Fed into Chat GPT to make art prompt: ' + ARTPrompt),  str('Chat GPT Prompt sent to Art AI: ' + Art2)]
        ArtistPoetInfo ='Poem Written By: ' + up.AI_Poet_Name + '      (' + 'Artwork by: ' + up.AI_ArtistName + ')'
        DigitalAssist.NamePoemSavePoem(self,  prompt, ArtPaths,Prompts_Used, ArtistPoetInfo)

    def SundayScary_Poem(self):
        DigitalAssist.Quick_Poem_v1(self,GPTPrompt = up.Sunday_Scaries_Poem_prompt)





    def Shake_Poem_v1(self):
        ArtPaths= []

        DigitalAssist.makeQuickPoem(self)
        prompt = DigitalAssist.ChatGPTDA(self, temp=0.4, MakeArt=True, Prompt=(up.Shake_Poem_v1_prompt + ' "' + self.Words + '"'))
        try:
            ArtPath = DigitalAssist.makeArt(self, prompt)
            ArtPaths.append(ArtPath)
        except:
            DigitalAssist.speak(self, 'Could not make the Art due to an error')

        Art2 = DigitalAssist.ChatGPTDA(self, temp=0.5, MakeArt=True, Prompt=( up.Poem_Art_2_prompt + '"' + prompt + '"'))

        try:
            ArtPath2 = DigitalAssist.makeArt(self, Art2)
            ArtPaths.append(ArtPath2)
        except:
            DigitalAssist.speak(self, 'Could not make the Art due to an error')

        Prompts_Used = ['Original Prompt: '+ up.Shake_Poem_v1_prompt+ + '"' + self.Words + '"' , 'ChatGPT to DALL-E (poem sent as is): ' + prompt,'User Prompt to get description' + up.Poem_Art_2_prompt + 'poem' ,'Chat-GPT to DALL-E (describing poem): ' + Art2 ]
        ArtistPoetInfo = str('Poem Written By: ' + up.AI_Poet_Name + '      (' + 'Artwork by: ' + up.AI_ArtistName + ')')
        DigitalAssist.NamePoemSavePoem(self, prompt, ArtPaths,Prompts_Used, ArtistPoetInfo)

    def Shake_Poem_v2(self):
        ArtPaths= []
        DigitalAssist.makeQuickPoem(self)
        prompt = DigitalAssist.ChatGPTDA(self, temp=0.4, MakeArt=True, Prompt=( up.Shake_Poem_v2_prompt + '"' + self.Words + '"'))
        try:
            ArtPath = DigitalAssist.makeArt(self, prompt)
            ArtPaths.append(ArtPath)
        except:
            DigitalAssist.speak(self, 'Could not make the Art due to an error')

        Art2 = DigitalAssist.ChatGPTDA(self, temp=0.5, MakeArt=True, Prompt=( up.Poem_Art_2_prompt + '"' + prompt + '"'))
        try:
            ArtPath2 = DigitalAssist.makeArt(self, Art2)
            ArtPaths.append(ArtPath2)
        except:
            DigitalAssist.speak(self, 'Could not make the Art due to an error')

        Prompts_Used = ['Original Prompt: '+ up.Shake_Poem_v2_prompt+  '"' + self.Words + '"' , 'ChatGPT to DALL-E (poem sent as is): ' + prompt,'User Prompt to get description' + up.Poem_Art_2_prompt + 'poem' ,'Chat-GPT to DALL-E (describing poem): ' + Art2 ]
        ArtistPoetInfo =str('Poem Written By: ' + up.AI_Poet_Name + '      (' + 'Artwork by: ' + up.AI_ArtistName + ')')
        DigitalAssist.NamePoemSavePoem(self, prompt, ArtPaths,Prompts_Used, ArtistPoetInfo)





    def Sampler(self, Mode = 'ALL'):
        sample_w = 'Not Techno Mode - N/A'
        words = ''
        if Mode == 'Techno':
            sample_w = up.Techno_Sample_Question2
            words += DigitalAssist.ChatGPTDA(self, temp=0.5, MakeArt=True, Prompt=(sample_w), ConfirmBot=False)

        sample_p = up.Techno_Sample_Question

        Samples = DigitalAssist.ChatGPTDA(self, temp=0.5, MakeArt=True, Prompt=(sample_p), ConfirmBot=False)
        Prompts_Used = [str('Prompt1: ' + sample_w),str('Prompt2: ' + sample_p) ]
        ArtistPoetInfo = 'Lyrics Written By: ' + up.Song_Writer
        Title = 'Sampler Man'
        prompt = Samples + '\n'+ '\n'+ '\n' + words
        DigitalAssist.NamePoemSavePoem(self, prompt, [], Prompts_Used, ArtistPoetInfo, title=Title,FolderPath=up.AI_Music_Path)




    def Make_a_Chorus(self, Mode = 'Random Song'):
        ArtPaths = []
        if Mode == 'Random Song':
            Verse_p = up.verse_prompt
            Chorus_p = up.Chorus_prompt
            Bridge_p = up.Bridge_prompt

        elif Mode == 'Techno':
            Verse_p = up.Techno_Story_prompt
            Chorus_p = up.Techno_Chorus_prompt
            Bridge_p = up.Techno_Bridge_prompt

        else:
            print('No Mode provided')
            Verse_p = up.Rap_Story_prompt
            Chorus_p = up.Rap_Chorus_prompt
            Bridge_p = up.Rap_Bridge_prompt

        chorus = DigitalAssist.ChatGPTDA(self, temp=0.6, MakeArt=True, Prompt=(Chorus_p),ConfirmBot = False)

        chorus2 = chorus[0:500]


        bridge = DigitalAssist.ChatGPTDA(self, temp=0.4, MakeArt=True, Prompt=(Bridge_p + '"' +  chorus2 + '"'),ConfirmBot = False)

        if Mode == 'Techno':
            bridge2 = chorus2
        else:
            bridge2 = bridge[0:500]
            Verse_p = Verse_p + chorus2
        combo = str(bridge2)
        combo2 = combo[0:900]


        try:
            if combo2 =='':
                if chorus != '':
                    combo2 = 'Make a work of art inspired by the following song lyrics: ' +  chorus
                    ArtPath = DigitalAssist.makeArt(self, combo2)
                    ArtPaths.append(ArtPath)
                else:
                    combo2 = 'Make a random work of art, be creative and unique, make it visually and aesthetic pleasing to the viewer'
                    ArtPath = ''
            else:
                ArtPath = DigitalAssist.makeArt(self, combo2)
                ArtPaths.append(ArtPath)
        except:
            ArtPath = ''
        Art2 = DigitalAssist.ChatGPTDA(self, temp=0.5, MakeArt=True, Prompt=(up.Chorus_Art_2_prompt + bridge2), ConfirmBot=False)

        try:
            if Art2 == '':
                if chorus != '':
                    Art2 = 'Make a work of art inspired by the following song lyrics: ' + chorus
                    ArtPath2 = DigitalAssist.makeArt(self, Art2)
                    ArtPaths.append(ArtPath2)
                else:
                    Art2 = 'Make a random work of art, be creative and unique, make it visually and aesthetic pleasing to the viewer'
            else:
                ArtPath2 = DigitalAssist.makeArt(self, Art2)
                ArtPaths.append(ArtPath2)
        except:
            ArtPath2 = ''

        Title = DigitalAssist.ChatGPTDA(self, temp=0.6, MakeArt=True, Prompt=(up.Song_Title_Prompt +  chorus2 ), ConfirmBot=False)
        verses = DigitalAssist.ChatGPTDA(self, temp=0.6, MakeArt=True, Prompt=( Verse_p ),ConfirmBot = False)
        prompt = ('verses: ' + verses + "\n" + "\n" +'bridge: ' + bridge+  "\n" + "\n" + 'chorus: ' +  chorus)
        Prompts_Used = [str('Song Prompt: ' + up.Chorus_prompt), 'Song_prompt: Art Generated using song directly', str('Prompt Fed into Chat GPT to make art prompt: ' + up.Chorus_Art_2_prompt),str('Chat GPT Prompt sent to Art AI: ' + Art2), 'Bridge Prompt: ' + up.Bridge_prompt]
        ArtistPoetInfo ='Lyrics Written By: ' + up.Song_Writer + '      (' + 'Artwork by: ' + up.AI_ArtistName + ')'
        DigitalAssist.NamePoemSavePoem(self, prompt, ArtPaths,Prompts_Used, ArtistPoetInfo, title = Title, FolderPath = up.AI_Music_Path + '\\' + Mode,  ArtType = 'Song Lyrics' + Mode)


    def Make_a_Song(self, Mode = 'Random Song'):
        ArtPaths = []
        song = DigitalAssist.ChatGPTDA(self, temp=0.6, MakeArt=True, Prompt=(up.Song_prompt2),ConfirmBot = False)
        song2 = song[0:500]
        bridge = DigitalAssist.ChatGPTDA(self, temp=0.6, MakeArt=True, Prompt=(up.Bridge_prompt + '"'  + song2 + '"'),ConfirmBot = False)
        bridge2 = bridge[0:500]
        try:
            if bridge =='':
                if song !='':
                    bridge = 'Make a work of art inspired by the following song lyrics: ' +  song
                    ArtPath = DigitalAssist.makeArt(self, bridge)
                    ArtPaths.append(ArtPath)
                else:
                    bridge = 'Make a random work of art, be creative and unique, make it visually and aesthetic pleasing to the viewer'
                    ArtPath = ''
            else:
                ArtPath = DigitalAssist.makeArt(self, bridge)
                ArtPaths.append(ArtPath)
        except:
            ArtPath = ''

        Art2 = DigitalAssist.ChatGPTDA(self, temp=0.5, MakeArt=True, Prompt=(up.Chorus_Art_2_prompt + bridge2), ConfirmBot=False)
        try:
            if Art2 == '':
                if song != '':
                    Art2 = 'Make a work of art inspired by the following song lyrics: ' + song
                    ArtPath2 = DigitalAssist.makeArt(self, Art2)
                    ArtPaths.append(ArtPath2)
                else:
                    Art2 = 'Make a random work of art, be creative and unique, make it visually and aesthetic pleasing to the viewer'
                    ArtPath2 = ''

            else:
                ArtPath2 = DigitalAssist.makeArt(self, Art2)
                ArtPaths.append(ArtPath2)
        except:
            ArtPath2 = ''

        Title = DigitalAssist.ChatGPTDA(self, temp=0.5, MakeArt=True,   Prompt=(up.Song_Title_Prompt + bridge2  + '. ' + song2), ConfirmBot=False)
        prompt = ('bridge: ' + bridge + "\n" + "\n" + 'song: ' + song +  "\n" + "\n" )
        Prompts_Used = [str('Song Prompt: ' + up.Chorus_prompt), str('Song_prompt: ' + up.Chorus_Art_2_prompt), str('Prompt Fed into Chat GPT to make art prompt: ' + up.Chorus_Art_2_prompt), str('Chat GPT Prompt sent to Art AI: ' + Art2), 'Bridge Prompt: ' + up.Bridge_prompt]
        ArtistPoetInfo ='Lyrics Written By: ' + up.Song_Writer + '      (' + 'Artwork by: ' + up.AI_ArtistName + ')'
        DigitalAssist.NamePoemSavePoem(self, prompt, ArtPaths,Prompts_Used, ArtistPoetInfo, title = Title, FolderPath = up.AI_Music_Path + '\\' + Mode,  ArtType = 'Song Lyrics' + Mode)

    def Make_a_Rap(self, Mode = 'Rap'):
        ArtPaths = []
        if Mode == 'Rap':
            Verse_p = up.Rap_Story_prompt
            Chorus_p = up.Rap_Chorus_prompt
            Bridge_p = up.Rap_Bridge_prompt

        elif Mode == 'Raggae':
            Verse_p = up.Reggae_Story_prompt
            Chorus_p = up.Reggae_Rap_Chorus_prompt
            Bridge_p = up.Reggae_Bridge_prompt


        elif Mode == 'SadRap':
            Verse_p = up.Rap_Story_prompt2
            Chorus_p = up.Rap_Chorus_prompt2
            Bridge_p = up.Rap_Bridge_prompt2


        else:
            print('No Mode provided')
            Verse_p = up.Rap_Story_prompt
            Chorus_p = up.Rap_Chorus_prompt
            Bridge_p = up.Rap_Bridge_prompt

        verses = DigitalAssist.ChatGPTDA(self, temp=0.4, MakeArt=True,Prompt=(  Verse_p),ConfirmBot=False)
        verses2 = verses[0:500]
        print('Verse Complete, Chorus below')
        chorus = DigitalAssist.ChatGPTDA(self, temp=0.6, MakeArt=True, Prompt=(Chorus_p),  ConfirmBot=False)
        chorus2 = chorus[0:500]
        print('Chorus Complete, Title below')
        Title = DigitalAssist.ChatGPTDA(self, temp=0.5, MakeArt=True, Prompt=(up.Song_Title_Prompt  + chorus2 ), ConfirmBot=False)
        print('Title Complete, bridge below')
        bridge = DigitalAssist.ChatGPTDA(self, temp=0.5, MakeArt=True,Prompt=(Bridge_p), ConfirmBot=False)
        bridge2 = bridge[0:500]
        print('bridge Complete, Make Art below')
        combo = Title + bridge2
        combo2 = combo[0:500]


        try:
            if combo2 == '':
                if chorus != '':
                    combo2 = 'Make a work of art inspired by the following song lyrics: ' + chorus[0:200]
                    ArtPath = DigitalAssist.makeArt(self, combo2)
                    ArtPaths.append(ArtPath)
                else:
                    combo2 = 'Make a random work of art, be creative and unique, make it visually and aesthetic pleasing to the viewer'
                    ArtPath = ''
            else:
                ArtPath = DigitalAssist.makeArt(self, combo2)
                ArtPaths.append(ArtPath)

        except:
            ArtPath = ''
        Art2 = DigitalAssist.ChatGPTDA(self, temp=0.5, MakeArt=True, Prompt=(up.Chorus_Art_2_prompt + bridge2), ConfirmBot=False)
        try:
            if Art2 == '':
                if chorus != '':
                    Art2 = 'Make a work of art inspired by the following song lyrics: ' +  chorus[0:200]
                    ArtPath2 = DigitalAssist.makeArt(self, Art2)
                    ArtPaths.append(ArtPath2)
                else:
                    Art2 = 'Make a random work of art, be creative and unique, make it visually and aesthetic pleasing to the viewer'
                    ArtPath2 = ''

            else:
                ArtPath2 = DigitalAssist.makeArt(self, Art2)
                ArtPaths.append(ArtPath2)
        except:
            ArtPath2 = ''
        print('Art2 Script Complete, Make Art2 below')

        prompt = ('verses: ' + verses + "\n" + "\n" +'bridge: ' + bridge+  "\n" + "\n" +'chorus: ' +  chorus)
        Prompts_Used = [str('Song Prompt: ' + up.Rap_Story_prompt), 'Song_prompt: ' + Art2 , str('Prompt Fed into Chat GPT to make art prompt: ' + up.Chorus_Art_2_prompt),str('Chat GPT Prompt sent to Art AI: ' + Art2), 'Bridge Prompt: ' + up.Bridge_prompt]
        ArtistPoetInfo = str('Lyrics Written By: ' + up.Song_Writer + '      (' + 'Artwork by: ' + up.AI_ArtistName + ')')
        DigitalAssist.NamePoemSavePoem(self, prompt, ArtPaths,Prompts_Used, ArtistPoetInfo, title = Title, FolderPath = up.AI_Music_Path + '\\' + Mode,  ArtType = 'Song Lyrics' + Mode)



    def Wiki4PenNames(self):
        current_time1 = datetime.datetime.now()
        current_time2 = current_time1.strftime('%m-%d-%Y_%H.%M.%S')
        f2 = up.AI_Poetry_Path
        SavePath1 = f2
        FileName = 'S.L. Rose & Sage Pixel '
        Filename = '\\' + FileName + "_"
        Title = SavePath1 + Filename + current_time2

        Script1 = DigitalAssist.ChatGPTDA(self, temp=0.5, MakeArt=False, Prompt=(up.Custom_prompt_1), ConfirmBot=False)

        CoverArt1 = DigitalAssist.ChatGPTDA(self, temp=0.5, MakeArt=False, Prompt=(up.Art_prompt_1 +  '"' + Script1[0:200] + '"'), ConfirmBot=False)
        try:
            CoverArtPath1 = DigitalAssist.makeArt(self, CoverArt1)
        except:
            CoverArtPath1 = ''
        Script2 = DigitalAssist.ChatGPTDA(self, temp=0.6, MakeArt=False, Prompt=(up.Custom_prompt_2), ConfirmBot=False)
        CoverArt2 = DigitalAssist.ChatGPTDA(self, temp=0.5, MakeArt=False, Prompt=(up.Art_prompt_2 +  '"' + Script2[0:200] + '"'),ConfirmBot=False)
        try:
            CoverArtPath2 = DigitalAssist.makeArt(self, CoverArt2)
        except:
            CoverArtPath2 = ''
        Script3 = DigitalAssist.ChatGPTDA(self, temp=0.6, MakeArt=False, Prompt=(up.Custom_prompt_3), ConfirmBot=False)
        CoverArt3 = DigitalAssist.ChatGPTDA(self, temp=0.5, MakeArt=False, Prompt=(up.Art_prompt_3 +  '"' + Script3[0:200] + '"'),ConfirmBot=False)

        try:
            CoverArtPath3 = DigitalAssist.makeArt(self, CoverArt3)
        except:
            CoverArtPath3 = ''


        ArtPaths = [CoverArtPath1, CoverArtPath2, CoverArtPath3]


        Prompts = (up.Custom_prompt_1+ '|' + up.Art_prompt_1+'|' + up.Custom_prompt_2+ '|' +up.Art_prompt_2+ '|' + up.Custom_prompt_3+'|' + up.Art_prompt_3)


        document = Document()
        document.add_heading('Pen Name info: S.L. Rose & Sage Pixel', 0)
        document.add_heading('Wiki for S L Rose  ', 4)
        p = document.add_paragraph()
        r = p.add_run()
        r.add_text(Script1)

        document.add_heading('Potential Profile pictures', 4)
        p = document.add_paragraph()
        r = p.add_run()

        for i in ArtPaths:
            r.add_picture(i)
        document.add_heading('Wiki for Sage Pixel 2', 4)
        p = document.add_paragraph()
        r = p.add_run()
        r.add_text(Script2)
        document.add_heading('Wiki for Dark Poet', 4)
        p = document.add_paragraph()
        r = p.add_run()
        r.add_text(Script3)
        document.save(Title + '.docx')

        poem = str('SJ Rose: ' + Script1 + 'Sage Pixel: ' + Script2 + 'Macabre Artist: ' + Script3)
        Tag = DigitalAssist.YayorNay(self)

        DigitalAssist.Add2MasterLyrics(self, current_time2, 'Wiki Page and Social Media', 'Wiki for SJ Rose, Pixel Art and Macabre artist', self.AssistantName, poem, Tag, SavePath1, Prompts)




    def makeBlogPost(self, GPTprompt = up.Random_Blog_Topic, ARTprompt = up.Art_1_prompt, BlogAuthor =  up.AI_Blogger):
        current_time1 = datetime.datetime.now()
        current_time2 = current_time1.strftime('%m-%d-%Y')
        Script1 = DigitalAssist.ChatGPTDA(self, temp=0.5, MakeArt=False, Prompt=(GPTprompt), ConfirmBot=False)
        CoverArt1 = DigitalAssist.ChatGPTDA(self, temp=0.5, MakeArt=False, Prompt=(ARTprompt + '"' + Script1[0:200] + '"'), ConfirmBot=False)

        try:
            CoverArtPath1 = DigitalAssist.makeArt(self, CoverArt1)
        except:
            CoverArtPath1 = ''

        Prompt1 = (GPTprompt + '|' + ARTprompt)
        ArtPaths = [CoverArtPath1]

        DigitalAssist.NamePoemSavePoem(self, Script1, ArtPaths, Prompt1, BlogAuthor,title='Blog_'+Script1[0:15]+'_' + current_time2, FolderPath=up.AI_Blog_Path, ArtType='AI Blog')





    def YayorNay(self):
        self.Mode = 'Yay or Nay -  Menu'

        DigitalAssist.speak(self, self.Mode + ' What did you think ' + self.UserName)

        print( self.Mode + ' What did you think ' + self.UserName)

        s2 = True
        while (s2 == True):
            DigitalAssist.Add2Transcript(self,' \n')
            query = DigitalAssist.takeCommand(self).lower()


            #Say Quick Art or Auto Art or Basic Art
            if 'awesome'  in query or 'love'  in query or 'yay' in query or  'yes'  in query or 'best' in query:
                rr = '- Sounded Great'
                s2 = False
                continue

            elif 'good'  in query or 'love'  in query or 'yay' in query or  'yes'  in query:
                rr = '- Good - POTENTIAL HERE'
                s2 = False
                continue

            elif 'ok' in query or 'okay' in query or 'not bad' in query or 'so so' in query:
                rr = ' - So So'
                s2 = False

            elif 'no' in query or 'hate' in query or 'bad' in query or 'terrible' in query or 'yuck' in query:
                rr = ' - BAD'
                s2 = False
            else:
                rr = ''
                s2 = False

        return rr

    def NamePoemSavePoem(self,poem, ArtPaths,Prompts_Used,ArtistPoetInfo, title = '', FolderPath = up.AI_Poetry_Path , ArtType = 'Poem'):
        dfPrompts = ''
        Tag = ''
        if ArtType == 'Poem':
            title_p = up.Poem_Title_Prompt
        else:
            title_p= up.Song_Title_Prompt

        if title =='':
            Title = DigitalAssist.ChatGPTDA(self, temp=0.5, MakeArt=True, Prompt=(title_p  +'"' + poem + '"'), ConfirmBot= False)
        else:
            Title = title
        current_time1 = datetime.datetime.now()
        current_time2 = current_time1.strftime('%m-%d-%Y_%H.%M.%S')
        f2 = FolderPath
        SavePath1 = f2
        invalidCharRemoved = re.sub(r"[^a-zA-Z0-9 ]", "", Title)
        folder = str(invalidCharRemoved)
        if len(folder)>40:
            folder = folder[0:40]
        Title1 = '\\' + str(invalidCharRemoved) + "_"
        if len(Title1)>180:
            Title1 = Title1[0:180]
        SavePath2 = SavePath1 +'\\' +folder
        Title2 = SavePath2 + Title1 + current_time2
        Tag = DigitalAssist.YayorNay(self)
        Title2 = Title2 + Tag


        isExist = os.path.exists(SavePath1)
        if not isExist:
            # Create a new directory because it does not exist
            os.makedirs(SavePath1)

        isExist = os.path.exists(SavePath2)
        if not isExist:
            # Create a new directory because it does not exist
            os.makedirs(SavePath2)

        document = Document()
        document.add_heading(Title, 0)
        document.add_heading(ArtistPoetInfo, 1)
        p = document.add_paragraph()
        r = p.add_run()
        r.add_text(ArtistPoetInfo)
        p = document.add_paragraph()
        r = p.add_run()
        r.add_text(poem)
        p = document.add_paragraph()
        r = p.add_run()
        image_counter= 0
        for i in ArtPaths:
            image_counter +=1
            r.add_picture(i)
            original = i
            target = SavePath2 + Title1  + str(image_counter) +current_time2 + '.png'
            shutil.copyfile(original, target)

        document.add_heading('AI Prompts used: ', 7)

        for i in Prompts_Used:
            p = document.add_paragraph()
            r = p.add_run()
            r.add_text(i)
            dfPrompts += i + '|'
        document.save(Title2+  '_Details.docx')
        document2 = Document()
        document2.add_heading(Title, 1)
        for i in ArtistPoetInfo:
            document2.add_heading(i, 1)
        p = document2.add_paragraph()
        r = p.add_run()
        r.add_text(poem)
        document2.save(Title2 + '.docx')
        DigitalAssist.Add2MasterLyrics(self,current_time2, ArtType, title, ArtistPoetInfo, poem, Tag, SavePath2, dfPrompts)


    def SaveLiveArt(self, ArtPaths, Title = 'Live Art'):
        current_time1 = datetime.datetime.now()
        self.current_time22 = current_time1.strftime('%m-%d-%Y')
        current_time3 = current_time1.strftime('%m-%d-%Y_%H.%M.%S')
        f2 = up.AI_Live_Art_Path
        SavePath1 = f2

        invalidCharRemoved = re.sub(r"[^a-zA-Z0-9 ]", "", Title)
        Title1 = '\\' + str(invalidCharRemoved) + "_"
        if len(Title1)>100:
            Title1 = Title1[0:100] + self.current_time22
        folder = str(invalidCharRemoved)
        if len(folder) > 40:
            folder = folder[0:40]

        SavePath2 = SavePath1 + '\\' + folder
        isExist = os.path.exists(SavePath1)
        if not isExist:
            # Create a new directory because it does not exist
            os.makedirs(SavePath1)

        isExist = os.path.exists(SavePath2)
        if not isExist:
            # Create a new directory because it does not exist
            os.makedirs(SavePath2)

        image_counter = 0
        for i in ArtPaths:
            image_counter += 1
            original = i
            target = SavePath2 + Title1  + current_time3 + '.png'
            shutil.copyfile(original, target)


#Blog Make Art Live mode
    def MakeArtLive(self):
        self.onlyMyWords = ''
        self.onlyMyWordsLatest2 = ''
        self.onlyMyWordsLatest = []
        self.Mode = 'Make Art Live Mode'
        Title = DigitalAssist.getfilename(self)

        #Ask user if you should have silent mode on?

        s2 = True
        while (s2 == True):

            DigitalAssist.speak(self, 'Do you Want Silent Mode activated for Chat GPT?')
            query = DigitalAssist.takeCommand(self).lower()
            # DigitalAssist.Add2Transcript(self,text2Add=query)
            if "yes" in query or "silent" in query:
                self.SilentMode = True
                s2 = False
                DigitalAssist.speak(self, 'Silent Mode Activated', voice = 1)
            elif "no" in query or "nope" in query or "loud" in query:
                DigitalAssist.speak(self, 'Interrupt Mode Activated', voice=1)
                self.SilentMode = False
                s2 = False
            else:
                print ("Gonna Have to Repeat That....")

        HowManyLinestoAdd = 7
        self.LineCount = 0
        LiveMode = True
        while (LiveMode ==True):
            print(self.onlyMyWordsLatest2)
            self.LineCount = self.LineCount + 1
            self.query1 = DigitalAssist.transcribe_Build_Query(self,1.44).lower()
            self.query1 = self.query1 + '.'+ '\n'
            self.onlyMyWords = self.onlyMyWords +  self.query1
            self.onlyMyWordsLatest2 = self.onlyMyWordsLatest2 + self.query1
            self.onlyMyWordsLatest.append( self.query1)

            if 'loud mode' in self.query1:
                self.SilentMode = False
            elif (("brick top" in self.query1 or "bricktop" in self.query1 or "show me" in self.query1 or "extra" in self.query1 or "additional" in self.query1) and ("option" in self.query1 or  "tool" in self.query1)) or ("edit mode" in self.query1) or  ("need to" in self.query1 and "make edit" in self.query1):
                s3 = True
                while (s3 == True):
                #this is where you say next steps (This can be a separate function)
                    query2 = DigitalAssist.transcribe_Build_Query_Pause(self).lower()
                    print(query2)

                    if 'chat' in query2 and ('gpt' in query2 or 'bot' in query2 or 'ai' in query2 or 'mode' in query2):
                        DigitalAssist.RunChatGPT(self)
                    elif ('take' in query2 or 'note' in query2 or 'write' in query2 or 'transcribe' in query2):
                        DigitalAssist.transcribe(self)
                    elif ('none' in query2 or 'no ' in query2 or 'mistake' in query2 or 'my bad' in query2 or 'go back' in query2):
                        s3 == True
                    elif ('stop' in query2 or 'cancel' in query2 or 'quit' in query2 or 'cut' in query2):
                        sys.exit()
                        continue
            else:
                SubjectMatter = DigitalAssist.Pick_Random_Lines(self,SetofLines=self.onlyMyWordsLatest   , LineCount = 2)
                if self.LineCount == HowManyLinestoAdd + 1:
                    DigitalAssist.LiveArt(self, SubjectMatter, Title= Title)
                    self.onlyMyWordsLatest = []
                    self.onlyMyWordsLatest2 = ''

        Script1 = 'Human Transcript: ' +  self.onlyMyWords + '\n' + 'Full Transcript: ' +    self.transcript_Final
        DigitalAssist.NamePoemSavePoem(self, Script1, [], '', 'Shane Donovan - MondeVert CEO', title=Title + '_' + self.current_time22, FolderPath=up.AI_Live_Art_Path,ArtType='MondeVert Podcast with Live Art')



    def LiveArt(self , SubjectMatter = 'Random subject of your choice, weirder the better', Title = 'Podcast UnNamed'):
        #LiveArt1 = DigitalAssist.ChatGPTDA(self, temp=0.5, MakeArt=True, Prompt=(up.MakeArtLive_prompt1), ConfirmBot=False)
        #ArtPath1 = DigitalAssist.makeArt(self, LiveArt1)
        ArtPath2 = DigitalAssist.makeArt(self, up.MakeArtLive_prompt2 + ': ' + SubjectMatter)
        #ArtPaths = [ArtPath1, ArtPath2]
        ArtPaths = [ ArtPath2]
        DigitalAssist.SaveLiveArt(self, ArtPaths, Title)

    def Pick_Random_Lines(self, SetofLines, LineCount = 2):
        Line1 = random.choices(SetofLines)
        Line1 = Line1[0]
        Line2 = Line1
        while Line2 == Line1:
            Line2 = random.choices(SetofLines)
            Line2 = Line2[0]
            continue
        return str(Line1 + Line2)


#Main Menu
    def Take_query(self):
        DigitalAssist.Hello(self)

        self.Mode = 'Main Menu'
       # DigitalAssist.Add2Transcript(self,text2Add= ( self.Mode + ':'))
        s2 = True
        while (s2 == True):
            self.Mode = 'Main Menu'
            DigitalAssist.speak(self,self.Mode)
            query = DigitalAssist.takeCommand(self).lower()
            #DigitalAssist.Add2Transcript(self,text2Add=query)
            if "company dashboard" in query or "open dashboard" in query :
                DigitalAssist.speak(self,"Opening MondDay Vert's Dasboard, good luck with your work!")

                # in the open method we just to give the link
                # of the website and it automatically open
                # it in your default browser
                Mondevert = "http://mondeVert.co"
                godaddy = 'https://dashboard.godaddy.com/'
                Mail = 'https://outlook.office.com/mail/'
                Square_store = 'https://square.online/app/home/'
                DigitalAssist.Open_Web(Mail)
                DigitalAssist.Open_Web(Square_store)
                DigitalAssist.Open_Web(Mondevert)
                DigitalAssist.Open_Web(godaddy)
                continue

            elif "pay bills" in query or "pay credit cards" in query:
                DigitalAssist.speak(self,"Bills Suck man, at least you are getting them done, Hope you have a nice day!")

                Liberty_Bay_CC = "https://www.myaccountaccess.com/onlineCard/login.do"
                Home_Depot_CC = 'https://citiretailservices.citibankonline.com/RSnextgen/svc/launch/index.action?siteId=PLCN_HOMEDEPOT#signon'
                Macys_CC = 'https://www.macys.com/my-credit/gateway/guest'
                Insurance = 'https://customer.concordgroupinsurance.com/ccp/login'
                National_Grid = 'https://login.nationalgridus.com/loginnationalgridus.onmicrosoft.com/oauth2/v2.0/authorize?cancel_uri=https%3A%2F%2Fmyaccount.nationalgrid.com%2Fservices%2Fauth%2Fsso%2FNGP_SignIn_MA_Electric_Home&client_id=88d004b4-3d39-4599-b410-093849907ee5&customer_type=home&p=B2C_1A_UWP_NationalGrid_convert_merge_signin&redirect_uri=https%3A%2F%2Fmyaccount.nationalgrid.com%2Fservices%2Fauthcallback%2FNGP_SignIn_MA_Electric_Home&region=masselec&response_type=code&scope=https%3A%2F%2Flogin.nationalgridus.com%2Fapi%2Fread+openid+profile+email+offline_access&state=CAAAAYUokN2YMDAwMDAwMDAwMDAwMDAwAAAA8Iy1JaCJ8hTkDIfqZ5i6KE2oIrQFv5bPH9zBI_P-5wWki9riiDhgLKrNlnVD4FeR38QBEzYH17vRm_CbsR8msY3J3fEDKFqtBsup_Bp9KwM-Flb38WYF6YqeqejsI-Iuj-yWpXfajw6-QKG_DMoOBtkqcL8zubs4c8KgKiuF-yIk3yoa1_vp_KqVAuBw0DSopTJrDwsYAZN5qW9zaUmwyW-wk6zbTNddqu24EJObkHkfEGKHEKQwrqArl40WukP3dxbKBfj7sSMhdqjTPa0ido8%3D&switch_uri=https%3A%2F%2Fmyaccount.nationalgrid.com%2Fs%2Flogin'
                LLBean_CC = 'https://citiretailservices.citibankonline.com/RSauth/signon?pageName=signon&siteId=PLCN_LLBEAN&langId=en_US'
                Amazon_CC = 'https://secure07ea.chase.com/web/auth/#/logon/logon/chaseOnline'
                DigitalAssist.Open_Web(Liberty_Bay_CC)
                DigitalAssist.Open_Web(Home_Depot_CC)
                DigitalAssist.Open_Web(Macys_CC)
                DigitalAssist.Open_Web(Insurance)
                DigitalAssist.Open_Web(National_Grid)
                DigitalAssist.Open_Web(LLBean_CC)
                DigitalAssist.Open_Web(Amazon_CC)
                continue


            elif ("chat" in query and "bot" in query) or  ("gpt" in query and "chat" in query) or  ("smart" in query and "assist" in query) or  ("extra" in query and "help" in query) or   ("artificial" in query and "intellig" in query):

                DigitalAssist.RunChatGPT(self)
                DigitalAssist.speak(self,"High Tech!")
                # prompt = DigitalAssist.ChatGPTDA(self)
                continue

            elif ("make" in query and "art" in query) or ("make" in query and "picture" in query) or ("ai" in query and "art" in query) or ("art" in query and "mode" in query) or ("super" in query and "shake" in query and 'art' in query):

                DigitalAssist.speak(self, "High Tech and sheek!")
                if 'super' in query or 'using' in query or 'extra' in query:
                    if 'poem' in query or 'shake' in query or 'poet' in query:
                        DigitalAssist.makeQuickPoem(self)
                        prompt = DigitalAssist.ChatGPTDA(self, MakeArt=True, Prompt=('make a prompt for openai  DALL-E program to create a work of art that is based on the following poem. '+'\n'+'\n' + self.Words))
                        DigitalAssist.makeArt(self, prompt)
                    else:
                        prompt = DigitalAssist.ChatGPTDA(self, MakeArt=True)
                        DigitalAssist.makeArt(self, prompt)

                else:

                    if 'poem' in query or 'shake' in query or 'poet' in query:
                        DigitalAssist.makeQuickPoem(self)
                        DigitalAssist.makeArt(self, self.Words)
                    else:
                        DigitalAssist.makeArt(self)
                        continue



            elif "stop" in query or "all set" in query or "quit" in query or "m done" in query or "thank" in query or "peace out" in query:
                DigitalAssist.speak(self,"Happy to help, Peace out brothah man")
                s2 = False

            elif ("record" in query or "listen up" in query or "can you right" in query or "record mode"  in query  or  "transcribe" in query or "could you write" in query or "transcribe" in query or "can you write" in query) and 'stop' not in query :
                DigitalAssist.speakSweet(self,"Dictated but not Red Mode Activated")
                DigitalAssist.transcribe(self)
                continue


            elif ( "blog" in query )and ("create" in query or "together" in query or "put"in query or "make"in query ) or ("add" in query and "blog" in query):
                DigitalAssist.speakSweet(self,"Mon Day Vert is lucky to have a great CEO this is going to be a great post!")
                DigitalAssist.CompileBlog(self)
                DigitalAssist.speakSweet(self,"Blog Consolidation Complete Shane D")
                continue

            elif (("consolidate" in query or "compile" in query  or "collect" in query )and ("file" in query or "folder" in query or "list"in query )) or ("add" in query and "master" in query):
                DigitalAssist.speakSweet(self,"Consolidation in progress")
                DigitalAssist.compileFiles(self)
                DigitalAssist.speakSweet(self,"Consolidation Complete my dude")
                continue

            elif "make list" in query or " list" in query or "to do" in query or "tracker" in query or "checklist" in query  or  "make a plan" in query or "quick note" in query  or "quick plan" in query or "take note" in query:
                #DigitalAssist.speak("Sure thing, One moment while I prepare")
                DigitalAssist.makelist(self)
                continue


            elif "poetic" in query or "swoon me" in query or "me a poem" in query  or  "poem" in query or "while I think you" in query or "poetry will help me think" in query or "write me a poem" in query or "we need some inspiration" in query or "some art" in query or "be creative" in query or "inspire me" in query:
                x22 = s.PoemBot(1, 1, 1, 1, 40, 3)
                x22.ReloadModel("model.h5")
                x22.setupdata()
                self.Words = x22.shakesbot_DA()
                print (self.Words)
                DigitalAssist.speakSweet(self,self.Words)
                continue




            elif ("play" in query or "script" in query) and ("make" in query or "write" in query or "create" in query or "mode" in query):
                # Call for the save file name after listing all of the subjects to the user
                self.Words = ''
                self.Correction_Comment = ''
                self.Type = 'AI Script Break'
                self.Significance = '**AI content attached'
                self.FileName = 'AI ShakeScript'
                self.Subject = 'AI ShakeScript'
                DigitalAssist.speak(self,'I like your style, my man')

                DigitalAssist.Make_Script(self)

                DigitalAssist.speakSweet(self,self.Words)
                s3 = pd.DataFrame()
                s3['Script'] = [self.Words]
                s3['Script - Corrected'] = [self.AI_Corrected_Content]
                s3['Correction Comment'] = [self.Correction_Comment]
                DigitalAssist.SaveText(self,s3, self.FileName, self.Subject)
                continue

            elif ("go" in query or 'try' in query or "mode" in query) and ("greek" in query or "myth" in query ):
                # Call for the save file name after listing all of the subjects to the user
                self.Words = ''
                self.Type = 'Greek Script Break'
                self.Significance = '**AI content attached'
                self.FileName = 'AI Going Greek'
                self.Subject = 'AI Going Greek'
                DigitalAssist.speak(self,'I like your style, my man... Dont eat too much tzatziki!')
                self.pb = s.PoemBot(1, 1, 1, 1, 40, 3)
                self.pb.ReloadModel("model_Greek.h5")
                self.pb.setupdata()
                size = 3
                trd = numpy.empty(size, dtype=str)
                w = numpy.empty(size, dtype=str)
                for i in range(0, size - 1):
                    globals()[f"trd{i}"] = threading.Thread(target=DigitalAssist.generateGreek(self))
                    w[i] = globals()[f"trd{i}"].start()  # starting the thread i

                for i in range(0, size - 1):
                    threading.active_count()
                    threading.enumerate()
                    if i == 0:
                        self.Words += w[i]
                    else:
                        i2 = i - 1
                        globals()[f"trd{i2}"].join()
                        globals()[f"trd{i}"].join()
                        self.Words += w[i]
                        print(self.Words)

                DigitalAssist.speak(self,self.Words)
                s3 = pd.DataFrame()
                s3['Script'] = [self.Words]
                DigitalAssist.SaveText(self,s3, self.FileName, self.subject)
                continue



            elif "social media" in query or "marketing" in query:
                DigitalAssist.speak(self,"Will do")

                # in the open method we just to give the link
                # of the website and it automatically open
                # it in your default browser
                Facebook = "https://www.facebook.com/"
                Insta = 'https://www.instagram.com/'
                TikTok = 'https://www.tiktok.com/'
                Twitter = 'https://twitter.com/home'
                DigitalAssist.Open_Web(Facebook)
                DigitalAssist.Open_Web(TikTok)
                DigitalAssist.Open_Web(Twitter)
                DigitalAssist.Open_Web(Insta)
                continue

            elif "open google" in query:
                DigitalAssist.speak(self,"Opening Google ")
                url = "http://google.com"
                DigitalAssist.Open_Web(url)
                continue



            elif   (("voice" in query or "person" in query or "bot" in query or "ai" in query) and ( "change" in query or  "update" in query or "switch" in query or "set" in query )):

                if "quick" in query or "manual" in query:
                    DigitalAssist.Setvoices(self, Quick = True)
                    print('Quick mode activated')
                else:
                    DigitalAssist.Setvoices(self)
                continue

            elif "s the date" in query:
                DigitalAssist.tellDay(self)
                continue

            elif "the time" in query or "what time" in query or "clock mode" in query:
                DigitalAssist.tellTime(self)
                continue

            # this will exit and terminate the program
            elif "bye" in query:
                DigitalAssist.speak(self,"Peace Out!")
                s2 = False

            elif "from wikipedia" in query or "wikipedia" in query:
                DigitalAssist.speak(self,"Checking the wikipedia ")
                query = query.replace("wikipedia", "")
                result = wikipedia.summary(query, sentences=4)
                DigitalAssist.speak(self,"According to wikipedia")
                DigitalAssist.speak(self,result)



            elif "tell me your name" in query or  "your name" in query or "introduce yourself" in query or "introduce your self" in query:
                DigitalAssist.speak(self,"I am Big Master Funk the fourth also known as Brick top, AND i AM  Your personal desktop Assistant")


            elif "shut up" in query:
                DigitalAssist.speak(self,"my bad I am tripping, I will leave you be Shane D")
                s2 = False

            else:
                #uk = cs.DA_Unknown_Command()
                #DigitalAssist.speak(uk)
                #print(uk)
                #Event().wait(1)
                continue




    def transcribe(self):
        self.entry = 1

        DigitalAssist.createDF(self)
        self.subject = 'Default'
        self.FileName = 'Default'
        s2 = True
        DigitalAssist.getsubject(self)
        self.visualList = 'Subject: ' + self.subject + ' \n'
        self.visualList2 = 'Subject*: ' + self.subject + ' \n'
        DigitalAssist.speak(self,"Please prepare for recording")
        print("Please prepare for recording...")
        while (s2 ==True):
            xx = ''
            self.query1 = ''
            self. new_Query = ''
            self.Type = ''
            self.Significance = ''
            self.Words = ''
            self.Notes = ''
            self.Added2TextFile = ''
            self.Completed = ''
            print(self.visualList)
            print(self.visualList2)
            self.query1 = DigitalAssist.transcribe_Build_Query(self,1.44).lower()
            self.query1 = self.query1 + '.'
            self.Type = 'Original'
            self.Significance = 'Default'



            if "stop recording" in self.query1 or "m all set" in self.query1 or "end list" in self.query1 or "stop listening" in self.query1 or "save this file" in self.query1 or "save this recording" in self.query1:
                # Call for the save file name after listing all of the subjects to the user
                DigitalAssist.getfilename(self)
                DigitalAssist.SaveText(self,self.transcript, self.FileName, self.subject)
                DigitalAssist.add2Master(self.transcript)
                s2 = False
                DigitalAssist.saveTranscript(self)
                sys.exit()
                continue


            elif (("brick top" in self.query1 or "bricktop" in self.query1 or "show me" in self.query1 or "extra" in self.query1 or "additional" in self.query1) and ("option" in self.query1 or  "tool" in self.query1)) or ("edit mode" in self.query1) or  ("need to" in self.query1 and "make edit" in self.query1):
                s3 = True
                while (s3 == True):
                #this is where you say next steps (This can be a separate function)
                    query2 = DigitalAssist.transcribe_Build_Query_Pause(self).lower()
                    print(query2)

                    if  "continue" in query2 or "just thinking" in query2 or "yes" in query2 or "chill" in query2 or "just thinking" in query2 or "keep going" in query2:
                        DigitalAssist.speak(self,'no worries, still here')
                        s3 = False
                        #continue



                    elif "change subject" in self.query1 or "update subject" in self.query1 or "switch category" in self.query1 or "change category" in self.query1:
                        # transcript2[self.subject] = [str(self.visualList)]
                        # DigitalAssist.SaveText(self.transcript2, self.FileName + 'v2' + self.subject, self.subject)
                        DigitalAssist.getsubject(self)
                        s3 = False
                        #continue


                    elif  "change" in query2 or "fix" in query2 or "edit" in query2 or "fix that" in query2 or "not quite" in query2 or "review last" in query2 or "back to me" in query2 or "read that back" in query2:
                        # Call for the save file name after listing all of the subjects to the user
                        self.Type = 'Edited'
                        self.Significance = 'Review edits'
                        self.new_Query = DigitalAssist.editentry(self,self.query1)
                        s3 = False
                        #continue


                    elif   "while I think you" in query2 or "poetry will help me think" in query2 or "write me a poem" in query2 or "we need some inspiration" in query2 or "some art" in query2 or "be creative" in query2 or "inspire me" in query2 or ( "poem" in query2 and  ("write" in query2 or "create" in query2 or "shake" in query2)):
                        # Call for the save file name after listing all of the subjects to the user
                        self.Type = 'Poetry Break'
                        self.Significance = '**AI content attached'
                        DigitalAssist.speakSweet(self,'I like your style, Shane you are brilliant and beautiful')
                        x = s.PoemBot(1, 1, 1, 1, 40, 3)
                        x.ReloadModel("model.h5")
                        x.setupdata()
                        self.Words += x.shakesbot_DA()
                        print (self.Words)
                        DigitalAssist.speakSweet(self,self.Words)
                        s3 = False
                        #continue

                    elif ("chat" in query2 and "bot" in query2) or ("gpt" in query2 and "chat" in query2) or ("smart" in query2 and "assist" in query2) or ("extra" in query2 and "help" in query2):
                        DigitalAssist.speak(self, "High Tech!")
                        prompt = DigitalAssist.ChatGPTDA(self)
                        self.Type = 'AI Conversation with Chat GPT'
                        self.Words += prompt
                        self.Significance = '**AI content attached'
                        continue

                    elif ("make" in query2 and "art" in query2) or ("make" in query2 and "picture" in query2) or ("ai" in query2 and "art" in query2) or ("art" in query2 and "mode" in query2) or ("super" in query2 and "shake" in query2 and 'art' in query2):

                        DigitalAssist.speak(self, "High Tech and sheek!")
                        if 'super' in query2 or 'using' in query2 or 'extra' in query2:
                            if 'poem' in query2 or 'shake' in query2 or 'poet' in query2:
                                DigitalAssist.makeQuickPoem(self)
                                prompt = DigitalAssist.ChatGPTDA(self, MakeArt=True, Prompt=('Using the poem that follows make a prompt to create a work of art that describes the feelings or the scene that fits the poem. '+'\n'+'\n' + self.Words))
                                DigitalAssist.makeArt(self, prompt)
                            else:
                                prompt = DigitalAssist.ChatGPTDA(self, MakeArt=True)
                                DigitalAssist.makeArt(self, prompt)
                            self.Type = 'AI Made Art - Add to BLOG?'
                            self.Words += prompt
                            self.Significance = '**AI content Please see image that goes with prompt'
                        else:

                            if 'poem' in query2 or 'shake' in query2  or 'poet' in query2:
                                DigitalAssist.makeQuickPoem(self)
                                DigitalAssist.makeArt(self, self.Words)
                            else:
                                DigitalAssist.makeArt(self)
                        continue

                    elif ( "play" in query2 or "script" in query2 )and ("make" in query2 or "write" in query2 or "create" in query2 or "mode" in query2 ):
                        # Call for the save file name after listing all of the subjects to the user
                        self.Type = 'AI Script Break'
                        self.Significance = '**AI content attached'

                        DigitalAssist.speak(self,'I like your style, my man')


                        DigitalAssist.Make_Script(self)
                        DigitalAssist.speakSweet(self,self.Words)

                        s3 = False
                        #continue
                    elif ("go" in query2 or 'try' in query2 or "mode" in query2) and ("greek" in query2 or "myth" in query2):
                        # Call for the save file name after listing all of the subjects to the user
                        self.Words = ''
                        self.Type = 'Greek Script Break'
                        self.Significance = '**AI content attached'
                        self.FileName = 'AI Going Greek'
                        self.Subject = 'AI Going Greek'
                        DigitalAssist.speak(self,'I like your style, my man... Dont eat too much tzatziki!')
                        self.pb = s.PoemBot(1, 1, 1, 1, 40, 3)
                        self.pb.ReloadModel("model_Greek.h5")
                        self.pb.setupdata()
                        size = 3
                        trd = numpy.empty(size, dtype=str)
                        w = numpy.empty(size, dtype=str)
                        for i in range(0, size - 1):
                            globals()[f"trd{i}"] = threading.Thread(target=DigitalAssist.generateGreek(self))
                            w[i] = globals()[f"trd{i}"].start()  # starting the thread i

                        for i in range(0, size - 1):
                            threading.active_count()
                            threading.enumerate()
                            if i == 0:
                                self.Words += w[i]
                            else:
                                i2 = i - 1
                                globals()[f"trd{i2}"].join()
                                globals()[f"trd{i}"].join()
                                self.Words += w[i]
                                print(self.Words)

                        DigitalAssist.speak(self,self.Words)

                        continue
                    elif "stop recording" in query2 or "m all set" in query2 or "no thanks" in query2 or "stop recording" in query2 or "save" in query2 or "no" in query2:
                        #Call for the save file name after listing all of the subjects to the user
                        print(self.subject)
                        DigitalAssist.getfilename(self)
                        DigitalAssist.SaveText(self,self.transcript, self.FileName, self.subject)
                        DigitalAssist.add2Master(self.transcript)
                        s3 = False
                        s2 = False
                        continue


                    elif query2 == 'none':
                        zz = 1
                        #Event().wait(1)

                    else:
                        zz = 1
                        #Event().wait(1)

            elif self.query1 == 'none':
                #Event().wait(1)
                continue
            else:

                self.AI_Corrected_Text = DigitalAssist.cleanText(self, self.query1)
                DigitalAssist.adddata2DF(self)
                self.visualList += str(self.entry) + '). ' + self.query1 + ' \n'
                self.visualList2 += str(self.entry) + '). ' + self.AI_Corrected_Text + ' \n'
                self.entry += 1

            #Do the dataframe stuff here




    def makelist(self):
         self.entry = 1
         self.transcript2 = pd.DataFrame()
         #self.transcript = pd.DataFrame()
         DigitalAssist.createDF(self)
         self.subject = 'To Do'
         self.FileName = 'To Do '
         s2 = True
         DigitalAssist.getsubject(self)
         DigitalAssist.speak(self,"Please prepare for recording")
         print("Please prepare for recording...")
         self.visualList = ''
         self.visualList2 = ''
         self.Type = 'To Do list'
         self.Significance = 'High'
         while (s2 == True):
             self.xx = ''
             self.query1 = ''
             self.new_Query = ''
             self.Words = ''
             self.Notes = ''
             self.Added2TextFile = ''
             self.Completed = ''
             print(self.visualList)
             print(self.visualList2)
             self.query1 = DigitalAssist.transcribe_Build_Query(self,1.3).lower()
             self.query1 = self.query1 + '.'




             if "stop recording" in self.query1 or "m all set" in self.query1 or "end list" in self.query1 or "stop listening" in self.query1 or "save this file" in self.query1 or "save this recording" in self.query1:
                 # Call for the save file name after listing all of the subjects to the user
                 DigitalAssist.SaveText(self,self.transcript,self.FileName,self.subject)
                 DigitalAssist.add2Master(self.transcript)
                 s2 = False
                 sys.exit()
                 continue

             elif "change subject" in self.query1 or "update subject" in self.query1 or "switch category" in self.query1 or"change category" in self.query1:
                 #transcript2[self.subject] = [str(self.visualList)]
                 #DigitalAssist.SaveText(self.transcript2, self.FileName + 'v2' + self.subject, self.subject)
                 DigitalAssist.getsubject(self)

                 continue
             elif  self.query1 == 'none' :
                 continue

             else:
                 self.AI_Corrected_Text = DigitalAssist.cleanText(self, self.query1)
                 DigitalAssist.adddata2DF(self)
                 self.visualList += str(self.entry) + '). ' + self.query1 + ' \n'
                 self.visualList2 += str(self.entry) + '). ' + self.AI_Corrected_Text + ' \n'
                 self.entry += 1








if __name__ == '__main__':
    # main method for executing
    # the functions
    Record = ''

    x = DigitalAssist(1)
    x.Take_query()


    #x.saveTranscript()
    atexit.register(x.saveTranscript)

   #Turn this back on after issue is understood

    # try:
    #     x = DigitalAssist(1)
    #     x.Take_query()
    #     x.saveTranscript()
    # except:
    #     x.saveTranscript()
    #     atexit.register(x.saveTranscript)
