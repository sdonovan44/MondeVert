from __future__ import unicode_literals
# from ffmpeg import Progress
# from ffmpeg.asyncio import FFmpeg
import threading
import whisper

import platform,subprocess
import time
from email.mime.text import MIMEText
from email.mime.multipart import MIMEMultipart
from email.mime.application import MIMEApplication
from email.header import Header
import smtplib
import smtplib
import email.mime.multipart
import email.mime.text
import email.mime.base
import os
# import FFProbe
# import FFProbe.utility
import random
from MondeVert_IP.SHAINE_MonderVert import Instagram_Posts as IP
from MondeVert_IP.SHAINE_MonderVert.SHAINE_WIZARD_PROMPTS import Long_User_Prompts as lup, Social_Media_SHAINE as sms,StoryMode_Wizard as StoryMode, StoryPrompts as SP
import shutil
from docx import Document
import pandas as pd
import pyttsx3
import datetime
from pdf2image import convert_from_path
import re
from boto3 import Session
from botocore.exceptions import BotoCoreError, ClientError
from contextlib import closing
import os
import sys
from tempfile import gettempdir
from MondeVert_IP.SHAINE_MonderVert.SHAINE_WIZARD_PROMPTS import Stories_For_Audio_Files as SAF
import MondeVert_IP.SHAINE_MonderVert.SHAINE_WIZARD_PROMPTS.User_Prefs as up
from MondeVert_IP.SHAINE_MonderVert.Utilities import DoNotCommit as DNC

from MondeVert_IP.SHAINE_MonderVert import SHAINE as GPT
from MondeVert_IP.SHAINE_MonderVert.SHAINE_WIZARD_PROMPTS import Stories_For_Audio_Files as sfa

from pydub import AudioSegment
import openai
from MondeVert_IP.SHAINE_MonderVert.Utilities.DoNotCommit import API_Key
import requests
import argparse
from pytube import YouTube


import youtube_dl
current_time1 = datetime.datetime.now()
current_time = current_time1.strftime('%m-%d-%Y_%H.%M')

from moviepy.editor import *



import speech_recognition as sr
import os
from pydub import AudioSegment
from pydub.silence import split_on_silence

# create a speech recognition object





def DownloadYoutubeMovie(video_url):
    VIDEO_SAVE_DIRECTORY = up.AI_Audio_Transcript + '\\' + "YoutubeVideos"
    Check_Folder_Exists(VIDEO_SAVE_DIRECTORY)




    try:
        try:
            video = YouTube(video_url,use_oauth=True,allow_oauth_cache=True)
            video = video.streams.get_highest_resolution()
            x = video.download(VIDEO_SAVE_DIRECTORY)
        except:
            yt = YouTube(video_url)
            yt = yt.streams.filter(progressive=True, file_extension='mp4').order_by('resolution').desc().first()
            if not os.path.exists(VIDEO_SAVE_DIRECTORY):
                os.makedirs(VIDEO_SAVE_DIRECTORY)
            x = yt.download(VIDEO_SAVE_DIRECTORY)

        print("video was downloaded successfully")
    except:
        print("Failed to download video")


    MovieSubtitles(x)



def AI_Task_Audio(Text):
    NewText = Basic_GPT_Query2(Line2_Role = SP.Task_Role, Line3_Format=SP.Task_Format, Line4_Task= "Use the following Text as your source Text   Source Text: " + Text, Line5_Task=SP.Task_Task, Big=True)

    return NewText

def ReWrite_Transcript(Text):
    NewText = Basic_GPT_Query2(Line2_Role = SP.ReWriteTranscript_Role, Line3_Format=SP.ReWriteTranscript_Format, Line4_Task= "Use the following Text as your source Text   Source Text: " + Text, Line5_Task=SP.ReWriteTranscript_Task, Big=True)

    return NewText

def MovieSubtitles(File, Origin="English", Output=["Spanish"], Rewrite=False, AI_Task = False):
    AudioFile = Movie2Audio(File)
    TranscribeAudio(AudioFile = AudioFile,File= File, Origin = Origin, Output = Output, Rewrite = Rewrite, AI_Task = AI_Task)

def TranscribeAudio(File,AudioFile, Origin = "English", Output = ["Spanish"], Rewrite = False,AI_Task = False):
    d = 1

    SavePath = up.AI_Audio_Transcript + '\\' + 'Movie2Audio'
    Check_Folder_Exists(SavePath)
    x = File.rfind('.')
    x2 = File.rfind('\\')
    FileName = File[x2:x]

    FileName = CleanFileName(FileName)

    SavePath = SavePath +'\\'+ FileName
    print(FileName)
    Check_Folder_Exists(SavePath)



    Transcript = transcribe_Large_audio(File,Origin = Origin)

    try:
        file_stats = os.stat(AudioFile)
        file_size = file_stats.st_size
        filesize2 = file_size/1000000
        print(f"File Size in MegaBytes is {filesize2}")
    except FileNotFoundError:
        print("File not found.")
    except OSError:
        print("OS error occurred.")

    Transcript2 = ''
    ReWrite_Transcript2 = ''
    ReWrite_Transcript1 = ''
    try:
        if filesize2 < 38:
            model = whisper.load_model("tiny")
        elif filesize2 < 74:
            model = whisper.load_model("base")
        elif filesize2 < 244:
            model = whisper.load_model("small")
        elif filesize2 < 768:
            model = whisper.load_model("medium")
        elif filesize2 >=768:
            model = whisper.load_model("large")



        Text= model.transcribe(AudioFile)
        Transcript2  = Text["text"]
        if Rewrite ==True:
            ReWrite_Transcript2 = ReWrite_Transcript(Transcript2)
            SaveCSV(Text=ReWrite_Transcript2, SavePath=SavePath, Title=FileName + '_Transcript_Revised_AI')
        SaveCSV(Text=Transcript2, SavePath=SavePath, Title=FileName + '_Transcript_AI')

    except:
        print("Error Using Whisper")

    if Rewrite == True:
        ReWrite_Transcript1 = ReWrite_Transcript(Transcript)
        SaveCSV(Text=ReWrite_Transcript1, SavePath=SavePath, Title=FileName + '_Transcript_Revised')

    SaveCSV(Text=Transcript, SavePath=SavePath, Title=FileName + '_Transcript')


    if ReWrite_Transcript2 == '' and Rewrite == True:
        Final_Text = ReWrite_Transcript2
    elif Transcript2 == '':
        Final_Text = Transcript2
    else:
        Final_Text = Transcript

    if AI_Task ==True:

        if Rewrite == True:
            Request = Final_Text

        elif Rewrite ==False:
            if Transcript2 == '':
                Request = Transcript
            else:
                Request = Transcript2

        try:

            Response = AI_Task_Audio(Request)
            SaveCSV(Text=Response, SavePath=SavePath, Title=FileName + '_Response')
        except:
            print ('Did not provide answer')

    SaveText2Audio(SavePath=SavePath, FileName=FileName,
                              Neural='Neural',
                              Mode='AUDIOBOOK', Chunk_Limit=1300, Translate=Output,
                              Text=Final_Text, Origin_Language= Origin)

    try:
        shutil.copyfile(AudioFile, SavePath + FileName + '.mp3')
    except:
        dn = 100
        print('Error moving audio file')


def Audio_Transcript(File):
    d = 1

def Movie2Audio(File):
    # Load the mp4 file
    video = VideoFileClip(File)

    x = File.rfind('.')

    nFile =  File[:x] + '.mp3'

    # Extract audio from video
    video.audio.write_audiofile(nFile)

    return nFile


# a function to recognize speech in the audio file
# so that we don't repeat ourselves in in other functions
def transcribe_audio(path, Origin = 'English'):
    # use the audio file as the audio source
    if Origin == "French":
        Origin_Recognize = "fr-FR"
    elif Origin =="Italian":
        Origin_Recognize = "it-IT"
    elif Origin =="Japanese":
        Origin_Recognize = "ja"

    else:
        Origin_Recognize = "en-US"

    r = sr.Recognizer()
    with sr.AudioFile(path) as source:
        audio_listened = r.record(source)
        # try converting it to text
        text = r.recognize_google(audio_listened, language=Origin_Recognize)
    return text

# a function that splits the audio file into chunks on silence
# and applies speech recognition
def transcribe_Large_audio(path, Origin = "English"):

    # open the audio file using pydub
    sound = AudioSegment.from_file(path)
    # split audio sound where silence is 500 miliseconds or more and get chunks
    chunks = split_on_silence(sound,
        # experiment with this value for your target audio file
        min_silence_len = 400,
        # adjust this per requirement
        silence_thresh = sound.dBFS-14,
        # keep the silence for 1 second, adjustable as well
        keep_silence=400,
    )

    # chunk_length_ms = 60000 # pydub calculates in millisec
    # chunks = make_chunks(sound, chunk_length_ms) #Make chunks of ten sec


    folder_name = "audio-chunks"
    # create a directory to store the audio chunks
    if not os.path.isdir(folder_name):
        os.mkdir(folder_name)
    whole_text = ""
    # process each chunk
    for i, audio_chunk in enumerate(chunks, start=1):
        # export audio chunk and save it in
        # the `folder_name` directory.
        chunk_filename = os.path.join(folder_name, f"chunk{i}.wav")
        audio_chunk.export(chunk_filename, format="wav")
        # recognize the chunk
        try:
            text = transcribe_audio(chunk_filename, Origin = Origin)
        except sr.UnknownValueError as e:
            print("Error:", str(e))
            print('Using Whisper')
            model = whisper.load_model("base")
            text2 = ''
            text1 = model.transcribe(chunk_filename)
            text = text1["text"]
            print(text)
        finally:
            if text != '':
                text = f"{text.capitalize()}. "
                print(chunk_filename, ":", text)
                whole_text += text
    # return the text for all chunks detected
    return whole_text


def Translator(Line1 = SAF.Translate_Sys,Origin_Language = 'English', Language_Final = 'French', Text = 'Sorry no text was sent but translate this instead, Monde Vert apologizes for the inconvenience', Line2 = SAF.Translate_Line2, Line3= SAF.Translate_Line3,Line4= SAF.Translate_Line4, SavePath = up.AI_AudioBook_Path, File_Name = 'Italian Translation'):

    SAF.Pick_Voice(Language = Language_Final)


    Line4 = Line4 + Text + "] From ["+ Origin_Language +"] to [" +Language_Final + "] " #+ Line4

    File_Name = File_Name + '_' +Origin_Language + ' to ' + Language_Final+'_Translation'
    Translated_Text = ''
    if Text != ' ' and Text != '  ' and Text != '':
        Translated_Text = Basic_GPT_Query2(Line2_Role= Line1 , Line3_Format=Line3, Line4_Task= Line2, Line5_Task= Line4)
    return Translated_Text

def PDF2JPG():
    # import module


    # Store Pdf with convert_from_path function
    images = convert_from_path('example.pdf')

    for i in range(len(images)):

    # Save pages as images in the pdf
        images[i].save('page' + str(i) + '.jpg', 'JPEG')


def Check_Folder_Exists(SavePath):
    try:
        isExist = os.path.exists(SavePath)
        if not isExist:
            # Create a new directory because it does not exist
            os.makedirs(SavePath)
            print("Created new filePath: " + SavePath)
    except:
        dn = True





def speak(self, audio, Add2T=True, voice=''):
    if voice == '':
        voice = self.voice

    if Add2T == True:
        Add2Transcript(self, self.AssistantName + ': ' + audio)
    engine = pyttsx3.init()
    # getter method(gets the current value
    # of engine property)
    voices = engine.getProperty('voices')
    engine.setProperty('voice', voices[voice].id)
    # Method for the speaking of the assistant
    if self.SilentMode == False:
        engine.say(audio + "")
        # Blocks while processing all the currently
        # queued commands
        engine.runAndWait()
    else:
        print('Silent Mode Activated: ' + audio)




def Add2Transcript(self, text2Add):
    transcript_Final = ''
    transcript_Final += text2Add + ' \n'
    current_time1 = datetime.datetime.now()
    current_time = current_time1.strftime('%m-%d-%Y_%H.%M.%S')

    data = [(current_time, text2Add)]
    print(data)
    try:
        df1 = pd.DataFrame(data, columns=["TimeStamp", "Transcript"])
        add2Master2(df1)
    except:
        print('Error could not add to transcript')

def add2Master2(df1):
     f2 = up.getMF2Path()
     df2 = pd.read_excel(f2)
     df3 = pd.concat([df1, df2])
     df3.iloc[:, 1:]
     df3.drop_duplicates()

     # creating a new excel file and save the data
     df3.to_excel(f2, index=False)





def CheckFileLength(SavePath,Title1,current_time_f, FileType):
    isExist = True
    Title2 = SavePath + '\\' + Title1 + current_time_f + FileType
    Add_Dupe_File = 0

    while len(Title2) > 250:
        if len(SavePath) >= 240 and len(Title2) > 250:
            Title2 = SavePath + '\\' + 'BU' + FileType

        elif len(Title2) > 250:
            TitleLen = len(Title2)
            TrimNum = round(TitleLen * .1) + 1
            Title2 = SavePath + '\\' + Title1[TrimNum:] + current_time_f + FileType

    while isExist == True:

        isExist = os.path.exists(Title2)
        if isExist == True:
            Add_Dupe_File += 1
            Title2 = Title2[:-4] + '_' + str(Add_Dupe_File) + FileType

    return Title2

def SaveCSV(Text, Title, SavePath, AddTimeStamp = True, FileType = '.txt'):


    if AddTimeStamp==True:
        Text1 = 'Time: ' + str(current_time) + '| Title:' + Title+'| Text:' + Text
        current_time_f = '_' + str(current_time)
    else:
        Text1 = Text
        current_time_f = ''


    Title1  = CleanFileName(Title)
    Title2 = SavePath + '\\' + Title1 + '_' + current_time_f + FileType
    if len(Title2) > 250 :
        Title2 = CheckFileLength(SavePath=SavePath,Title1=Title1, current_time_f=current_time_f,FileType=FileType)
    # folder = str(Title1)
    # if len(folder) > 44:
    #     folder = folder[0:44]

    Title1 =  str(Title1)
    Check_Folder_Exists(SavePath)
    Data1 = [(Text1)]
    try:
        df1 = pd.DataFrame(data=Data1)
        df1.to_csv(Title2)
        print(Title1 + ' - File Saved in the following location: ' + Title2)
    except:
        print('Review Error File did not save ')
    try:

        # MondeVert_IP.SaveText(self,df1,'MondeVert_IP Assistant', 'Full Transcript')
        add2Master2(df1)

    except:
        print('Review  did not save to master df2 file ')



    return Title2


# ********************************
"""Getting Started Example for Python 2.7+/3.3+"""

# Create a client using the credentials and region defined in the [adminuser]
# section of the AWS credentials file (~/.aws/credentials).

def makeArt( Prompt='', SavePath=up.AI_Art_Path, OpenFile=False):
    try:


        Mode = 'AI Art Mode'



        Promptc = 'True'

        if Prompt == '':
            Promptc = 'True'
        else:
            Promptc = 'False'

        if Prompt == '':
            print('Prompt is NULL: ' + Promptc)



        else:
            prompt = Prompt
            print(prompt)


        # Set up the OpenAI API client :
        openai.api_key = API_Key
        print("Sending to OpenAI...")
        print()

        try:
            response = openai.Image.create(
                prompt=prompt,
                n=1,
                size="1024x1024"
            )

            image_url = response['data'][0]['url']

            print(f"Image URL: {image_url}")
            print()

            # URL of the image to be downloaded is defined as image_url
            r = requests.get(image_url)  # create HTTP response object

            # send a HTTP request to the server and save
            # the HTTP response in a response object called r

            # Commented out for testing
            FileName = prompt
            print('Length of File Name: ' + str(len(FileName)))
            if len(FileName) >= 80:
                FileName = FileName[-80:]
                print('Length of File Name: ' + str(len(FileName)))

            FileName = CleanFileName(FileName)

            # fname = f"{''.join([c for c in FileName.strip().replace(' ', '_') if c.isalnum() or c == '_'])}.png"
            #
            #
            #
            # if os.path.isfile(os.path.join(SavePath, '\'', fname)):
            #     fname = fname.split(".")[0] + f".{''.join(random.choice(string.ascii_letters) for x in range(5))}.png"

            fname_only = FileName
            # fname = os.path.join(SavePath, '\'',FileName)

            fname = SavePath + '\\' + FileName + '.png'
            print(f"Filename: {fname}")
            with open(fname, 'wb') as f:

                f.write(r.content)

            try:
                if OpenFile == True:
                    if platform.system() == 'Darwin':  # macOS
                        subprocess.call(('open', fname))
                    elif platform.system() == 'Windows':  # Windows
                        os.startfile(fname)
                    else:  # linux variants
                        subprocess.call(('xdg-open', fname))
            except:
                dn = 100
            # this line saves the
            # IP.png2JPG(fname, up.PNGPath,fname_only, up.PNGPath_Archive, Del = False )
        except:
            fname = 'File not saved error occurred'

    except:
        fname = 'File not saved error occurred'

    try:

        Text1 = 'Original Prompt: ' + prompt + 'Path to Photo: ' + fname

        fnameText = fname_only
        SaveCSV(Text=Text1, Title=fnameText, SavePath=SavePath)
    except:
        dn = 100

        # data = ([Text1])
        # df1 = pd.DataFrame(data = data)

        # cu.add2Master2(self, Text1)
    return fname






def Chop4Art(Text, SavePath, FileName,FilePath = '',  Chunk_Limit = 1500,  Chunk_Replaces = ['.','?'], Chunk_Delimiter = '!',Artist_Persona = 'embrace the spirit and culture of the following text describe it to be illustrated by an artist', aPrompt = StoryMode.ArtPrompt_Story):
    Length_Text = len(Text)
    Chunk_Limit = Chunk_Limit -10
    Text_Chunks = []
    Translate_Folder = r'\Translated'
    Text_New = Text
    SavePath_Pics = SavePath + r'\SHAINE - Art'

    length = len(Text)
    Audio_File_Count= 1
    if 'Audio Chunks' in SavePath:
        SavePath_Chunks= SavePath
    else:
        SavePath_Chunks = SavePath + r'\SHAINE - Audio'
    Check_Folder_Exists(SavePath)
    Check_Folder_Exists(SavePath_Chunks)
    Check_Folder_Exists(SavePath_Pics)
    RunningCheck = 0
    LastPunc = 100
    #New_Text_Translated=[]

    Voices = []
    countl = 0

    x = GPT.MondeVert()
    TranslateLanguages = []
    FilePaths = []
    Audio_File_Count = 0
    New_Text_Translated = ''
    Text_New = Text



    countl = 0
    LastPunc = Chunk_Limit
    while  LastPunc != -1 and len(Text_New)> 0:
        countl += 1
        Text_Chunk = Text_New[:Chunk_Limit]
        Audio_File_Count += 1
        Text_Chunk_Info_only= Text_Chunk
        for i in Chunk_Replaces:
            Text_Chunk_Info_only = Text_Chunk_Info_only.replace(i, Chunk_Delimiter)

        LastPunc =Text_Chunk_Info_only.rfind(Chunk_Delimiter)


        print('Starting Chars:')
        print(len(Text))
        print('Chars Remaining:')
        print(len(Text_New))
        print('Last Punc Position(New):')
        print(LastPunc)

        #LastPunc = Text_Chunk_Info_only.rfind(Chunk_Delimiter)
        if LastPunc == -1 and len(Text_New) > Chunk_Limit:

            LastPunc = Text_Chunk_Info_only.rfind('\\n')
            print('Last Punc Position(using new Char)')
            print(LastPunc)
            if LastPunc == -1 and len(Text_New) > Chunk_Limit:
                LastPunc = Text_Chunk_Info_only.rfind(' ')
                if LastPunc == -1 and len(Text_New) > Chunk_Limit:

                    if Chunk_Limit < 1700:
                        Chunk_Limit_new = round(Chunk_Limit * (1.33))
                        if Chunk_Limit_new > 1700:
                            Chunk_Limit_new == 1700

                        print('New Chunk Limit')
                        print(Chunk_Limit_new)
                        Text2Send = Text_New[:Chunk_Limit_new]

                        try:
                            # length = len(Text_New)
                            New_pos = (Chunk_Limit_new + 1)
                            Text_New = Text_New[New_pos:]
                            # length = len(Text_New)
                        except:
                            print('Review Error Triggered, could be just how this has to work')

                    else:
                        Text_Chunk_Final = Text_New[:LastPunc]
                        New_pos = (LastPunc + 1)
                        Text_New = Text_New[New_pos:]

                elif LastPunc != -1 and len(Text_New) > Chunk_Limit:
                    Text_Chunk_Final = Text_New[:LastPunc]
                    try:
                        # length = len(Text_New)
                        New_pos = (LastPunc + 1)
                        Text_New = Text_New[New_pos:]
                        # length = len(Text_New)
                    except:
                        print('Review Error Triggered, could be just how this has to work')
                elif LastPunc == -1 and len(Text_New) <= Chunk_Limit:
                    Text_Chunk_Final = Text_New
                    Text_New = ''


        elif LastPunc > -1 and LastPunc <= Chunk_Limit:

            Text_Chunk_Final = Text_New[:LastPunc]
            try:
                # length = len(Text_New)
                New_pos = (LastPunc + 1)
                Text_New = Text_New[New_pos:]
                # length = len(Text_New)
            except:
                print('Review Error Triggered, could be just how this has to work')

        elif LastPunc == -1 and len(Text_New) <= Chunk_Limit:
            Text_Chunk_Final = Text_New
            Text_New = ''


        FileName_Chunk = FileName + '_Chunk_' + str(Audio_File_Count)

        if 1== 1:


            if len(Text_Chunk_Final) > 44:
                try:
                    ArtPrompt = x.GPTArt2( User_Subject='Create a prompt for an artist to create a work of art based on the following text: '  +Text_Chunk_Final, ArtFormat=aPrompt)
                except:
                    ArtPrompt = Artist_Persona + Text_Chunk_Final
                    ArtPrompt[:313]

                try:
                    PicPath = makeArt(Prompt='Using the following details: ' + Artist_Persona + " Art Prompt:  " + ArtPrompt)
                    print('successfully made a work of art in foreign language')
                    newPicPath = SavePath_Pics + '\\'+FileName_Chunk + '.png'
                    try:

                        shutil.copy(PicPath,newPicPath)
                    except:
                        dn = 100
                except:
                    print('Art not made, not an issue for now, maybe down the road')

        #print('Error Could not parse files together, see chunks for now')


    return FilePaths



def Combine_Splitsof_audio(AudioFiles_ordered,FilePath,FileName = 'SHAINE - Audio Combined', SavePath = up.SavePath + 'r/Orphan Audio Files',Opening_Sound = up.MondeVertIntro, Voice = random.choices(SAF.Original_List_of_Voices_English)[0]):
    Check_Folder_Exists(SavePath)
    Merged_Audio = Opening_Sound
    FilePath_merge = SavePath + '\\' + FileName + '.mp3'
    counter = 1
    len_Audio = len(AudioFiles_ordered)
    for x in range (0,len_Audio):
        if x ==0:
            #print(Opening_Sound)
            #Opening = AudioSegment.from_mp3(Opening_Sound)
            sound2add = AudioSegment.from_mp3(AudioFiles_ordered[x])
            Merged_Audio = sound2add

        else:
            sound2add = AudioSegment.from_mp3(AudioFiles_ordered[x])
            Merged_Audio += sound2add

    try:
    # simple export
         FilePath_new = Merged_Audio.export(FilePath_merge, format="mp3")
    except:
         print('could not save file audio file')
    #     FilePath = 'No File Saved'

    return FilePath_merge

#on character list Have Narrator Scene (stern and manly) and then you have preferred Narrators to choose from (to do all other narrations)
#Split on Space and :
#also look for Narrator, for now we should mostly ignore whats in the parenthesis so I will work to remove these
#Setting and other parts are Narrator1, and then narrator2
#def SplitScriptbyPart(Script,Character_Voice_Relationship,FileName,SavePath )
#remove (voiceover), otherwise have the system remove anywhere that has the parenthesis as its not important for audio
#(in person) (on the phone)
#Grab the Left most item on the line
#Def Clean_Text_For_Script

#Remove - FADE OUT.
#(Impressed) dip in the voice idk
#(defeated)/(Nervous) make quieter -2/(relieved) quiet but happy

#(Excited).(Urgently)/(upset)  - speed up maybe --> Make it louder

#other parenthesis should be spoken by one of the narrators or just the one

# not sure what to do maybe add a pause or something like that(smiles)
#Replace with onomonapea
#(sighs)
#(Laughs)


#Remove the reduandant parts like out loud and voiceover etc. eventually I can key off this but dont need it
#if line above is NULL or " " then you look for : if no : then ignore the line completely
#if the line says Music: send to GPT to come up with a quick prompt to relate it to the story/scene
#Make sure to add results to a script so we can use it again (Call it on the fly Revisions)
#if line says Narrator, remove parenthesis and just talk for the most part
#If its not a scenario above it is a character so we identify the name prior to the : and after the line before that is NULL or " "
#We eventually try to map the name to a voice and then we use the same voice for the character (for now we can live with random and maybe assign it and remove it from the list)
#We will eventually have the narrator say the lines that are in parenthesis, we can have the words split up or something


def Resume():
    Response = Basic_GPT_Query2( SavePath=up.AI_Task_Path + '\\' + "Resume", Line1_System_Rule=up.system_TextJoaT,
                            Line2_Role=lup.Test_Role_Resume,
                            Line5_Task="Use the following Resume for your edits:" + lup.Resume_Text, Line4_Task=lup.Test_Task_Resume + lup.Test_Special_Resume,
                             Line3_Format=lup.Test_Format_Resume,
                           SaveFile=True, MakeArt=True, Mode = 'Resume')

def Basic_GPT_Query2(   Line2_Role  , Line5_Task,Line3_Format,Line4_Task,Big = False,Model = "gpt-3.5-turbo",upgradeLimit = 3000, Special = '',Line1_System_Rule = StoryMode.system_TextJoaT_quick, crazy = .5, Subject= '', Outline = '', Allowed_Fails = 8, SaveFile = False,MakeArt = False, Mode = 'SHAINE SAYS', SavePath= ''):#use this to create art style for the work


    if Subject != '':
        Line2_Role = Line2_Role + """Your role and subject matter expertise should fit the following Subject and or style and mood in the {Text} provided by the user Text:###""" + Subject + """###"""



    Line1_System_Rule = Line2_Role
    #Line2_Role = Line2_Role + Special
    Full_User_Prompt = """User Inputs to Chat GPT: 
    1). """ + Line1_System_Rule +"""
    2).""" + Line3_Format+ """
    3).""" + Line4_Task + """
    4).""" + Line5_Task

    if len(Full_User_Prompt) > upgradeLimit:
        Model = "gpt-3.5-turbo-16k-0613"
    elif len(Full_User_Prompt) < upgradeLimit:
        Model = "gpt-3.5-turbo"



    if Big ==True:
        Model = "gpt-3.5-turbo-16k-0613"
    KeepGoing = False
    KillSwitch = 0
    while KeepGoing == False and KillSwitch < Allowed_Fails:


        try:



            # This is for the result if you let the AI describe project and details and then make the response
            response = openai.ChatCompletion.create(
                model=Model,
                messages=[
                    {"role": "system", "content": Line1_System_Rule},
                    {"role": "user", "content": Line3_Format},
                    {"role": "user", "content": Line4_Task },
                    {"role": "user", "content": Line5_Task},
                ]
                , temperature=crazy
            )
            GPT_Response = str(response.choices[0].message.content)

            KeepGoing = True
        except:
            print(' Error ChatGPT failed, trying to rerun prompt now.... if this happens too many times we will kill the script')
            KillSwitch += 1
            print(Full_User_Prompt)
            #as written this does not run but if I got rid of +1 it could
            if KillSwitch == Allowed_Fails+1:
                print('could not create a writer persona, redoing it now')
                GPT_Response = Basic_GPT_Query2( Line2_Role='You are a skilled writer', Line3_Format=Line3_Format, Line4_Task=Line4_Task, Line1_System_Rule=up.system_Text_ScreenPlay)

            continue


        #print(up.breakupOutput)

        # self.UserPromptsCount+=1
        #
        # self.UserPrompts += 'User Input #' + str(UserPromptsCount)
        #
        # self.UserPrompts += Full_User_Prompt + up.breakupOutput2 + up.breakupOutput2


        Title = Mode + '_' + current_time
        if SaveFile ==True:
            SaveText = current_time+ up.breakupOutput2 + Full_User_Prompt + up.breakupOutput2 + 'SHAINE SAYS: '+  GPT_Response
            print(SaveText)
            SaveCSV(Text=SaveText, SavePath=SavePath, Title = Title)
        if MakeArt ==True:
            ArtPrompt = GPTArt2(User_Subject = GPT_Response)
            print(ArtPrompt)
            originalFilepath = makeArt(Prompt=ArtPrompt)
            PicNewPath1 = SavePath + '\\' + Mode
            Check_Folder_Exists(PicNewPath1)
            PicNewPath =PicNewPath1 + '\\' + Title + '.png'
            shutil.copyfile(originalFilepath, PicNewPath)
        return GPT_Response


def GPTArt2( crazy=.5,  Model = "gpt-3.5-turbo",prompt=sms.ArtPrompt_Clean_Social_Media_Post_Line2_Prompt,
             User_Subject='Pick a random subject and medium go wild and make it exciting, beautiful and shocking',
             ArtFormat=sms.ArtPrompt_Clean_Social_Media_Post_Line3,
             sys_prompt=sms.ArtPrompt_Sys):

    Full_User_Prompt = sys_prompt + prompt + User_Subject + ArtFormat
    if len(Full_User_Prompt) > 6000:
        Model = "gpt-3.5-turbo-16k-0613"
    elif len(Full_User_Prompt) < 6000:
        Model = "gpt-3.5-turbo"

    keepgoing = True
    GPTARTPROMPT = User_Subject
    while keepgoing == True:
          try:

              Art_Prompt1 = openai.ChatCompletion.create(
          model=Model,
          messages = [
                         {"role": "system", "content": sys_prompt},
                         {"role": "user", "content": prompt + User_Subject},
                         {"role": "user", "content": ArtFormat}
                     ], temperature = crazy
              )

              GPTARTPROMPT = Art_Prompt1.choices[0].message.content
              keepgoing == False

          except:
              print('GPT error waiting 13 seconds and trying again...')

              try:
                  time.sleep(13)
              except:
                  dn = 100

          return GPTARTPROMPT

def Split_Audio2(Text, SavePath, FileName,FilePath = '', OpeningSound = up.MondeVertIntro, Chunk_Limit = 1500, Voice = random.choices(SAF.Original_List_of_Voices_English)[0], Translate =["English"], Chunk_Replaces = ['.','?'], Chunk_Delimiter = '!',Artist_Persona = 'embrace the spirit and culture of the following text describe it to be illustrated by an artist', aPrompt = StoryMode.ArtPrompt_Story,Origin_Language = 'English'):
    Length_Text = len(Text)
    Chunk_Limit = Chunk_Limit -10
    Text_Chunks = []
    Translate_Folder = r'\Translated'
    Text_New = Text
    SavePath_Pics = SavePath + r'\SHAINE - Art'

    length = len(Text)
    Audio_File_Count= 1
    if 'Audio Chunks' in SavePath:
        SavePath_Chunks= SavePath
    else:
        SavePath_Chunks = SavePath + r'\SHAINE - Audio'
    Check_Folder_Exists(SavePath)
    Check_Folder_Exists(SavePath_Chunks)
    Check_Folder_Exists(SavePath_Pics)
    RunningCheck = 0
    LastPunc = 100
    #New_Text_Translated=[]

    Voices = []
    countl = 0
    # print('All Languages to translate')
    # print(Translate)
    x = GPT.MondeVert()
    TranslateLanguages = []
    for l in Translate:

        FilePaths = []
        Audio_File_Count = 0
        New_Text_Translated = ''
        Text_New = Text
        if l == Origin_Language:
            v = Voice
        else:
            v = SAF.Pick_Voice(Language=l)

        Voices.append(v)
        #New_Text_Translated.append('Voice: ' + str(Voices[countl]) + '   ' + str(l) + '_Translation: ')
        FileName_Language_Voice = FileName + '_' + l + '_' + v
        countl = 0
        LastPunc = Chunk_Limit
        while  LastPunc != -1 and len(Text_New)> 0:
            countl += 1
            Text_Chunk = Text_New[:Chunk_Limit]
            Audio_File_Count += 1
            Text_Chunk_Info_only= Text_Chunk
            for i in Chunk_Replaces:
                Text_Chunk_Info_only = Text_Chunk_Info_only.replace(i, Chunk_Delimiter)

            LastPunc =Text_Chunk_Info_only.rfind(Chunk_Delimiter)


            #
            # print(l + 'Voice: ' + v)
            # print('Starting Chars:')
            # print(len(Text))
            # print('Chars Remaining:')
            # print(len(Text_New))
            # print('Last Punc Position(New):')
            # print(LastPunc)

            #LastPunc = Text_Chunk_Info_only.rfind(Chunk_Delimiter)
            if LastPunc == -1 and len(Text_New) > Chunk_Limit:

                LastPunc = Text_Chunk_Info_only.rfind('\\n')
                # print('Last Punc Position(using new Char)')
                # print(LastPunc)
                if LastPunc == -1 and len(Text_New) > Chunk_Limit:
                    LastPunc = Text_Chunk_Info_only.rfind(' ')
                    if LastPunc == -1 and len(Text_New) > Chunk_Limit:

                        if Chunk_Limit < 1700:
                            Chunk_Limit_new = round(Chunk_Limit * (1.33))
                            if Chunk_Limit_new > 1700:
                                Chunk_Limit_new == 1700
                            #
                            # print('New Chunk Limit')
                            # print(Chunk_Limit_new)
                            Text2Send = Text_New[:Chunk_Limit_new]

                            try:
                                # length = len(Text_New)
                                New_pos = (Chunk_Limit_new + 1)
                                Text_New = Text_New[New_pos:]
                                # length = len(Text_New)
                            except:
                                print('Review Error Triggered, could be just how this has to work')
                            try:
                                Split_Audio2(Text=Text2Send, Translate=[l], SavePath=SavePath,
                                             FileName=FileName + '_' + str(Audio_File_Count),
                                             Chunk_Limit=Chunk_Limit_new, Chunk_Delimiter=Chunk_Delimiter,
                                             Artist_Persona=Artist_Persona, Voice=v, aPrompt = aPrompt, Origin_Language = Origin_Language)

                                New_Text_Translated += '****Review, there is missing text that was captured elsewhere'
                            except:
                                print('Error trying to split files with new Chunk Limit')
                                continue
                        else:
                            Text_Chunk_Final = Text_New[:LastPunc]
                            New_pos = (LastPunc + 1)
                            Text_New = Text_New[New_pos:]

                    elif LastPunc != -1 and len(Text_New) > Chunk_Limit:
                        Text_Chunk_Final = Text_New[:LastPunc]
                        try:
                            # length = len(Text_New)
                            New_pos = (LastPunc + 1)
                            Text_New = Text_New[New_pos:]
                            # length = len(Text_New)
                        except:
                            print('Review Error Triggered, could be just how this has to work')
                    elif LastPunc == -1 and len(Text_New) <= Chunk_Limit:
                        Text_Chunk_Final = Text_New
                        Text_New = ''


            elif LastPunc > -1 and LastPunc <= Chunk_Limit:

                Text_Chunk_Final = Text_New[:LastPunc]
                try:
                    # length = len(Text_New)
                    New_pos = (LastPunc + 1)
                    Text_New = Text_New[New_pos:]
                    # length = len(Text_New)
                except:
                    print('Review Error Triggered, could be just how this has to work')

            elif LastPunc == -1 and len(Text_New) <= Chunk_Limit:
                Text_Chunk_Final = Text_New
                Text_New = ''

            FileName_Language_Voice = FileName + '_' + l + '_' + v
            FileName_Chunk = FileName + '_' + l + '_' + v + '_Chunk_' + str(Audio_File_Count)
            # print(FileName_Chunk)
            # print(SavePath_Pics)
            # print(l)
            # print(countl)
            # print(Voices[countl])
            # print(New_Text_Translated[countl])
            New_Text_Translated_Temp = ''

            #print(Text_Chunk_Final)

            if l != Origin_Language:
                d= 100
                tempTranslate = Translator(Text = Text_Chunk_Final, Language_Final=l, Origin_Language=Origin_Language)
                Translate_Voice = v
                New_Text_Translated+= tempTranslate
                translate2 = [l]
                try:

                    Chunk_Limit_Translate = len(tempTranslate) + 10
                    FilePathc = SaveText2Audio(Text=tempTranslate, SavePath=SavePath_Chunks+Translate_Folder,
                                          FileName=FileName_Chunk,
                                          Voice=Translate_Voice, Chunk_Mode=True, FilePath=FilePath, Translate=Translate, Chunk_Limit=Chunk_Limit_Translate, Artist_Persona=Artist_Persona,  aPrompt = aPrompt, Origin_Language = Origin_Language)
                    FilePaths.append(FilePathc)
                    # print('successfully made an audio file in ' + l)
                except:
                    print('Error could not create audio for this text')


                if len(tempTranslate) > 44:
                    try:
                        ArtPrompt = x.GPTArt2( User_Subject='Create a prompt for an artist to create a work of art based on the following text (Use the original language for inspiration/context/Location of artwork): '  +tempTranslate, ArtFormat=StoryMode.StoryArt)
                    except:
                        ArtPrompt = tempTranslate
                        ArtPrompt[:313]

                    try:
                        PicPath = x.makeArt(Prompt='Using the following details: ' + Artist_Persona + " Art Prompt: " + ArtPrompt)
                            #x.makeArt(Prompt="Using the following Art Sytles/Details from the text create a work of art pased on the following Art Prompt:  " + ArtPrompt)
                        # print('successfully made a work of art in foreign language')
                        newPicPath = SavePath_Pics + '\\'+FileName_Chunk + '.png'
                        try:

                            shutil.copy(PicPath,newPicPath)
                        except:
                            dn = 100
                    except:
                        print('Art not made, not an issue for now, maybe down the road')



            else:
                # print('Text_Chunk_Final: ' + str(len(Text_Chunk_Final)))
                translate2 = [l]
                FilePathc = SaveText2Audio(Text = Text_Chunk_Final, SavePath = SavePath_Chunks, FileName=FileName +'_' + l +'_' + Voice +'_Chunk_' + str(Audio_File_Count), Voice = v , Chunk_Mode=True, FilePath=FilePath, Translate=[l], Chunk_Delimiter=Chunk_Delimiter,Chunk_Replaces=Chunk_Replaces,Chunk_Limit=Chunk_Limit + 10 ,Artist_Persona=Artist_Persona,  aPrompt = aPrompt, Origin_Language = Origin_Language)
                FilePaths.append(FilePathc)
                if len(Text_Chunk_Final)>44:
                    ArtPrompt = x.GPTArt2(User_Subject='Create a prompt for an artist to create a work of art based on the following text: ' + Text_Chunk_Final, ArtFormat=aPrompt)
                    try:
                        PicPath = x.makeArt(Prompt='Using the following details: ' + Artist_Persona + " Art Prompt: " + ArtPrompt)
                        newPicPath = SavePath_Pics + '\\' +FileName_Chunk + '.png'
                        try:
                            shutil.copy(PicPath, newPicPath)
                        except:
                            dn = 100
                    except:
                        print('Art not made, not an issue for now, maybe down the road')



            Text_Chunks.append(Text_Chunk_Final)
        #RunningCheck +=len(Text_Chunk_Final)
        if l != Origin_Language:
            SaveCSV(SavePath=SavePath+Translate_Folder, Title=FileName +'_' +  l + '_Translated', Text=New_Text_Translated)
            SavePath_temp = SavePath+Translate_Folder
        else:
            SavePath_temp = SavePath

        #FilePath2 = FilePaths[0]

            #print(Text_New)
            #print(len(Text_New))
        try:

            FilePath2 = Combine_Splitsof_audio(AudioFiles_ordered = FilePaths, FilePath=FilePath,SavePath= SavePath_temp, FileName= FileName_Language_Voice, Voice = v)
        except:
            FilePath2 = FilePaths[0]

            #print('Error Could not parse files together, see chunks for now')
    return FilePath2

def CleanLyrics4audio( text,Chunk_Delimiter_right = ':', Chunk_Delimiter_left= '\\n', RemoveChars = [':', 'Chorus', 'Verse', 'Bridge', 'Sample']):
    findChars = True
    Text_Final = ''
    Text_Pre = text

    if Chunk_Delimiter_right  in text:
        while findChars ==True:
            rPoint = Text_Pre.rfind(Chunk_Delimiter_right)
            Text_Pre_L = Text_Pre[:rPoint]
            lPoint = Text_Pre_L.rfind(Chunk_Delimiter_left)
            if rPoint == -1:
                findChars= False
                continue
            elif lPoint>=rPoint:
                #Text_Final = Text_Pre[rPoint+1:] + lup.NewLine  + Text_Final
                print('how did this happen - review?')
                #Text_Pre = Text_Pre[:(lPoint)]
                continue
            elif lPoint < rPoint:
                Text_Final= Text_Pre[rPoint+1:] + lup.NewLine + Text_Final
                Text_Pre = Text_Pre[:lPoint]
                continue
        for r in RemoveChars:
            Text_Final = Text_Final.replace(r,'')

    else:
        Text_Final= ''

    if Text_Final =='':
        Text_Final = text
    return Text_Final

def CleanFileName(Text):
    Text = re.sub(r"[^a-zA-Z0-9 ]", "", Text)
    return Text
def SaveText2Audio( MultiThread = True,Translate = ["English"],Text = '', SavePath =up.SavePath, FileName="Text2Audio", FilePath = '', Neural='Neural', Mode='Text2Audio', Chunk_Limit = 1500,Voice = random.choices(SAF.Original_List_of_Voices_English)[0], Chunk_Replaces = ['.','?','\n'], Chunk_Delimiter = '!',Artist_Persona = '', Chunk_Mode = False,  aPrompt = sms.ArtPrompt_Clean_Social_Media_Post_Line2_Prompt, Origin_Language = 'English'):
    #Voice = CleanFileName(Voice)

    Check_Folder_Exists(SavePath)




    if Chunk_Mode == True:
        print('Save text to audio for : '+ FileName + ', should be saving to : ' + SavePath)
        FilePath_m = SavePath + '\\' + FileName + '.mp3'
    else:
        if Text == '' and os.path.exists(FilePath) and FilePath != '':

            FilePath_m = FilePath[:-4] + '.mp3'
            SavePath_index = FilePath.rfind('\\')
            FStart = (SavePath_index+1)
            FileName = FilePath[FStart:-4]
            SavePath = FilePath[:SavePath_index]
            # print('SavePath: ' + SavePath)
            # print('FileName: ' + FileName)

            with open(FilePath) as f:
                Text = f.read()

        else:
            FilePath_m = SavePath + '\\' + FileName + '.mp3'

    if Artist_Persona == 'embrace the spirit/symbolism and culture of the following text describe it to be illustrated by an artist' or Artist_Persona == '':
        arttext = Text
        try:
            arttext = Text[:1700]
        except:
            dn = 100
        try:
            x = GPT.MondeVert()
            Artist_Persona = x.GPTArt2(User_Subject=arttext,ArtFormat = lup.artDetailsFormat, prompt=lup.artDetailsPrompt)
            # print(Artist_Persona)
        except:
            Artist_Persona = 'embrace the spirit/symbolism and culture of the following text describe it to be illustrated by an artist'

    #polly = Session(region_name='eu-west-2').client('polly')

    # print(Translate)

    if MultiThread == True and len(Translate)>1:
        for l in Translate:
            ll = [l]
            t = threading.Thread(target=SaveText2Audio, args=(MultiThread,ll,Text,SavePath,FileName,FilePath,Neural,Mode,Chunk_Limit,Voice,Chunk_Replaces,Chunk_Delimiter,Artist_Persona,Chunk_Mode,aPrompt, Origin_Language)).start()

            #FilePath = '', Neural='Neural', Mode='Text2Audio', Chunk_Limit = 1500,Voice = random.choices(SAF.Original_List_of_Voices_English)[0], Chunk_Replaces = ['.','?','\n'], Chunk_Delimiter = '!',Artist_Persona = '', Chunk_Mode = False
            time.sleep(23)
            FilePathNew = 'Used Multithreading so multiple files created'
    else:

        if len(str(Text)) >= Chunk_Limit or (( len(Translate) > 1 or Translate[0] != 'English') and  Chunk_Mode == False):
            #try:
            FilePathNew = Split_Audio2(Text = Text,SavePath = SavePath, FileName=FileName, FilePath=FilePath_m, Chunk_Limit= Chunk_Limit, Voice = Voice, Chunk_Replaces = Chunk_Replaces, Chunk_Delimiter =Chunk_Delimiter, Artist_Persona = Artist_Persona, Translate=Translate, aPrompt=aPrompt,Origin_Language=Origin_Language)
            #except:
                #print('Error trying to create audio file try again later')

        else:
            #FilePaths = ''

            # print(FilePath)
            # print(len(str(Text)))
            if Voice not in FileName:
                FileName = FileName + '_' + Voice
                FilePath1 = 'No File Saved'

            try:
                session = Session(aws_access_key_id=DNC.aws_access_key_id, aws_secret_access_key=DNC.aws_secret_access_key,
                                  region_name='us-west-2')
                polly = session.client("polly")
                response = polly.synthesize_speech(Text=Text, OutputFormat="mp3",
                                                   VoiceId=Voice, Engine='neural', )

                # # Request speech synthesis #StartSpeechSynthesisTask #synthesize_speech
                # response = polly.start_speech_synthesis_task(Text=Text, OutputFormat="mp3",
                #                                    VoiceId=Voice, Engine='neural',
                #                                 OutputS3KeyPrefix= "InputAudio",
                #                                 OutputS3BucketName='s3bucketname',)
            except (BotoCoreError, ClientError) as error:
                # The service returned an error, exit gracefully
                print(error)
                sys.exit(-1)

            # Access the audio stream from the response
            if "AudioStream" in response:
                # Note: Closing the stream is important because the service throttles on the
                # number of parallel connections. Here we are using contextlib.closing to
                # ensure the close method of the stream object will be called automatically
                # at the end of the with statement's scope.
                with closing(response["AudioStream"]) as stream:
                    output = os.path.join(gettempdir(), FilePath_m)


                    try:
                        # Open a file for writing the output as a binary stream
                        with open(output, "wb") as file:
                            file.write(stream.read())
                            #FilePath = FilePath
                        FilePathNew = FilePath_m
                    except IOError as error:
                        # Could not write to file, exit gracefully
                        print(error)
                        sys.exit(-1)

            else:
                # The response didn't contain audio data, exit gracefully
                print("Could not stream audio as file is large, trying alternate ways")
               # Monitor_Audio_File(FileName, FilePath)
                sys.exit(-1)
            #
            # Play the audio using the platform's default player
            if sys.platform == "win32":
                os.startfile(output)
            # else:
            #     # The following works on macOS and Linux. (Darwin = mac, xdg-open = linux).
            #     opener = "open" if sys.platform == "darwin" else "xdg-open"
            #     subprocess.call([opener, output])
            # #


    return FilePathNew


#
#
# import boto3
# def Monitor_Audio_File(FileName,FilePath):
#
#     session = Session(aws_access_key_id=DNC.aws_access_key_id, aws_secret_access_key=DNC.aws_secret_access_key,
#                       region_name='us-west-2')
#     polly = session.client("polly")
#     s3 = boto3.client('s3', aws_access_key_id=DNC.aws_access_key_id, aws_secret_access_key=DNC.aws_secret_access_key)
#     #somhow download whatever this is lol
#     while FileDone ==False:
#         polly.get_speech_synthesis_task()
#         try:
#             s3.get_object(
#                 #Bucket=self._bucket,
#                 Key=FileName,
#             )
#             FileDone =  True
#         except s3.exceptions.NoSuchKey:
#             FileDone =  False
#             print('File not there, waiting 15 seconds')
#             time.sleep(15)
#     polly.download_file('your_bucket',FileName,FilePath)





def add2Master_Persona(df1 = '', Text = 'No Data Provided'):
    try:
        if df1 != '':
            df1 = df1
        else:
            data = ([Text])
            df1 = pd.DataFrame(data)

        f2 =  up.MasterFilePersona
        df2 = pd.read_excel(f2)
        df3 = pd.concat([df1, df2])
        df3.drop_duplicates()
        df3.iloc[:, 1:]
        # creating a new excel file and save the data
        df3.to_excel(f2, index=False)
    except:
        dn = 1000



def add2Master_Character(df1 = '', Text = 'No Data Provided'):

    if df1 !='':
        df1 = df1
    else:
        data = ([Text])
        df1 = pd.DataFrame(data)

    f2 =  up.MasterFile_Character
    df2 = pd.read_excel(f2)
    df3 = pd.concat([df1, df2])
    df3.drop_duplicates()
    df3.iloc[:, 1:]
    # creating a new excel file and save the data
    df3.to_excel(f2, index=False)





def add2Master(df1):
    f2 = up.getMFPath()
    df2 = pd.read_excel(f2)
    df3 = pd.concat([df1, df2])
    df3.drop_duplicates()
    df3.iloc[:, 1:]
    # creating a new excel file and save the data
    df3.to_excel(f2, index=False)

def add2Master2(df1):
    f2 = up.getMF2Path()
    df2 = pd.read_excel(f2)
    df3 = pd.concat([df1, df2])
    df3.iloc[:, 1:]
    df3.drop_duplicates()

    # creating a new excel file and save the data
    df3.to_excel(f2, index=False)

def add2Master3(df1):
    f2 = up.MasterFile3
    df2 = pd.read_excel(f2)
    df3 = pd.concat([df1, df2])
    df3.iloc[:, 1:]
    df3.drop_duplicates()

    # creating a new excel file and save the data
    df3.to_excel(f2, index=False)



def SaveText( df, FileName, tabname):

    text1 = df
    f2 = up.getPath()
    SavePath1 = f2
    invalidCharRemoved = re.sub(r"[^a-zA-Z0-9 ]", "", FileName)
    Filename = '\\' + str(invalidCharRemoved) + "_"
    SavePath2 = SavePath1 + Filename + ".xlsx"
    # SavePath2 = r'D:\ShakeBot Testing\ShaKeBotTest for DA - 12-20-2022.xlsx'
    try:
        with pd.ExcelWriter(SavePath2) as writer:
            text1.to_excel(writer, sheet_name=str(tabname))
            #MondeVert_IP.speak(self, 'File Saved')
            print('File Saved')
    except:
        Exception




def saveTranscript(Transcript,current_time):


    Filename = 'SHAINE Transcript v1'
    Filename = '\\' + Filename + "_"
    f2 = up.getPath()
    SavePath1 = f2
    data = [(current_time, Transcript)]
    print(data)
    try:
        df1 = pd.DataFrame(data, columns=["TimeStamp", "Transcript"])
        # df1 = pd.DataFrame(data)
        df1.to_csv(SavePath1 + Filename + current_time + '.csv')

        # MondeVert_IP.SaveText(self,df1,'MondeVert_IP Assistant', 'Full Transcript')
        add2Master2(df1)
    except:
        print('Review Error')


def Add2MasterLyrics(self, current_time2, Mode, title, ArtistPoetInfo, poem, Tag, SavePath2, Prompts):
    data = [(current_time2, Mode, title, ArtistPoetInfo, poem, Tag, SavePath2, Prompts)]
    df = pd.DataFrame(data, columns=['Date_Time_Added', 'Art_Type', 'Title', 'Poet_Artist_Info', 'Poem_Song_Lyrics',
                                     'Quality', 'Folder_Path', 'Prompts_Used'])
    add2Master3(df)






def NamePoemSavePoem(self, poem, ArtPaths, Prompts_Used, ArtistPoetInfo, title='', SavePath=up.AI_Poetry_Path,
                     Mode='SHAINE', current_time = current_time ):


    try:
        dfPrompts = ''
        Tag = ''
        if Mode == 'Poem':
            title_p = up.Poem_Title_Prompt
        else:
            title_p = up.Song_Title_Prompt
        try:
            if title == '':
                x = GPT.MondeVert()
                Title = x.Get_Title_GPT(  Text= poem[:1500] )
            else:
                Title = title

        except:
            print('error could not get title')

        if Title == '':
            Title = 'AI - No Name Project'

        Title1 = CleanFileName(Title)

        Title1 = '\\' + str(Title1) + "_"
        if len(Title1) > 60:
            Title1 = Title1[0:60]


        Title2 = SavePath + Title1

        Check_Folder_Exists(SavePath)





        data = Title +  ArtistPoetInfo + poem +  SavePath+  Mode +  Tag
        # print(dat

        Text = 'Time: ' + current_time + '| Title:' + Title + '| Text:' + data

        Data1 = [(Text)]

        try:
            df1 = pd.DataFrame(Text, columns=["TEXT"])

            SaveCSV(Text = Text, Title= Title , SavePath=SavePath)
            # print(df1)
            # print(Title2)
            # df1.to_csv(Title2 + '.txt')
            Add2MasterLyrics(self, current_time, Mode, title, ArtistPoetInfo, poem, Tag, SavePath,dfPrompts)
            # MondeVert_IP.SaveText(self,df1,'MondeVert_IP Assistant', 'Full Transcript')
            try:
                #add2Master2(df1)
                dn = True
            except:
                print('Did not add to master, error')
        except:
            print('Review Error')

            try:
    #Word Doc Creation below, not quite what I need so I do text file too above
                document = Document()
                document.add_heading(Title, 0)
                document.add_heading(ArtistPoetInfo, 1)
                p = document.add_paragraph()
                r = p.add_run()
                # r.add_text(ArtistPoetInfo)
                p = document.add_paragraph()
                r = p.add_run()
                r.add_text(poem)
                p = document.add_paragraph()
                r = p.add_run()
                image_counter = 0
                for i in ArtPaths:
                    image_counter += 1
                    r.add_picture(i)
                    original = i
                    target = SavePath + Title1 + str(image_counter)  + '.png'


                    try:
                        shutil.copyfile(original, target)
                    except:
                        dn = 100
                    try:
                        IP.png2JPG(original, up.PNGPath, Title1, up.PNGPath_Archive, Del=False)

                    except:
                        print('instagram failed, as always!')
                document.add_heading('AI Prompts used: ', 7)

                for i in Prompts_Used:
                    p = document.add_paragraph()
                    r = p.add_run()
                    r.add_text(i)
                    dfPrompts += i + '|'
                document.save(Title2 + '_Details.docx')

                document2 = Document()
                document2.add_heading(Title, 0)
                document2.add_heading(ArtistPoetInfo, 1)

                p = document2.add_paragraph()
                r = p.add_run()
                r.add_text(poem)
                document2.save(Title2 + '.docx')
            except:
                print('Error saving word doc')
    except:
        print('Error did not save file')

def MultiThread(self, Functions, Args):
    thread_list = []
    print("Start")
    print(Args)

    t = threading.Thread(target=Functions, args=(Args))
    thread_list.append(t)

    # Starts threads
    for thread in thread_list:
        print("New Thread Started")
        thread.start()


    # This blocks the calling thread until the thread whose join() method is called is terminated.
    # From http://docs.python.org/2/library/threading.html#thread-objects
    for thread in thread_list:
        thread.join()

# Demonstrates that the main process waited for threads to complete
    print("Done")



def send_email_no_attachment_outlook( password = DNC.Outlookpassword, sender = DNC.OutlookEmail, body= DNC.subject + current_time, filename = [up.MasterFilePersona], to= DNC.to, subject= DNC.subject + current_time, smtpHost = 'smtp.office365.com', smtpPort = 587):
    try:

    #smtp

        # Add the From: and To: headers at the start!
        message = ("From: %s\r\nTo: %s\r\nSubject: %s\r\n\r\n"
               % (sender, ", ".join(to), subject))
        message += """Sent by SHAINE - MondeVert"""
        print (message)
        try:
            smtpObj = smtplib.SMTP(smtpHost, smtpPort)
            #smtpObj.set_debuglevel(1)
            smtpObj.ehlo()
            smtpObj.starttls()
            smtpObj.ehlo()
            smtpObj.login(sender,password)
            smtpObj.sendmail(sender, to, message)
            smtpObj.quit()
            print ("Successfully sent email")
        except:
            print ("Error: unable to send email")
    except:
        print("Error: unable to send email")

def send_email_w_attachment_gmail(sender = DNC.gmail_user, password = DNC.gmail_pass, body= DNC.subject + current_time, filename = [up.MasterFilePersona], to= DNC.to, subject= DNC.subject + current_time,fType = "txt"):
    # create message object
    #to = 'sdonovan@mondevert.co'#; RichardDX44@gmail.com'
    try:

        message = MIMEMultipart()
        # add in header
        message['From'] = Header(sender)
        message['To'] = Header(to)
        message['Subject'] = Header(subject)
        # attach message body as MIMEText
        message.attach(MIMEText(body, 'plain', 'utf-8'))

        # locate and attach desired attachments

        for n in filename:
            att_name = os.path.basename(filename[n])
            _f = open(filename[n], 'rb')
            att = MIMEApplication(_f.read(), _subtype=fType)
            _f.close()
            att.add_header('Content-Disposition', 'attachment', filename=att_name)
            message.attach(att)

        # setup email server
        server = smtplib.SMTP_SSL(DNC.host, DNC.port)
        server.login(sender, password)

        # send email and quit server
        server.sendmail(sender, to, message.as_string())
        server.quit()
    except:
        print("Error: unable to send email")


def send_email_no_attachment_gmail(sender = DNC.gmail_user, password = DNC.gmail_pass, body= DNC.subject + current_time, filename = [up.MasterFilePersona], to= DNC.to, subject= DNC.subject + current_time,fType = "txt"):
    # create message object
    #to = 'sdonovan@mondevert.co'#; RichardDX44@gmail.com'
    try:

        message = MIMEMultipart()
        # add in header
        message['From'] = Header(sender)
        message['To'] = Header(to)
        message['Subject'] = Header(subject)

        # attach message body as MIMEText
        message.attach(MIMEText(body, 'plain', 'utf-8'))

        # setup email server
        server = smtplib.SMTP_SSL(DNC.host, DNC.port)
        server.login(sender, password)

        # send email and quit server
        server.sendmail(sender, to, message.as_string())
        server.quit()
    except:
        print("Error: unable to send email")

# !/usr/bin/python
def send_email_w_attachment_outlook( password = DNC.Outlookpassword, sender = DNC.OutlookEmail, body= DNC.subject + current_time, filename = [up.MasterFilePersona], to= DNC.to, subject= DNC.subject + current_time, server_host = 'smtp.office365.com', server_port = 587,contype = 'application/octet-stream'):

    # set sender email and password

    try:
        # set receivers
        receivers = to
        # set attachment file

        # set outlook smtp server host and port

        # create email text content
        # create MIMEMultipart object
        main_msg = email.mime.multipart.MIMEMultipart()
        # create a MIMEText object, it is the text content of email
        text_msg = email.mime.text.MIMEText(body)
        # add MIMEText object to MIMEMultipart object
        main_msg.attach(text_msg)
        # create MIMEBase object

        maintype, subtype = contype.split('/', 1)
        # read attachment content
        for f in filename:
            data = open(f, 'rb')
            file_msg = email.mime.base.MIMEBase(maintype, subtype)
            file_msg.set_payload(data.read())
            data.close()
            # file_msg is content of attachment
            email.encoders.encode_base64(file_msg)
            # attachment header
            basename = os.path.basename(f)
            file_msg.add_header('Content-Disposition',
                                'attachment', filename=basename)
        # add attachment to MIMEMultipart object
        main_msg.attach(file_msg)
        # set email format
        main_msg['From'] = sender
        main_msg['To'] = ", ".join(receivers)
        main_msg['Subject'] = "This attachment sent from SHAINE - MondeVert"
        main_msg['Date'] = email.utils.formatdate()
        # full content of email
        fullText = main_msg.as_string()
        # send email by outlook smtp
        server = smtplib.SMTP(server_host, server_port)
        try:
            server.ehlo()
            server.starttls()
            server.ehlo()
            server.login(sender, password)

            server.sendmail(sender, receivers, fullText)
            print("Successfully sent email")
        except:
            print("Error: unable to send email")
        finally:
            server.quit()

    except:
        print("Error: unable to send email")
#
# import SMTPException
# import smtplib.SMTPException

def CutUP_long_String_to_list(self, Text, SavePath, FileName, Art_Style_details,Key_Words = ['Characters','Character', 'Page', 'Illustration Details'], Key_Char = ':',Delimiter = '|', crazy = .5):
    dummy =1
    ArtPaths = []
    if dummy ==1:
        try:
            Art_Prompt_Chop = Text
        except:
            print('Error Trying to make original Painting')
            Art_PromptCharo =  Text

        Art_PromptCharoo = Art_Prompt_Chop
        # ArtPrompt_Mondevert = + Result
        # ArtPrompt_Mondevert = ArtFormatAdvance + Result

        for Key_Word in Key_Words:
            try:
                if Delimiter in Art_Prompt_Chop:
                    Art_Prompt_Chop = Art_Prompt_Chop.split(Delimiter)
                else:
                    Art_Prompt_Chop = Art_Prompt_Chop.replace(Key_Word + Key_Char, Delimiter)
                    Art_Prompt_Chop = Art_Prompt_Chop.replace(Key_Word, Delimiter)
                    if Delimiter in Art_Prompt_Chop:
                        Art_Prompt_Chop = Art_Prompt_Chop.split(Delimiter)
                    else:
                        Art_Prompt_Chop = Art_Prompt_Chop.replace(Key_Char, Delimiter)
                        Art_Prompt_Chop = Art_Prompt_Chop.split(Delimiter)




                # Test_Skel1 = Result.replace("Page:", "#")
                # Test_Skel1 = Test_Skel1.replace("Page", "#")

               # print('Art_PromptCharo')
                #print(Art_PromptCharo)



            except:
                print('Could not properly split out the string for multiple illustrations')

        return Art_Prompt_Chop




def Create_Art_from_String_to_list(self, Text, SavePath, FileName, Art_Style_details,Key_Words = ['Characters','Character', 'Page', 'Illustration Details'], Key_Char = ':',Delimiter = '|', crazy = .5):
    dummy = 1
    ArtPaths = []
    if dummy == 1:
        try:
            Art_Prompt_Chop = Text
        except:
            print('Error Trying to make original Painting')
            Art_PromptCharo = Text

        Art_PromptCharoo = Art_Prompt_Chop
        # ArtPrompt_Mondevert = + Result
        # ArtPrompt_Mondevert = ArtFormatAdvance + Result

        for Key_Word in Key_Words:
            try:
                if Delimiter in Art_Prompt_Chop:
                    Art_Prompt_Chop = Art_Prompt_Chop.split(Delimiter)
                else:
                    Art_Prompt_Chop = Art_Prompt_Chop.replace(Key_Word + Key_Char, Delimiter)
                    Art_Prompt_Chop = Art_Prompt_Chop.replace(Key_Word, Delimiter)

                if Delimiter in Art_Prompt_Chop:
                    Art_Prompt_Chop = Art_Prompt_Chop.split(Delimiter)
                else:
                    Art_Prompt_Chop = Art_Prompt_Chop.replace(Key_Char, Delimiter)
                    Art_Prompt_Chop = Art_Prompt_Chop.split(Delimiter)

                # Test_Skel1 = Result.replace("Page:", "#")
                # Test_Skel1 = Test_Skel1.replace("Page", "#")

            # print('Art_PromptCharo')
            # print(Art_PromptCharo)

            except:
                print('Could not properly split out the string for multiple illustrations')

        #ore sending as art I should do one more clean up of the language/prompt to make it best

        KillSwitch = 0
        KeepGoing = False
        while KeepGoing == False and KillSwitch < 6:
            try:
                Art_PromptCharoo + 'Revised Prompts: '
                Characters = []
                x = len(Art_Prompt_Chop)
                for c in range(0, x):
                    # print(c)
                    #Do the change here before setting the value or keep both
                    if len(Art_Prompt_Chop[c])>22:
                        Characters.append(Art_Prompt_Chop[c])
                        Clean_Prompt_User_GPT_Input = GPT.MondeVert.Basic_GPT_Query(self,   Line2_Role = lup.Clean_Role_after_Delimit + c , Line3_Format = lup.Clean_Format_after_Delimit,Line4_Task= lup.Clean_Task_after_Delimit ,Special = 'Use the following Text for the styles to embody/artist persona to embody###  ' + Art_Style_details+ c + '###',Line1_System_Rule = up.system_Text_Art, crazy = .5, Subject= '')
                        Characters.append(Art_Prompt_Chop[Clean_Prompt_User_GPT_Input])

                        Art_PromptCharoo +=Clean_Prompt_User_GPT_Input


                #print('Characters')
                #print(Characters)

                for c in Characters:
                    print(' ')
                    print('-------------------------------------------')
                    print('-------------------------------------------')
                    print("Character: " +  c)
                    print('-------------------------------------------')
                    # print("Character Art1: " + up.Character_Art1 + c)
                    if len(Art_Prompt_Chop[c]) > 22:
                        ArtPath2 = GPT.MondeVert.makeArt(self, Art_Prompt_Chop + c)
                        ArtPaths.append(ArtPath2)
                    #print(ArtPath2)


            except:
                print('Error Trying to create multiple paintings')
                KillSwitch += 1
            KeepGoing = True

    return ArtPaths #after its returned I append this to the folder with all relevant data